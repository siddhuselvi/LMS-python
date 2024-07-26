from django.shortcuts import render
from rest_framework.response import Response
from rest_framework import generics
from .models import appinfo,course_schedule,user_profile,course_trainer_feedback,student_request,interview_result,interview_schedule,interview_master,training_attendance_sheet, course_content_feedback, attendance_master,job_master,event_master,company_master, announcement_master, trainer_skill_map, trainer_availability, test_question_map,skill_type,login,tests_candidates_answers,tests_candidates_map,content_detail,course_candidates_map,college_admin,test_master,course_contents,question_master,candidate_master,content_master,skills_master,courses_master,topic_master,test_type,question_type,college_master,department_master,trainer_master, rules, question_paper_master, compiler_output, job_master, event_master, company_master, interview_master, interview_schedule, interview_result, student_request, training_attendance_sheet, course_trainer_feedback
from .serializers import candidateSerializerImport,UserProfileSerializer,UserSerializer,studentRequestSerializer,interviewResultSerilaizer,interviewScheduleSerializer,interviewMasterSerializer,trainerfeedbackSerializer,trainingreportsheetSerializer,loginSerializer,eventSerializer,companySerializer,jobSerializer,appinfoSerializer,tests_candidates_answerSerializer,courseScheduleSerializer, courseContentFeedbackSerializer, attendanceMasterSerializer, announcementSerializer, trainerSkillMapSerializer, trainerAvailabilitySerializer, testQuestionMapSerializer,testcandidatemapSerializers,contentdetailSerializer,coursecandidatesmapSerializer,collegeadminSerializers,coursecontentsSerializer,testsSerializers,questionsSerializer,skilltypeSerializer,contentSerializers,collegeSerializers,departmentSerializers,questiontypeSerializers,testtypeSerializers,topicSerializers,courseSerializer,skillSerializer,candidatesSerializer,trainerSerializer, testsSerializersImport, questionsSerializerImport, ruleSerializers, testsSerializersAddUpdate, questionsSerializerMasterData, loginSerializerStu, questionsSerializerCodeImport, candidatesoneSerializer, questionsPaperSerializer, questionsSerializer_IO, candidateuserSerializerImport, loginSerializerStuser, questionsSerializer_code, NonDbTestAssignSerializer, compilerSerializer, jobSerializer, eventSerializer, companySerializer, interviewMasterSerializer, interviewResultSerilaizer, interviewScheduleSerializer, studentRequestSerializer, trainingreportsheetSerializer, trainerfeedbackSerializer
from rest_framework.decorators import api_view
from .forms import QuestionForm, QuestionCodeForm
from django.http import HttpResponse 
import base64


from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from django.http.response import JsonResponse
from rest_framework.views import APIView
from rest_framework import status
import pandas as pd
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Avg, Sum, Subquery, OuterRef
from django.http import JsonResponse
from docx.oxml.ns import qn
from .models import question_master, question_paper_master

import requests
import io
import docx
import json
import re
import base64
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.db.models import Q
from thefuzz import fuzz
import urllib.parse
import math
import random
from django.db.models import Count, Value, F, DateTimeField
from django.db.models import Avg, Sum, Subquery, OuterRef, Max

from django.db.models import Count
import io,math,random,datetime,json,string,collections,itertools,functools
import sys
import builtins
from django.core.exceptions import ObjectDoesNotExist
from datetime import date
from io import StringIO
import datetime
import logging
import subprocess
import os,re,shutil
import uuid
import io,math,random,datetime,json,string,collections,itertools,functools
import sys
from django.conf import settings
from datetime import datetime
import json
from django.contrib.auth.models import User
from .models import user_profile
from .serializers import UserProfileSerializer, TestCandidateMapSerializerNeedInfo, testsSerializersAdd, testcandidatemapSerializersupdate
from django.utils.timezone import localtime
from django.utils.dateformat import format as django_format_date
from django.db.models.functions import ExtractMonth
import calendar
import pytz
# from .models import StudentAttendance
from django.utils import timezone
from django.db.models.functions import Coalesce
from django.db.models import Case, When
from django.utils.timezone import make_aware

# Get an instance of a logger
logger = logging.getLogger('app')
from django.db import transaction

# Create your views here.


class appinfogetAPIView(generics.ListAPIView):
    queryset = appinfo.objects.filter(deleted=0)

    
    serializer_class = appinfoSerializer


class appinfoAPIView(generics.CreateAPIView):
    queryset = appinfo.objects.all()
    serializer_class = appinfoSerializer

class appinfoRetrieveUpdateDestroyAPIView(generics.RetrieveUpdateDestroyAPIView):
    queryset = appinfo.objects.all()
    serializer_class = appinfoSerializer
#_____________________________Master Tables_____________________________________________#

class skilltypegetAPIView(generics.ListAPIView):
    queryset = skill_type.objects.filter(deleted=0)
    serializer_class = skilltypeSerializer

    def get(self, request, *args, **kwargs):
       
        return super().get(request, *args, **kwargs)


class skilltypecreateAPIView(generics.CreateAPIView):
    queryset = skill_type.objects.all()
    serializer_class = skilltypeSerializer

    def create(self, request, *args, **kwargs):
       
        return super().create(request, *args, **kwargs)


class skilltypeRetrieveUpdateDestroyAPIView(generics.RetrieveUpdateAPIView):
    queryset = skill_type.objects.all()
    serializer_class = skilltypeSerializer

    def update(self, request, *args, **kwargs):
       # ##logger.debug('Updating skill type with id: %s', kwargs.get('pk'))
        return super().update(request, *args, **kwargs)

    def retrieve(self, request, *args, **kwargs):
       # ##logger.debug('Retrieving skill type with id: %s', kwargs.get('pk'))
        return super().retrieve(request, *args, **kwargs)


@api_view(['PUT', 'PATCH'])
def delete_skill_type(request, pk):
   # ##logger.debug('Entering delete_skill_type function with id: %s', pk)
    try:
        skilltype = skill_type.objects.get(id=pk)
      #  ##logger.debug('Skill type found: %s', skilltype)
    except skill_type.DoesNotExist:
        logger.error('Skill type with id %s not found', pk)
        return JsonResponse("Skill type not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    skilltype.deleted = 1
    skilltype.save()
   # ##logger.debug('Skill type marked as deleted: %s', skilltype)

    return JsonResponse("Skill 'deleted' field updated successfully", safe=False)


#------------------------------Test-type-------------------------------------------------------------#


class test_type_listView(generics.ListAPIView):
    queryset = test_type.objects.filter(deleted=0)
    serializer_class = testtypeSerializers

    def get(self, request, *args, **kwargs):
       # #logger.info("Fetching test types where deleted=0")
        return super().get(request, *args, **kwargs)

class test_type_create(generics.CreateAPIView):
    queryset = test_type.objects.all()
    serializer_class = testtypeSerializers

    def post(self, request, *args, **kwargs):
      #  #logger.info("Creating a new test type")
        return super().post(request, *args, **kwargs)

class test_type_Update(generics.RetrieveUpdateAPIView):
    queryset = test_type.objects.all()
    serializer_class = testtypeSerializers

    def put(self, request, *args, **kwargs):
      #  #logger.info(f"Updating test type with id {kwargs.get('pk')}")
        return super().put(request, *args, **kwargs)

    def patch(self, request, *args, **kwargs):
       # #logger.info(f"Partially updating test type with id {kwargs.get('pk')}")
        return super().patch(request, *args, **kwargs)

@api_view(['PUT', 'PATCH'])
def delete_test_type(request, pk):
    try:
       # #logger.info(f"Attempting to mark test type with id {pk} as deleted")
        test_typevar = test_type.objects.get(id=pk)
    except test_type.DoesNotExist:
        logger.error(f"Test type with id {pk} not found")
        return JsonResponse("test_type not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    test_typevar.deleted = 1
    test_typevar.save()

   # #logger.info(f"Marked test type with id {pk} as deleted")
    return JsonResponse("test_type 'deleted' field updated successfully", safe=False)

#------------------------------------Question_type--------------------------------------#

class question_type_listView(generics.ListAPIView):
    serializer_class = questiontypeSerializers

    def get_queryset(self):
       # #logger.info("Fetching question types where deleted=0")
        queryset = question_type.objects.filter(deleted=0)
        ##logger.info("Fetched question types successfully")
        return queryset

class question_type_create(generics.CreateAPIView):
    queryset = question_type.objects.all()
    serializer_class = questiontypeSerializers

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new question type")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new question type successfully")
        return response

class question_type_Update(generics.RetrieveUpdateAPIView):
    queryset = question_type.objects.all()
    serializer_class = questiontypeSerializers

    def put(self, request, *args, **kwargs):
       # #logger.info(f"Updating question type with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
      #  #logger.info(f"Updated question type with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
       # #logger.info(f"Partially updating question type with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
       # #logger.info(f"Partially updated question type with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_question_type(request, pk):
   # #logger.info(f"Attempting to mark question type with id {pk} as deleted")
    try:
        question_typevar = question_type.objects.get(id=pk)
    except question_type.DoesNotExist:
        logger.error(f"Question type with id {pk} not found")
        return JsonResponse("question_type not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    question_typevar.deleted = 1
    question_typevar.save()

   # #logger.info(f"Marked question type with id {pk} as deleted successfully")
    return JsonResponse("question_type 'deleted' field updated successfully", safe=False)


#------------------------------college----------------------------------------#


class CollegeListView(generics.ListAPIView):
    serializer_class = collegeSerializers

    def get_queryset(self):
        
        queryset = college_master.objects.filter(deleted=0)
       
        return queryset

class collegeCreateView(generics.CreateAPIView):
    queryset = college_master.objects.all()
    serializer_class = collegeSerializers

    def post(self, request, *args, **kwargs):
        
        response = super().post(request, *args, **kwargs)
       
        return response

class collegeUpdateView(generics.RetrieveUpdateAPIView):
    queryset = college_master.objects.all()
    serializer_class = collegeSerializers

    def put(self, request, *args, **kwargs):
       
        response = super().put(request, *args, **kwargs)
       
        return response

    def patch(self, request, *args, **kwargs):
        
        response = super().patch(request, *args, **kwargs)
       
        return response

@api_view(['PUT', 'PATCH'])
def delete_college(request, pk):
   
    try:
        college = college_master.objects.get(id=pk)
    except college_master.DoesNotExist:
        logger.error(f"College with id {pk} not found")
        return Response("college not found", status=404)

    college.deleted = 1
    college.save()

   

    return Response("college 'deleted' field updated successfully")

#----------------------------------------department---------------------------------------------#

class DepartmentListView(generics.ListAPIView):
    serializer_class = departmentSerializers

    def get_queryset(self):
        
        return department_master.objects.filter(deleted=0)

class departmentCreateView(generics.CreateAPIView):
    queryset = department_master.objects.all()
    serializer_class = departmentSerializers

    def post(self, request, *args, **kwargs):
       
        response = super().post(request, *args, **kwargs)
      
        return response

class departmentUpdateView(generics.RetrieveUpdateAPIView):
    queryset = department_master.objects.all()
    serializer_class = departmentSerializers

    def put(self, request, *args, **kwargs):
       
        response = super().put(request, *args, **kwargs)
       
        return response

    def patch(self, request, *args, **kwargs):
        
        response = super().patch(request, *args, **kwargs)
       
        return response

@api_view(['PUT', 'PATCH'])
def delete_department(request, pk):
   
    try:
        department = department_master.objects.get(id=pk)
    except department_master.DoesNotExist:
        logger.error(f"department with id {pk} not found")
        return Response("department not found", status=404)

    department.deleted = 1
    department.save()

   


    return Response("department 'deleted' field updated successfully")


#__________________________topic_master___________________________________#


class topicListView(generics.ListAPIView):
    serializer_class = topicSerializers

    def get_queryset(self):
       
        return topic_master.objects.filter(deleted=0)

class topicCreateView(generics.CreateAPIView):
    queryset = topic_master.objects.all()
    serializer_class = topicSerializers

    def post(self, request, *args, **kwargs):
       
        response = super().post(request, *args, **kwargs)
       
        return response

class topicUpdateView(generics.RetrieveUpdateAPIView):
    queryset = topic_master.objects.all()
    serializer_class = topicSerializers

    def put(self, request, *args, **kwargs):
       
        response = super().put(request, *args, **kwargs)
        
        return response

    def patch(self, request, *args, **kwargs):
       
        response = super().patch(request, *args, **kwargs)
       
        return response

@api_view(['PUT', 'PATCH'])
def delete_topic(request, pk):
   
    try:
        topic = topic_master.objects.get(id=pk)
    except topic_master.DoesNotExist:
        return Response("topic not found", status=404)

    topic.deleted = 1
    topic.save()

   


    return Response("topic 'deleted' field updated successfully")
#___________________________SKILL_master_________________________________________
class skillsAPIView(generics.ListAPIView):
    queryset = skills_master.objects.all()
    serializer_class = skillSerializer

   

class skillscreateAPIView(generics.CreateAPIView):
    queryset = skills_master.objects.all()
    serializer_class = skillSerializer

    def post(self, request, *args, **kwargs):
        
        response = super().post(request, *args, **kwargs)
        
        return response

class skillsRetrieveUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = skills_master.objects.all()
    serializer_class = skillSerializer

    def put(self, request, *args, **kwargs):
       
        response = super().put(request, *args, **kwargs)
       
        return response

    def patch(self, request, *args, **kwargs):
       
        response = super().patch(request, *args, **kwargs)
       
        return response

@api_view(['PUT', 'PATCH'])
def delete_skills(request, pk):
   
    try:
        print("Entering Function..")
        skills = skills_master.objects.get(id=pk)

        print("skill: ", skills)
    except skills_master.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    skills.deleted = 1
    skills.save()

   

    print("skill: ", skills)

    return JsonResponse("skill 'deleted' field updated successfully", safe=False)


#_________________________ Courses_master______________________________________



@api_view(['GET'])
def get_course(request):
    courseset=courses_master.objects.filter(deleted=0)
   
   
    course_data=[]
    for courses in  courseset:
       
        topic_id=None
        sub_topic=None
        skill_id=None
       
        if courses.topic_id:
           topic_id=courses.topic_id.topic 
           sub_topic=courses.topic_id.sub_topic
        if courses.skill_id:
           skill_id=courses.skill_id.skill_name
          
        course_data.append({
               'id': courses.id,
            'course_name': courses.course_name,
           
            'topic_id':topic_id,
            'sub_topic':sub_topic,
            'skill_id':skill_id,
            'dtm_start': courses.dtm_start,
            'dtm_end': courses.dtm_end,
           
            'course_count': courses.course_count,
            
            'total_enrollment': courses.total_enrollment,
            'is_active':courses.is_active,
           
       
          

            })
    return Response(course_data)



class coursecreateAPIView(generics.CreateAPIView):
    queryset = courses_master.objects.all()
    serializer_class =courseSerializer

    def post(self, request, *args, **kwargs):
       
        response = super().post(request, *args, **kwargs)
       
        return response

class courseUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = courses_master.objects.all()
    serializer_class =courseSerializer

    def put(self, request, *args, **kwargs):
       
        response = super().put(request, *args, **kwargs)
       
        return response

    def patch(self, request, *args, **kwargs):
      
        response = super().patch(request, *args, **kwargs)
       
        return response

@api_view(['PUT', 'PATCH'])
def delete_course(request, pk):
   
    try:
        print("Entering Function..")
        course=courses_master.objects.get(id=pk)

        print("course: ",course)
    except courses_master.DoesNotExist:
        return JsonResponse("courses not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    course.deleted = 1
    course.save()

    

    print("course: ",course)

    return JsonResponse("course 'deleted' field updated successfully", safe=False)

#___________________________________candidate_master________________________________________

@api_view(['GET'])
def get_candidate(request):
    candidatelist=candidate_master.objects.filter(deleted=0)
   
    candidate_data=[]
    for candidate in  candidatelist:
        college_id=None
        collegeID=None
       
        department_id=None
        departmentID=None
        if candidate.college_id:
           college_id=candidate.college_id.college
           collegeID=candidate.college_id.id
      
        if candidate.department_id:
           department_id=candidate.department_id.department
           departmentID=candidate.department_id.id
       
           candidate_data.append({
              'id': candidate.id,
               'college_id': college_id,
               'college_name_id': collegeID,
              'students_name':candidate.students_name,
            'user_name':candidate.user_name,
            'registration_number':candidate.registration_number,
            'gender':candidate.gender,
           #'skill_id':candidate.skill_id,
            'email_id':candidate.email_id,
            'mobile_number':candidate.mobile_number,
            'year':candidate.year,
            'cgpa':candidate.cgpa,
            'department_id': department_id,
            'department_name_id': departmentID,
            'marks_10th':candidate.marks_10th,
            'marks_12th': candidate.marks_12th,
            'marks_semester_wise':candidate.marks_semester_wise,
            'history_of_arrears':candidate. history_of_arrears,
            'standing_arrears':candidate.standing_arrears,
            'number_of_offers':candidate.number_of_offers,
            'text': candidate.text,
            'it_of_offers':candidate.it_of_offers,
            'core_of_offers':candidate.core_of_offers



            })
    return Response(candidate_data)



@api_view(['GET'])
def get_candidate_all_OLD(request):
    try:
        candidatelist = candidate_master.objects.filter(deleted=0).select_related('college_id', 'department_id').values(
            'id', 
            'college_id__college',  # Get college name
            'students_name', 
            'user_name', 
            'registration_number', 
            'gender',
            'email_id', 
            'mobile_number', 
            'year', 
            'cgpa', 
            'department_id__department',  # Get department name
            'marks_10th',
            'marks_12th', 
            'marks_semester_wise', 
            'history_of_arrears', 
            'standing_arrears',
            'number_of_offers', 
            'text', 
            'it_of_offers', 
            'core_of_offers', 
            'skill_id',
        )
        candidate_data = list(candidatelist)
        return Response(candidate_data)
    except Exception as e:
        return Response({'error': str(e)}, status=500)


@api_view(['GET'])
def get_candidate_all(request):
    try:
        # Fetch candidates with related college and department, and prefetch skills
        candidatelist = candidate_master.objects.filter(deleted=0).select_related('college_id', 'department_id').prefetch_related('skill_id').values(
            'id', 
            'college_id__college',  # Get college name
            'students_name', 
            'user_name', 
            'registration_number', 
            'gender',
            'email_id', 
            'mobile_number', 
            'year', 
            'cgpa', 
            'department_id__department',  # Get department name
            'marks_10th',
            'marks_12th', 
            'marks_semester_wise', 
            'history_of_arrears', 
            'standing_arrears',
            'number_of_offers', 
            'text', 
            'it_of_offers', 
            'core_of_offers'
        ).distinct()

        # Create a list to store the final output
        candidate_data = []
        
        for candidate in candidatelist:
            skills = candidate_master.objects.get(id=candidate['id']).skill_id.all()
            skill_ids = [skill.id for skill in skills]
            candidate['skill_id'] = skill_ids
            candidate_data.append(candidate)
        
        return Response(candidate_data)
    except Exception as e:
        return Response({'error': str(e)}, status=500)




class candidatescreateAPIView(generics.CreateAPIView):
    queryset = candidate_master.objects.all()
    serializer_class = candidatesSerializer

    def post(self, request, *args, **kwargs):
        
        response = super().post(request, *args, **kwargs)
       
        return response

class candidates_Select_Update(generics.RetrieveUpdateAPIView):
    queryset = candidate_master.objects.all()
    serializer_class = candidatesSerializer

    def put(self, request, *args, **kwargs):
        
        response = super().put(request, *args, **kwargs)
        
        return response

    def patch(self, request, *args, **kwargs):
       
        response = super().patch(request, *args, **kwargs)
        
        return response

@api_view(['PUT', 'PATCH'])
def delete_candidates(request, pk):
    
    try:
        print("Entering Function..")
        trainees = candidate_master.objects.get(id=pk)

        print("tests: ", trainees)
    except candidate_master.DoesNotExist:
        return JsonResponse("trainees not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    trainees.deleted = 1
    trainees.save()

    

    print("tests: ", trainees)

    return JsonResponse("trainees 'deleted' field updated successfully", safe=False)

#____________________________________content_master________________________________________
@api_view(['GET'])
def get_content_OLD2007(request):
    contentlist=content_master.objects.filter(deleted=0)
   
   
    content_data=[]
    for content in  contentlist:
        topic_id=None
        question_type_id=None
        skill_type_id=None
        sub_topic=None
      
        if content.question_type_id:
           question_type_id=content.question_type_id.question_type
       

        if content.skill_type_id:
            skill_type_id=content.skill_type_id.skill_type
        
        dtm_validity_formatted = django_format_date(localtime(content.dtm_validity), 'd-m-Y h:i A')
        
        content_data.append({
            'id': content.id,
           # 'content_name': content.content_name,
            'topic': content.topic,
            'content_type':content.content_type,
            'question_type_id':question_type_id,
            'skill_type_id':skill_type_id,
            'content_url':content.content_url,
            'actual_content':content.actual_content,
           
            'status':content.status,
           # 'added_by':content.added_by,
           # 'size':content.size,
           # 'guidelines':content.guidelines,
            'sub_topic': sub_topic,
            'dtm_active_from':content.dtm_active_from,
            'dtm_validity': dtm_validity_formatted,
           # 'feedback':content.feedback,
          


            })
    return Response(content_data)


@api_view(['GET'])
def get_content(request):
    contentlist=content_master.objects.filter(deleted=0)
   
   
    content_data=[]
    for content in  contentlist:
       # topic_id=None
        question_type_id=None
        skill_type_id=None
        #sub_topic=None
      
       # if content.topic_id:
          # topic_id=content.topic_id.topic
           #sub_topic=content.topic_id.sub_topic
        if content.question_type_id:
           question_type_id=content.question_type_id.question_type
       

        if content.skill_type_id:
            skill_type_id=content.skill_type_id.skill_type
       # if content.sub_topic_id:
        #   sub_topic_id=content.sub_topic_id.sub_topic
        
        dtm_validity_formatted = django_format_date(localtime(content.dtm_validity), 'd-m-Y h:i A')
        
        content_data.append({
            'id': content.id,
           # 'content_name': content.content_name,
            'topic': content.topic,
            'content_type':content.content_type,
            'question_type_id':question_type_id,
            'skill_type_id':skill_type_id,
            'content_url':content.content_url,
            'actual_content':content.actual_content,
           
            'status':content.status,
           # 'added_by':content.added_by,
           # 'size':content.size,
           # 'guidelines':content.guidelines,
            'sub_topic': content.sub_topic,
            'dtm_active_from':content.dtm_active_from,
            'dtm_validity': dtm_validity_formatted,
           # 'feedback':content.feedback,
          


            })
    return Response(content_data)





class contentcreateAPIView(generics.CreateAPIView):
    queryset = content_master.objects.all()
    serializer_class =contentSerializers

    def post(self, request, *args, **kwargs):
       
        response = super().post(request, *args, **kwargs)
       
        return response

class contentUpdateAPIView(generics.RetrieveUpdateDestroyAPIView):
    queryset = content_master.objects.all()
    serializer_class =contentSerializers

    def put(self, request, *args, **kwargs):
        
        response = super().put(request, *args, **kwargs)
       
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating Content with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated Content with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_content(request, pk):
    ##logger.info(f"Attempting to mark Content with id {pk} as deleted")
    try:
        print("Entering Function..")
        content=content_master.objects.get(id=pk)

        print("content: ",content)
    except content_master.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    content.deleted = 1
    content.save()

    ##logger.info(f"Marked Content with id {pk} as deleted successfully")

    print("content: ",content)

    return JsonResponse("content 'deleted' field updated successfully", safe=False)
#_______________________________________trainer_master__________________________________________________________


class TrainerListAPIView(generics.ListAPIView):
    queryset = trainer_master.objects.filter(deleted=0)
    ##logger.info("Fetching Trainer where deleted=0")
    serializer_class = trainerSerializer

class TrainerCreateAPIView(generics.CreateAPIView):
    queryset = trainer_master.objects.all()
    serializer_class = trainerSerializer

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new Trainer")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new Trainer successfully")
        return response

class TrainerRetrieveUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = trainer_master.objects.all()
    serializer_class = trainerSerializer

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating Trainer with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated Trainer with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating Trainer with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated Trainer with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_trainer(request, pk):
    ##logger.info(f"Attempting to mark Trainer with id {pk} as deleted")
    try:
        print("Entering Function..")
        trainer=trainer_master.objects.get(id=pk)

        print("trainer: ",trainer)
    except trainer_master.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    trainer.deleted = 1
    trainer.save()

    ##logger.info(f"Marked Trainer with id {pk} as deleted successfully")

    print("content: ",trainer)

    return JsonResponse("trainer 'deleted' field updated successfully", safe=False)



#_____________________________________Test_________________________________________________________#



@api_view(['GET'])
def get_test(request):
    testset=test_master.objects.filter(deleted=0)
    ##logger.info("Fetching Test where deleted=0")
   
    test_data=[]
    for test in  testset:
        test_type_id = None
        test_type_categories=None
        question_type_id = None
        skill_type_id=None
       

        if test.test_type_id:
           test_type_id=test.test_type_id.test_type
           test_type_categories=test.test_type_id.test_type_categories
        if test.question_type_id:
           question_type_id=test.question_type_id.question_type
       

        if test.skill_type_id:
            skill_type_id=test.skill_type_id.skill_type
          
        test_data.append({
               'id': test.id,
            'test_name': test.test_name,
            'test_type_id': test_type_id,
            'test_type_categories':test_type_categories,
            'question_type_id': question_type_id, 
            'skill_type_id':skill_type_id,
       
            })
    return Response(test_data)

class testcreateAPIView(generics.CreateAPIView):
    queryset = test_master.objects.all()
    serializer_class = testsSerializersAdd

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new Test")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new Test successfully")
        return response

class testsRetrieveUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = test_master.objects.all()
    serializer_class = testsSerializersAddUpdate

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating Test with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated Test with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating Test with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated Test with id {kwargs.get('pk')} successfully")
        return response

#class testsRetrievedeleteAPIView(generics.DestroyAPIView):
   # queryset = tests_master.objects.all()
   # serializer_class = testsSerializers

@api_view(['PUT', 'PATCH'])
def delete_tests(request, pk):
    ##logger.info(f"Attempting to mark Test with id {pk} as deleted")
    try:
        print("Entering Function..")
        test = test_master.objects.get(id=pk)

      
    except test_master.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    test.deleted = 1
    test.save()

    ##logger.info(f"Marked Test with id {pk} as deleted successfully")

    print("skill1: ",test)

    return JsonResponse("skill 'deleted' field updated successfully", safe=False)
#______________________________________Question_master________________________________________
@api_view(['GET'])
def get_questionsOLD03(request):
    questionset=question_master.objects.filter(deleted=0)
    ##logger.info("Fetching Questions where deleted=0")
   
    question_data=[]
    for question in  questionset:
        topic_id=None
        sub_topic=None
      
        if question.topic_id:
           topic_id=question.topic_id.topic
           
           sub_topic=question.topic_id.sub_topic
           question_data.append({
              'id': question.id,
              'question_name':question.question_name,
              'question_text':question.question_text,
              'option_a':question.option_a,
              'option_b':question.option_b,
              'option_c':question.option_c,
              'option_d':question.option_d,
              'view_hint':question.view_hint,
              'mark':question.mark,
              'explain_answer':question.explain_answer,
             'topic_id':topic_id,
             'sub_topic':sub_topic,
              'answer':question.answer,
              'negative_mark':question.negative_mark,
             'input_format': question.input_format,


            })
    return Response(question_data)


@api_view(['GET'])
def get_questions(request):
    questionset=question_master.objects.filter(deleted=0)
    ##logger.info("Fetching Questions where deleted=0")
   
    question_data=[]
    for question in  questionset:
        question_id=None
        #sub_topic=None
      
        if question.question_name_id:
          question_id=question.question_name_id.question_paper_name
           
        question_data.append({
              'id': question.id,
              'question_id':question_id,
              'question_text':question.question_text,
              'option_a':question.option_a,
              'option_b':question.option_b,
              'option_c':question.option_c,
              'option_d':question.option_d,
              'view_hint':question.view_hint,
              'mark':question.mark,
              'explain_answer':question.explain_answer,
              'answer':question.answer,
              'negative_mark':question.negative_mark,
             'input_format': question.input_format,


            })
    return Response(question_data)

class questionscreateAPIView(generics.CreateAPIView):
    queryset = question_master.objects.all()
    serializer_class = questionsSerializer

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new Questions")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new Questions successfully")
        return response

class questionsRetrieveUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = question_master.objects.all()
    serializer_class = questionsSerializer 

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating Questions with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated Questions with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating Questions with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated Questions with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_question(request, pk):
    ##logger.info(f"Attempting to mark Questions with id {pk} as deleted")
    try:
        print("Entering Function..")
        question = question_master.objects.get(id=pk)

        print("skill: ", question)
    except question_master.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    question.deleted = 1
    question.save()

    ##logger.info(f"Marked Questions with id {pk} as deleted successfully")

    print("skill: ", question)

    return JsonResponse("skill 'deleted' field updated successfully", safe=False)
#_______________________________________COLLEGE_admin____________________


@api_view(['GET'])
def get_collegeadmin(request):
    adminset = college_admin.objects.filter(deleted=0)  # Accessing the skill model directly
    ##logger.info("Fetching College-admin where deleted=0") 

    skill_data = []
    for admin in adminset:
        college_id = None
        if admin.college_id:
            college_id= admin.college_id.college
            skill_data.append({
                'id': admin.id,
                'admin_name': admin.admin_name,
                'college_id': college_id,
            })
    return Response(skill_data)

class collegeadmincreateAPIView(generics.CreateAPIView):
    queryset = college_admin.objects.all()
    serializer_class = collegeadminSerializers

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new College-admin")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new college-admin successfully")
        return response

class collegeadminRetrieveUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = college_admin.objects.all()
    serializer_class = collegeadminSerializers

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating College-admin with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated College-admin with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating College-admin with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated College-admin with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_college_admin(request, pk):
    ##logger.info(f"Attempting to mark College-admin with id {pk} as deleted")

    try:
        print("Entering Function..")
        college_ad = college_admin.objects.get(id=pk)

       
    except college_admin.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    college_ad.deleted = 1
    college_ad.save()

    ##logger.info(f"Marked College-admin with id {pk} as deleted successfully")

   

    return JsonResponse("collegeadmin 'deleted' field updated successfully", safe=False)


#__________________________course_contents____________________________________#


@api_view(['GET'])
def get_course_content(request):
    coursecntent=course_contents.objects.filter(deleted=0)
    ##logger.info("Fetching course-content where deleted=0")
   
    content_course_data=[]
    for contents in  coursecntent:
        
        course_id=None
        content_id=None
        topic_id=None
        sub_topic=None
        content_name = None
       
        if contents.content_id:
           content_id=contents.content_id.id
           content_name = contents.content_id.content_name
       
        if contents.course_id:
           course_id=contents.course_id.course_name
        
        if contents.topic_id:
           topic_id=contents.topic_id.topic
           sub_topic=contents.topic_id.sub_topic
       
      
       
        content_course_data.append({
            'id':contents.id,
            'content_name': content_name,
            
            'content_id':content_id,
            'course_id':course_id,
            'topic':topic_id,
            'sub_topic':sub_topic,
           
              })
    return Response(content_course_data)



class coursecontentscreateAPIView(generics.CreateAPIView):
    queryset = course_contents.objects.all()
    serializer_class =coursecontentsSerializer

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new course-content")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new course-content successfully")
        return response

class coursecontentsUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = course_contents.objects.all()
    serializer_class =coursecontentsSerializer

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating course-content with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated course-content with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating course-content with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated course-content with id {kwargs.get('pk')} successfully")
        return response
    
@api_view(['PUT', 'PATCH'])
def delete_course_contents(request, pk):
    ##logger.info(f"Attempting to mark course-content with id {pk} as deleted")

    try:
        print("Entering Function..")
        contents=course_contents.objects.get(id=pk)

       
    except course_contents.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    contents.deleted = 1
    contents.save()

    ##logger.info(f"Marked course-content with id {pk} as deleted successfully")

   

    return JsonResponse("tests_result 'deleted' field updated successfully", safe=False)


#_______________________course_candidates_map______________________________



@api_view(['GET'])
def get_course_candidates_map(request):
    coursecan=course_candidates_map.objects.filter(deleted=0)
    ##logger.info("Fetching course_candidates_map where deleted=0")
   
    course_map_data=[]
    for content in  coursecan:
         # Initialize question_type variable
        course_id=None
        student_id=None
        college_id=None
        if content.college_id:
            college_id= content.college_id.college
       
        if content.student_id:
           student_id=content.student_id.students_name
       
        if content.course_id:
           course_id=content.course_id.course_name
      
       
        course_map_data.append({
               'id':content.id,
               
               
             # Add question_type field to the response
            'student_id':student_id,
            'course_id':course_id,
            'collge_id':college_id,
            'dt_enrolled':content.dt_enrolled,
            'dt_validity':content.dt_validity,
            'status':content.status,
           
              })
    return Response(course_map_data)



class coursecandidatesmapcreateAPIView(generics.CreateAPIView):
    queryset = course_candidates_map.objects.all()
    serializer_class =coursecandidatesmapSerializer

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new course_candidates_map")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new course_candidates_map successfully")
        return response

class coursecandidatesmapUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = course_candidates_map.objects.all()
    serializer_class =coursecandidatesmapSerializer

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating course_candidates_map with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated course_candidates_map with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating course_candidates_map with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated course_candidates_map with id {kwargs.get('pk')} successfully")
        return response


@api_view(['PUT', 'PATCH'])
def delete_course_can(request, pk):
    ##logger.info(f"Attempting to mark course_candidates_map with id {pk} as deleted")

    try:
        print("Entering Function..")
        contnt=course_candidates_map.objects.get(id=pk)

       
    except course_candidates_map.DoesNotExist:
        return JsonResponse("candidate map for course not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    contnt.deleted = 1
    contnt.save()

    ##logger.info(f"Marked course_candidates_map with id {pk} as deleted successfully")

   

    return JsonResponse("candidate map 'deleted' field updated successfully", safe=False)


#_________________________________content_detail_______________________________________________________________#

@api_view(['GET'])
def get_content_detail(request):
    cntent=content_detail.objects.filter(deleted=0)
    ##logger.info("Fetching content_detail where deleted=0")
   
    contents_data=[]
    for contentsdetail in  cntent:
         # Initialize question_type variable
       
        content_id=None
        
       
        if contentsdetail.content_id:
           content_id=contentsdetail.content_id.content_url
       
       
       
        contents_data.append({
               'id':contentsdetail.id,
             # Add question_type field to the response
            'content_id':content_id,
           
            'actual_content':contentsdetail.actual_content,
            'status':contentsdetail.status,
            'content_url':contentsdetail.content_url
           
              })
    return Response(contents_data)



class contentdetailcreateAPIView(generics.CreateAPIView):
    queryset = content_detail.objects.all()
    serializer_class =contentdetailSerializer

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new content_detail")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new content_detail successfully")
        return response

class coursecontentdetailUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = content_detail.objects.all()
    serializer_class =contentdetailSerializer

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating content_detail with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated content_detail with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating content_detail with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated content_detail with id {kwargs.get('pk')} successfully")
        return response


@api_view(['PUT', 'PATCH'])
def delete_content_detail(request, pk):
    ##logger.info(f"Attempting to mark content_detail with id {pk} as deleted")

    try:
        print("Entering Function..")
        contentdtil=content_detail.objects.get(id=pk)

       
    except content_detail.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    contentdtil.deleted = 1
    contentdtil.save()

    ##logger.info(f"Marked content_detail with id {pk} as deleted successfully")

   

    return JsonResponse("tests_result 'deleted' field updated successfully", safe=False)



#____________________________test_candidate-map________________________________#
@api_view(['GET'])
def get_tests_candidates_map_OLD(request):
    tests_candidates=tests_candidates_map.objects.filter(deleted=0)
    ##logger.info('Fetching test-candidate-map')
   
    test_candidate_map_data=[]
    for testing in  tests_candidates:
        test_id = None
        test_name = None
        question_id= None  # Initialize question_type variable
        topic_name=None
        student_id=None
        student_name=None
        user_name=None
        college_id=None
        department_id=None
        department_name=None
        test_type=None
        rules_id=None
        instruction=None
        skill_type=None

        if testing.rules_id:
           rules_id=testing.rules_id.rule_name
           instruction=testing.rules_id.instruction
        if testing.test_id:
           test_id=testing.test_id.id
           test_name=testing.test_id.test_name
           skill_type=skilltypeSerializer(testing.test_id.skill_type_id).data
           skill_type_value = skill_type.get('skill_type')
           test_type = testtypeSerializers(testing.test_id.test_type_id).data
           test_type_value = test_type.get('test_type')

        if testing.department_id:
           department_id=testing.department_id.department
      
        if testing.question_id:
           question_id=testing.question_id.id
           question_name=testing.question_id.question_name_id
           topic_name = topicSerializers(testing.question_id.topic_id).data
           topic_name_val = topic_name.get('topic')


        if testing.student_id:
           student_id=testing.student_id.id
           student_name=testing.student_id.students_name
           user_name=testing.student_id.user_name
        if testing.college_id:
            college_id= testing.college_id.college

        test_candidate_map_data.append({
            'id': testing.id,
            'test_id':test_id,
            'test_name':test_name,
            'test_type': test_type_value,
            'college_id':college_id,
            'department_id':department_id,
            'question_id':question_id, 
            'question_name':question_name, 
            'student_id':student_id,
            'student_name':student_name,
            'user_name':user_name,
            'dtm_start': testing.dtm_start,
            'dtm_end': testing.dtm_end,
           'skill_type_id':skill_type_value,
            'attempt_count':testing.attempt_count,
            'is_actual_test':testing.is_actual_test,
            'is_active': testing.is_active,
            'duration':testing.duration,
            'duration_type':testing.duration_type,
            'year':testing.year,
            'rules':rules_id,
            'instruction':instruction,
            'need_candidate_info':testing.need_candidate_info,
            'topic': topic_name_val,
            'total_score': testing.total_score
       
       
          

            })
    return Response(test_candidate_map_data)



@api_view(['GET'])
def get_tests_candidates_map(request):
    tests_candidates=tests_candidates_map.objects.filter(deleted=0)
    ##logger.info('Fetching test-candidate-map')
   
    test_candidate_map_data=[]
    for testing in  tests_candidates:
        question_id= None  # Initialize question_type variable
        topic_name=None
        student_id=None
        student_name=None
        user_name=None
        college_id=None
        college_id_id=None
        department_id=None
        department_id_id=None
        department_name=None
        test_type=None
        rules_id=None
        rules_id_id = None
        instruction=None
        skill_type=None
        question_paper_name=None
        test_type=None

        if testing.rules_id:
           rules_id_id=testing.rules_id.id
           rules_id=testing.rules_id.rule_name
           instruction=testing.rules_id.instruction
        
        if testing.department_id:
           department_id=testing.department_id.department
           department_id_id=testing.department_id.id
      
        if testing.question_id:
           question_id=testing.question_id.id
           question_paper_name=testing.question_id.question_paper_name
           test_type=testing.question_id.test_type


        if testing.student_id:
           student_id=testing.student_id.id
           student_name=testing.student_id.students_name
           user_name=testing.student_id.user_name
        if testing.college_id:
            college_id= testing.college_id.college
            college_id_id= testing.college_id.id

        test_candidate_map_data.append({
            'id': testing.id,
            'test_name':testing.test_name,
            'college_id_id':college_id_id,
            'college_id':college_id,
            'department_id_id':department_id_id,
            'department_id':department_id,
            'question_id':question_id, 
            'question_paper_name':question_paper_name,
            'test_type': test_type,
            'student_id':student_id,
            'student_name':student_name,
            'user_name':user_name,
            'dtm_start': testing.dtm_start,
            'dtm_end': testing.dtm_end,
           
            'attempt_count':testing.attempt_count,
            'is_actual_test':testing.is_actual_test,
            'is_active': testing.is_active,
            'duration':testing.duration,
            'duration_type':testing.duration_type,
            'year':testing.year,
            'rules_id':rules_id_id,
            'rules':rules_id,
            'instruction':instruction,
            'need_candidate_info':testing.need_candidate_info,
            'total_score': testing.total_score,
            'avg_mark': testing.avg_mark,
            'dtm_created': testing.dtm_created,
       
       
          

            })
    return Response(test_candidate_map_data)



@api_view(['GET'])
def get_tests_candidates_map_Update_ID(request, id):
    tests_candidates=tests_candidates_map.objects.filter(deleted=0, id=id)
    ##logger.info('Fetching test-candidate-map')
   
    test_candidate_map_data=[]
    for testing in  tests_candidates:
        question_id= None  # Initialize question_type variable
        topic_name=None
        student_id=None
        student_name=None
        user_name=None
        college_id=None
        department_id=None
        department_name=None
        test_type=None
        rules_id=None
        rules_name=None
        instruction=None
        skill_type=None
        question_paper_name=None
        test_type=None

        if testing.rules_id:
           rules_id=testing.rules_id.id
           rules_name=testing.rules_id.rule_name
           instruction=testing.rules_id.instruction
        
        if testing.department_id:
           department_id=testing.department_id.department
      
        if testing.question_id:
           question_id=testing.question_id.id
           question_paper_name=testing.question_id.question_paper_name
           test_type=testing.question_id.test_type


        if testing.student_id:
           student_id=testing.student_id.id
           student_name=testing.student_id.students_name
           user_name=testing.student_id.user_name
        if testing.college_id:
            college_id= testing.college_id.college

        
        dtm_start_formatted = django_format_date(localtime(testing.dtm_start), 'd-m-Y h:i A')
        dtm_end_formatted = django_format_date(localtime(testing.dtm_end), 'd-m-Y h:i A')


        test_candidate_map_data.append({
            'id': testing.id,
            'test_name':testing.test_name,
            'college_id':college_id,
            'department_id':department_id,
            'question_id':question_id, 
            'question_paper_name':question_paper_name,
            'test_type': test_type,
            'student_id':student_id,
            'student_name':student_name,
            'user_name':user_name,
            'dtm_start': dtm_start_formatted,
            'dtm_end': dtm_end_formatted,
           
            'attempt_count':testing.attempt_count,
            'is_actual_test':testing.is_actual_test,
            'is_active': testing.is_active,
            'duration':testing.duration,
            'duration_type':testing.duration_type,
            'year':testing.year,
            'rules_id':rules_id,
            'rules_name': rules_name,
            'instruction':instruction,
            'need_candidate_info':testing.need_candidate_info,
            'total_score': testing.total_score,
            'avg_mark': testing.avg_mark
       
       
          

            })
    return Response(test_candidate_map_data)



@api_view(['GET'])
def get_tests_candidates_map_Reports(request):
    print('Entering reports functiom......')

    tests_candidates=tests_candidates_map.objects.filter(deleted=0)
    print('filtered candidates: ', tests_candidates)

    ##logger.info('Fetching test-candidate-map')
   
    test_candidate_map_data=[]
    for testing in  tests_candidates:
        question_name= None 
        student_id=None
        student_name=None
        user_name=None
        college_id=None
        department_id=None
        rules_id=None
        instruction=None       
        test_type=None
        email_id=None
        mobile_number=None
        gender=None
        registration_number=None

        if testing.rules_id:
           rules_id=testing.rules_id.rule_name
           instruction=testing.rules_id.instruction
        
        if testing.department_id:
           department_id=testing.department_id.department
      
       

        if testing.question_id:
           question_name=testing.question_id.question_paper_name
          # print("id",question_name)
          # question_id_value=question_name.get('question_paper_name')
        if testing.student_id:
           student_id=testing.student_id.id
           student_name=testing.student_id.students_name
           user_name=testing.student_id.user_name
           email_id=testing.student_id.email_id
           mobile_number=testing.student_id.mobile_number
           gender=testing.student_id.gender
           registration_number=testing.student_id.registration_number
        if testing.college_id:
            college_id= testing.college_id.college

        test_candidate_map_data.append({
            'id': testing.id,
            'test_name':testing.test_name,
            'college_id':college_id,
            'department_id':department_id,
            #'question_id':question_id, 
            'question_id':question_name,
           # 'test_type': test_type,
            'student_id':student_id,
            'registration_number':registration_number,
            'email_id':email_id,
            'mobile_number':mobile_number,
            'gender':gender,
            'student_name':student_name,
            'user_name':user_name,
            'dtm_start': testing.dtm_start,
            'dtm_end': testing.dtm_end,
           
            'attempt_count':testing.attempt_count,
            'is_actual_test':testing.is_actual_test,
            'is_active': testing.is_active,
            'duration':testing.duration,
            'duration_type':testing.duration_type,
            'year':testing.year,
            'rules':rules_id,
            'instruction':instruction,
            'need_candidate_info':testing.need_candidate_info,
            #'topic': topic_name_val,
            'total_score': testing.total_score,
            'avg_mark': testing.avg_mark
       
       
          

            })
    print('test_candidate_map_data: ', test_candidate_map_data)
    return Response(test_candidate_map_data)






@api_view(['GET'])
def get_tests_Reports(request):
    tests_candidates = tests_candidates_map.objects.filter(deleted=0)
    ##logger.info('Fetching test-candidate-map')
    
    test_candidate_map_data = []
    for testing in tests_candidates:
        question_name = None
        student_id = None
        student_name = None
        user_name = None
        college_id = None
        department_id = None
        test_type = None
        rules_id = None
        instruction = None
        email_id = None
        mobile_number = None
        gender = None
        registration_number = None

        if testing.rules_id:
            rules_id = testing.rules_id.rule_name
            instruction = testing.rules_id.instruction

        if testing.department_id:
            department_id = testing.department_id.department

        if testing.question_id:
            question_name = testing.question_id.question_paper_name

        if testing.student_id:
            student_id = testing.student_id.id
            student_name = testing.student_id.students_name
            user_name = testing.student_id.user_name
            email_id = testing.student_id.email_id
            mobile_number = testing.student_id.mobile_number
            gender = testing.student_id.gender
            registration_number = testing.student_id.registration_number

        if testing.college_id:
            college_id = testing.college_id.college

        # Format dates
        dtm_start_formatted = django_format_date(localtime(testing.dtm_start), 'd-m-Y h:i A')
        dtm_end_formatted = django_format_date(localtime(testing.dtm_end), 'd-m-Y h:i A')

        # Check if total_score is null and set to 'AA' if so
        total_score_display = testing.total_score if testing.total_score is not None else 'AA'

        test_candidate_map_data.append({
            'id': testing.id,
            'test_name': testing.test_name,
            'college_id': college_id,
            'department_id': department_id,
            'question_id': question_name,
            'student_id': student_id,
            'registration_number': registration_number,
            'email_id': email_id,
            'mobile_number': mobile_number,
            'gender': gender,
            'student_name': student_name,
            'user_name': user_name,
            'dtm_start': dtm_start_formatted,
            'dtm_end': dtm_end_formatted,
            'attempt_count': testing.attempt_count,
            'is_actual_test': testing.is_actual_test,
            'is_active': testing.is_active,
            'duration': testing.duration,
            'duration_type': testing.duration_type,
            'year': testing.year,
            'rules': rules_id,
            'instruction': instruction,
            'need_candidate_info': testing.need_candidate_info,
            'total_score': total_score_display,
            'avg_mark': testing.avg_mark
        })
    return Response(test_candidate_map_data)




class testcandidatemapcreateAPIView(generics.CreateAPIView):
    queryset = tests_candidates_map.objects.all()
    serializer_class =testcandidatemapSerializers

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new test-candidate-map")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new test-candidate-map successfully")
        return response



class testcandidatemapUpdateAPIView_OLD(generics.UpdateAPIView):
    serializer_class = testcandidatemapSerializers

    def get_queryset(self):
        test_name = self.request.data.get('testName')
        print(f"Querying for test_name: {test_name}")  # Add logging
        return tests_candidates_map.objects.filter(test_name=test_name)

    def put(self, request, *args, **kwargs):
        print(f"Request data: {request.data}")  # Log the entire request data
        queryset = self.get_queryset()
        if not queryset.exists():
            print("Test not found")  # Add logging
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)
        
        response_data = []
        for instance in queryset:
            serializer = self.get_serializer(instance, data=request.data, partial=False)
            serializer.is_valid(raise_exception=True)
            self.perform_update(serializer)
            response_data.append(serializer.data)
        
        print("Update successful")  # Add logging
        return Response(response_data, status=status.HTTP_200_OK)

    def patch(self, request, *args, **kwargs):
        print(f"Request data: {request.data}")  # Log the entire request data
        queryset = self.get_queryset()
        if not queryset.exists():
            print("Test not found")  # Add logging
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)
        
        response_data = []
        for instance in queryset:
            serializer = self.get_serializer(instance, data=request.data, partial=True)
            serializer.is_valid(raise_exception=True)
            self.perform_update(serializer)
            response_data.append(serializer.data)
        
        print("Partial update successful")  # Add logging
        return Response(response_data, status=status.HTTP_200_OK)

    def perform_update(self, serializer):
        # Parse the datetime fields to ensure correct timezone handling
        dtm_start = self.request.data.get('dtm_start')
        dtm_end = self.request.data.get('dtm_end')

        if dtm_start:
            dtm_start = datetime.fromisoformat(dtm_start)
            serializer.validated_data['dtm_start'] = dtm_start

        if dtm_end:
            dtm_end = datetime.fromisoformat(dtm_end)
            serializer.validated_data['dtm_end'] = dtm_end

        serializer.save()



class testcandidatemapUpdateAPIView(generics.UpdateAPIView):
    serializer_class = testcandidatemapSerializersupdate

    def get_queryset(self):
        test_name = self.request.data.get('testName')
        print(f"Querying for test_name: {test_name}")
        return tests_candidates_map.objects.filter(test_name=test_name)

    def update(self, request, *args, **kwargs):
        current_date_time = datetime.now()
        print(f"Request data: {request.data}")
        queryset = self.get_queryset()
        print('QuerySet: ', queryset)
        if not queryset.exists():
            print("Test not found")
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)

        response_data = []
        for instance in queryset:
            request.data['dtm_created'] = current_date_time
            serializer = self.get_serializer(instance, data=request.data, partial=False)
            try:
                serializer.is_valid(raise_exception=True)
                print(f"Validated data before saving: {serializer.validated_data}")
                serializer.save()
                response_data.append(serializer.data)
            except Exception as e:
                print(f"Validation error: {e}")
                return Response({"detail": str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        print("Update successful")
        return Response(response_data, status=status.HTTP_200_OK)

    def partial_update(self, request, *args, **kwargs):
        current_date_time = datetime.now()
        print(f"Request data: {request.data}")
        queryset = self.get_queryset()
        if not queryset.exists():
            print("Test not found")
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)

        response_data = []
        for instance in queryset:
            request.data['dtm_created'] = current_date_time
            serializer = self.get_serializer(instance, data=request.data, partial=True)
            try:
                serializer.is_valid(raise_exception=True)
                print(f"Validated data before saving: {serializer.validated_data}")
                serializer.save()
                response_data.append(serializer.data)
            except Exception as e:
                print(f"Validation error: {e}")
                return Response({"detail": str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        print("Partial update successful")
        return Response(response_data, status=status.HTTP_200_OK)







class test_master_UpdateAPIView(generics.UpdateAPIView):
    serializer_class = testsSerializersAddUpdate

    def get_queryset(self):
        test_name = self.request.data.get('testName')
        print(f"Querying for test_name: {test_name}")  # Add logging
        return test_master.objects.filter(test_name=test_name)  # Use correct model

    def put(self, request, *args, **kwargs):
        print(f"Request data: {request.data}")  # Log the entire request data
        queryset = self.get_queryset()
        if not queryset.exists():
            print("Test not found")  # Add logging
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)
        
        response_data = []
        for instance in queryset:
            serializer = self.get_serializer(instance, data=request.data, partial=False)
            serializer.is_valid(raise_exception=True)
            self.perform_update(serializer)
            response_data.append(serializer.data)
        
        print("Update successful")  # Add logging
        return Response(response_data, status=status.HTTP_200_OK)

    def patch(self, request, *args, **kwargs):
        print(f"Request data: {request.data}")  # Log the entire request data
        queryset = self.get_queryset()
        if not queryset.exists():
            print("Test not found")  # Add logging
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)
        
        response_data = []
        for instance in queryset:
            serializer = self.get_serializer(instance, data=request.data, partial=True)
            serializer.is_valid(raise_exception=True)
            self.perform_update(serializer)
            response_data.append(serializer.data)
        
        print("Partial update successful")  # Add logging
        return Response(response_data, status=status.HTTP_200_OK)




@api_view(['PUT', 'PATCH'])
def delete_testcandidatemap(request, test_name):
    try:
        print("Entering Function..")
        tests_candidates = tests_candidates_map.objects.filter(test_name=test_name)
        if not tests_candidates.exists():
            return JsonResponse({"error": "tests not found"}, status=404)

        for candidate in tests_candidates:
            candidate.deleted = 1
            candidate.save()
            print("Updated candidate: ", candidate)

        return JsonResponse("tests_candidates_map 'deleted' field updated successfully", safe=False)
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)


@api_view(['PUT', 'PATCH'])
def update_testcandidatemap_is_active(request, pk):
    try:
        print("Entering Function..")
        tests_candidates=tests_candidates_map.objects.get(id=pk)
        ##logger.info('Fetching test-candidate-map ')

        print("tests_candidates: ",tests_candidates)
    except tests_candidates_map.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    tests_candidates.is_active = True
    tests_candidates.save()
    ##logger.info('test-candidates updated')
    print("tests_candidates: ",tests_candidates)

    return JsonResponse("tests_candidates_map 'deleted' field updated successfully", safe=False)



#__________________________test_candidate_answer____________________________
@api_view(['GET'])
def get_tests_candidates_answer(request):
    tests_candidates_ans=tests_candidates_answers.objects.filter(deleted=0).order_by('-id')
    ##logger.info("Fetching test-candidate-answer where deleted=0")

    test_candidate_answer_data=[]
    for testans in  tests_candidates_ans:
         # Initialize question_type variable
        student_id=None
        student_name=None
        question_id=None
        question_name=None
    
        if testans.student_id:
           student_id=testans.student_id.id
           student_name=testans.student_id.students_name
        
        if testans.question_id:
           question_id=testans.question_id.id
           question_name=testans.question_id.question_text
       
        test_candidate_answer_data.append({
               'id': testans.id,
             # Add question_type field to the response
            'student_id': student_id,
            'student_name':student_name,
            'question_id': question_id,
            'question_name':question_name,
            'test_name':testans.test_name,
            'answer': testans.answer,
            'result':testans.result,
            'dtm_start': testans.dtm_start,
            'dtm_end': testans.dtm_end,
            'submission_compile_code': testans.submission_compile_code,
            'compile_code_editor': testans.compile_code_editor,
            })
    return Response(test_candidate_answer_data)



@api_view(['POST'])
def test_candidates_answer_viewOLD(request, ques_id, ans, code, format=None):
    print("Parameters: ", ques_id, ans, code)
    
    try:
        question = question_master.objects.get(id=ques_id)
    except question_master.DoesNotExist:
        return Response({'error': 'Question not found'}, status=status.HTTP_404_NOT_FOUND)

    if question.answer == ans:
        result = question.mark
    else:
        similarity_score = fuzz.ratio(question.explain_answer, code)
        result = (similarity_score/100) * question.mark

    print("Result: ", result)
    
    test_candidate_answer_data = {
        'test_name': request.data.get('test_name'),
        'question_id': request.data.get('question_id'),
        'student_id': request.data.get('student_id'),
        'submission_compile_code': ans,
        'compile_code_editor': code,
        'result': result,
        'dtm_start': request.data.get('dtm_start'),
        'dtm_end': request.data.get('dtm_end'),
    }

    serializer = tests_candidates_answerSerializer(data=test_candidate_answer_data)
    print("Serializer: ", serializer)
    
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    else:
        print(serializer.errors)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)



@csrf_exempt
@api_view(['POST'])
def test_candidates_answer_view(request, format=None):
    ques_id =  request.data.get('question_id')
    ans =  request.data.get('ans')
    code =  request.data.get('code')

    # code = urllib.parse.unquote(code)
    print("Parameters: ", ques_id, ans, code)
    
    try:
        question = question_master.objects.get(id=ques_id)
    except question_master.DoesNotExist:
        return JsonResponse({'error': 'Question not found'}, status=404)

    if question.answer == ans:
        result = question.mark
    else:
        # Handling multi-line strings
        similarity_score = fuzz.ratio(question.explain_answer.strip(), code.strip())
        result = round((similarity_score / 100) * question.mark)


    # Round the result to the nearest integer
    result = round(result)
    print("Result: ", result)
    
    test_candidate_answer_data = {
        'test_name': request.data.get('test_name'),
        'question_id': ques_id,
        'student_id': request.data.get('student_id'),
        'submission_compile_code': ans,
        'compile_code_editor': code,
        'result': result,
        'dtm_start': request.data.get('dtm_start'),
        'dtm_end': request.data.get('dtm_end'),
    }

    serializer = tests_candidates_answerSerializer(data=test_candidate_answer_data)
    print("Serializer: ", serializer)
    
    if serializer.is_valid():
        serializer.save()
        print('Test Answer is added Successfully')
        return JsonResponse(serializer.data, status=201)
    else:
        print(serializer.errors)
        return JsonResponse(serializer.errors, status=400)




@csrf_exempt
@api_view(['POST'])
def test_candidates_answer_view_Submit(request, format=None):
    ques_id =  request.data.get('question_id')
    ans =  request.data.get('ans')
    code =  request.data.get('code')

    # code = urllib.parse.unquote(code)
    print("Parameters: ", ques_id, ans, code)
    
    try:
        question = question_master.objects.get(id=ques_id)
    except question_master.DoesNotExist:
        return JsonResponse({'error': 'Question not found'}, status=404)

    if question.answer == ans:
        result = question.mark
    else:
        # Handling multi-line strings
        similarity_score = fuzz.ratio(question.explain_answer.strip(), code.strip())
        result = round((similarity_score / 100) * question.mark)


    # Round the result to the nearest integer
    result = round(result)
    print("Result: ", result)
    
    test_candidate_answer_data = {
        'test_name': request.data.get('test_name'),
        'question_id': ques_id,
        'student_id': request.data.get('student_id'),
        'answer': ans,
        'compile_code_editor': code,
        'result': result,
        'dtm_start': request.data.get('dtm_start'),
        'dtm_end': request.data.get('dtm_end'),
    }

    serializer = tests_candidates_answerSerializer(data=test_candidate_answer_data)
    print("Serializer: ", serializer)
    
    if serializer.is_valid():
        serializer.save()
        print('Test Answer is added Successfully')
        return JsonResponse(serializer.data, status=201)
    else:
        print(serializer.errors)
        return JsonResponse(serializer.errors, status=400)




class testcandidateanscreateAPIView(generics.CreateAPIView):
    queryset = tests_candidates_answers.objects.all()
    serializer_class =tests_candidates_answerSerializer

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new test-candidate-answer")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new test-candidate-answer successfully")
        return response

class testcandidateansUpdateAPIView(generics.RetrieveUpdateDestroyAPIView):
    queryset = tests_candidates_answers.objects.all()
    serializer_class =tests_candidates_answerSerializer

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating test-candidate-answer with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated test-candidate-answer with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating test-candidate-answer with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated test-candidate-answer with id {kwargs.get('pk')} successfully")
        return response



@api_view(['PUT', 'PATCH'])
def delete_testcandidateanswer(request, pk):
    ##logger.info(f"Attempting to mark test-candidate-answer with id {pk} as deleted")

    try:
        print("Entering Function..")
        tests_candidates_ans=tests_candidates_answers.objects.get(id=pk)

        print("tests_candidates_ans: ",tests_candidates_ans)
    except tests_candidates_answers.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    tests_candidates_ans.deleted = 1
    tests_candidates_ans.save()

    ##logger.info(f"Marked test-candidate-answer with id {pk} as deleted successfully")

    print("tests_candidates: ",tests_candidates_ans)

    return JsonResponse("tests_candidates_ans 'deleted' field updated successfully", safe=False)



#-----------------------------Login-----------------------------------------#


class login_create(generics.CreateAPIView):
    queryset = login.objects.all()
    serializer_class = loginSerializer

    def create(self, request, *args, **kwargs):
        ##logger.debug('login_create view called')
        
        try:
            response = super().create(request, *args, **kwargs)
            #logger.debug(f'Login created with data: {request.data}')
            #logger.debug(f'Response: {response.data}')
            return response
        except Exception as e:
            logger.error(f'Error occurred in login_create view: {e}')
            return Response({'error': 'An error occurred while creating login'}, status=500)


@api_view(['GET'])
def get_login_OLD0506(request):
    #logger.debug('get_login function called')
    
    try:
        login_set = login.objects.select_related('college_id').all()
        #logger.debug(f'Fetched {login_set.count()} login records')
        
        login_data = []
        for logins in login_set:
            login_data.append({
                'id': logins.id,
                'email_id': logins.email_id,
                'user_name': logins.user_name,
                'password': logins.password,
                'role': logins.role,
                'college_id': logins.college_id.id,
                'college_name': logins.college_id.college
            })
        
        #logger.debug('Login data successfully serialized')
        return Response(login_data)
    except Exception as e:
        logger.error(f'Error occurred in get_login function: {e}')
        return Response({'error': 'An error occurred'}, status=500)


@api_view(['GET'])
def get_login(request):
    #logger.debug('get_login function called')
    
    try:
        login_set = login.objects.select_related('college_id').all()
        #logger.debug(f'Fetched {login_set.count()} login records')
        
        login_data = []
        for logins in login_set:
            college_id = None
            college_name = None
            if logins.college_id:
                college_id = logins.college_id.id
                college_name = logins.college_id.college
                
            login_data.append({
                'id': logins.id,
                'email_id': logins.email_id,
                'user_name': logins.user_name,
                'password': logins.password,
                'role': logins.role,
                'college_id': college_id,
                'college_name': college_name
            })
        
        ##logger.debug('Login data successfully serialized')
        return Response(login_data)
    except Exception as e:
        logger.error(f'Error occurred in get_login function: {e}')
        return Response({'error': 'An error occurred'}, status=500)


class update_login(generics.RetrieveUpdateAPIView):
    queryset = login.objects.all()
    serializer_class = loginSerializer



@api_view(['PUT', 'PATCH'])
def delete_login(request, pk):
    try:
        print("Entering Function..")
        logins = login.objects.get(id=pk)

        print("skilltyp: ", logins)
    except login.DoesNotExist:
        return JsonResponse("login not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    logins.deleted = 1
    logins.save()
    print("logins: ", logins)

    return JsonResponse("logins 'deleted' field updated successfully", safe=False)


#-----------------------------Course Schedule-----------------------------------------#

class add_course_schedule(generics.CreateAPIView):
    queryset = course_schedule.objects.all()
    serializer_class = courseScheduleSerializer


@api_view(['GET'])
def get_course_schedule_OLD(request):
    cs_set = course_schedule.objects.select_related('college_id','topic_id' ,'department_id', 'trainer_id').filter(deleted=0)

    course_schedule_data=[]
    for cs in  cs_set:
           course_schedule_data.append({
              'id': cs.id,
              'college_id': cs.college_id.college,
              'department_id': cs.department_id.department,
              'topic_id':cs.topic_id.topic,
              'sub_topic':cs.topic_id.sub_topic,
              'year':cs.year,
              'trainer_id': cs.trainer_id.trainer_name,
              'dtm_start_student': cs.dtm_start_student,
              'dtm_end_student': cs.dtm_end_student,
              'dtm_start_trainer': cs.dtm_start_trainer,
              'dtm_end_trainer': cs.dtm_end_trainer,
            
              'dtm_of_training': cs.dtm_of_training,
            
            })
    return Response(course_schedule_data)


@api_view(['GET'])
def get_course_schedule(request):
    candidatelist=course_schedule.objects.filter(deleted=0)
   
    candidate_data=[]
    for candidate in  candidatelist:
        college_id=None
        collegeID=None
        student_id=None
        department_id=None
        departmentID=None
        topic_id=None
        sub_topic=None
        trainer_id=None
        if candidate.college_id:
           college_id=candidate.college_id.college
           collegeID=candidate.college_id.id
        if candidate.student_id:
           student_id=candidate.student_id.students_name
      
        if candidate.department_id:
           department_id=candidate.department_id.department
           departmentID=candidate.department_id.id
        if candidate.topic_id:
           topic_id=candidate.topic_id.topic
           sub_topic=candidate.topic_id.sub_topic
      
        if candidate.trainer_id:
           trainer_id=candidate.trainer_id.trainer_name
          
      
           candidate_data.append({
              'id': candidate.id,
               'college_id': college_id,
               'college_name_id': collegeID,
              'student_id':student_id,
              'year':candidate.year,
              'department_id': department_id,
            'department_name_id': departmentID,
            'sub_topic':sub_topic,
              'year':candidate.year,
               'topic_id':topic_id,
             
              'trainer_id':trainer_id,
              'dtm_start_student': candidate.dtm_start_student,
              'dtm_end_student': candidate.dtm_end_student,
             'dtm_start_trainer': candidate.dtm_start_trainer,
              'dtm_end_trainer': candidate.dtm_end_trainer,
            
              'dtm_of_training': candidate.dtm_of_training,
          


            })
    return Response(candidate_data)




class update_course_schedule(generics.UpdateAPIView):
    queryset = course_schedule.objects.all()
    serializer_class = courseScheduleSerializer



@api_view(['PUT', 'PATCH'])
def delete_course_schedule(request, pk):
    try:
        print("Entering Function..")
        course = course_schedule.objects.get(id=pk)

        print("course: ", course)
    except course_schedule.DoesNotExist:
        return JsonResponse("course not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    course.deleted = 1
    course.save()
    print("logins: ", course)

    return JsonResponse("course 'deleted' field updated successfully", safe=False)



#-----------------------------Course Content Feedback-----------------------------------------#

class add_course_content_feedback(generics.CreateAPIView):
    queryset = course_content_feedback.objects.all()
    serializer_class = courseContentFeedbackSerializer

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new Course content feedback")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new Course content feedback successfully")
        return response


@api_view(['GET'])
def get_course_content_feedback(request):
    course_set = course_content_feedback.objects.select_related( 'student_id', 'topic_id', 'trainer_id','department_id').filter(deleted=0)

    course_content_data=[]
    for cc in  course_set:
           course_content_data.append({
              'id': cc.id,
             # 'course_id': cc.course_id.course_name,
              'student_id': cc.student_id.students_name,
              'topic_id': cc.topic_id.topic,
               'sub_topic': cc.topic_id.sub_topic,
              'dtm_session': cc.dtm_session,
              'trainer_id': cc.trainer_id.trainer_name,
              'feedback': cc.feedback,
              'department_id':cc.department_id.department,
            })
    return Response(course_content_data)

class update_course_content_feedback(generics.UpdateAPIView):
    queryset = course_content_feedback.objects.all()
    serializer_class = courseContentFeedbackSerializer

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating course content feedback with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated course content feedback with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating course content feedback with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated course content feedback with id {kwargs.get('pk')} successfully")
        return response



@api_view(['PUT', 'PATCH'])
def delete_course_content_feedback(request, pk):
    ##logger.info(f"Attempting to mark course content feedback with id {pk} as deleted")

    try:
        print("Entering Function..")
        course = course_content_feedback.objects.get(id=pk)

        print("course: ", course)
    except course_content_feedback.DoesNotExist:
        return JsonResponse("course not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    course.deleted = 1
    course.save()

    ##logger.info(f"Marked course content feedback with id {pk} as deleted successfully")

    print("course: ", course)

    return JsonResponse("course 'deleted' field updated successfully", safe=False)


#-----------------------------Attendance Master-----------------------------------------#

class add_attendance_master(generics.CreateAPIView):
    queryset = attendance_master.objects.all()
    serializer_class = attendanceMasterSerializer

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new Attendance")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new Attendance successfully")
        return response


@api_view(['GET'])
def get_attendance_master(request):
    attend_set = attendance_master.objects.select_related('student_id', 'course_id', 'test_id').filter(deleted=0)
    
    ##logger.info("Fetching Attendance where deleted=0")

    attendance_data=[]
    for attend in  attend_set:
           attendance_data.append({
              'id': attend.id,
              'student_id': attend.student_id.students_name,
              'course_id': attend.course_id.course_name,
              'test_id': attend.test_id.test_name,
              'dtm_attendance': attend.dtm_attendance
            })
    return Response(attendance_data)

class update_attendance_master(generics.UpdateAPIView):
    queryset = attendance_master.objects.all()
    serializer_class = attendanceMasterSerializer

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating Attendance with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        #logger.info(f"Updated Attendance with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        #logger.info(f"Partially updating Attendance with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        #logger.info(f"Partially updated Attendance with id {kwargs.get('pk')} successfully")
        return response



@api_view(['PUT', 'PATCH'])
def delete_attendance_master(request, pk):
    #logger.info(f"Attempting to mark Attendance with id {pk} as deleted")

    try:
        print("Entering Function..")
        attend = attendance_master.objects.get(id=pk)

        print("attend: ", attend)
    except attendance_master.DoesNotExist:
        return JsonResponse("course not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    attend.deleted = 1
    attend.save()

    #logger.info(f"Marked Attendance with id {pk} as deleted successfully")

    print("attend: ", attend)

    return JsonResponse("attend 'deleted' field updated successfully", safe=False)



#-----------------------------Announcement Master-----------------------------------------#

class add_announcement_master(generics.CreateAPIView):
    queryset = announcement_master.objects.all()
    serializer_class = announcementSerializer


@api_view(['GET'])
def get_announcement_master(request):
    announce_set = announcement_master.objects.select_related('college_id').filter(deleted=0)

    announce_data=[]
    for announce in announce_set:
        announce_data.append({
            'id': announce.id,
            'college_id': announce.college_id.college,
           # 'trainer_id': announce.trainer_id.trainer_name,
            'dtm_start': announce.dtm_start,
            'dtm_end': announce.dtm_end,
            'content': announce.content,
            'is_active': announce.is_active
        })
    return Response(announce_data)

class update_announcement_master(generics.UpdateAPIView):
    queryset = announcement_master.objects.all()
    serializer_class = announcementSerializer



@api_view(['PUT', 'PATCH'])
def delete_announcement_master(request, pk):
    try:
        print("Entering Function..")
        announce = announcement_master.objects.get(id=pk)

        print("announce: ", announce)
    except announcement_master.DoesNotExist:
        return JsonResponse("announce not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    announce.deleted = 1
    announce.save()
    print("announce: ", announce)

    return JsonResponse("announce 'deleted' field updated successfully", safe=False)


#-----------------------------Trainer Skill Map-----------------------------------------#

class add_trainer_skill_map(generics.CreateAPIView):
    queryset = trainer_skill_map.objects.all()
    serializer_class = trainerSkillMapSerializer


@api_view(['GET'])
def get_trainer_skill_map(request):
    mapping = trainer_skill_map.objects.select_related('trainer_id', 'skill_id').filter(deleted=0)

    map_data=[]
    for maps in  mapping:
           map_data.append({
              'id': maps.id,
              'trainer_id': maps.trainer_id.trainer_name,
              'skill_id': maps.skill_id.skill_name,
              'skill_level': maps.skill_level,
              'dt_skill_from': maps.dt_skill_from,
              'is_handson': maps.is_handson,
              'last_session': maps.last_session
            })
    return Response(map_data)

class update_trainer_skill_map(generics.UpdateAPIView):
    queryset = trainer_skill_map.objects.all()
    serializer_class = trainerSkillMapSerializer



@api_view(['PUT', 'PATCH'])
def delete_trainer_skill_map(request, pk):
    try:
        print("Entering Function..")
        skill_map = trainer_skill_map.objects.get(id=pk)

        print("skill_map: ", skill_map)
    except trainer_skill_map.DoesNotExist:
        return JsonResponse("skill_map not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    skill_map.deleted = 1
    skill_map.save()
    print("skill_map: ", skill_map)

    return JsonResponse("skill_map 'deleted' field updated successfully", safe=False)


#-----------------------------Trainer Skill Availability-----------------------------------------#

class add_trainer_availability(generics.CreateAPIView):
    queryset = trainer_availability.objects.all()
    serializer_class = trainerAvailabilitySerializer


@api_view(['GET'])
def get_trainer_availability(request):
    available = trainer_availability.objects.select_related('trainer_id', 'college_id', 'skill_id').filter(deleted=0)

    available_data=[]
    for data in available:
        available_data.append({
            'id': data.id,
            'trainer_id': data.trainer_id.trainer_name,
            'is_available': data.is_available,
            'dtm_start': data.dtm_start,
            'dtm_stop': data.dtm_stop,
            'college_id': data.college_id.college_name,
            'skill_id': data.skill_id.skill_name
        })
    return Response(available_data)

class update_trainer_availability(generics.UpdateAPIView):
    queryset = trainer_availability.objects.all()
    serializer_class = trainerAvailabilitySerializer



@api_view(['PUT', 'PATCH'])
def delete_trainer_availability(request, pk):
    try:
        print("Entering Function..")
        available = trainer_availability.objects.get(id=pk)

        print("available: ", available)
    except trainer_availability.DoesNotExist:
        return JsonResponse("available not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    available.deleted = 1
    available.save()
    print("available: ", available)

    return JsonResponse("available 'deleted' field updated successfully", safe=False)



#-----------------------------Test Question Map-----------------------------------------#

class add_testQuestion_map(generics.CreateAPIView):
    queryset = test_question_map.objects.all()
    serializer_class = testQuestionMapSerializer


@api_view(['GET'])
def get_testQuestion_map(request):
    mapping = test_question_map.objects.select_related('test_id', 'question_id').filter(deleted=0)

    mapping_data=[]
    for data in mapping:
        mapping_data.append({
            'id': data.id,
            'test_id': data.test_id.test_name,
            'question_id': data.question_id.question_name
        })
    return Response(mapping_data)

class update_testQuestion_map(generics.UpdateAPIView):
    queryset = test_question_map.objects.all()
    serializer_class = testQuestionMapSerializer



@api_view(['PUT', 'PATCH'])
def delete_testQuestion_map(request, pk):
    try:
        print("Entering Function..")
        mapping = test_question_map.objects.get(id=pk)

        print("mapping: ", mapping)
    except test_question_map.DoesNotExist:
        return JsonResponse("mapping not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    mapping.deleted = 1
    mapping.save()
    print("mapping: ", mapping)

    return JsonResponse("mapping 'deleted' field updated successfully", safe=False)





#--------------------------------------Import Questions-----------------------------------#


class ExcelImportView_Tests(APIView):
    def post(self, request, format=None):
        print('Files:', request.FILES)
        
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())  # Print the first few rows of the DataFrame
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        # Replace question_type and skill_type values with their corresponding foreign key IDs
        for index, row in df.iterrows():
            test_type_name = row['test_type_id']
            question_type_name = row['question_type_id']
            skill_type_id_name = row['skill_type_id']
           # college_name = row['college_id']
            #department_name = row['department_id']
           # rules_name = row['rules_id']
            
            id_test_type = test_type.objects.filter(test_type=test_type_name).first().id
            print("testtype",id_test_type)
            id_question_type = question_type.objects.filter(question_type=question_type_name).first().id
            print("questiontype",id_question_type)
            id_skill_type_name = skill_type.objects.filter(skill_type=skill_type_id_name).first().id
            print("skilltype",id_skill_type_name)
           # id_college_name = college_master.objects.filter(college=college_name).first().id
           # id_department_name = department_master.objects.filter(department=department_name).first().id
           # id_rules_name = rules.objects.filter(rule_name=rules_name).first().id

            df.at[index, 'test_type_id'] = id_test_type
            df.at[index, 'question_type_id'] = id_question_type
            df.at[index, 'skill_type_id'] = id_skill_type_name
          #  df.at[index, 'college_id'] = id_college_name
           # df.at[index, 'department_id'] = id_department_name
           # df.at[index, 'rules_id'] = id_rules_name
        
        records = df.to_dict(orient='records')
        
        serializer = testsSerializersImport(data=records, many=True)
        
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        



class ExcelImportView_CandidateLAST_one(APIView):
    def post(self, request, format=None):
        print('Files:', request.FILES)
        
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())  # Print the first few rows of the DataFrame
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        
        # Mapping of Excel header names to actual column names
        header_mapping = {
            'College Name**': 'college_id',
            'Student Name': 'students_name',
            'User Name**': 'user_name',
            'Reg No': 'registration_number',
            'Gender': 'gender',
            'Email ID': 'email_id',
            'Mobile Number': 'mobile_number',
            'Year**': 'year',
            'CGPA': 'cgpa',
            'Department**': 'department_id',
            '10th Mark': 'marks_10th',
            '12th Mark': 'marks_12th',
            'Semaster Wise': 'marks_semester_wise',
            'History Of Arrears': 'history_of_arrears',
            'Standing Arrears': 'standing_arrears',
            'No.Of.IT Offers': 'it_of_offers',
            'No.Of.Core Offers': 'core_of_offers',
            'No.Of.Offers': 'number_of_offers',
            'Password**': 'password'
        }
        
        # Rename columns in DataFrame to match expected column names
        df.rename(columns=header_mapping, inplace=True)
        
        
        # Check for mandatory columns
        mandatory_columns = ['college_id', 'user_name', 'department_id', 'year', 'password']
        for column in mandatory_columns:
            if column not in df.columns:
                return Response({'error': f'Mandatory column {column} is missing'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Ensure mandatory columns are not null
        if df[mandatory_columns].isnull().any().any():
            return Response({'error': 'Mandatory fields cannot be null'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Replace college_id and department_id values with their corresponding foreign key IDs
        for index, row in df.iterrows():
            college_name = row['college_id']
            department_name = row['department_id']

            college_instance = college_master.objects.filter(college=college_name).first()
            department_instance = department_master.objects.filter(department=department_name).first()
            
            if not college_instance or not department_instance:
                return Response({'error': 'Invalid college or department name'}, status=status.HTTP_400_BAD_REQUEST)

            id_college_name = college_instance.id
            id_department_name = department_instance.id

            df.at[index, 'college_id'] = id_college_name
            df.at[index, 'department_id'] = id_department_name

        # Add default role
        df['role'] = 'Student'
        
        # Handle optional fields by setting them to pd.NA if they are empty
        optional_columns = ['students_name', 'registration_number', 'gender', 'email_id', 'mobile_number',
                            'cgpa', 'marks_10th', 'marks_12th', 'marks_semester_wise',
                            'history_of_arrears', 'standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers', 'text']
        
        for column in optional_columns:
            if column not in df.columns:
                df[column] = pd.NA  # Add the column with pd.NA values if it doesn't exist
            else:
                df[column] = df[column].fillna(pd.NA)  # Fill NaN values with pd.NA
        
        # Ensure email addresses are valid, add placeholder email if missing or invalid
        df['email_id'] = df['email_id'].apply(lambda x: x if pd.notna(x) and '@' in str(x) else 'placeholder@example.com')
        
        # Fill numeric fields with default values (e.g., 0) if they are NaN
        numeric_fields = ['standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers']
        for field in numeric_fields:
            df[field] = df[field].fillna(0)
        
        # Prepare login data
        login_data = df[['email_id', 'user_name', 'password', 'college_id', 'role']]
        login_records = login_data.to_dict(orient='records')
        
        # Serialize login data
        login_serializer = loginSerializerStu(data=login_records, many=True)
        
        if not login_serializer.is_valid():
            print('Login serializer errors:', login_serializer.errors)
            return Response(login_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        # Save login data
        login_serializer.save()

        # Prepare candidate data, ensure user_name is correctly included
        candidate_data = df[['college_id', 'students_name', 'registration_number', 'gender', 'email_id', 'mobile_number',
                             'department_id', 'year', 'cgpa', 'marks_10th', 'marks_12th', 'marks_semester_wise',
                             'history_of_arrears', 'standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers', 'user_name', 'text']]
        print("Candidate_data: ", candidate_data)

        candidate_records = candidate_data.to_dict(orient='records')
        print("Candidate_records: ", candidate_records)
        
        # Serialize candidate data
        candidate_serializer = candidateSerializerImport(data=candidate_records, many=True)
        
        if not candidate_serializer.is_valid():
            print('Candidate serializer errors:', candidate_serializer.errors)
            return Response(candidate_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        # Save candidate data
        candidate_serializer.save()

        print('login_data: ', login_serializer.data)
        print('Candidate_data: ', candidate_serializer.data)
        
        return Response({'login_data': login_serializer.data, 'candidate_data': candidate_serializer.data}, status=status.HTTP_201_CREATED)



class ExcelImportView_CandidateLASTollll(APIView):
    def post(self, request, format=None):
        if 'file' not in request.FILES:
            print("No file uploaded")
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            print("File is not in Excel format")
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print("Excel file read successfully")
        except Exception as e:
            print(f"Error reading Excel file: {str(e)}")
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        header_mapping = {
            'College Name**': 'college_id',
            'Student Name': 'students_name',
            'User Name**': 'user_name',
            'Reg No': 'registration_number',
            'Gender': 'gender',
            'Email ID': 'email_id',
            'Mobile Number': 'mobile_number',
            'Year**': 'year',
            'CGPA': 'cgpa',
            'Department**': 'department_id',
            '10th Mark': 'marks_10th',
            '12th Mark': 'marks_12th',
            'Semaster Wise': 'marks_semester_wise',
            'History Of Arrears': 'history_of_arrears',
            'Standing Arrears': 'standing_arrears',
            'No.Of.IT Offers': 'it_of_offers',
            'No.Of.Core Offers': 'core_of_offers',
            'No.Of.Offers': 'number_of_offers',
            'Password**': 'password'
        }
        
        df.rename(columns=header_mapping, inplace=True)
        print("Columns renamed according to header_mapping")

        # Trim all string columns
        for column in df.select_dtypes(include=['object']).columns:
            df[column] = df[column].str.strip()

        mandatory_columns = ['college_id', 'user_name', 'department_id', 'year', 'password']
        for column in mandatory_columns:
            if column not in df.columns:
                print(f"Mandatory column {column} is missing")
                return Response({'error': f'Mandatory column {column} is missing'}, status=status.HTTP_400_BAD_REQUEST)
        
        if df[mandatory_columns].isnull().any().any():
            print("Mandatory fields cannot be null")
            return Response({'error': 'Mandatory fields cannot be null'}, status=status.HTTP_400_BAD_REQUEST)

        email_regex = re.compile(r'^\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b$')

        user_names = df['user_name']
        existing_users = set(User.objects.filter(username__in=user_names).values_list('username', flat=True))

        if existing_users:
            duplicates = [name for name in user_names if name in existing_users]
            print(f"Duplicate usernames found: {', '.join(duplicates)}")
            return Response({'error': f'Usernames already exist: {", ".join(duplicates)}'}, status=status.HTTP_400_BAD_REQUEST)

        processed_usernames = set()
        for index, row in df.iterrows():
            username = row['user_name']
            if username in processed_usernames:
                print(f"Duplicate username within the file: {username} in row {index + 2}")
                return Response({'error': f'Duplicate username within the file: {username} in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            processed_usernames.add(username)
            print(f"Processing row {index + 2}")
            college_instance = college_master.objects.filter(college=row['college_id']).first()
            department_instance = department_master.objects.filter(department=row['department_id']).first()
            
            if not college_instance:
                print(f"College not found in row {index + 2}")
                return Response({'error': f'College not found in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
            elif not department_instance:
                print(f"Department not found in row {index + 2}")
                return Response({'error': f'Department not found in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            year_value = str(row['year'])
            if year_value not in ['1', '2', '3', '4']:
                print(f"Year is mismatch in row {index + 2}")
                return Response({'error': f'Year is mismatch in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            if pd.notna(row['email_id']):
                if not isinstance(row['email_id'], str) or not email_regex.match(row['email_id']):
                    print(f"Email ID is not in correct format in row {index + 2}")
                    return Response({'error': f'Email ID is not in correct format in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
        
            if pd.notna(row['mobile_number']):
                if len(str(int(row['mobile_number']))) != 10:
                    print(f"Mobile number is not 10 digits in row {index + 2}")
                    return Response({'error': f'Mobile number is not 10 digits in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            # Convert and validate integer fields
            integer_fields = ['history_of_arrears', 'standing_arrears', 'core_of_offers', 'it_of_offers', 'number_of_offers']
            for field in integer_fields:
                if pd.notna(row[field]):
                    try:
                        value = float(row[field])
                        if value.is_integer():
                            row[field] = int(value)
                        else:
                            print(f"{field.replace('_', ' ').title()} must be an integer in row {index + 2}")
                            return Response({'error': f'{field.replace("_", " ").title()} must be an integer in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
                    except ValueError:
                        print(f"{field.replace('_', ' ').title()} must be an integer in row {index + 2}")
                        return Response({'error': f'{field.replace("_", " ").title()} must be an integer in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
            
            # Convert and validate float fields
            float_fields = ['cgpa', 'marks_10th', 'marks_12th']
            for field in float_fields:
                if pd.notna(row[field]):
                    try:
                        row[field] = float(row[field])
                    except ValueError:
                        print(f"{field.replace('_', ' ').title()} must be a number in row {index + 2}")
                        return Response({'error': f'{field.replace("_", " ").title()} must be a number in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
    
            df.at[index, 'user_name'] = username
            df.at[index, 'college_id'] = college_instance.id
            df.at[index, 'department_id'] = department_instance.id

        df['role'] = 'Student'
        print("Role set to 'Student' for all records")
        
        optional_columns = ['students_name', 'registration_number', 'gender', 'email_id', 'mobile_number',
                            'cgpa', 'marks_10th', 'marks_12th', 'marks_semester_wise',
                            'history_of_arrears', 'standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers', 'text']
        
        # Replace NaN or None with default values
        for column in optional_columns:
            if column in df.columns:
                if df[column].dtype == 'float64':  # For float columns
                    df[column] = df[column].fillna(0.0)
                elif df[column].dtype == 'int64':  # For integer columns
                    df[column] = df[column].fillna(0)
                else:  # For other columns
                    df[column] = df[column].fillna(pd.NA)
            else:
                if column in ['cgpa', 'marks_10th', 'marks_12th']:
                    df[column] = 0.0
                elif column in ['history_of_arrears', 'standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers']:
                    df[column] = 0
                else:
                    df[column] = pd.NA
        
        df['email_id'] = df['email_id'].apply(lambda x: x if pd.notna(x) and '@' in str(x) else 'placeholder@example.com')
        
        # numeric_fields = ['standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers']
        # for field in numeric_fields:
        #     df[field] = df[field].fillna(0)
        
        # Replace NaNs with None to avoid serialization issues
        df = df.where(pd.notnull(df), None)
        
        login_data = df[['email_id', 'user_name', 'password', 'college_id', 'role']]
        login_records = login_data.to_dict(orient='records')
        print(f"Login data prepared: {login_records}")

        login_serializer = loginSerializerStu(data=login_records, many=True)
        
        if not login_serializer.is_valid():
            print(f"Login serializer errors: {login_serializer.errors}")
            return Response(login_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        candidate_data = df[['college_id', 'students_name', 'registration_number', 'gender', 'email_id', 'mobile_number',
                             'department_id', 'year', 'cgpa', 'marks_10th', 'marks_12th', 'marks_semester_wise',
                             'history_of_arrears', 'standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers', 'user_name', 'text']]
        candidate_records = candidate_data.to_dict(orient='records')
        print(f"Candidate data prepared: {candidate_records}")
        
        candidate_serializer = candidateSerializerImport(data=candidate_records, many=True)
        
        if not candidate_serializer.is_valid():
            print(f"Candidate serializer errors: {candidate_serializer.errors}")
            return Response(candidate_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        # Save data in a transaction to ensure atomicity
        with transaction.atomic():
            login_serializer.save()
            print("Login data saved successfully")
            candidate_serializer.save()
            print("Candidate data saved successfully")

        return Response({'message': 'Excel file imported successfully'}, status=status.HTTP_201_CREATED)
class ExcelImportView_CandidateLAST(APIView):
    def post(self, request, format=None):
        if 'file' not in request.FILES:
            print("No file uploaded")
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            print("File is not in Excel format")
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print("Excel file read successfully")
        except Exception as e:
            print(f"Error reading Excel file: {str(e)}")
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        header_mapping = {
            'College Name**': 'college_id',
            'Student Name': 'students_name',
            'User Name**': 'user_name',
            'Reg No': 'registration_number',
            'Gender': 'gender',
            'Email ID': 'email_id',
            'Mobile Number': 'mobile_number',
            'Year**': 'year',
            'CGPA': 'cgpa',
            'Department**': 'department_id',
            '10th Mark': 'marks_10th',
            '12th Mark': 'marks_12th',
            'Semaster Wise': 'marks_semester_wise',
            'History Of Arrears': 'history_of_arrears',
            'Standing Arrears': 'standing_arrears',
            'No.Of.IT Offers': 'it_of_offers',
            'No.Of.Core Offers': 'core_of_offers',
            'No.Of.Offers': 'number_of_offers',
            'Password**': 'password'
        }
        
        df.rename(columns=header_mapping, inplace=True)
        print("Columns renamed according to header_mapping")

        # Trim all string columns
        for column in df.select_dtypes(include=['object']).columns:
            df[column] = df[column].str.strip()

        mandatory_columns = ['college_id', 'user_name', 'department_id', 'year', 'password']
        for column in mandatory_columns:
            if column not in df.columns:
                print(f"Mandatory column {column} is missing")
                return Response({'error': f'Mandatory column {column} is missing'}, status=status.HTTP_400_BAD_REQUEST)
        
        if df[mandatory_columns].isnull().any().any():
            print("Mandatory fields cannot be null")
            return Response({'error': 'Mandatory fields cannot be null'}, status=status.HTTP_400_BAD_REQUEST)

        email_regex = re.compile(r'^\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b$')

        user_names = df['user_name']
        existing_users = set(User.objects.filter(username__in=user_names).values_list('username', flat=True))

        if existing_users:
            duplicates = [name for name in user_names if name in existing_users]
            print(f"Duplicate usernames found: {', '.join(duplicates)}")
            return Response({'error': f'Usernames already exist: {", ".join(duplicates)}'}, status=status.HTTP_400_BAD_REQUEST)

        processed_usernames = set()
        for index, row in df.iterrows():
            username = row['user_name']
            if username in processed_usernames:
                print(f"Duplicate username within the file: {username} in row {index + 2}")
                return Response({'error': f'Duplicate username within the file: {username} in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            processed_usernames.add(username)
            print(f"Processing row {index + 2}")
            college_instance = college_master.objects.filter(college=row['college_id']).first()
            department_instance = department_master.objects.filter(department=row['department_id']).first()
            
            if not college_instance:
                print(f"College not found in row {index + 2}")
                return Response({'error': f'College not found in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
            elif not department_instance:
                print(f"Department not found in row {index + 2}")
                return Response({'error': f'Department not found in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            year_value = str(row['year'])
            if year_value not in ['1', '2', '3', '4']:
                print(f"Year is mismatch in row {index + 2}")
                return Response({'error': f'Year is mismatch in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            if pd.notna(row['email_id']):
                if not isinstance(row['email_id'], str) or not email_regex.match(row['email_id']):
                    print(f"Email ID is not in correct format in row {index + 2}")
                    return Response({'error': f'Email ID is not in correct format in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
        
            if pd.notna(row['mobile_number']):
                if len(str(int(row['mobile_number']))) != 10:
                    print(f"Mobile number is not 10 digits in row {index + 2}")
                    return Response({'error': f'Mobile number is not 10 digits in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            # Convert and validate integer fields
            # integer_fields = ['history_of_arrears', 'standing_arrears', 'core_of_offers', 'it_of_offers', 'number_of_offers']
            # for field in integer_fields:
            #     if pd.notna(row[field]):
            #         try:
            #             value = float(row[field])
            #             if value.is_integer():
            #                 row[field] = int(value)
            #             else:
            #                 print(f"{field.replace('_', ' ').title()} must be an integer in row {index + 2}")
            #                 return Response({'error': f'{field.replace("_", " ").title()} must be an integer in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
            #         except ValueError:
            #             print(f"{field.replace('_', ' ').title()} must be an integer in row {index + 2}")
            #             return Response({'error': f'{field.replace("_", " ").title()} must be an integer in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
            
            # Convert and validate integer fields
            integer_fields = ['history_of_arrears', 'standing_arrears', 'core_of_offers', 'it_of_offers', 'number_of_offers']
            for field in integer_fields:
                if pd.notna(row[field]):
                    try:
                        value = float(row[field])
                        if value.is_integer() and 0 <= value <= 100:
                            row[field] = int(value)
                        else:
                            print(f"{field.replace('_', ' ').title()} must be an integer between 0 and 100 in row {index + 2}")
                            return Response({'error': f'{field.replace("_", " ").title()} must be an integer between 0 and 100 in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
                    except ValueError:
                        print(f"{field.replace('_', ' ').title()} must be an integer between 0 and 100 in row {index + 2}")
                        return Response({'error': f'{field.replace("_", " ").title()} must be an integer between 0 and 100 in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)


            # Convert and validate float fields
            float_fields = ['cgpa', 'marks_10th', 'marks_12th']
            for field in float_fields:
                if pd.notna(row[field]):
                    try:
                        row[field] = float(row[field])
                    except ValueError:
                        print(f"{field.replace('_', ' ').title()} must be a number in row {index + 2}")
                        return Response({'error': f'{field.replace("_", " ").title()} must be a number in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
    
            df.at[index, 'user_name'] = username
            df.at[index, 'college_id'] = college_instance.id
            df.at[index, 'department_id'] = department_instance.id

        df['role'] = 'Student'
        print("Role set to 'Student' for all records")
        
        optional_columns = ['students_name', 'registration_number', 'gender', 'email_id', 'mobile_number',
                            'cgpa', 'marks_10th', 'marks_12th', 'marks_semester_wise',
                            'history_of_arrears', 'standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers', 'text']
        
        # Replace NaN or None with default values
        for column in optional_columns:
            if column in df.columns:
                if df[column].dtype == 'float64':  # For float columns
                    df[column] = df[column].fillna(0.0)
                elif df[column].dtype == 'int64':  # For integer columns
                    df[column] = df[column].fillna(0)
                else:  # For other columns
                    df[column] = df[column].fillna(pd.NA)
            else:
                if column in ['cgpa', 'marks_10th', 'marks_12th']:
                    df[column] = 0.0
                elif column in ['history_of_arrears', 'standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers']:
                    df[column] = 0
                else:
                    df[column] = pd.NA
        
        df['email_id'] = df['email_id'].apply(lambda x: x if pd.notna(x) and '@' in str(x) else 'placeholder@example.com')
        
        # numeric_fields = ['standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers']
        # for field in numeric_fields:
        #     df[field] = df[field].fillna(0)
        
        # Replace NaNs with None to avoid serialization issues
        df = df.where(pd.notnull(df), None)
        
        login_data = df[['email_id', 'user_name', 'password', 'college_id', 'role']]
        login_records = login_data.to_dict(orient='records')
        print(f"Login data prepared: {login_records}")

        login_serializer = loginSerializerStu(data=login_records, many=True)
        
        if not login_serializer.is_valid():
            print(f"Login serializer errors: {login_serializer.errors}")
            return Response(login_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        candidate_data = df[['college_id', 'students_name', 'registration_number', 'gender', 'email_id', 'mobile_number',
                             'department_id', 'year', 'cgpa', 'marks_10th', 'marks_12th', 'marks_semester_wise',
                             'history_of_arrears', 'standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers', 'user_name', 'text']]
        candidate_records = candidate_data.to_dict(orient='records')
        print(f"Candidate data prepared: {candidate_records}")
        
        candidate_serializer = candidateSerializerImport(data=candidate_records, many=True)
        
        if not candidate_serializer.is_valid():
            print(f"Candidate serializer errors: {candidate_serializer.errors}")
            return Response(candidate_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        # Save data in a transaction to ensure atomicity
        with transaction.atomic():
            login_serializer.save()
            print("Login data saved successfully")
            candidate_serializer.save()
            print("Candidate data saved successfully")

        return Response({'message': 'Excel file imported successfully'}, status=status.HTTP_201_CREATED)



class ExcelImportView_CandidateLAST_2607(APIView):
    def post(self, request, format=None):
        if 'file' not in request.FILES:
            print("No file uploaded")
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            print("File is not in Excel format")
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print("Excel file read successfully")
        except Exception as e:
            print(f"Error reading Excel file: {str(e)}")
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        header_mapping = {
            'College Name**': 'college_id',
            'Student Name': 'students_name',
            'User Name**': 'user_name',
            'Reg No': 'registration_number',
            'Gender': 'gender',
            'Email ID': 'email_id',
            'Mobile Number': 'mobile_number',
            'Year**': 'year',
            'CGPA': 'cgpa',
            'Department**': 'department_id',
            '10th Mark': 'marks_10th',
            '12th Mark': 'marks_12th',
            'Semaster Wise': 'marks_semester_wise',
            'History Of Arrears': 'history_of_arrears',
            'Standing Arrears': 'standing_arrears',
            'No.Of.IT Offers': 'it_of_offers',
            'No.Of.Core Offers': 'core_of_offers',
            'No.Of.Offers': 'number_of_offers',
            'Password**': 'password'
        }
        
        df.rename(columns=header_mapping, inplace=True)
        print("Columns renamed according to header_mapping")

        # Trim all string columns
        for column in df.select_dtypes(include=['object']).columns:
            df[column] = df[column].str.strip()

        mandatory_columns = ['college_id', 'user_name', 'department_id', 'year', 'password']
        for column in mandatory_columns:
            if column not in df.columns:
                print(f"Mandatory column {column} is missing")
                return Response({'error': f'Mandatory column {column} is missing'}, status=status.HTTP_400_BAD_REQUEST)
        
        if df[mandatory_columns].isnull().any().any():
            print("Mandatory fields cannot be null")
            return Response({'error': 'Mandatory fields cannot be null'}, status=status.HTTP_400_BAD_REQUEST)

        email_regex = re.compile(r'^\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b$')

        user_names = df['user_name']
        existing_users = set(User.objects.filter(username__in=user_names).values_list('username', flat=True))

        if existing_users:
            duplicates = [name for name in user_names if name in existing_users]
            print(f"Duplicate usernames found: {', '.join(duplicates)}")
            return Response({'error': f'Usernames already exist: {", ".join(duplicates)}'}, status=status.HTTP_400_BAD_REQUEST)

        processed_usernames = set()
        for index, row in df.iterrows():
            username = row['user_name']
            if username in processed_usernames:
                print(f"Duplicate username within the file: {username} in row {index + 2}")
                return Response({'error': f'Duplicate username within the file: {username} in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            processed_usernames.add(username)
            print(f"Processing row {index + 2}")
            college_instance = college_master.objects.filter(college=row['college_id']).first()
            department_instance = department_master.objects.filter(department=row['department_id']).first()
            
            if not college_instance:
                print(f"College not found in row {index + 2}")
                return Response({'error': f'College not found in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
            elif not department_instance:
                print(f"Department not found in row {index + 2}")
                return Response({'error': f'Department not found in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            year_value = str(row['year'])
            if year_value not in ['1', '2', '3', '4']:
                print(f"Year is mismatch in row {index + 2}")
                return Response({'error': f'Year is mismatch in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            if pd.notna(row['email_id']):
                if not isinstance(row['email_id'], str) or not email_regex.match(row['email_id']):
                    print(f"Email ID is not in correct format in row {index + 2}")
                    return Response({'error': f'Email ID is not in correct format in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
        
            if pd.notna(row['mobile_number']):
                if len(str(int(row['mobile_number']))) != 10:
                    print(f"Mobile number is not 10 digits in row {index + 2}")
                    return Response({'error': f'Mobile number is not 10 digits in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
        
            if pd.notna(row['history_of_arrears']):
                try:
                    # Convert to float first
                    value = float(row['history_of_arrears'])
                    # Check if it's an integer by comparing it to its integer representation
                    if value.is_integer():
                        row['history_of_arrears'] = int(value)
                    else:
                        # If it's not an integer, raise an error
                        print(f"History of arrears must be an integer in row {index + 2}")
                        return Response({'error': f'History of arrears must be an integer in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
                except ValueError:
                    # Handle any conversion errors
                    print(f"History of arrears must be an integer in row {index + 2}")
                    return Response({'error': f'History of arrears must be an integer in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
    
        
            if pd.notna(row['standing_arrears']):
                try:
                    # Convert to float first
                    value = float(row['standing_arrears'])
                    # Check if it's an integer by comparing it to its integer representation
                    if value.is_integer():
                        row['standing_arrears'] = int(value)
                    else:
                        # If it's not an integer, raise an error
                        print(f"Standing of arrears must be an integer in row {index + 2}")
                        return Response({'error': f'Standing of arrears must be an integer in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
                except ValueError:
                    # Handle any conversion errors
                    print(f"Standing of arrears must be an integer in row {index + 2}")
                    return Response({'error': f'Standing of arrears must be an integer in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
    
            df.at[index, 'user_name'] = username
            df.at[index, 'college_id'] = college_instance.id
            df.at[index, 'department_id'] = department_instance.id

        df['role'] = 'Student'
        print("Role set to 'Student' for all records")
        
        optional_columns = ['students_name', 'registration_number', 'gender', 'email_id', 'mobile_number',
                            'cgpa', 'marks_10th', 'marks_12th', 'marks_semester_wise',
                            'history_of_arrears', 'standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers', 'text']
        
        for column in optional_columns:
            if column not in df.columns:
                df[column] = pd.NA
            else:
                df[column] = df[column].fillna(pd.NA)
        
        df['email_id'] = df['email_id'].apply(lambda x: x if pd.notna(x) and '@' in str(x) else 'placeholder@example.com')

        login_data = df[['email_id', 'user_name', 'password', 'college_id', 'role']]
        candidate_data = df[['college_id', 'students_name', 'registration_number', 'gender', 'email_id', 'mobile_number',
                             'department_id', 'year', 'cgpa', 'marks_10th', 'marks_12th', 'marks_semester_wise',
                             'history_of_arrears', 'standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers', 'user_name', 'text']]
        
        login_records = login_data.to_dict(orient='records')
        print(f"Login data prepared: {login_records}")

        login_serializer = loginSerializerStu(data=login_records, many=True)
        
        if not login_serializer.is_valid():
            print(f"Login serializer errors: {login_serializer.errors}")
            return Response({'error': 'Login data validation errors', 'details': login_serializer.errors}, status=status.HTTP_400_BAD_REQUEST)
        
        candidate_records = candidate_data.to_dict(orient='records')
        print(f"Candidate data prepared: {candidate_records}")

        # Custom error messages
        def get_error_message(field):
            error_messages = {
                'cgpa': 'CGPA must be a valid number',
                'marks_10th': '10th Marks must be a valid number',
                'marks_12th': '12th Marks must be a valid number',
                'history_of_arrears': 'History of Arrears must be a valid integer'
            }
            return error_messages.get(field, 'Invalid value')

        candidate_serializer = candidateSerializerImport(data=candidate_records, many=True)
        
        if not candidate_serializer.is_valid():
            error_details = candidate_serializer.errors
            formatted_errors = {}
            for idx, errors in enumerate(error_details):
                formatted_errors[idx] = {field: get_error_message(field) for field in errors}
            print(f"Candidate serializer errors: {formatted_errors}")
            return Response({'error': 'Candidate data validation errors', 'details': formatted_errors}, status=status.HTTP_400_BAD_REQUEST)
        
        # Save data in a transaction to ensure atomicity
        with transaction.atomic():
            login_serializer.save()
            print("Login data saved successfully")
            candidate_serializer.save()
            print("Candidate data saved successfully")
        
        return Response({'login_data': login_serializer.data, 'candidate_data': candidate_serializer.data}, status=status.HTTP_201_CREATED)





class ExcelImportView_Candidate(APIView):
    def post(self, request, format=None):
        print('Files:', request.FILES)
        
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())  # Print the first few rows of the DataFrame
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        # Replace college_id and department_id values with their corresponding foreign key IDs
        for index, row in df.iterrows():
            college_name = row['college_id']
            department_name = row['department_id']

            college_instance = college_master.objects.filter(college=college_name).first()
            department_instance = department_master.objects.filter(department=department_name).first()
            
            if not college_instance or not department_instance:
                return Response({'error': 'Invalid college or department name'}, status=status.HTTP_400_BAD_REQUEST)

            id_college_name = college_instance.id
            id_department_name = department_instance.id

            df.at[index, 'college_id'] = id_college_name
            df.at[index, 'department_id'] = id_department_name

        
        # Add default role
        df['role'] = 'Student'
        
        # Prepare login data
        df['user_name'] = df['registration_number']  # Set user_name to registration_number for loginSerializerStu
        login_data = df[['email_id', 'user_name', 'password', 'college_id', 'role']]
        login_records = login_data.to_dict(orient='records')
        
        # Serialize login data
        login_serializer = loginSerializerStu(data=login_records, many=True)
        
        if not login_serializer.is_valid():
            return Response(login_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        # Save login data
        login_serializer.save()

        # Prepare candidate data, ensure user_name is correctly included
        candidate_data = df[['college_id', 'students_name', 'registration_number', 'gender', 'email_id', 'mobile_number',
                             'department_id', 'year', 'cgpa', 'marks_10th', 'marks_12th', 'marks_semester_wise',
                             'history_of_arrears', 'standing_arrears', 'number_of_offers', 'text']]
        print("Candidate_data: ", candidate_data)

        candidate_records = candidate_data.to_dict(orient='records')
        print("Candidate_records: ", candidate_records)
        
        # Serialize candidate data
        candidate_serializer = candidateSerializerImport(data=candidate_records, many=True)
        
        if not candidate_serializer.is_valid():
            return Response(candidate_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        # Save candidate data
        candidate_serializer.save()

        print('login_data: ', login_serializer.data)
        print('Candidate_data: ', candidate_serializer.data)
        
        return Response({'login_data': login_serializer.data, 'candidate_data': candidate_serializer.data}, status=status.HTTP_201_CREATED)



class ExcelImportView_CandidateuserOLD0406(APIView):
    def post(self, request, format=None):
        print('Files:', request.FILES)
        
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())  # Print the first few rows of the DataFrame
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        # Ensure user_name is not null
        if df['user_name'].isnull().any():
            return Response({'error': 'User name cannot be null'}, status=status.HTTP_400_BAD_REQUEST)

        # Add default role
        df['role'] = 'Student'
        
        # Drop unnecessary columns
        if 'college_id' in df.columns:
            df.drop(columns=['college_id'], inplace=True)
        if 'department_id' in df.columns:
            df.drop(columns=['department_id'], inplace=True)
        
        # Prepare login data with default values for college_id and department_id
        login_data = df[['user_name', 'password', 'role']]
        login_records = login_data.to_dict(orient='records')
        
        # Serialize login data
        login_serializer = loginSerializerStuser(data=login_records, many=True)
        
        if not login_serializer.is_valid():
            return Response(login_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        # Save login data
        login_serializer.save()

        # Prepare candidate data, ensure user_name is correctly included
        candidate_data = df[['user_name']]
        candidate_records = candidate_data.to_dict(orient='records')
        
        # Serialize candidate data
        candidate_serializer = candidateuserSerializerImport(data=candidate_records, many=True)
        
        if not candidate_serializer.is_valid():
            return Response(candidate_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        # Save candidate data
        candidate_serializer.save()

        print('login_data: ', login_serializer.data)
        print('Candidate_data: ', candidate_serializer.data)
        
        return Response({'login_data': login_serializer.data, 'candidate_data': candidate_serializer.data}, status=status.HTTP_201_CREATED)


class ExcelImportView_Candidateuser(APIView):
    def post(self, request, format=None):
        print('Files:', request.FILES)
        
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())  # Print the first few rows of the DataFrame
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        # Ensure user_name is not null
        if 'user_name' not in df.columns or df['user_name'].isnull().any():
            return Response({'error': 'User name cannot be null'}, status=status.HTTP_400_BAD_REQUEST)

        # Ensure password column is present
        if 'password' not in df.columns:
            return Response({'error': 'Password column is missing'}, status=status.HTTP_400_BAD_REQUEST)

        # Prepare login data with default values for role
        login_data = df[['user_name', 'password']]
        login_data['role'] = 'Student'  # Default role
        login_records = login_data.to_dict(orient='records')

        print('login records: ', login_records)
        
        # Serialize login data
        login_serializer = loginSerializerStuser(data=login_records, many=True)
        
        if not login_serializer.is_valid():
            return Response(login_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        # Save login data
        login_serializer.save()
        
       
        # Prepare candidate data
        candidate_data = df[['user_name']]
        candidate_records = candidate_data.to_dict(orient='records')
        print('candidate records: ', candidate_records)

        # Serialize candidate data
        candidate_serializer = candidateuserSerializerImport(data=candidate_records, many=True)
        
        if not candidate_serializer.is_valid():
            return Response(candidate_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        # Save candidate data
        candidate_serializer.save()

        print('login_data: ', login_serializer.data)
        print('Candidate_data: ', candidate_serializer.data)
        
        return Response({'login_data': login_serializer.data, 'candidate_data': candidate_serializer.data}, status=status.HTTP_201_CREATED)



class ExcelImportView_QuestionsOLD(APIView):
    def post(self, request, format=None):
        print('Files:', request.FILES)
        
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())  # Print the first few rows of the DataFrame
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        # Replace question_type and skill_type values with their corresponding foreign key IDs
        for index, row in df.iterrows():
            topic_name = row['topic_id']
            #question_type_name = row['question_type_id']
            #skills_name = row['skill_id']
            
            id_topic_name = topic_master.objects.filter(topic=topic_name).first().id
            #id_question_type_name = question_type.objects.filter(question_type=question_type_name).first().id
            #id_skill_name = skills_master.objects.filter(skill_name=skills_name).first().id

            df.at[index, 'topic_id'] = id_topic_name
            #df.at[index, 'question_type_id'] = id_question_type_name
            #df.at[index, 'skill_id'] = id_skill_name
        
        records = df.to_dict(orient='records')
        
        serializer = questionsSerializerImport(data=records, many=True)
        print("Serializer: ", serializer)
        print("Entered of Saving....")
        
        if serializer.is_valid():
            print("Before Saving..")
            serializer.save()
            print("Afetr Saving..")
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        


class ExcelImportView_Questions_OLD3105(APIView):
    def post(self, request, format=None):
        print('Files:', request.FILES)
        
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())  # Print the first few rows of the DataFrame
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        # Replace question_type and skill_type values with their corresponding foreign key IDs
        for index, row in df.iterrows():
            topic_name = row['topic_id']
            
            try:
                topic_obj = topic_master.objects.filter(topic=topic_name).first()
                if not topic_obj:
                    return Response({'error': f'Topic "{topic_name}" not found in topic_master table'}, status=status.HTTP_400_BAD_REQUEST)
                df.at[index, 'topic_id'] = topic_obj.id
            except Exception as e:
                return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        records = df.to_dict(orient='records')
        
        serializer = questionsSerializerImport(data=records, many=True)
        print("Serializer: ", serializer)
        
        if serializer.is_valid():
            print("Before Saving..")
            serializer.save()
            print("After Saving..")
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        else:
            print("Validation Errors: ", serializer.errors)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)




class ExcelImportView_Questions_Code_OLD3105(APIView):
    def post(self, request, format=None):
        print('Files:', request.FILES)
        
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())  # Print the first few rows of the DataFrame
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        # Replace question_type and skill_type values with their corresponding foreign key IDs
        for index, row in df.iterrows():
            topic_name = row['topic_id']
            #question_type_name = row['question_type_id']
            #skills_name = row['skill_id']
            
            id_topic_name = topic_master.objects.filter(topic=topic_name).first().id
            #id_question_type_name = question_type.objects.filter(question_type=question_type_name).first().id
            #id_skill_name = skills_master.objects.filter(skill_name=skills_name).first().id

            df.at[index, 'topic_id'] = id_topic_name
            #df.at[index, 'question_type_id'] = id_question_type_name
            #df.at[index, 'skill_id'] = id_skill_name
        
        records = df.to_dict(orient='records')
        
        serializer = questionsSerializerCodeImport(data=records, many=True)
        print("Serializer: ", serializer)
        
        if serializer.is_valid():
            print("Before Saving..")
            serializer.save()
            print("Afetr Saving..")
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)



class ExcelImportView_Questions(APIView):
    def post(self, request, format=None):
        # Extract data for the question_paper_master
        question_paper_name = request.data.get('question_paper_name')
        duration_of_test = request.data.get('duration_of_test')
        topic = request.data.get('topic')
        sub_topic = request.data.get('sub_topic')
        no_of_questions = request.data.get('no_of_questions')
        upload_type = request.data.get('upload_type')
        test_type = request.data.get('test_type')

        # Validate required fields
        if not all([question_paper_name, duration_of_test, topic, sub_topic, no_of_questions, upload_type, test_type]):
            return Response({'error': 'Missing fields for question_paper_master'}, status=status.HTTP_400_BAD_REQUEST)

        # Create a new question_paper_master entry
        question_paper_data = {
            'question_paper_name': question_paper_name,
            'duration_of_test': duration_of_test,
            'topic': topic,
            'sub_topic': sub_topic,
            'no_of_questions': no_of_questions,
            'upload_type': upload_type,
            'test_type': test_type
        }

        question_paper_serializer = questionsPaperSerializer(data=question_paper_data)
        if question_paper_serializer.is_valid():
            question_paper_instance = question_paper_serializer.save()
            question_paper_id = question_paper_instance.id
        else:
            return Response(question_paper_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        print('Files:', request.FILES)

        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)

        file = request.FILES['file']

        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())  # Print the first few rows of the DataFrame
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        # Mapping of Excel header names to actual column names
        header_mapping = {
            'Questions**': 'question_text',
            'Option A**': 'option_a',
            'Option B**': 'option_b',
            'Option C**': 'option_c',
            'Option D**': 'option_d',
            'Answer**': 'answer',
            'Mark**': 'mark',
            'Explain Answer': 'explain_answer',
        }
        
        # Rename columns in DataFrame to match expected column names
        df.rename(columns=header_mapping, inplace=True)
        

        # List of mandatory columns
        mandatory_columns = [
            'question_text',
            'option_a',
            'option_b',
            'option_c',
            'option_d',
            'answer',
            'mark'
        ]

        # Check for empty values in mandatory columns
        empty_columns = [col for col in mandatory_columns if df[col].isnull().any()]

        if empty_columns:
            error_message = f'Mandatory columns have null values: {", ".join(empty_columns)}'
            return Response({'error': error_message}, status=status.HTTP_400_BAD_REQUEST)

        
        records = df.to_dict(orient='records')

        for record in records:
            record['question_name_id'] = question_paper_id
            for key in record:
                if isinstance(record[key], str):
                    record[key] = record[key].strip()

        # Serialize and save the records
        serializer = questionsSerializerImport(data=records, many=True)
        print("Serializer: ", serializer)

        if serializer.is_valid():
            print("Before Saving..")
            serializer.save()
            print("After Saving..")
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        else:
            print("Serializer Errors: ", serializer.errors)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)



class ExcelImportView_Questions_Code1(APIView):
    def post(self, request, format=None):
        print('Files:', request.FILES)
        
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())  # Print the first few rows of the DataFrame
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        # Mapping of Excel header names to actual column names
        header_mapping = {
            'Questions**': 'question_name_id',
            'Answer**': 'answer',
            'Mark**': 'mark',
            'Explain Answer**': 'explain_answer',
            'Input Format': 'input_format'
        }
        
        # Rename columns in DataFrame to match expected column names
        df.rename(columns=header_mapping, inplace=True)
        

       # Remove 'topic_id' column from the DataFrame
        # if 'topic_id' in df.columns:
        #     df.drop(columns=['topic_id'], inplace=True)
        
        records = df.to_dict(orient='records')

        try:
            last_added_paper = question_paper_master.objects.filter(deleted=0).order_by('-id').first()
            if not last_added_paper:
                return JsonResponse({'message': 'No question papers available'}, status=404)
            
            question_name_id = last_added_paper.id

            for record in records:
                record['question_name_id'] = question_name_id
                # Ensure datetime fields are timezone-aware if they exist in the records
                if 'dtm_created' in record:
                    record['dtm_created'] = timezone.make_aware(record['dtm_created'], timezone.get_default_timezone())


        
            serializer = questionsSerializerCodeImport(data=records, many=True)
            print("Serializer: ", serializer)

            if serializer.is_valid():
                print("Before Saving..")
                serializer.save()
                print("After Saving..")
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            else:
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        except question_paper_master.DoesNotExist:
            return JsonResponse({'message': 'No question papers available'}, status=404)     


class ExcelImportView_Questions_Code_OLD_2607(APIView):
    def post(self, request, format=None):
        print('Files:', request.FILES)
        
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())  # Print the first few rows of the DataFrame
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        # Mapping of Excel header names to actual column names
        header_mapping = {
            'Questions**': 'question_text',
            'Answer**': 'answer',
            'Mark**': 'mark',
            'Explain Answer**': 'explain_answer',
            'Input Format': 'input_format'
        }
        
        # Rename columns in DataFrame to match expected column names
        df.rename(columns=header_mapping, inplace=True)

        # List of mandatory columns
        mandatory_columns = [
            'question_text',
            'answer',
            'mark',
            'explain_answer'
        ]

        # Check for empty values in mandatory columns
        empty_columns = [col for col in mandatory_columns if df[col].isnull().any()]

        if empty_columns:
            error_message = f'Mandatory columns have null values: {", ".join(empty_columns)}'
            return Response({'error': error_message}, status=status.HTTP_400_BAD_REQUEST)

        
        # Remove 'topic_id' column from the DataFrame if it exists
        # if 'topic_id' in df.columns:
        #     df.drop(columns=['topic_id'], inplace=True)
        
        records = df.to_dict(orient='records')

        question_paper_name = request.data.get('question_paper_name', None)
        if not question_paper_name:
            print('Question paper name is required')
            return Response({'error': 'Question paper name is required'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            last_added_paper = question_paper_master.objects.filter(question_paper_name=question_paper_name, deleted=0).order_by('-id').first()
            if not last_added_paper:
                print('No question papers available')
                return JsonResponse({'message': 'No question papers available'}, status=404)
            
            question_name_id = last_added_paper.id

            for record in records:
                record['question_name_id'] = question_name_id

                # Trim whitespace for each value in the record
                for key in record:
                    if isinstance(record[key], str):
                        record[key] = record[key].strip()


            serializer = questionsSerializerCodeImport(data=records, many=True)
            print("Serializer: ", serializer)

            if serializer.is_valid():
                print("Before Saving..")
                serializer.save()
                print("After Saving..")
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            else:
                print("Serializer Errors: ", serializer.errors)
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        except question_paper_master.DoesNotExist:
            return JsonResponse({'message': 'No question papers available'}, status=404)

class ExcelImportView_Questions_Code(APIView):
    def post(self, request, format=None):
        # Extract data for the question_paper_master
        question_paper_name = request.data.get('question_paper_name')
        duration_of_test = request.data.get('duration_of_test')
        topic = request.data.get('topic')
        sub_topic = request.data.get('sub_topic')
        no_of_questions = request.data.get('no_of_questions')
        upload_type = request.data.get('upload_type')
        test_type = request.data.get('test_type')

        # Validate required fields
        if not all([question_paper_name, duration_of_test, topic, sub_topic, no_of_questions, upload_type]):
            return Response({'error': 'Missing fields for question_paper_master'}, status=status.HTTP_400_BAD_REQUEST)

        # Create a new question_paper_master entry
        question_paper_data = {
            'question_paper_name': question_paper_name,
            'duration_of_test': duration_of_test,
            'topic': topic,
            'sub_topic': sub_topic,
            'no_of_questions': no_of_questions,
            'upload_type': upload_type,
            'test_type': test_type
        }

        question_paper_serializer = questionsPaperSerializer(data=question_paper_data)
        if question_paper_serializer.is_valid():
            question_paper_instance = question_paper_serializer.save()
            question_paper_id = question_paper_instance.id
        else:
            return Response(question_paper_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        # Process the uploaded file
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)

        file = request.FILES['file']

        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        # Map Excel columns to the expected format
        header_mapping = {
            'Questions**': 'question_text',
            'Answer**': 'answer',
            'Mark**': 'mark',
            'Explain Answer**': 'explain_answer',
            'Input Format': 'input_format'
        }

        df.rename(columns=header_mapping, inplace=True)

        mandatory_columns = [
            'question_text',
            'answer',
            'mark',
            'explain_answer'
        ]

        empty_columns = [col for col in mandatory_columns if df[col].isnull().any()]

        if empty_columns:
            error_message = f'Mandatory columns have null values: {", ".join(empty_columns)}'
            return Response({'error': error_message}, status=status.HTTP_400_BAD_REQUEST)

        # Convert DataFrame to records and associate with the newly created question_paper_master
        records = df.to_dict(orient='records')

        for record in records:
            record['question_name_id'] = question_paper_id
            for key in record:
                if isinstance(record[key], str):
                    record[key] = record[key].strip()

        # Serialize and save the records
        serializer = questionsSerializerCodeImport(data=records, many=True)
        print("Serializer: ", serializer)

        if serializer.is_valid():
            print("Before Saving..")
            serializer.save()
            print("After Saving..")
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        else:
            print("Serializer Errors: ", serializer.errors)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)





#------------------------------Rules-------------------------------------------------------------#


class rules_listView(generics.ListAPIView):
    queryset = rules.objects.filter(deleted=0)

    
    serializer_class = ruleSerializers

class rules_create(generics.CreateAPIView):
    queryset = rules.objects.all()
    serializer_class = ruleSerializers

class rules_Update(generics.RetrieveUpdateAPIView):
    queryset = rules.objects.all()
    serializer_class = ruleSerializers

@api_view(['PUT', 'PATCH'])
def delete_rules(request, pk):
    try:
        print("Entering Function..")
        rule = rules.objects.get(id=pk)

       
    except rules.DoesNotExist:
        return JsonResponse("rules not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    rule.deleted = 1
    rule.save()
 

    return JsonResponse("rules 'deleted' field updated successfully", safe=False)



#------------------------Dashboard---------------------------------#






def attendance_excel_data1(request):
    if request.method == 'GET':
        # Define URLs for each sheet
        sheet_urls = {
            'AIDS': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=257135364',
            'IT': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=747238312',
            'CSE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=1208914504',
            'ECE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=2128791655'
        }
        
        # Initialize an empty dictionary to store data for each sheet
        all_data = {}

        # Iterate over each sheet URL
        for sheet_name, url in sheet_urls.items():
            # Fetch the CSV data from the URL
            response = requests.get(url)

            # Read the CSV data into a Pandas DataFrame, skipping the first 8 rows
            df = pd.read_csv(io.StringIO(response.text), skiprows=7)

            # Transform the DataFrame into the desired format
            transformed_data = []
            for _, row in df.iterrows():
                transformed_row = {
                    'Department': row.get('DePArtment ', ''),  # Handle missing column gracefully
                    'Register Number': row.get('Register Number', ''),  # Handle missing column gracefully
                    'Name of the Student': row.get('NAme of the Student', '')
                }
                # Iterate over date columns and add FN and AN values
                for i in range(4, len(df.columns), 2):  # Start from index 4, increment by 2
                    col = df.columns[i]
                    date = col.split()[0]  # Extract the date
                    fn = row.get(col, '')  # Handle missing column gracefully
                    an_col_index = i + 1  # Get the index of the next column (AN)
                    if an_col_index < len(df.columns):
                        an = row.get(df.columns[an_col_index], '')  # Handle missing column gracefully
                    else:
                        an = ''  # If AN column does not exist, set it to empty string
                    transformed_row[date] = {'FN': fn, 'AN': an}
                transformed_data.append(transformed_row)

            # Store the transformed data in the dictionary
            all_data[sheet_name] = transformed_data

        return JsonResponse(all_data, safe=False)
    else:
        return JsonResponse({'error': 'Method not allowed'}, status=405)



#---------------------------------Password Reset----------------------------#


@api_view(['PUT'])
def update_login(request, user_name=None):
    login_to_update = login.objects.get(user_name=user_name)
    serializer = loginSerializer(instance=login_to_update, data=request.data, partial=True)
    
    if serializer.is_valid():
        serializer.save()
        print("Updated.")
        return JsonResponse("Login Updated Successfully", safe=False)
    return JsonResponse("Failed to Update Login")


@api_view(['PUT'])
def update_totalScore_test_candidate_map(request, pk=None):
    total_score_update = tests_candidates_map.objects.get(id=pk)
    serializer = testcandidatemapSerializers(instance=total_score_update, data=request.data, partial=True)

    if serializer.is_valid():
        serializer.save()
        print('Total score updated')
        return JsonResponse("Total score Updated Successfully", safe=False)
    return JsonResponse("Failed to Update total score")


@api_view(['PUT'])
def update_Avg_Marks_test_candidate_map(request, pk=None):
    avg_mark_update = tests_candidates_map.objects.get(id=pk)
    serializer = testcandidatemapSerializers(instance=avg_mark_update, data=request.data, partial=True)

    if serializer.is_valid():
        serializer.save()
        print('Avg mark is updated')
        return JsonResponse("Avg mark is  Updated Successfully", safe=False)
    return JsonResponse("Failed to Update Avg mark")




#---------------------------------Batch wise Test assign-------------------------------#


class TestCandidatesMapView_OLD(APIView):
    def post(self, request, batch_no, format=None):
        print("Batch No: ", batch_no)
        students = candidate_master.objects.filter(batch_no=batch_no)
        print("Students: ", students)
        data = []
        for student in students:
            print("Student id: ", student.id)
            test_candidate_data = {
                'test_id': request.data.get('test_id'),  # Assuming 'test_id' is provided in the request
                'question_id': request.data.get('question_id'),  # Assuming 'question_id' is provided in the request
                'student_id': student.id,
                'college_id': request.data.get('college_id'),
                'department_id': request.data.get('department_id'),
                'dtm_start': request.data.get('dtm_start'),
                'dtm_end': request.data.get('dtm_end'),
                'is_actual_test': request.data.get('is_actual_test'),
                'duration': request.data.get('duration'),
                'year': request.data.get('year'),
                'rules_id': request.data.get('rules_id'),
                'need_candidate_info': request.data.get('need_candidate_info'),
                # Add other fields here
            }

            serializer = testcandidatemapSerializers(data=test_candidate_data)
            print("Serializer: ", serializer)
            if serializer.is_valid():
                serializer.save()
                data.append(serializer.data)
                print("Data.append: ", data)
            else:
                print(serializer.errors)
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        print("Data: ", data)
        return Response(data, status=status.HTTP_201_CREATED)



class TestCandidatesMapView(APIView):
    def post(self, request, format=None):
        college_id = request.data.get('college_id', [])
        department_id = request.data.get('department_id', [])
        year = request.data.get('year')

        if not college_id or not department_id or not year:
            return Response({'error': 'college_id, department_id, and year are required fields.'}, status=status.HTTP_400_BAD_REQUEST)

        print("Batch No: ", year, college_id, department_id)

        # Filter candidates based on multiple college_ids and department_ids
        students = candidate_master.objects.filter(
            college_id__in=college_id,
            department_id__in=department_id,
            year=year
        )

        print("Students: ", students)
        data = []
        updated_candidates = []  # To store IDs of updated candidates

        current_date_and_time = datetime.now()  # Get the current date and time

        for student in students:
            print("Student id: ", student.id)
            test_candidate_data = {
                'test_name': request.data.get('test_name'),  # Assuming 'test_name' is provided in the request
                'question_id': request.data.get('question_id'),  # Assuming 'question_id' is provided in the request
                'student_id': student.id,
                'college_id': student.college_id.id,  # Use student.college_id.id for pk
                'department_id': student.department_id.id,  # Use student.department_id.id for pk
                'dtm_start': request.data.get('dtm_start'),
                'dtm_end': request.data.get('dtm_end'),
                'is_actual_test': request.data.get('is_actual_test'),
                'duration': request.data.get('duration'),
                'duration_type': request.data.get('duration_type'),
                'year': year,
                'rules_id': request.data.get('rules_id'),
                'need_candidate_info': request.data.get('need_candidate_info'),
                'dtm_created': current_date_and_time  # Add current date and time
            }

            serializer = testcandidatemapSerializers(data=test_candidate_data)
            print("Serializer: ", serializer)
            if serializer.is_valid():
                with transaction.atomic():  # Use transaction to ensure atomicity
                    serializer.save()
                    data.append(serializer.data)
                    print("Data.append: ", data)
    
                    # Check and update candidate_master if needed
                    if student.need_candidate_info is None and request.data.get('need_candidate_info') is True:
                        student.need_candidate_info = request.data.get('need_candidate_info')
                        student.save(update_fields=['need_candidate_info'])
                        updated_candidates.append(student.id)
            else:
                print(serializer.errors)
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        print("Data: ", data)
        return Response(data, status=status.HTTP_201_CREATED)





#---------------Non database test assign-------------------#

class NonDbTestAssign_OLD(APIView):
    def post(self, request):
        students = candidate_master.objects.filter(
            Q(college_id__isnull=True),
            Q(department_id__isnull=True) ,
            Q(year__isnull=True) | Q(year=''),
            Q(batch_no__isnull=True) | Q(batch_no=''),
            Q(students_name__isnull=True) | Q(students_name=''),
            Q(registration_number__isnull=True) | Q(registration_number=''),
            Q(gender__isnull=True) | Q(gender=''),
            Q(email_id__isnull=True) | Q(email_id=''),
            Q(mobile_number__isnull=True) | Q(mobile_number=''),
            Q(cgpa__isnull=True) | Q(cgpa=''),
            Q(marks_10th__isnull=True) | Q(marks_10th=''),
            Q(marks_12th__isnull=True) | Q(marks_12th=''),
            Q(marks_semester_wise__isnull=True) | Q(marks_semester_wise=''),
            Q(history_of_arrears__isnull=True) | Q(history_of_arrears=''),
            Q(standing_arrears__isnull=True),
            Q(number_of_offers__isnull=True),
            Q(text__isnull=True) | Q(text=''),
            ~Q(user_name__isnull=True) & ~Q(user_name=''),
        )
        
        print("Students: ", students)
        data = []
        updated_candidates = []  # To store IDs of updated candidates

        current_date_and_time = datetime.now() 

        for student in students:
            print("Student id: ", student.id)
            test_candidate_data = {
                'test_name': request.data.get('test_name'),  # Assuming 'test_id' is provided in the request
                'question_id': request.data.get('question_id'),  # Assuming 'question_id' is provided in the request
                'student_id': student.id,
                'dtm_start': request.data.get('dtm_start'),
                'dtm_end': request.data.get('dtm_end'),
                'is_actual_test': request.data.get('is_actual_test'),
                'duration': request.data.get('duration'),
                'duration_type': request.data.get('duration_type'),
                'rules_id': request.data.get('rules_id'),
                'need_candidate_info': request.data.get('need_candidate_info'),
                'dtm_created': current_date_and_time  # Add current date and time
                # Add other fields here
            }

            serializer = NonDbTestAssignSerializer(data=test_candidate_data)
            print("Serializer: ", serializer)
            if serializer.is_valid():
                with transaction.atomic():  # Use transaction to ensure atomicity
                    serializer.save()
                    data.append(serializer.data)
                    print("Data.append: ", data)

                    # Check and update candidate_master if needed
                    if student.need_candidate_info is None and request.data.get('need_candidate_info') is True:
                        student.need_candidate_info = request.data.get('need_candidate_info')
                        student.save(update_fields=['need_candidate_info'])
                        updated_candidates.append(student.id)
            else:
                print(serializer.errors)
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        print("Data: ", data)
        return Response(data, status=status.HTTP_201_CREATED)


class NonDbTestAssign(APIView):
    def post(self, request):
        # Use Q objects to build the filter query dynamically
        queries = [
            Q(college_id__isnull=True),
            Q(department_id__isnull=True),
            Q(year__isnull=True) | Q(year=''),
            Q(batch_no__isnull=True) | Q(batch_no=''),
            Q(students_name__isnull=True) | Q(students_name=''),
            Q(registration_number__isnull=True) | Q(registration_number=''),
            Q(gender__isnull=True) | Q(gender=''),
            Q(email_id__isnull=True) | Q(email_id=''),
            Q(mobile_number__isnull=True) | Q(mobile_number=''),
            Q(history_of_arrears__isnull=True),
            Q(standing_arrears__isnull=True),
            Q(number_of_offers__isnull=True),
            Q(text__isnull=True) | Q(text=''),
            ~Q(user_name__isnull=True) & ~Q(user_name=''),
        ]

        # Handle decimal fields
        decimal_fields = ['cgpa', 'marks_10th', 'marks_12th']
        for field in decimal_fields:
            queries.append(Q(**{f"{field}__isnull": True}) | Q(**{field: None}))

        # Filter the students based on the constructed query
        students = candidate_master.objects.filter(*queries)
        
        data = []
        updated_candidates = []  # To store IDs of updated candidates
        current_date_and_time = datetime.now()

        for student in students:
            test_candidate_data = {
                'test_name': request.data.get('test_name'),
                'question_id': request.data.get('question_id'),
                'student_id': student.id,
                'dtm_start': request.data.get('dtm_start'),
                'dtm_end': request.data.get('dtm_end'),
                'is_actual_test': request.data.get('is_actual_test'),
                'duration': request.data.get('duration'),
                'duration_type': request.data.get('duration_type'),
                'rules_id': request.data.get('rules_id'),
                'need_candidate_info': request.data.get('need_candidate_info'),
                'dtm_created': current_date_and_time
            }

            # Clean the input data
            for field in decimal_fields:
                if test_candidate_data.get(field) == '':
                    test_candidate_data[field] = None

            serializer = NonDbTestAssignSerializer(data=test_candidate_data)
            if serializer.is_valid():
                with transaction.atomic():  # Use transaction to ensure atomicity
                    serializer.save()
                    data.append(serializer.data)

                    # Check and update candidate_master if needed
                    if student.need_candidate_info is None and request.data.get('need_candidate_info') is True:
                        student.need_candidate_info = request.data.get('need_candidate_info')
                        student.save(update_fields=['need_candidate_info'])
                        updated_candidates.append(student.id)
            else:
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        return Response(data, status=status.HTTP_201_CREATED)







#-----------------------Get Unique Batch---------------------#

@api_view(['GET'])
def batch_list(request):
    batch_numbers = candidate_master.objects.exclude(batch_no__isnull=True).values_list('batch_no', flat=True).distinct()
    return Response({'batch_numbers': list(batch_numbers)})

@api_view(['GET'])
def question_name_list(request):
    question_names = question_master.objects.exclude(question_name__isnull=True).values_list('question_name', flat=True).distinct()
    return Response({'question_new': list(question_names)})

@api_view(['GET'])
def topic_list(request):
    topics = content_master.objects.filter(deleted=0).exclude(topic__isnull=True).values_list('topic', flat=True).distinct()
    return Response({'topics': list(topics)})


@api_view(['GET'])
def unique_test_type(request):
    test_types = test_type.objects.filter(deleted=0).values_list('test_type', flat=True).distinct()
    return Response({'test_type': list(test_types)})


@api_view(['GET'])
def MCQ_test_type(request):
    test_types = test_type.objects.filter(deleted=0, test_type='MCQ Test').values_list('test_type_categories', flat=True).distinct()

    return Response({"test_type": list(test_types)})

@api_view(['GET'])
def Coding_test_type(request):
    test_types = test_type.objects.filter(deleted=0, test_type='Coding Test').values_list('test_type_categories', flat=True).distinct()

    return Response({"test_type": list(test_types)})


@api_view(['GET'])
def unique_question_type(request):
    question_types = question_type.objects.filter(deleted=0).values_list('question_type', flat=True).distinct()

    return Response({"question_types": list(question_types)})




class sidebarMenu(generics.ListAPIView):
    queryset = appinfo.objects.filter(deleted=0, info_code='MAIN_MENU')
    serializer_class = appinfoSerializer


@api_view(['GET'])
def main_menu_Training(request):
    menus = appinfo.objects.filter(deleted=0, info_code='MAIN_MENU').values_list('info_value', flat=True).distinct()

    return Response({"MAIN_MENU": list(menus)})




def fetch_csv_with_retries(url, retries=3, delay=5):
    for attempt in range(retries):
        try:
            response = requests.get(url)
            response.raise_for_status()  # Raise an HTTPError for bad responses
            return response.text
        except requests.exceptions.RequestException as e:
            if attempt < retries - 1:  # If not the last attempt
                time.sleep(delay)
            else:
                raise e  # Re-raise the last exception

def attendance_excel_data(request):
    if request.method == 'GET':
        # Define URLs for each sheet
        sheet_urls = {
            'AIDS': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=257135364',
            'IT': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=747238312',
            'CSE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=1208914504',
            'ECE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=2128791655'
        }
        
        # Initialize an empty dictionary to store data for each sheet
        all_data = {}

        # Iterate over each sheet URL
        for sheet_name, url in sheet_urls.items():
            try:
                # Fetch the CSV data with retries
                csv_data = fetch_csv_with_retries(url)
                # Read the CSV data into a Pandas DataFrame, skipping the first 8 rows
                df = pd.read_csv(io.StringIO(csv_data), skiprows=7)

                # Transform the DataFrame into the desired format
                transformed_data = []
                for _, row in df.iterrows():
                    transformed_row = {
                        'Department': row.get('DePArtment ', ''),  # Handle missing column gracefully
                        'Register Number': row.get('Register Number', ''),  # Handle missing column gracefully
                        'Name of the Student': row.get('NAme of the Student', '')
                    }
                    # Iterate over date columns and add FN and AN values
                    for i in range(4, len(df.columns), 2):  # Start from index 4, increment by 2
                        col = df.columns[i]
                        date = col.split()[0]  # Extract the date
                        fn = row.get(col, '')  # Handle missing column gracefully
                        an_col_index = i + 1  # Get the index of the next column (AN)
                        if an_col_index < len(df.columns):
                            an = row.get(df.columns[an_col_index], '')  # Handle missing column gracefully
                        else:
                            an = ''  # If AN column does not exist, set it to empty string
                        transformed_row[date] = {'FN': fn, 'AN': an}
                    transformed_data.append(transformed_row)

                # Store the transformed data in the dictionary
                all_data[sheet_name] = transformed_data

            except requests.exceptions.RequestException as e:
                # Log the error and add an error message to the response
                all_data[sheet_name] = {'error': str(e)}
                print(f"Error fetching data for sheet {sheet_name}: {e}")

        return JsonResponse(all_data, safe=False)
    else:
        return JsonResponse({'error': 'Method not allowed'}, status=405)


class candidates_one_Update(generics.RetrieveUpdateAPIView):
    queryset = candidate_master.objects.all()
    serializer_class = candidatesoneSerializer


#------------------------------------------Question Paper Master----------------------------------#

@api_view(['GET'])
def get_question_paper(request):
    questionset=question_paper_master.objects.filter(deleted=0).order_by(
            Case(
                When(dtm_created__isnull=True, then=1),
                default=0,
                output_field=DateTimeField()
            ),
            '-dtm_created'
        )
    #logger.info("Fetching question paper where deleted=0")
   
    questions_data=[]
    for questions in  questionset:
        
        questions_data.append({
            'id': questions.id,
            'question_paper_name': questions.question_paper_name,
           
            'topic': questions.topic,
            'sub_topic':questions.sub_topic,
            'test_type':questions.test_type,
            'duration_of_test': questions.duration_of_test,
            'no_of_questions': questions.no_of_questions,
           
            'upload_type': questions.upload_type,
            'dtm_cretaed': questions.dtm_created,

            })
    
    # Sort the questions_data list in descending order based on 'id'
    # questions_data = sorted(questions_data, key=lambda x: x['id'], reverse=True)

    return Response(questions_data)



class questionpapercreateAPIView(generics.CreateAPIView):
    queryset = question_paper_master.objects.all()
    serializer_class =questionsPaperSerializer

    def post(self, request, *args, **kwargs):
        print("Request data:", request.data)  
        #logger.info("Creating a new question paper")
        response = super().post(request, *args, **kwargs)
        #logger.info("Created a new question paper successfully")
        return response

class questionpaperUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = question_paper_master.objects.all()
    serializer_class =questionsPaperSerializer

    def put(self, request, *args, **kwargs):
        #logger.info(f"Updating Question Paper with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        #logger.info(f"Updated Question Paper with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        #logger.info(f"Partially updating Question Paper with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        #logger.info(f"Partially updated Question Paper with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_question_paper(request, pk):
    #logger.info(f"Attempting to mark Courses with id {pk} as deleted")
    try:
        print("Entering Function..")
        questions=question_paper_master.objects.get(id=pk)

        print("questions: ",questions)
    except question_paper_master.DoesNotExist:
        return JsonResponse("Question Paper not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    questions.deleted = 1
    questions.save()

    #logger.info(f"Marked Courses with id {pk} as deleted successfully")

    print("questions: ",questions)

    return JsonResponse("questions paper 'deleted' field updated successfully", safe=False)



def get_last_added_question_paper(request):
    try:
        last_added_paper = question_paper_master.objects.order_by('-id').first()
        serializer = questionsPaperSerializer(last_added_paper)
        return JsonResponse(serializer.data)
    except question_paper_master.DoesNotExist:
        return JsonResponse({'message': 'No question papers available'}, status=404)



#---------------------Questions Master with images--------------#

@api_view(['GET'])
def get_questions_IO_OLD(request):
    questionset=question_master.objects.filter(deleted=0)
    #logger.info("Fetching Questions where deleted=0")
   
    question_data=[]
    for question in  questionset:
        question_name_id = None
        question_paper_name = None
      
        if question.question_name_id:
           question_name_id=question.question_name_id.id           
           question_paper_name = question.question_name_id.question_paper_name

           question_data.append({
              'id': question.id,
              'question_name_id':question_name_id,
              'question_paper_name':question_paper_name,
              'question_text':question.question_text,
              'question_image_data': question.question_image_data,
              'option_a_image_data': question.option_a_image_data,
              'option_b_image_data': question.option_b_image_data,
              'option_c_image_data': question.option_c_image_data,
              'option_d_image_data': question.option_d_image_data,
              'option_a':question.option_a,
              'option_b':question.option_b,
              'option_c':question.option_c,
              'option_d':question.option_d,
              'view_hint':question.view_hint,
              'mark':question.mark,
              'explain_answer':question.explain_answer,
              'answer':question.answer,
              'negative_mark':question.negative_mark,
             'input_format': question.input_format,


            })
    return Response(question_data)


class questionscreateAPIView_IO_OLD(generics.CreateAPIView):
    queryset = question_master.objects.all()
    serializer_class = questionsSerializer_IO

    def post(self, request, *args, **kwargs):
        #logger.info("Creating a new Questions")
        response = super().post(request, *args, **kwargs)
        #logger.info("Created a new Questions successfully")
        return response

class questionsRetrieveUpdateAPIView_IO_OLD(generics.RetrieveUpdateAPIView):
    queryset = question_master.objects.all()
    serializer_class = questionsSerializer_IO


    def update(self, request, *args, **kwargs):
        instance = self.get_object()
        response = super().update(request, *args, **kwargs)
        
        # Update the dtm_created field of the related question_paper_master
        question_paper = instance.question_name_id  # Adjust this based on your ForeignKey field name
        question_paper.dtm_created = datetime.now()
        question_paper.save()

        return response

    def put(self, request, *args, **kwargs):
        # logger.info(f"Updating Questions with id {kwargs.get('pk')}")
        response = self.update(request, *args, **kwargs)
        # logger.info(f"Updated Questions with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        # logger.info(f"Partially updating Questions with id {kwargs.get('pk')}")
        response = self.update(request, *args, **kwargs)
        # logger.info(f"Partially updated Questions with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_question_IO_OLD(request, pk):
    #logger.info(f"Attempting to mark Questions with id {pk} as deleted")
    try:
        print("Entering Function..")
        question = question_master.objects.get(id=pk)

        print("skill: ", question)
    except question_master.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    question.deleted = 1
    question.save()

    #logger.info(f"Marked Questions with id {pk} as deleted successfully")

    print("skill: ", question)

    return JsonResponse("skill 'deleted' field updated successfully", safe=False)



#-------------------------------------------------------------#


#---------------------Questions Master with images--------------#
def get_questions_IO(request):
    questionset = question_master.objects.filter(deleted=0)
    #logger.info("Fetching Questions where deleted=0")
   
    question_data = []
    for question in questionset:
        question_name_id = None
        question_paper_name = None
      
        if question.question_name_id:
            question_name_id = question.question_name_id.id           
            question_paper_name = question.question_name_id.question_paper_name

        question_image_data = None
        if question.question_image_data:
            question_image_data = base64.b64encode(question.question_image_data).decode('utf-8')
        
        option_a_image_data = None
        if question.option_a_image_data:
            option_a_image_data = base64.b64encode(question.option_a_image_data).decode('utf-8')

        option_b_image_data = None
        if question.option_b_image_data:
            option_b_image_data = base64.b64encode(question.option_b_image_data).decode('utf-8')

        option_c_image_data = None
        if question.option_c_image_data:
            option_c_image_data = base64.b64encode(question.option_c_image_data).decode('utf-8')

        option_d_image_data = None
        if question.option_d_image_data:
            option_d_image_data = base64.b64encode(question.option_d_image_data).decode('utf-8')

        question_data.append({
            'id': question.id,
            'question_name_id': question_name_id,
            'question_paper_name': question_paper_name,
            'question_text': question.question_text,
            'question_image_data': question_image_data,
            'option_a_image_data': option_a_image_data,
            'option_b_image_data': option_b_image_data,
            'option_c_image_data': option_c_image_data,
            'option_d_image_data': option_d_image_data,
            'option_a': question.option_a,
            'option_b': question.option_b,
            'option_c': question.option_c,
            'option_d': question.option_d,
            'view_hint': question.view_hint,
            'mark': question.mark,
            'explain_answer': question.explain_answer,
            'answer': question.answer,
            'negative_mark': question.negative_mark,
            'input_format': question.input_format,
            
        })
    random.shuffle(question_data)
    return JsonResponse(question_data, safe=False)


@csrf_exempt
def upload_question(request):
    print('***1')
    if request.method == 'POST':
        print('****2')
        print('POST data:', request.POST)
        print('FILES data:', request.FILES)

        form = QuestionForm(request.POST, request.FILES)
        print('****3')
        if form.is_valid():
            print('*****4')
            question = form.save(commit=False)
            if 'question_image_data' in request.FILES:
                question.question_image_data = request.FILES['question_image_data'].read()
            if 'option_a_image_data' in request.FILES:
                question.option_a_image_data = request.FILES['option_a_image_data'].read()
            if 'option_b_image_data' in request.FILES:
                question.option_b_image_data = request.FILES['option_b_image_data'].read()
            if 'option_c_image_data' in request.FILES:
                question.option_c_image_data = request.FILES['option_c_image_data'].read()
            if 'option_d_image_data' in request.FILES:
                question.option_d_image_data = request.FILES['option_d_image_data'].read()
            question.save()
            print('Question Created Successfully')
            return HttpResponse("Question created successfully")
    else:
        print('***1')
        form = QuestionForm()
        print('******2')
    return render(request, 'upload_image.html', {'form': form})



def update_question(request, question_id):
    question = get_object_or_404(question_master, id=question_id, deleted=0)

    if request.method == 'POST':
        form = QuestionForm(request.POST, request.FILES, instance=question)
        if form.is_valid():
            question = form.save(commit=False)
            if 'question_image_data' in request.FILES:
                question.question_image_data = request.FILES['question_image_data'].read()
            if 'option_a_image_data' in request.FILES:
                question.option_a_image_data = request.FILES['option_a_image_data'].read()
            if 'option_b_image_data' in request.FILES:
                question.option_b_image_data = request.FILES['option_b_image_data'].read()
            if 'option_c_image_data' in request.FILES:
                question.option_c_image_data = request.FILES['option_c_image_data'].read()
            if 'option_d_image_data' in request.FILES:
                question.option_d_image_data = request.FILES['option_d_image_data'].read()
            else:
                question.question_image_data = None
            question.save()
            return HttpResponse("Question updated successfully")
    else:
        form = QuestionForm(instance=question)

    return render(request, 'update_image.html', {'form': form})

def delete_question_IO(request, pk):
    #logger.info(f"Attempting to mark Questions with id {pk} as deleted")
    try:
        question = question_master.objects.get(id=pk)
    except question_master.DoesNotExist:
        return JsonResponse("Question not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    question.deleted = 1
    question.save()

    #logger.info(f"Marked Questions with id {pk} as deleted successfully")

    return JsonResponse("Question 'deleted' field updated successfully", safe=False)

def get_last_added_question_paper(request):
    try:
        last_added_paper = question_paper_master.objects.order_by('-id').first()
        serializer = questionsPaperSerializer(last_added_paper)
        return JsonResponse(serializer.data)
    except question_paper_master.DoesNotExist:
        return JsonResponse({'message': 'No question papers available'}, status=404)
    #QUESTION_Code_master_#

def get_questions_Code(request):
    questionset = question_master.objects.filter(deleted=0)
    #logger.info("Fetching Questions where deleted=0")
   
    question_data = []
    for question in questionset:
        question_name_id = None
        question_paper_name = None
      
        if question.question_name_id:
            question_name_id = question.question_name_id.id           
            question_paper_name = question.question_name_id.question_paper_name

        question_image_data = None
        if question.question_image_data:
            question_image_data = base64.b64encode(question.question_image_data).decode('utf-8')
        
        question_data.append({
            'id': question.id,
            'question_name_id': question_name_id,
            'question_paper_name': question_paper_name,
            'question_text': question.question_text,
            'question_image_data': question_image_data,

            
            'view_hint': question.view_hint,
            'mark': question.mark,
            'explain_answer': question.explain_answer,
            'answer': question.answer,
            'negative_mark': question.negative_mark,
            'input_format': question.input_format,
        })
    random.shuffle(question_data)
    return JsonResponse(question_data, safe=False)

def upload_question_code(request):
    if request.method == 'POST':
        form = QuestionCodeForm(request.POST, request.FILES)
        if form.is_valid():
            question = form.save(commit=False)
            if 'question_image_data' in request.FILES:
                question.question_image_data = request.FILES['question_image_data'].read()
            else:
                question.question_image_data = None
                question.save()
            return HttpResponse("Question created successfully")
    else:
        form = QuestionCodeForm()
    return render(request, 'upload_image.html', {'form': form})

def update_question_code(request, question_id):
    question = get_object_or_404(question_master, id=question_id, deleted=0)

    if request.method == 'POST':
        form = QuestionCodeForm(request.POST, request.FILES, instance=question)
        if form.is_valid():
            question = form.save(commit=False)
            if 'question_image_data' in request.FILES:
                question.question_image_data = request.FILES['question_image_data'].read()
            question.save()
            return HttpResponse("Question updated successfully")
    else:
        form = QuestionCodeForm(instance=question)

    return render(request, 'update_image.html', {'form': form})



#-----------------------Getting All questions where question_paper_id--------------------------#



@api_view(['GET'])
def get_questions_Qp_id(request, question_name_id):
    questionset = question_master.objects.filter(deleted=0, question_name_id=question_name_id)

    question_data = []
    for question in questionset:
        question_name_id = None
        question_paper_name = None
      
        if question.question_name_id:
            question_name_id = question.question_name_id.id           
            question_paper_name = question.question_name_id.question_paper_name

        question_image_data = None
        if question.question_image_data:
            question_image_data = base64.b64encode(question.question_image_data).decode('utf-8')
        
        option_a_image_data = None
        if question.option_a_image_data:
            option_a_image_data = base64.b64encode(question.option_a_image_data).decode('utf-8')

        option_b_image_data = None
        if question.option_b_image_data:
            option_b_image_data = base64.b64encode(question.option_b_image_data).decode('utf-8')

        option_c_image_data = None
        if question.option_c_image_data:
            option_c_image_data = base64.b64encode(question.option_c_image_data).decode('utf-8')

        option_d_image_data = None
        if question.option_d_image_data:
            option_d_image_data = base64.b64encode(question.option_d_image_data).decode('utf-8')

        question_data.append({
            'id': question.id,
            'question_name_id': question_name_id,
            'question_paper_name': question_paper_name,
            'question_text': question.question_text,
            'question_image_data': question_image_data,
            'option_a_image_data': option_a_image_data,
            'option_b_image_data': option_b_image_data,
            'option_c_image_data': option_c_image_data,
            'option_d_image_data': option_d_image_data,
            'option_a': question.option_a,
            'option_b': question.option_b,
            'option_c': question.option_c,
            'option_d': question.option_d,
            'view_hint': question.view_hint,
            'mark': question.mark,
            'explain_answer': question.explain_answer,
            'answer': question.answer,
            'negative_mark': question.negative_mark,
            'input_format': question.input_format,
            
        })
    return JsonResponse(question_data, safe=False)



class questionsRetrieveUpdateAPIView_IO_code(generics.RetrieveUpdateAPIView):
    queryset = question_master.objects.all()
    serializer_class = questionsSerializer_code


    def update(self, request, *args, **kwargs):
        instance = self.get_object()
        response = super().update(request, *args, **kwargs)
        
        # Update the dtm_created field of the related question_paper_master
        question_paper = instance.question_name_id  # Adjust this based on your ForeignKey field name
        print('question_paper: ', question_paper)
        question_paper.dtm_created = datetime.now()
        question_paper.save()
        print('question_paper_saved....')

        return response

    def put(self, request, *args, **kwargs):
        # logger.info(f"Updating Questions with id {kwargs.get('pk')}")
        response = self.update(request, *args, **kwargs)
        # logger.info(f"Updated Questions with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        # logger.info(f"Partially updating Questions with id {kwargs.get('pk')}")
        response = self.update(request, *args, **kwargs)
        # logger.info(f"Partially updated Questions with id {kwargs.get('pk')} successfully")
        return response



class questionscreateAPIView_IO_code(generics.CreateAPIView):
    queryset = question_master.objects.all()
    serializer_class = questionsSerializer_code

    def post(self, request, *args, **kwargs):
        #logger.info("Creating a new Questions")
        response = super().post(request, *args, **kwargs)
        #logger.info("Created a new Questions successfully")
        return response




@api_view(['PUT'])
def update_Need_Candidate_info(request, studentID=None):
    n_info_list = tests_candidates_map.objects.filter(student_id=studentID)
    print('request.data: ', request.data)

    if not n_info_list.exists():
        return JsonResponse("No candidate info found for the given student ID", safe=False)

    for n_info in n_info_list:
        serializer = testcandidatemapSerializers(instance=n_info, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            print('serializer: ', serializer)
        else:
            return JsonResponse("Failed to Update Need Candidate info, clg, dept, year", safe=False)

    print('Need Candidate info, clg, dept, year updated')
    return JsonResponse("Need Candidate info, clg, dept, year Updated Successfully", safe=False)


@api_view(['PUT'])
def update_clg_login(request, userName=None):
    clg = login.objects.get(user_name=userName)
    print('request.data: ', request.data)
    serializer = loginSerializer(instance=clg, data=request.data, partial=True)

    if serializer.is_valid():
        serializer.save()
        print('serializer: ', serializer)
        print('College updated')
        return JsonResponse("College Updated Successfully", safe=False)
    return JsonResponse("Failed to Update College. ")




@api_view(['GET'])
def get_candidate_login(request):
    candidatelist = candidate_master.objects.filter(deleted=0)
    logins = login.objects.all()  # Remove the filter for 'deleted' since the login model doesn't have this field

    #logger.info("Fetching candidates where deleted=0")

    candidate_data = []
    for candidate in candidatelist:
        # Filter the logins for the specific candidate's username
        candidate_login = logins.filter(user_name=candidate.user_name).first()
        password = candidate_login.password if candidate_login else None

        candidate_data.append({
            'student_id': candidate.id,
            'user_name': candidate.user_name,
            'student_name': candidate.students_name,
            'password': password
        })

    return Response(candidate_data)






class TestAssignFor_Selected(APIView):
    def post(self, request):
        stu_id = request.data.get('stu_id', [])  # Fetching stu_id from request data
        if not isinstance(stu_id, list):
            return Response({"error": "stu_id must be a list"}, status=status.HTTP_400_BAD_REQUEST)

        data = []
        for s_id in stu_id:
            try:
                student = candidate_master.objects.get(id=s_id)
            except candidate_master.DoesNotExist:
                return Response({"error": f"Student with id {s_id} does not exist"}, status=status.HTTP_400_BAD_REQUEST)

            print("Students: ", student)
            print("Student id: ", student.id)
            
            # Preparing test candidate data
            test_candidate_data = {
                'test_name': request.data.get('test_name'),  
                'question_id': request.data.get('question_id'), 
                'college_id': student.college_id.id if student.college_id else None,
                'department_id': student.department_id.id if student.department_id else None,
                'year': student.year if student.year else None, 
                'student_id': student.id,
                'dtm_start': request.data.get('dtm_start'),
                'dtm_end': request.data.get('dtm_end'),
                'is_actual_test': request.data.get('is_actual_test'),
                'duration': request.data.get('duration'),
                'duration_type': request.data.get('duration_type'),
                'rules_id': request.data.get('rules_id'),
                'need_candidate_info': student.college_id is None or student.department_id is None or student.year is None,
                'dtm_created': request.data.get('dtm_created')
            }

            serializer = testcandidatemapSerializers(data=test_candidate_data)
            print("Serializer: ", serializer)
            if serializer.is_valid():
                serializer.save()
                data.append(serializer.data)
                print("Data.append: ", data)
            else:
                print(serializer.errors)
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        print("Data: ", data)
        return Response(data, status=status.HTTP_201_CREATED)




@api_view(['GET'])
def get_group_test_name_OLD(request):
    # Join tests_candidates_map with test_master and filter by test_name and deleted=0
    tests_candidates = (
        tests_candidates_map.objects.filter(deleted=0)
        .values('test_name', 'question_id', 'question_id__question_paper_name', 'dtm_start', 'dtm_end', 'dtm_created')
        .annotate(student_count=Count('student_id'))
        .order_by('-dtm_created')
    )

    test_candidate_map_data = []
    for testing in tests_candidates:
        # Check if test_master with the same test_name exists
        if test_master.objects.filter(test_name=testing['test_name']).exists():
            # Get the test_master record to access its ID
            test_master_record = test_master.objects.get(test_name=testing['test_name'])
            
            # Format dates
            dtm_start_formatted = django_format_date(localtime(testing['dtm_start']), 'd-m-Y h:i A')
            dtm_end_formatted = django_format_date(localtime(testing['dtm_end']), 'd-m-Y h:i A')

            test_candidate_map_data.append({
                'test_name': testing['test_name'],
                'question_id': testing['question_id'],
                'question_paper_name': testing['question_id__question_paper_name'],
                'dtm_start': dtm_start_formatted,
                'dtm_end': dtm_end_formatted,
                'student_count': testing['student_count'],
                'test_master_id': test_master_record.id,
                'dtm_created': testing['dtm_created']
            })

    # Sort the results by test_master_id
    # test_candidate_map_data = sorted(test_candidate_map_data, key=lambda x: -x['test_master_id'])

    return Response(test_candidate_map_data)



@api_view(['GET'])
def get_group_test_name(request):
    # Join tests_candidates_map with test_master and filter by test_name and deleted=0
    tests_candidates = (
        tests_candidates_map.objects.filter(deleted=0)
        .values('test_name', 'question_id', 'question_id__question_paper_name', 'dtm_start', 'dtm_end', 'dtm_created')
        .annotate(student_count=Count('student_id'))
        .order_by(
            Case(
                When(dtm_created__isnull=True, then=1),
                default=0,
                output_field=DateTimeField()
            ),
            '-dtm_created'
        )
    )

    test_candidate_map_data = []
    for testing in tests_candidates:
        # Format dates
        dtm_start_formatted = django_format_date(localtime(testing['dtm_start']), 'd-m-Y h:i A')
        dtm_end_formatted = django_format_date(localtime(testing['dtm_end']), 'd-m-Y h:i A')

        test_candidate_map_data.append({
            'test_name': testing['test_name'],
            'question_id': testing['question_id'],
            'question_paper_name': testing['question_id__question_paper_name'],
            'dtm_start': dtm_start_formatted,
            'dtm_end': dtm_end_formatted,
            'student_count': testing['student_count'],
            'dtm_created': testing['dtm_created']
        })

    return Response(test_candidate_map_data)






#_____________________________________________Company_master______________________________________________


class company_listView(generics.ListAPIView):
    queryset = company_master.objects.filter(deleted=0)
    serializer_class = companySerializer

    def get(self, request, *args, **kwargs):
        #logger.info("Fetching test types where deleted=0")
        return super().get(request, *args, **kwargs)

class company_create(generics.CreateAPIView):
    queryset = company_master.objects.all()
    serializer_class = companySerializer

    def post(self, request, *args, **kwargs):
        #logger.info("Creating a new test type")
        return super().post(request, *args, **kwargs)
class company_master_delete(generics.RetrieveDestroyAPIView):
    queryset = company_master.objects.all()
    serializer_class = companySerializer

   


class company_master_Update(generics.RetrieveUpdateAPIView):
    queryset = company_master.objects.all()
    serializer_class = companySerializer

    def put(self, request, *args, **kwargs):
        #logger.info(f"Updating test type with id {kwargs.get('pk')}")
        return super().put(request, *args, **kwargs)

    def patch(self, request, *args, **kwargs):
        #logger.info(f"Partially updating test type with id {kwargs.get('pk')}")
        return super().patch(request, *args, **kwargs)

@api_view(['PUT', 'PATCH'])
def delete_company_master(request, pk):
    try:
        #logger.info(f"Attempting to mark test type with id {pk} as deleted")
        company_mastervar = company_master.objects.get(id=pk)
    except company_master.DoesNotExist:
        logger.error(f"Test type with id {pk} not found")
        return JsonResponse("company_master not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    company_mastervar.deleted = 1
    company_mastervar.save()

    #logger.info(f"Marked test type with id {pk} as deleted")
    return JsonResponse("company_master 'deleted' field updated successfully", safe=False)

#______________________________________________________job_master______________________________________________#


@api_view(['GET'])
def get_job(request):
    jobset=job_master.objects.filter(deleted=0)
    #logger.info("Fetching events where deleted=0")
   
    job_data=[]
    for events in  jobset:
       
        company_id=None
        company_profile=None
        department_id=None
        skill_id=None
        college_id=None
        if events.company_id:
           company_id=events.company_id.company_name
           company_profile=events.company_id.company_profile
           
        if events.skill_id:
           skill_id=events.skill_id.skill_name
        if events.department_id:
           department_id=events.department_id.department
        if events.college_id:
           college_id=events.college_id.college
            
         
     
        job_data.append({
               'id': events.id,
            'company_id': company_id,
           'company_profile':company_profile,
            'college_id':college_id,
            'department_id':department_id,
            'skill_id':skill_id,
            'intern_fulltime': events.intern_fulltime,
            'on_off_campus': events.on_off_campus,
           
            'marks_10th': events.marks_10th,
             'marks_12th': events.marks_12th,
            'cgpa': events.cgpa,
           
            'gender': events.gender,
             'history_of_arrears': events.history_of_arrears,
            'standing_arrears': events.standing_arrears,
           
            'location': events.location,
            
          
           
       
          

            })
    return Response(job_data)



class jobcreateAPIView(generics.CreateAPIView):
    queryset = job_master.objects.all()
    serializer_class =jobSerializer

    def post(self, request, *args, **kwargs):
        #logger.info("Creating a new job")
        response = super().post(request, *args, **kwargs)
        #logger.info("Created a new job successfully")
        return response

class jobUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = job_master.objects.all()
    serializer_class =jobSerializer

    def put(self, request, *args, **kwargs):
        #logger.info(f"Updating job with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        #logger.info(f"Updated job with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        #logger.info(f"Partially updating job with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        #logger.info(f"Partially updated job with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_job(request, pk):
    #logger.info(f"Attempting to mark job with id {pk} as deleted")
    try:
        print("Entering Function..")
        jobs=job_master.objects.get(id=pk)

        print("job: ",jobs)
    except job_master.DoesNotExist:
        return JsonResponse("job not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    jobs.deleted = 1
    jobs.save()

    #logger.info(f"Marked job with id {pk} as deleted successfully")

    print("course: ",jobs)

    return JsonResponse("course 'deleted' field updated successfully", safe=False)


#______________________________________________________Event_master______________________________________________#


@api_view(['GET'])
def get_event(request):
    eventset=event_master.objects.filter(deleted=0)
    #logger.info("Fetching events where deleted=0")
   
    event_data=[]
    for events in  eventset:
       
       
        department_id=None

        college_id=None
        
        if events.department_id:
           department_id=events.department_id.department
        if events.college_id:
           college_id=events.college_id.college
            
         
     
        event_data.append({
               'id': events.id,
           
            'college_id':college_id,
            'department_id':department_id,
           
            'event_name': events.event_name,
            'event_desc': events.event_desc,
           
            'dtm_start': events.dtm_start,
           
          

            })
    return Response(event_data)



class eventcreateAPIView(generics.CreateAPIView):
    queryset = event_master.objects.all()
    serializer_class =eventSerializer

    def post(self, request, *args, **kwargs):
        #logger.info("Creating a new event")
        response = super().post(request, *args, **kwargs)
        #logger.info("Created a new event successfully")
        return response

class eventUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = event_master.objects.all()
    serializer_class =eventSerializer

    def put(self, request, *args, **kwargs):
        #logger.info(f"Updating events with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        #logger.info(f"Updated events with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        #logger.info(f"Partially updating events with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        #logger.info(f"Partially updated events with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_events(request, pk):
    #logger.info(f"Attempting to mark events with id {pk} as deleted")
    try:
        print("Entering Function..")
        eventss=event_master.objects.get(id=pk)

        print("events: ",eventss)
    except event_master.DoesNotExist:
        return JsonResponse("events not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    eventss.deleted = 1
    eventss.save()

    #logger.info(f"Marked event with id {pk} as deleted successfully")

    print("course: ",eventss)

    return JsonResponse("event 'deleted' field updated successfully", safe=False)

#----------------------------Training admin Dashboard datas-------------------#

# Get Total Test Count

def get_distinct_test_name_count(request, college_id):
    # Perform the query to count distinct test names for the given college_id
    distinct_test_name_count = tests_candidates_map.objects.filter(college_id=college_id).values('test_name').distinct().count()
    
    # Return the count as a JSON response
    return JsonResponse({'distinct_test_name_count': distinct_test_name_count})



# Avg Score of Aptitude

def get_avg_score_by_department(request, college_id, dtm_start):
    try:
        # Convert dtm_start to a datetime object
        dtm_start_date = datetime.datetime.strptime(dtm_start, '%Y-%m-%d').date()

        # Get all departments
        departments = department_master.objects.filter(deleted=0).values('id', 'department')

        # Calculate the average score for each department
        results = []
        for department in departments:
            avg_score = tests_candidates_map.objects.filter(
                college_id=college_id,
                dtm_start__date=dtm_start_date,
                question_id__test_type='MCQ Test',
                department_id=department['id']
            ).aggregate(avg_score=Avg('total_score'))['avg_score'] or 0

            results.append({
                'department_name': department['department'],
                'avg_score': avg_score
            })

        # Return the results as a JSON response
        return JsonResponse(results, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


# Avg Score of Coding

def avg_score_by_department_Coding(request, college_id, dtm_start):
    try:
        # Convert dtm_start to a datetime object
        dtm_start_date = datetime.datetime.strptime(dtm_start, '%Y-%m-%d').date()

        # Get all departments
        departments = department_master.objects.filter(deleted=0).values('id', 'department')

        # Calculate the average score for each department
        results = []
        for department in departments:
            avg_score = tests_candidates_map.objects.filter(
                college_id=college_id,
                dtm_start__date=dtm_start_date,
                question_id__test_type='Coding Test',
                department_id=department['id']
            ).aggregate(avg_score=Avg('total_score'))['avg_score'] or 0

            results.append({
                'department_name': department['department'],
                'avg_score': avg_score
            })

        # Return the results as a JSON response
        return JsonResponse(results, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


#  College Topper


def get_max_score_by_department(request, college_id):
    try:
        # First, create a subquery to get the maximum total score per department
        subquery = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='MCQ Test',
            total_score__isnull=False,
            department_id=OuterRef('department_id')
        ).values(
            'department_id'
        ).annotate(
            max_total_score=Max('total_score')
        ).values('max_total_score')

        # Then, filter the original query using this subquery to get the corresponding student names
        results = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='MCQ Test',
            total_score__isnull=False,
            total_score=Subquery(subquery)  # Match the max score in each department
        ).values(
            'student_id__students_name',
            'department_id__department',
            'total_score'
        ).distinct()  # Ensure unique records in case of ties

        # Format the results as a list of dictionaries
        data = [
            {
                'student_name': result['student_id__students_name'],
                'department': result['department_id__department'],
                'max_total_score': result['total_score']
            }
            for result in results
        ]

        # Return the results as a JSON response
        return JsonResponse(data, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)



def get_max_score_by_department_coding(request, college_id):
    try:
        # First, create a subquery to get the maximum total score per department
        subquery = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='Coding Test',
            total_score__isnull=False,
            department_id=OuterRef('department_id')
        ).values(
            'department_id'
        ).annotate(
            max_total_score=Max('total_score')
        ).values('max_total_score')

        # Then, filter the original query using this subquery to get the corresponding student names
        results = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='Coding Test',
            total_score__isnull=False,
            total_score=Subquery(subquery)  # Match the max score in each department
        ).values(
            'student_id__students_name',
            'department_id__department',
            'total_score'
        ).distinct()  # Ensure unique records in case of ties

        # Format the results as a list of dictionaries
        data = [
            {
                'student_name': result['student_id__students_name'],
                'department': result['department_id__department'],
                'max_total_score': result['total_score']
            }
            for result in results
        ]

        # Return the results as a JSON response
        return JsonResponse(data, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


# Count of company


def count_company_names(request):
    try:
        # Perform the query to count company_name entries
        result = company_master.objects.aggregate(
            count_company_name=Count('company_name')
        )

        # Return the count as a JSON response
        return JsonResponse(result, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


#------------------------------Placement admin Dashboard datas-----------------------#


def get_distinct_test_name_count_today(request):
    try:
        # Get today's date
        today_date = date.today()

        # Perform the query to count distinct test_name where dtm_start is today's date
        distinct_test_name_count = tests_candidates_map.objects.filter(
            dtm_start__date=today_date
        ).values('test_name').distinct().count()

        # Return the count as JSON response
        return JsonResponse({'distinct_test_name_count': distinct_test_name_count}, status=200)
    
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)







# URLs for the CSV files
sheet_urls = {
    'AIDS': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=257135364',
    'IT': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=747238312',
    'CSE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=1208914504',
    'ECE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=2128791655'
}

# Function to process each sheet
def process_sheet(url):
    response = requests.get(url)
    csv_data = response.content.decode('utf-8')
    df = pd.read_csv(StringIO(csv_data), skiprows=7)  # Skip first 7 rows to reach the headers
    df = df[2:]  # Skip the next row (8th and 9th rows)
    df.reset_index(drop=True, inplace=True)  # Reset index after slicing
    df.columns = df.columns.str.strip()  # Strip any leading/trailing spaces from columns
    return df

@api_view(['GET'])
def get_avg_total_present(request):
    try:
        department_data = {}

        for department, url in sheet_urls.items():
            df = process_sheet(url)
            print(f"Headers for {department}: {df.columns}")  # Print headers for debugging

            # Check if 'Total Present' column exists
            if 'Total Present' in df.columns:
                # Calculate the average of 'Total Present'
                avg_total_present = df['Total Present'].astype(float).mean()  # Ensure conversion to float
                department_data[department] = round(avg_total_present, 1)
            else:
                department_data[department] = 'Total Present column not found'

        return Response(department_data)
    except Exception as e:
        return Response({'error': str(e)}, status=400)



@api_view(['GET'])
def get_avg_total_absent(request):
    try:
        department_data = {}

        for department, url in sheet_urls.items():
            df = process_sheet(url)
            print(f"Headers for {department}: {df.columns}")  # Print headers for debugging

            # Check if 'Total Present' column exists
            if 'Total Present' in df.columns:
                # Calculate the average of 'Total Present'
                avg_total_absent = df['Total Abscent'].astype(float).mean()  # Ensure conversion to float
                department_data[department] = round(avg_total_absent, 1)
            else:
                department_data[department] = 'Total absent column not found'

        return Response(department_data)
    except Exception as e:
        return Response({'error': str(e)}, status=400)







#---------------------------------------Students Dashboards-----------------------------------#



def get_events_by_college_and_department(request, college_id, department_id):
    try:
        # Filter the events based on college_id and department_id
        events = event_master.objects.filter(
            college_id=college_id,
            department_id=department_id
        ).values('event_name', 'event_desc', 'dtm_start')

        # Convert the queryset to a list of dictionaries
        data = list(events)

        # Return the results as a JSON response
        return JsonResponse(data, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


#__________________________________IMport Questions____________________________#
from django.shortcuts import render, redirect
from django.http import HttpResponse
from .forms import QuestionImportForm
from .models import question_master_temp, question_paper_master, topic_master,question_master
import docx
import json
import re
import base64
from docx import Document

def extract_text_and_images_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = []
    images_binary = {}
    image_count = 0

    for para in doc.paragraphs:
        text.append(para.text)

    for rel in doc.part.rels:
        target = doc.part.rels[rel].target_ref
        if "image" in target:
            image_count = str(target).split('.')[0][-1]
            image = doc.part.rels[rel].target_part.blob
            images_binary[f"image_{image_count}"] = base64.b64encode(image).decode('utf-8')

    return '\n'.join(text), images_binary

def remove_serial_number(question_text):
    return re.sub(r'^\d+\s?\)\s*', '', question_text)

def remove_option_choice(option_text):
    return re.sub(r'^[a-d]\s*?\)\s*', '', option_text)

def create_json_structure(text, images_binary):
    questions = []
    lines = text.split('\n')
    current_question = {}
    options = {}
    image_keys = list(images_binary.keys())
    image_index = 1
    unmatched_lines = []

    question_pattern = re.compile(r'^\s*\d+\s?\)')
    option_a_pattern = re.compile(r'^\s*a\s?\)')
    option_b_pattern = re.compile(r'^\s*b\s?\)')
    option_c_pattern = re.compile(r'^\s*c\s?\)')
    option_d_pattern = re.compile(r'^\s*d\s?\)')
    answer_pattern = re.compile(r'^\s*c_ans:')
    explanation_pattern = re.compile(r'^\s*e_ans:')
    mark_pattern = re.compile(r'^\s*mark:')
    negative_mark_pattern = re.compile(r'^\s*n_m:')

    for line in lines:
        line = line.strip()
        if question_pattern.match(line):
            if current_question:
                current_question['options'] = options
                questions.append(current_question)
                current_question = {}
                options = {}
            question_text = remove_serial_number(line)
            if question_text == '' :
                current_question['question_text'] =  [f"data:image/png;base64,{images_binary[f'image_{image_index}']}", True]
                image_index += 1
            else:
                current_question['question_text'] = [question_text, False]
        elif option_a_pattern.match(line):
            option_a_text = remove_option_choice(line)
            if option_a_text == '':
                options['a'] =  [f"data:image/png;base64,{images_binary[f'image_{image_index}']}", True]
                image_index += 1
            else:  
                options['a'] = [option_a_text,False] 
        elif option_b_pattern.match(line):
            option_b_text = remove_option_choice(line)
            if option_b_text == '':
                options['b'] =  [f"data:image/png;base64,{images_binary[f'image_{image_index}']}", True]
                image_index += 1
            else:
                options['b'] = [option_b_text, False]
        elif option_c_pattern.match(line):
            option_c_text = remove_option_choice(line)
            if option_c_text == '':
                options['c'] =  [f"data:image/png;base64,{images_binary[f'image_{image_index}']}", True]
                image_index += 1
            else:
                options['c'] = [option_c_text, False]
        elif option_d_pattern.match(line):
            option_d_text = remove_option_choice(line)
            if option_d_text == '':
                options['d'] =  [f"data:image/png;base64,{images_binary[f'image_{image_index}']}", True]
                image_index += 1
            else:
                options['d'] = [option_d_text, False]
        elif answer_pattern.match(line):
            current_question['answer'] = line.split(':', 1)[1].strip()
        elif explanation_pattern.match(line):
            current_question['explanation'] = line.split(':', 1)[1].strip()
        elif mark_pattern.match(line):
            current_question['marks'] = int(line.split(':', 1)[1].strip())
        elif negative_mark_pattern.match(line):
            current_question['negative_marks'] = int(line.split(':', 1)[1].strip())
        else:
            if line:
                unmatched_lines.append((line, "Unmatched pattern"))

    if current_question:
        current_question['options'] = options
        current_question['image_based'] = any(isinstance(opt[0], str) and 'data:image' in opt[0] for opt in options.values())
        questions.append(current_question)

    return {"questions": questions, "unmatched_lines": unmatched_lines}

def import_questions_from_word(request):
    if request.method == 'POST':
        form = QuestionImportForm(request.POST, request.FILES)
        if form.is_valid():
            docx_file = request.FILES['docx_file']
            text, images_binary = extract_text_and_images_from_docx(docx_file)
            data = create_json_structure(text, images_binary)

            for question_data in data['questions']:
                question_master.objects.create(
                    question_text=question_data.get('question_text')[0] if not question_data.get('question_text')[1] else None,
                    question_image_data=base64.b64decode(question_data.get('question_text')[0].split(',')[1]) if question_data.get('question_text')[1] else None,
                    option_a=question_data.get('options', {}).get('a')[0] if not question_data.get('options', {}).get('a')[1] else None,
                    option_a_image_data=base64.b64decode(question_data.get('options', {}).get('a')[0].split(',')[1]) if question_data.get('options', {}).get('a')[1] else None,
                    option_b=question_data.get('options', {}).get('b')[0] if not question_data.get('options', {}).get('b')[1] else None,
                    option_b_image_data=base64.b64decode(question_data.get('options', {}).get('b')[0].split(',')[1]) if question_data.get('options', {}).get('b')[1] else None,
                    option_c=question_data.get('options', {}).get('c')[0] if not question_data.get('options', {}).get('c')[1] else None,
                    option_c_image_data=base64.b64decode(question_data.get('options', {}).get('c')[0].split(',')[1]) if question_data.get('options', {}).get('c')[1] else None,
                    option_d=question_data.get('options', {}).get('d')[0] if not question_data.get('options', {}).get('d')[1] else None,
                    option_d_image_data=base64.b64decode(question_data.get('options', {}).get('d')[0].split(',')[1]) if question_data.get('options', {}).get('d')[1] else None,
                    answer=question_data.get('answer'),
                    explain_answer=question_data.get('explanation'),
                    mark=question_data.get('marks'),
                    negative_mark=question_data.get('negative_marks')
                )

            return HttpResponse("Questions imported successfully.")
    else:
        form = QuestionImportForm()

    return render(request, 'import_questions.html', {'form': form})

#___________________________________IMPORT CODING QUESTION___________________________________________________#
from django.shortcuts import render, redirect
from django.http import HttpResponse
from .forms import DocumentForm
from .models import question_master_temp, question_paper_master
import docx
import re

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return '\n'.join(text)

def remove_serial_number(question_text):
    return re.sub(r'^\d+\s?\)\s*', '', question_text)

def create_json_structure(text):
    questions = []
    lines = text.split('\n')
    current_question = {}
    unmatched_lines = []
    input_format_lines = []

    question_pattern = re.compile(r'^\d+\s?\)')
    answer_pattern = re.compile(r'^c_ans:')
    explanation_pattern = re.compile(r'^e_ans:')
    mark_pattern = re.compile(r'^mark:')
    negative_mark_pattern = re.compile(r'^n_m:')
    input_format_pattern = re.compile(r'^input_format:')

    def append_current_question():
        if current_question.get('question'):
            if input_format_lines:
                current_question['input_format'] = ' '.join(input_format_lines).strip()
                input_format_lines.clear()
            questions.append(current_question.copy())  # Append a copy to avoid modifying the original
            current_question.clear()  # Clear current_question for the next iteration

    for line in lines:
        line = line.strip()
        if question_pattern.match(line):
            append_current_question()
            current_question['question'] = remove_serial_number(line)
            current_question['answers'] = []  # Initialize answers list
        elif answer_pattern.match(line):
            current_question['answers'].append(line[len('c_ans:'):].strip())
        elif explanation_pattern.match(line):
            current_question['explanation'] = line[len('e_ans:'):].strip()
        elif mark_pattern.match(line):
            current_question['marks'] = line[len('mark:'):].strip()
        elif negative_mark_pattern.match(line):
            current_question['negative_marks'] = line[len('n_m:'):].strip()
        elif input_format_pattern.match(line):
            append_current_question()
            input_format_lines.append(line[len('input_format:'):].strip())
        elif input_format_lines:
            input_format_lines.append(line)
        else:
            unmatched_lines.append(line)

    append_current_question()  # Append the last question

    return {"questions": questions, "unmatched_lines": unmatched_lines}

def import_questions(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            docfile = request.FILES['docfile']
            text = extract_text_from_docx(docfile)
            data = create_json_structure(text)

            question_paper_master_instance = question_paper_master.objects.first()  # Adjust this as needed

            for ques in data["questions"]:
                if 'question' in ques:  # Ensure 'question' key exists
                    question_text = ques['question']
                    answers = ques.get('answers', [])
                    marks = ques.get('marks', '')
                    negative_mark = ques.get('negative_marks', '')
                    explain_answer = ques.get('explanation', '')
                    input_format = ques.get('input_format', '')

                    # Create the question
                    question = question_master(
                        question_name_id=question_paper_master_instance,
                        question_text=question_text,
                        input_format=input_format,
                        answer=', '.join(answers),
                        negative_mark=negative_mark,
                        mark=marks,
                        explain_answer=explain_answer
                    )
                    question.save()
                else:
                    return HttpResponse("Error: Missing 'question' key in data.")

            return HttpResponse("Questions imported successfully.")  # Redirect to a success page or elsewhere
    else:
        form = DocumentForm()

    return render(request, 'import_coding.html', {'form': form})


#_________________________Training_schedule_sheet________________________________________________________#


@api_view(['GET'])
def get_training_report(request):
    courseset = training_attendance_sheet.objects.filter(deleted=0)
    #logger.info("Fetching Courses where deleted=0")
   
    course_data = []
    for courses in courseset:
        year = None
        dtm_start = None
        dtm_end = None
        college_name = None
        department_name = None
        trainer_name = None
        topic = None
        sub_topic = None

        if courses.course_schedule_id:
            year = courses.course_schedule_id.year
            dtm_start = courses.course_schedule_id.dtm_start
            dtm_end = courses.course_schedule_id.dtm_end
            college_name = courses.course_schedule_id.college_id.college if courses.course_schedule_id.college_id else None
            department_name = courses.course_schedule_id.department_id.department if courses.course_schedule_id.department_id else None
            trainer_name = courses.course_schedule_id.trainer_id.trainer_name if courses.course_schedule_id.trainer_id else None
            topic = courses.course_schedule_id.topic
            sub_topic = courses.course_schedule_id.sub_topic

        course_data.append({
            'id': courses.id,
            'year': year,
            'dtm_start': dtm_start,
            'dtm_end': dtm_end,
            'college_name': college_name,
            'department_name': department_name,
            'trainer_name': trainer_name,
            'topic': topic,
            'sub_topic': sub_topic,
            'dtm_attendance': courses.dtm_attendance,
            'present': courses.present,
            'absent': courses.absent,
        })
    return Response(course_data)

class trainingreportcreateAPIView(generics.CreateAPIView):
    queryset = training_attendance_sheet.objects.all()
    serializer_class =trainingreportsheetSerializer

    def post(self, request, *args, **kwargs):
        #logger.info("Creating a new training_report")
        response = super().post(request, *args, **kwargs)
        #logger.info("Created a new training_report successfully")
        return response

class trainingreportUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = training_attendance_sheet.objects.all()
    serializer_class =trainingreportsheetSerializer

    def put(self, request, *args, **kwargs):
        #logger.info(f"Updating training_report with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        #logger.info(f"Updated training_report with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        #logger.info(f"Partially updating training_report with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        #logger.info(f"Partially updated training_report with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_training_report(request, pk):
    #logger.info(f"Attempting to mark training_report with id {pk} as deleted")
    try:
        print("Entering Function..")
        training_report=training_attendance_sheet.objects.get(id=pk)

        print("trainingreport: ",training_report)
    except training_attendance_sheet.DoesNotExist:
        return JsonResponse("training_report not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    training_report.deleted = 1
    training_report.save()

    #logger.info(f"Marked training_report with id {pk} as deleted successfully")

    print("training_report: ",training_report)

    return JsonResponse("training_report 'deleted' field updated successfully", safe=False)



#-----------------------------Course _ trainer_feedback-----------------------------------------#

class add_trainer_feedback(generics.CreateAPIView):
    queryset = course_trainer_feedback.objects.all()
    serializer_class = trainerfeedbackSerializer


@api_view(['GET'])
def get_trainer_feedback(request):
    cs_set = course_trainer_feedback.objects.select_related('college_id','topic_id' , 'trainer_id','department_id').filter(deleted=0)

    trainer_feedback_data=[]
    for cs in  cs_set:
           trainer_feedback_data.append({
              'id': cs.id,
              'college_id': cs.college_id.college,
               'department_id': cs.department_id.department,
              'topic_id':cs.topic_id.topic,
             # 'year':cs.year,
              'trainer_id': cs.trainer_id.trainer_name,
              'dtm_complete': cs.dtm_complete,
              
              'completion_status': cs.completion_status,
              'feedback': cs.feedback
            })
    return Response(trainer_feedback_data)

class update_trainer_feedback(generics.UpdateAPIView):
    queryset = course_trainer_feedback.objects.all()
    serializer_class = trainerfeedbackSerializer



@api_view(['PUT', 'PATCH'])
def delete_trainer_feedback(request, pk):
    try:
        print("Entering Function..")
        course = course_trainer_feedback.objects.get(id=pk)

        print("course: ", course)
    except course_trainer_feedback.DoesNotExist:
        return JsonResponse("course not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    course.deleted = 1
    course.save()
    print("logins: ", course)

    return JsonResponse("course 'deleted' field updated successfully", safe=False)

#------------------------Compiler---------------------#



# List of forbidden file operation functions
FORBIDDEN_FUNCTIONS = [
    r'\bfopen\b', r'\bfclose\b', r'\bfread\b', r'\bfwrite\b', r'\bfscanf\b', 
    r'\bfprintf\b', r'\bfgetc\b', r'\bfputc\b', r'\bfgets\b', r'\bfputs\b', 
    r'\bremove\b', r'\brename\b',r'\bfopen\b',r'\bfread\b',r'\bfwrite\b',r'\bstd::ifstream\b',r'\bstd::ofstream\b',r'\bstd::fstream\b'
]

def contains_forbidden_functions(code):
    """Check if the code contains any forbidden functions."""
    for func in FORBIDDEN_FUNCTIONS:
        if re.search(func, code):
            return True
    return False




def execute_python_code(code,inputs=""):

    try:
        

        addon="""
import builtins
# Define safe import and restricted open functions
def safe_import(name, globals=None, locals=None, fromlist=(), level=0):
    if name not in ['os',"globe","shutill",subprocess]:
        return original_import(name, globals, locals, fromlist, level)
    raise ImportError(f"Importing {name} is not allowed")

def restricted_open(*args, **kwargs):
    raise IOError("File opening is not allowed.")

# Save the original built-in functions
original_import = builtins.__import__
original_open = builtins.open

# Replace built-in functions with restricted versions
builtins.__import__ = safe_import
builtins.open = restricted_open

"""
        code=addon+"\n"+code
        #use subprocess to execute the code at terminal
        result = subprocess.run(
            ['python', '-c', code],
            capture_output=True,
            input=inputs,
            text=True,
            timeout=settings.CODE_EXECUTION_TIMEOUT #the program will raise timeout error if it exceeds settings.CODE_EXECUTION_TIMEOUTs of execution
        )
        return result.stdout or result.stderr
    except subprocess.TimeoutExpired:
        #time error is handled and the message is returned
        return {'error':'Execution timed out'}
    except Exception as e:
        return {'error': str(e)}




def execute_java_code(code,inputs=""):
    asset_path = settings.ASSET_DIR
    # Extract class name
    class_match = re.search(r'public\s+class\s+(\w+)', code)
    if not class_match:
        return 'No public class found'
    
    class_name = class_match.group(1)
    dir_path = os.path.abspath(os.path.join(asset_path,str(uuid.uuid4())))
    os.makedirs(dir_path)
    #file_path = os.path.join(dir_path, f'{class_name}.java')

    #using uuid to create temporary c file for execution
    temp_java_path = os.path.abspath(os.path.join(dir_path, f'{class_name}.java'))
    temp_exec_path = f'{class_name}'
    #writing the code into temporary file
    with open(temp_java_path, 'w') as f:
        f.write(code)
    try:
        #compiling the file using subprocess at terminal
        compile_result = subprocess.run(
            ['javac', temp_java_path],
            capture_output=True,
            cwd=dir_path,
            text=True
        )
        #if program has syntax error return the error message
        if compile_result.returncode != 0:
            print("syntax error occured")
            return compile_result.stderr
        #return the programs output
        run_result = subprocess.run(
            ['java', '-Djava.security.manager', '-Djava.security.policy==java.policy', temp_exec_path],
            capture_output=True,
            input=inputs,
            text=True,
            cwd=dir_path,
            timeout=settings.CODE_EXECUTION_TIMEOUT
        )
        return run_result.stdout or run_result.stderr
    except subprocess.TimeoutExpired:
        return 'Execution timed out'
    except Exception as e:
         return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    finally:
        shutil.rmtree(dir_path)
     

   
         
def execute_c_code(code,inputs=""):
    file_uuid=uuid.uuid4()
    asset_path = settings.ASSET_DIR
    #using uuid to create temporary c file for execution
    temp_c_path = os.path.abspath(os.path.join(asset_path, f'{file_uuid}.c'))
    temp_exec_path = os.path.abspath(os.path.join(asset_path, f'{file_uuid}'))
    # Check for forbidden file operations
    if contains_forbidden_functions(code):
        return "Error: Code contains forbidden file operations."
    #writing the code into temporary file
    with open(temp_c_path, 'w') as f:
        f.write(code)
    try:
        #compiling the file using subprocess at terminal
        compile_result = subprocess.run(
            ['gcc', temp_c_path, '-o', temp_exec_path],
            capture_output=True,
            #inputs=inputs,
            text=True
        )
        #if program has syntax error return the error message
        if compile_result.returncode != 0:
            print("syntax error occured")
            return compile_result.stderr
        #return the programs output
        # run_result = subprocess.run(
        #     [temp_exec_path],
        #     capture_output=True,
        #     text=True,
        #     timeout=settings.CODE_EXECUTION_TIMEOUT
        # )
        # Execute the compiled  program with inputs
        process = subprocess.Popen(
                [temp_exec_path],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
        )

            # Communicate with the process: send inputs and read outputs
        stdout, stderr = process.communicate(inputs)
        return stdout or stderr
    except subprocess.TimeoutExpired:
        return 'Execution timed out'
    except Exception as e:
         return f'error:{str(e)}'
    finally:
        os.remove(temp_c_path)
        os.remove(f"{temp_exec_path}.exe")
        #subprocess.run(['rm', temp_c_path, temp_exec_path])

def execute_cpp_code(code,inputs=""):
    file_uuid=uuid.uuid4()
    asset_path = settings.ASSET_DIR
    file_path = os.path.abspath(os.path.join(asset_path, f'{file_uuid}.cpp'))
    exec_path = os.path.abspath(os.path.join(asset_path, f'{file_uuid}'))
    # Check for forbidden file operations
    if contains_forbidden_functions(code):
        return "Error: Code contains forbidden file operations."
   
    try:
        with open(file_path, 'w') as f:
            f.write(code)
        compile_result = subprocess.run(
            ['g++', file_path, '-o', exec_path],
            capture_output=True,
            text=True
        )
        if compile_result.returncode != 0:
            return compile_result.stderr
        process = subprocess.Popen(
                [exec_path],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True)

            # Communicate with the process: send inputs and read outputs
        stdout, stderr = process.communicate(inputs)
        return stdout or stderr
    except subprocess.TimeoutExpired:
        return 'Execution timed out'
    finally:
        os.remove(file_path)
        os.remove(f"{exec_path}.exe")
        # subprocess.run(['rm', file_path, exec_path])






@api_view(['POST'])
def program_compiler(request):

    language_dict={
        "python":execute_python_code,
        "c":execute_c_code,
        "cpp":execute_cpp_code,
        "java":execute_java_code,
    }
    
    code = request.data.get('code')
    p_type = request.data.get('p_type')
    inputs=request.data.get('inputs')
    if p_type not in language_dict.keys():
        return Response({"result":"Unsupported language"})
    else:
        print(f"executing {p_type} code")
        return Response({"result":language_dict[p_type](code,inputs)})




def execute_code(p_type, code, inputs):
    language_dict = {
        "python": execute_python_code,
        "c": execute_c_code,
        "cpp": execute_cpp_code,
        "java": execute_java_code,
    }

    if p_type not in language_dict.keys():
        raise ValueError("Unsupported language")
    
    print('output: ', language_dict[p_type](code,inputs))

    return language_dict[p_type](code,inputs)

@csrf_exempt
@api_view(['POST'])
def test_candidates_answer_view_Com(request):
    stu_id =  request.data.get('student_id')
    
    # Delete all rows of the compiler_output model
    # compiler_output.objects.filter(student_id=stu_id).delete()

    ques_id = request.data.get('question_id')
    code = request.data.get('code')
    p_type = request.data.get('p_type')
    inputs = request.data.get('inputs')

    if not ques_id or not code or not p_type:
        return JsonResponse({'error': 'Missing required parameters'}, status=400)

    try:
        output = execute_code(p_type, code, inputs)
        # Update the output where student_id matches stu_id
        checking = compiler_output.objects.filter(student_id=stu_id).update(output=output)
        print('Checking: ', checking)

        # compiler_output_instance = compiler_output(output=output)
        # compiler_output_instance.save()
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

    # Get the last data of compiler_output and store it in ans
    last_output = compiler_output.objects.filter(student_id=stu_id).last()
    print('last_output: ', last_output)
    ans = last_output.output if last_output else ''
    print('ans: ', ans)


    print("Parameters: ", ques_id, ans, code)

    question = get_object_or_404(question_master, id=ques_id)

    if question.answer == ans:
        result = question.mark
    else:
        similarity_score = fuzz.ratio(question.explain_answer.strip(), code.strip())
        result = round((similarity_score / 100) * question.mark)

    result = round(result)
    print("Result: ", result)

    test_candidate_answer_data = {
        'test_name': request.data.get('test_name'),
        'question_id': ques_id,
        'student_id': request.data.get('student_id'),
        'submission_compile_code': ans,
        'compile_code_editor': code,
        'result': result,
        'dtm_start': request.data.get('dtm_start'),
        'dtm_end': request.data.get('dtm_end'),
    }

    serializer = tests_candidates_answerSerializer(data=test_candidate_answer_data)
    print("Serializer: ", serializer)

    if serializer.is_valid():
        serializer.save()
        print('Test Answer is added Successfully')
        return JsonResponse(serializer.data, status=201)
    else:
        print(serializer.errors)
        return JsonResponse(serializer.errors, status=400)




@csrf_exempt
@api_view(['POST'])
def test_candidates_answer_view_Submit_Com(request, format=None):
    stu_id =  request.data.get('student_id')


    # Delete all rows of the compiler_output model
    # compiler_output.objects.filter(student_id=stu_id).delete()

    ques_id =  request.data.get('question_id')
    # ans =  request.data.get('ans')
    code =  request.data.get('code')
    p_type = request.data.get('p_type')
    inputs = request.data.get('inputs')


    if not ques_id or not code or not p_type:
        return JsonResponse({'error': 'Missing required parameters'}, status=400)

    try:
        output = execute_code(p_type, code, inputs)
         # Update the output where student_id matches stu_id
        checking = compiler_output.objects.filter(student_id=stu_id).update(output=output)
        print('Checking: ', checking)

        #compiler_output_instance = compiler_output(student_id_id=stu_id, output=output)
        #compiler_output_instance.save()
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


    # code = urllib.parse.unquote(code)

    # Get the last data of compiler_output and store it in ans
    last_output = compiler_output.objects.filter(student_id=stu_id).last()
    print('last_output: ', last_output)
    ans = last_output.output if last_output else ''
    print('ans: ', ans)


    print("Parameters: ", ques_id, ans, code)
    
    try:
        question = question_master.objects.get(id=ques_id)
    except question_master.DoesNotExist:
        return JsonResponse({'error': 'Question not found'}, status=404)

    if question.answer == ans:
        result = question.mark
    else:
        # Handling multi-line strings
        similarity_score = fuzz.ratio(question.explain_answer.strip(), code.strip())
        result = round((similarity_score / 100) * question.mark)


    # Round the result to the nearest integer
    result = round(result)
    print("Result: ", result)
    
    test_candidate_answer_data = {
        'test_name': request.data.get('test_name'),
        'question_id': ques_id,
        'student_id': request.data.get('student_id'),
        'answer': ans,
        'compile_code_editor': code,
        'result': result,
        'dtm_start': request.data.get('dtm_start'),
        'dtm_end': request.data.get('dtm_end'),
    }

    serializer = tests_candidates_answerSerializer(data=test_candidate_answer_data)
    print("Serializer: ", serializer)
    
    if serializer.is_valid():
        serializer.save()
        print('Test Answer is added Successfully')

        compiler_output.objects.filter(student_id=stu_id).delete()
        print('Compiler Output Data are Deleted')
        return JsonResponse(serializer.data, status=201)
    else:
        print(serializer.errors)
        return JsonResponse(serializer.errors, status=400)

    


@api_view(['GET'])
def get_last_compiler_output(request, student_id):
    last_output = compiler_output.objects.filter(student_id=student_id).last()
    output = last_output.output if last_output else ""
    
    return Response({"last_output": output})




@api_view(['POST'])
def insert_empty_output_view(request, student_id):
    if not student_id:
        print('student_id is required')
        return Response({'error': 'student_id is required'}, status=400)
    
    try:
        student = candidate_master.objects.get(id=student_id)
        print('student id: ', student)
    except candidate_master.DoesNotExist:
        return Response({'error': 'CandidateMaster with the given id does not exist'}, status=404)
    
    existing_record = compiler_output.objects.filter(student_id=student)
    if not existing_record:
        compiler_output.objects.create(student_id=student, output='')
        print('Record created successfully')
        return Response({'message': 'Record created successfully'})
    else:
        print('Record already exists')
        return Response({'message': 'Record already exists'})


#----------------------------Training admin Dashboard datas-------------------#

# Get Total Test Count

def get_distinct_test_name_count(request, college_id):
    # Perform the query to count distinct test names for the given college_id
    distinct_test_name_count = tests_candidates_map.objects.filter(college_id=college_id, deleted=0).values('test_name').distinct().count()
    
    # Return the count as a JSON response
    return JsonResponse({'distinct_test_name_count': distinct_test_name_count})



# Avg Score of Aptitude

def get_avg_score_by_department(request, college_id, dtm_start):
    try:
        # Convert dtm_start to a datetime object
        dtm_start_date = datetime.datetime.strptime(dtm_start, '%Y-%m-%d').date()

        # Get all departments
        departments = department_master.objects.filter(deleted=0).values('id', 'department')

        # Calculate the average score for each department
        results = []
        for department in departments:
            avg_score = tests_candidates_map.objects.filter(
                college_id=college_id,
                dtm_start__date=dtm_start_date,
                question_id__test_type='MCQ Test',
                department_id=department['id'],
                deleted=0
            ).aggregate(avg_score=Avg('total_score'))['avg_score'] or 0

            results.append({
                'department_name': department['department'],
                'avg_score': avg_score
            })

        # Return the results as a JSON response
        return JsonResponse(results, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


# Avg Score of Coding

def avg_score_by_department_Coding(request, college_id, dtm_start):
    try:
        # Convert dtm_start to a datetime object
        dtm_start_date = datetime.datetime.strptime(dtm_start, '%Y-%m-%d').date()

        # Get all departments
        departments = department_master.objects.filter(deleted=0).values('id', 'department')

        # Calculate the average score for each department
        results = []
        for department in departments:
            avg_score = tests_candidates_map.objects.filter(
                college_id=college_id,
                dtm_start__date=dtm_start_date,
                question_id__test_type='Coding Test',
                department_id=department['id'],
                deleted=0
            ).aggregate(avg_score=Avg('total_score'))['avg_score'] or 0

            results.append({
                'department_name': department['department'],
                'avg_score': avg_score
            })

        # Return the results as a JSON response
        return JsonResponse(results, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


#  College Topper


def get_max_score_by_department(request, college_id):
    try:
        # First, create a subquery to get the maximum total score per department
        subquery = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='MCQ Test',
            total_score__isnull=False,
            department_id=OuterRef('department_id'),
            deleted=0
        ).values(
            'department_id'
        ).annotate(
            max_total_score=Max('total_score')
        ).values('max_total_score')

        # Then, filter the original query using this subquery to get the corresponding student names
        results = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='MCQ Test',
            total_score__isnull=False,
            total_score=Subquery(subquery),
            deleted=0
        ).values(
            'student_id__students_name',
            'department_id__department',
            'total_score'
        ).distinct()  # Ensure unique records in case of ties

        # Format the results as a list of dictionaries
        data = [
            {
                'student_name': result['student_id__students_name'],
                'department': result['department_id__department'],
                'max_total_score': result['total_score']
            }
            for result in results
        ]

        # Return the results as a JSON response
        return JsonResponse(data, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)



def get_max_score_by_department_coding(request, college_id):
    try:
        # First, create a subquery to get the maximum total score per department
        subquery = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='Coding Test',
            total_score__isnull=False,
            department_id=OuterRef('department_id'),
            deleted=0
        ).values(
            'department_id'
        ).annotate(
            max_total_score=Max('total_score')
        ).values('max_total_score')

        # Then, filter the original query using this subquery to get the corresponding student names
        results = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='Coding Test',
            total_score__isnull=False,
            total_score=Subquery(subquery),
            deleted=0
        ).values(
            'student_id__students_name',
            'department_id__department',
            'total_score'
        ).distinct()  # Ensure unique records in case of ties

        # Format the results as a list of dictionaries
        data = [
            {
                'student_name': result['student_id__students_name'],
                'department': result['department_id__department'],
                'max_total_score': result['total_score']
            }
            for result in results
        ]

        # Return the results as a JSON response
        return JsonResponse(data, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


# Count of company


def count_company_names(request):
    try:
        # Perform the query to count company_name entries
        result = company_master.objects.aggregate(
            count_company_name=Count('company_name')
        )

        # Return the count as a JSON response
        return JsonResponse(result, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


#------------------------------Placement admin Dashboard datas-----------------------#


def get_distinct_test_name_count_today(request):
    try:
        # Get today's date
        today_date = date.today()

        # Perform the query to count distinct test_name where dtm_start is today's date
        distinct_test_name_count = tests_candidates_map.objects.filter(
            dtm_start__date=today_date,
            deleted=0
        ).values('test_name').distinct().count()

        # Return the count as JSON response
        return JsonResponse({'distinct_test_name_count': distinct_test_name_count}, status=200)
    
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)







# URLs for the CSV files
sheet_urls = {
    'AIDS': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=257135364',
    'IT': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=747238312',
    'CSE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=1208914504',
    'ECE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=2128791655'
}

# Function to process each sheet
def process_sheet(url):
    response = requests.get(url)
    csv_data = response.content.decode('utf-8')
    df = pd.read_csv(StringIO(csv_data), skiprows=7)  # Skip first 7 rows to reach the headers
    df = df[2:]  # Skip the next row (8th and 9th rows)
    df.reset_index(drop=True, inplace=True)  # Reset index after slicing
    df.columns = df.columns.str.strip()  # Strip any leading/trailing spaces from columns
    return df

@api_view(['GET'])
def get_avg_total_present(request):
    try:
        department_data = []

        for department, url in sheet_urls.items():
            df = process_sheet(url)
            # print(f"Headers for {department}: {df.columns}")  # Print headers for debugging

            # Check if 'Total Present' column exists
            if 'Total Present' in df.columns:
                # Calculate the average of 'Total Present'
                avg_total_present = df['Total Present'].astype(float).mean()  # Ensure conversion to float
                department_data.append({
                    "department_name": department,
                    "total_present_avg": round(avg_total_present, 1)
                })
            else:
                department_data.append({
                    "department_name": department,
                    "total_present_avg": 0
                })

        return Response(department_data)
    except Exception as e:
        return Response({'error': str(e)}, status=400)



@api_view(['GET'])
def get_avg_total_absent(request):
    try:
        department_data = []

        for department, url in sheet_urls.items():
            df = process_sheet(url)
            # print(f"Headers for {department}: {df.columns}")  # Print headers for debugging

            # Check if 'Total Present' column exists
            if 'Total Abscent' in df.columns:
                # Calculate the average of 'Total Present'
                avg_total_present = df['Total Abscent'].astype(float).mean()  # Ensure conversion to float
                department_data.append({
                    "department_name": department,
                    "total_present_avg": round(avg_total_present, 1)
                })
            else:
                department_data.append({
                    "department_name": department,
                    "total_absent_avg": 0
                })

        return Response(department_data)
    except Exception as e:
        return Response({'error': str(e)}, status=400)







#---------------------------------------Students Dashboards-----------------------------------#



def get_events_by_college_and_department(request, college_id, department_id):
    try:
        # Filter the events based on college_id and department_id
        events = event_master.objects.filter(
            college_id=college_id,
            department_id=department_id,
            deleted=0
        ).values('event_name', 'event_desc', 'dtm_start')

        # Convert the queryset to a list of dictionaries and format the datetime
        data = []
        for event in events:
            dtm_start = event['dtm_start']
            formatted_dtm_start = dtm_start.strftime('%A %I:%M %p')
            data.append({
                'event_name': event['event_name'],
                'event_desc': event['event_desc'],
                'dtm_start': dtm_start,
                'day_time': formatted_dtm_start
            })

        # Return the results as a JSON response
        return JsonResponse(data, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)




#------------upcomming interview------------#

@api_view(['GET'])
def interview_schedule_view(request, college_id, department_id):
    # Get the students from the specified college
    students = candidate_master.objects.filter(college_id=college_id, department_id=department_id).values_list('id', flat=True)
    # Filter interview_schedule based on the students from the specified college
    results = interview_schedule.objects.filter(student_id__in=students) \
        .values('company_id__company_name', 'dtm_interview') \
        .annotate(student_count=Count('student_id')) \
        .order_by('company_id__company_name')

    return Response(results)



#-------------------offer status--------------------#

@api_view(['GET'])
def interview_status_count_view(request, college_id, company_id):
    # Get the students from the specified college
    students = candidate_master.objects.filter(college_id=college_id).values_list('id', flat=True)
    # Filter interview_schedule based on the students from the specified college
    results = interview_schedule.objects.filter(student_id__in=students, comapny_id=company_id) \
        .values('status') \
        .annotate(status_count=Count('status')) \
        .order_by('status')

    return Response(results)



@api_view(['GET'])
def interview_result_view(request, college_id):
    # Get the students from the specified college
    students = candidate_master.objects.filter(college_id=college_id).values_list('id', flat=True)
    # Filter interview_result based on the students from the specified college
    results = interview_result.objects.filter(student_id__in=students) \
        .select_related('company_id', 'student_id__department_id') \
        .values(
            'student_id__students_name',
            'company_id__company_name',
            'student_id__department_id__department',
            'dtm_offer',
            'package'
        ).order_by('student_id__students_name')

    return Response(results)


@api_view(['GET'])
def interview_result_view_emailAddress(request, college_id):
    # Get the students from the specified college
    students = candidate_master.objects.filter(college_id=college_id).values_list('id', flat=True)
    # Filter interview_result based on the students from the specified college
    results = interview_result.objects.filter(student_id__in=students) \
        .select_related('company_id', 'student_id__department_id') \
        .values(
            'student_id__students_name',
            'company_id__company_name',
            'student_id__email_id',
            'company_id__company_profile',
        ).order_by('student_id__students_name')

    return Response(results)


@api_view(['GET'])
def total_no_of_offers(request, college_id):
    # Get the students from the specified college
    student_ids = candidate_master.objects.filter(college_id=college_id).values_list('id', flat=True)
    # Count the number of interview results for these students
    offer_count = interview_result.objects.filter(student_id__in=student_ids).count()

    return Response({'total_offers': offer_count})


@api_view(['GET'])
def pending_requests_count(request):
    # Count the number of student_request records where status is 'pending'
    pending_count = student_request.objects.filter(status='pending').count()

    return Response({'pending_requests': pending_count})



#---------------Student Course Progress----------------------#


@api_view(['GET'])
def course_progress(request, student_id):
    # Count distinct topics in course_content_feedback for the specified student
    course_completed_count = course_content_feedback.objects.filter(student_id=student_id).values('topic_id').distinct().count()
    
    # Count distinct topics in content_master
    content_master_count = content_master.objects.values('id').distinct().count()
    
    # Calculate course in progress
    course_inprogress_count = content_master_count - course_completed_count
    
    results = {
        "Course_Inprogress": course_inprogress_count,
        "Course_Completed": course_completed_count
    }

    return Response(results)




@api_view(['GET'])
def get_tests_by_student(request, student_id):
    # Filter tests_candidates_map by student_id and select required fields
    tests = tests_candidates_map.objects.filter(student_id=student_id).values('test_name', 'total_score', 'is_active')
    
    return Response(list(tests))


#----------------Performance of students-------------------#


@api_view(['GET'])
def get_avg_total_score_by_month(request, student_id):
    # Perform the query
    result = tests_candidates_map.objects.filter(
        question_id__test_type='MCQ Test', 
        student_id=student_id  # Filter by student_id
    ).annotate(
        month=ExtractMonth('dtm_start')
    ).values('month').annotate(
        avg_total_score=Avg('total_score')
    ).order_by('month')
    
    # Convert the QuerySet to a dictionary with months as keys
    result_dict = {item['month']: item['avg_total_score'] for item in result}

    # Create the final list with all months
    result_list = [
        {
            'month': calendar.month_abbr[month],
            'avg_total_score': result_dict.get(month, 0)
        }
        for month in range(1, 13)
    ]

    return Response(result_list)

#-----------------------newly-- student attendance performance-------------------------------#



@api_view(['GET'])
def attendance_report_new(request, reg_no):
    try:
        department_data = {}

        for department, url in sheet_urls.items():
            df = process_sheet(url)
            
            # Filter data by registration number for the current department
            filtered_data = df[df['Register Number'] == reg_no]
            
            # Group by weekdays and calculate average performance for the current department
            weekdays = group_by_weekdays_NEW(filtered_data)
            avg_performance = calculate_avg_performance_NEW(weekdays)
            
            # Store results for each department
            department_data[department] = {
                'weeks': weekdays,
                'avg_performance': avg_performance
            }
        
        return Response(department_data)
    
    except Exception as e:
        return Response({'error': str(e)}, status=400)

def group_by_weekdays_NEW(data):
    weekdays = {'Monday': [], 'Tuesday': [], 'Wednesday': [], 'Thursday': [], 'Friday': [], 'Saturday': [], 'Sunday': []}
    for idx, row in data.iterrows():
        for column_name in data.columns[4:]:  # Assuming attendance starts from the 5th column
            value = row[column_name]
            if value == 'P':
                try:
                    date_obj = datetime.strptime(column_name, '%d-%b')
                    weekday = date_obj.strftime('%A')
                    weekdays[weekday].append(value)
                except ValueError:
                    # Handle invalid or empty dates if needed
                    pass
    return weekdays

def calculate_avg_performance_NEW(weekdays):
    avg_performance = {}
    for day, attendance in weekdays.items():
        total = len(attendance)
        present = attendance.count('P')
        avg_performance[day] = (present / total) * 100 if total else 0
    return avg_performance







#------------------------Django Login----------------------#




@api_view(['POST'])
def add_user_profile(request):
    if request.method == 'POST':
        serializer = UserProfileSerializer(data=request.data)
        print('serializer: ', serializer)

        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)




@api_view(['GET'])
def view_user_profile(request):
    try:
        user_profiles = user_profile.objects.all()  # Renamed variable to user_profiles
    except user_profile.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        serializer = UserProfileSerializer(user_profiles, many=True, context={'request': request})  # Updated to handle multiple objects
        return Response(serializer.data)


@api_view(['PUT'])
def update_user_profile(request, pk):
    try:
        user_profiles = user_profile.objects.get(pk=pk)
    except user_profile.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)

    if request.method == 'PUT':
        serializer = UserProfileSerializer(user_profiles, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@api_view(['DELETE'])
def delete_user_profile(request, pk):
    try:
        user_profile = user_profile.objects.get(pk=pk)
    except user_profile.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)

    if request.method == 'DELETE':
        user_profile.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)





@api_view(['GET'])
def custom_user_profiles(request):
    try:
        user_profiles = user_profile.objects.all()
    except user_profile.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)

    data = []
    for profile in user_profiles:
        college = None
        if profile.college_id:
            college = profile.college_id.college
        
        serialized_data = {
            "id": profile.id,
            "username": profile.user.username,
            "email": profile.user.email,
            "password": "****",  # Avoid exposing passwords directly
            "role": profile.role,
            "college_id": profile.college_id.id if profile.college_id else None,  # Use college_id instead of the object
            "college": college if college else None  # Serialize college_master object
        }
        data.append(serialized_data)

    return Response(data)

#------------Interview Master------------------#

@api_view(['GET'])
def get_interview_master(request):
    interviews = interview_master.objects.filter(deleted=0)
    # logger.info("Fetching interviews where deleted=0")
   
    interview_data=[]
    for interview in  interviews:
        company_id=None
        company_name=None
        job_id=None
        
        if interview.company_id:
           company_id=interview.company_id.id
           company_name=interview.company_id.company_name
        if interview.job_id:
           job_id=interview.job_id.id

        interview_data.append({
            'id': interview.id,
            'company_id':company_id,
            'company_name': company_name,
            'job_id':job_id,
            'post_name': interview.post_name,
            'dtm_interview': interview.dtm_interview,
            'interview_type': interview.interview_type,
            })
    return Response(interview_data)



class interview_master_createAPIView(generics.CreateAPIView):
    queryset = interview_master.objects.all()
    serializer_class =interviewMasterSerializer

    def post(self, request, *args, **kwargs):
        # logger.info("Creating a new event")
        response = super().post(request, *args, **kwargs)
        # logger.info("Created a new interview_master successfully")
        return response

class interview_master_UpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = interview_master.objects.all()
    serializer_class =interviewMasterSerializer

    def put(self, request, *args, **kwargs):
        # logger.info(f"Updating interview_master with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        # logger.info(f"Updated interview_master with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        # logger.info(f"Partially updating interview_master with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        # logger.info(f"Partially updated interview_master with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_interview_master(request, pk):
    # logger.info(f"Attempting to mark interview_master with id {pk} as deleted")
    try:
        print("Entering Function..")
        interview=interview_master.objects.get(id=pk)

        print("interview_master: ",interview)
    except interview_master.DoesNotExist:
        return JsonResponse("interview_master not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    interview.deleted = 1
    interview.save()

    # logger.info(f"Marked interview_master with id {pk} as deleted successfully")

    print("course: ",interview)

    return JsonResponse("interview_master 'deleted' field updated successfully", safe=False)



#---------------------------------Interview Schedule---------------------------#


@api_view(['GET'])
def get_interview_schedule(request):
    interviews = interview_schedule.objects.filter(deleted=0)
    # logger.info("Fetching interviews where deleted=0")
   
    interview_data=[]
    for interview in  interviews:
        company_id=None
        company_name=None
        student_id=None
        student_name = None
        post_name_id=None
        post_name=None
        
        if interview.company_id:
           company_id=interview.company_id.id
           company_name=interview.company_id.company_name
           
        if interview.student_id:
           student_id=interview.student_id.id
           student_name=interview.student_id.students_name

        if interview.post_name_id:
           post_name_id=interview.post_name_id.id
           post_name=interview.post_name_id.post_name

        interview_data.append({
            'id': interview.id,
            'company_id':company_id,
            'company_name': company_name,
            'student_id':student_id,
            'student_name': student_name,
            'post_name_id':post_name_id,
            'post_name': post_name,
            'dtm_interview': interview.dtm_interview,
            'status': interview.status,
            'interview_round': interview.interview_round,
            })
    return Response(interview_data)



class interview_schedule_createAPIView(generics.CreateAPIView):
    queryset = interview_schedule.objects.all()
    serializer_class =interviewScheduleSerializer

    def post(self, request, *args, **kwargs):
        # logger.info("Creating a new interview_schedule")
        response = super().post(request, *args, **kwargs)
        # logger.info("Created a new interview_schedule successfully")
        return response

class interview_schedule_UpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = interview_schedule.objects.all()
    serializer_class =interviewScheduleSerializer

    def put(self, request, *args, **kwargs):
        # logger.info(f"Updating interview_schedule with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        # logger.info(f"Updated interview_schedule with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        # logger.info(f"Partially updating interview_schedule with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        # logger.info(f"Partially updated interview_schedule with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_interview_schedule(request, pk):
    # logger.info(f"Attempting to mark interview_schedule with id {pk} as deleted")
    try:
        print("Entering Function..")
        interview=interview_schedule.objects.get(id=pk)

        print("interview_schedule: ",interview)
    except interview_schedule.DoesNotExist:
        return JsonResponse("interview_schedule not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    interview.deleted = 1
    interview.save()

    # logger.info(f"Marked interview_schedule with id {pk} as deleted successfully")

    print("course: ",interview)

    return JsonResponse("interview_schedule 'deleted' field updated successfully", safe=False)


#-------------------------Interview Result-------------------------#




@api_view(['GET'])
def get_interview_result(request):
    interviews = interview_result.objects.filter(deleted=0)
    # logger.info("Fetching interviews where deleted=0")
   
    interview_data=[]
    for interview in  interviews:
        company_id=None
        company_name=None
        student_id=None
        student_name = None
        post_name_id=None
        post_name=None
        
        if interview.company_id:
           company_id=interview.company_id.id
           company_name=interview.company_id.company_name
           
        if interview.student_id:
           student_id=interview.student_id.id
           student_name=interview.student_id.students_name

        if interview.post_name_id:
           post_name_id=interview.post_name_id.id
           post_name=interview.post_name_id.post_name

        interview_data.append({
            'id': interview.id,
            'company_id':company_id,
            'company_name': company_name,
            'student_id':student_id,
            'student_name': student_name,
            'post_name_id':post_name_id,
            'post_name': post_name,
            'dtm_interview': interview.dtm_interview,
            'status': interview.status,
            'interview_round': interview.interview_round,
            })
    return Response(interview_data)



class interview_result_createAPIView(generics.CreateAPIView):
    queryset = interview_result.objects.all()
    serializer_class =interviewResultSerilaizer

    def post(self, request, *args, **kwargs):
        # logger.info("Creating a new interview_result")
        response = super().post(request, *args, **kwargs)
        # logger.info("Created a new interview_result successfully")
        return response

class interview_result_UpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = interview_result.objects.all()
    serializer_class =interviewResultSerilaizer

    def put(self, request, *args, **kwargs):
        # logger.info(f"Updating interview_result with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        # logger.info(f"Updated interview_result with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        # logger.info(f"Partially updating interview_result with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        # logger.info(f"Partially updated interview_result with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_interview_result(request, pk):
    # logger.info(f"Attempting to mark interview_result with id {pk} as deleted")
    try:
        print("Entering Function..")
        interview=interview_result.objects.get(id=pk)

        print("interview_result: ",interview)
    except interview_result.DoesNotExist:
        return JsonResponse("interview_result not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    interview.deleted = 1
    interview.save()

    # logger.info(f"Marked interview_result with id {pk} as deleted successfully")

    print("course: ",interview)

    return JsonResponse("interview_result 'deleted' field updated successfully", safe=False)


#-------------------------------Student Request------------------------------------#


@api_view(['GET'])
def get_student_request(request):
    stu_queries = student_request.objects.filter(deleted=0)
    # logger.info("Fetching student_request where deleted=0")
   
    query_data=[]
    for stu_query in  stu_queries:
        student_id=None
        student_name=None
        
        if stu_query.student_id:
           student_id=stu_query.student_id.id
           student_name=stu_query.student_id.students_name

        query_data.append({
            'id': stu_query.id,
            'student_id':student_id,
            'student_name':student_name,
            'dtm_request': stu_query.dtm_request,
            'student_query':stu_query.student_query,
            'approved_by': stu_query.approved_by,
            'dtm_approved': stu_query.dtm_approved,
            'dtm_student_update': stu_query.dtm_student_update,
            'status': stu_query.status,
            })
    return Response(query_data)



class student_request_createAPIView(generics.CreateAPIView):
    queryset = student_request.objects.all()
    serializer_class =studentRequestSerializer

    def post(self, request, *args, **kwargs):
        # logger.info("Creating a new student_request")
        response = super().post(request, *args, **kwargs)
        # logger.info("Created a new student_request successfully")
        return response

class student_request_UpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = student_request.objects.all()
    serializer_class =studentRequestSerializer

    def put(self, request, *args, **kwargs):
        # logger.info(f"Updating student_request with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        # logger.info(f"Updated student_request with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        # logger.info(f"Partially updating student_request with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        # logger.info(f"Partially updated student_request with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_student_request(request, pk):
    # logger.info(f"Attempting to mark student_request with id {pk} as deleted")
    try:
        print("Entering Function..")
        stu_query=student_request.objects.get(id=pk)

        print("interview_master: ",stu_query)
    except student_request.DoesNotExist:
        return JsonResponse("student_request not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    stu_query.deleted = 1
    stu_query.save()

    # logger.info(f"Marked student_request with id {pk} as deleted successfully")

    print("course: ",stu_query)

    return JsonResponse("student_request 'deleted' field updated successfully", safe=False)




#----------------get LMS Updated id data-------------------#


@api_view(['GET'])
def update_LMS_id(request, id):
    # Get the content_master object by ID
    try:
        lms = content_master.objects.get(id=id)
    except content_master.DoesNotExist:
        return Response(status=404)

    # Serialize the content_master object
    serializer = contentSerializers(lms)
    
    return Response(serializer.data)



@api_view(['GET'])
def update_LMS_Topic_id(request, id):
    # Get the content_master object by ID
    try:
        lms = topic_master.objects.get(id=id)
    except topic_master.DoesNotExist:
        return Response(status=404)

    # Serialize the content_master object
    serializer = topicSerializers(lms)
    
    return Response(serializer.data)






@api_view(['GET'])
def get_test_type_category(request, test_name):
    try:
        test_master_instance = test_master.objects.get(test_name=test_name)
        test_type_category = test_master_instance.test_type_id.test_type_categories
        return Response({'test_type_category': test_type_category})
    except test_master.DoesNotExist:
        return Response({'error': 'Test name does not exist'}, status=404)



#------------Need Candidate info-----------------------#


@api_view(['GET'])
def get_test_candidates_NeedInfo(request, username):
    candidates = tests_candidates_map.objects.filter(student_id__user_name=username).select_related('student_id')
    serializer = TestCandidateMapSerializerNeedInfo(candidates, many=True)
    return Response(serializer.data)





#---------------------Student Dasboard New--------------------------#

#--Total Test Taken--------#

@api_view(['GET'])
def active_tests_count(request, student_id):
    count = tests_candidates_map.objects.filter(is_active=True, student_id=student_id).count()
    return Response({'count': count})


#---Total no.of offers------#


@api_view(['GET'])
def get_number_of_offers(request, candidate_id):
    try:
        candidate = candidate_master.objects.get(id=candidate_id)
        number_of_offers = candidate.number_of_offers
    except candidate_master.DoesNotExist:
        return Response({'error': 'Candidate not found'}, status=404)
    return Response({'number_of_offers': number_of_offers})


#----Request count---------#


@api_view(['GET'])
def count_student_requests(request, student_id):
    try:
        request_count = student_request.objects.filter(student_id=student_id).count()
    except student_request.DoesNotExist:
        return Response({'error': 'Student not found'}, status=404)
    return Response({'request_count': request_count})


#------Apditute Avg Score------#

@api_view(['GET'])
def get_monthly_avg_total_score_apditute(request, student_id):
    try:
        # Subquery to get the question_type id where question_type is 'Apditute'
        question_type_subquery = question_type.objects.filter(
            question_type='Aptitude'
        ).values('id')
        print('question_type_subquery: ', question_type_subquery)

        # Subquery to get the test_type id where test_type is 'MCQ Test'
        test_type_subquery = test_type.objects.filter(
            test_type='MCQ Test'
        ).values('id')
        print('test_type_subquery: ', test_type_subquery)

        # Subquery to get the test_name from test_master where the question_type and test_type match
        test_name_subquery = test_master.objects.filter(
            question_type_id__in=Subquery(question_type_subquery),
            test_type_id__in=Subquery(test_type_subquery)
        ).values('test_name')
        print('test_name_subquery: ', test_name_subquery)

        # Main query to get the average total score grouped by month
        monthly_avg_scores = tests_candidates_map.objects.filter(
            test_name__in=Subquery(test_name_subquery),
            student_id=student_id
        ).annotate(month=ExtractMonth('dtm_start')).values('month').annotate(avg_total_score=Avg('total_score')).order_by('month')

        print('monthly_avg_scores: ', monthly_avg_scores)

        # Convert month numbers to month names and create a full list with 0 defaults
        full_months = {i: {"month": calendar.month_abbr[i], "avg_score": 0} for i in range(1, 13)}
        
        print('full months: ', full_months)

        for entry in monthly_avg_scores:
            full_months[entry["month"]]["avg_score"] = entry["avg_total_score"]

        return Response(list(full_months.values()))
    except Exception as e:
        return Response({'error': str(e)}, status=500)




#------SoftSkills Avg Score------#

@api_view(['GET'])
def get_monthly_avg_total_score_softskill(request, student_id):
    try:
        # Subquery to get the question_type id where question_type is 'Apditute'
        question_type_subquery = question_type.objects.filter(
            question_type='Softskills'
        ).values('id')
        print('question_type_subquery: ', question_type_subquery)

        # Subquery to get the test_type id where test_type is 'MCQ Test'
        test_type_subquery = test_type.objects.filter(
            test_type='MCQ Test'
        ).values('id')
        print('test_type_subquery: ', test_type_subquery)

        # Subquery to get the test_name from test_master where the question_type and test_type match
        test_name_subquery = test_master.objects.filter(
            question_type_id__in=Subquery(question_type_subquery),
            test_type_id__in=Subquery(test_type_subquery)
        ).values('test_name')
        print('test_name_subquery: ', test_name_subquery)

        # Main query to get the average total score grouped by month
        monthly_avg_scores = tests_candidates_map.objects.filter(
            test_name__in=Subquery(test_name_subquery),
            student_id=student_id
        ).annotate(month=ExtractMonth('dtm_start')).values('month').annotate(avg_total_score=Avg('total_score')).order_by('month')

        print('monthly_avg_scores: ', monthly_avg_scores)

        # Convert month numbers to month names and create a full list with 0 defaults
        full_months = {i: {"month": calendar.month_abbr[i], "avg_score": 0} for i in range(1, 13)}
        
        print('full months: ', full_months)

        for entry in monthly_avg_scores:
            full_months[entry["month"]]["avg_score"] = entry["avg_total_score"]

        return Response(list(full_months.values()))
    except Exception as e:
        return Response({'error': str(e)}, status=500)




#------Technical Avg Score------#

@api_view(['GET'])
def get_monthly_avg_total_score_technical(request, student_id):
    try:
        # Subquery to get the question_type id where question_type is 'Apditute'
        question_type_subquery = question_type.objects.filter(
            question_type='Technical'
        ).values('id')
        print('question_type_subquery: ', question_type_subquery)

        # Subquery to get the test_type id where test_type is 'MCQ Test'
        test_type_subquery = test_type.objects.filter(
            test_type='MCQ Test'
        ).values('id')
        print('test_type_subquery: ', test_type_subquery)

        # Subquery to get the test_name from test_master where the question_type and test_type match
        test_name_subquery = test_master.objects.filter(
            question_type_id__in=Subquery(question_type_subquery),
            test_type_id__in=Subquery(test_type_subquery)
        ).values('test_name')
        print('test_name_subquery: ', test_name_subquery)

        # Main query to get the average total score grouped by month
        monthly_avg_scores = tests_candidates_map.objects.filter(
            test_name__in=Subquery(test_name_subquery),
            student_id=student_id
        ).annotate(month=ExtractMonth('dtm_start')).values('month').annotate(avg_total_score=Avg('total_score')).order_by('month')

        print('monthly_avg_scores: ', monthly_avg_scores)

        # Convert month numbers to month names and create a full list with 0 defaults
        full_months = {i: {"month": calendar.month_abbr[i], "avg_score": 0} for i in range(1, 13)}
        
        print('full months: ', full_months)

        for entry in monthly_avg_scores:
            full_months[entry["month"]]["avg_score"] = entry["avg_total_score"] if entry["avg_total_score"] is not None else 0

        return Response(list(full_months.values()))
    except Exception as e:
        return Response({'error': str(e)}, status=500)




#------Coding Avg Score------#


@api_view(['GET'])
def get_monthly_avg_total_score_Coding(request, student_id):
    try:
        # Subquery to get the test_type id where test_type is 'CodingTest'
        test_type_subquery = test_type.objects.filter(
            test_type='Coding Test'
        ).values('id')

        # Subquery to get the test_name from test_master where the test_type matches
        test_name_subquery = test_master.objects.filter(
            test_type_id__in=Subquery(test_type_subquery)
        ).values('test_name')

        # Main query to get the average total score grouped by month
        monthly_avg_scores = tests_candidates_map.objects.filter(
            test_name__in=Subquery(test_name_subquery),
            student_id=student_id
        ).annotate(month=ExtractMonth('dtm_start')).values('month').annotate(avg_total_score=Avg('total_score')).order_by('month')

        # Convert month numbers to abbreviated month names and create a full list with 0 defaults
        full_months = {i: {"month": calendar.month_abbr[i], "avg_score": 0} for i in range(1, 13)}

        for entry in monthly_avg_scores:
            full_months[entry["month"]]["avg_score"] = entry["avg_total_score"]

        return Response(list(full_months.values()))
    except Exception as e:
        return Response({'error': str(e)}, status=500)

#------------------Offer Chart where college id


@api_view(['GET'])
def get_candidate_all_college_id(request, college_id):
    try:
        candidatelist = candidate_master.objects.filter(deleted=0, college_id=college_id).select_related('college_id', 'department_id').values(
            'id', 
            'college_id__college',  # Get college name
            'students_name', 
            'user_name', 
            'registration_number', 
            'gender',
            'email_id', 
            'mobile_number', 
            'year', 
            'cgpa', 
            'department_id__department',  # Get department name
            'marks_10th',
            'marks_12th', 
            'marks_semester_wise', 
            'history_of_arrears', 
            'standing_arrears',
            'number_of_offers', 
            'text', 
            'it_of_offers', 
            'core_of_offers', 
            'skill_id',
        )
        candidate_data = list(candidatelist)
        return Response(candidate_data)
    except Exception as e:
        return Response({'error': str(e)}, status=500)

#------------------number of offers --college_id--------

@api_view(['GET'])
def get_number_of_offers_college_id(request, college_id):
    try:
        # Filter candidate_master records by college_id
        candidates = candidate_master.objects.filter(college_id=college_id)
        
        # If no candidates are found, return an appropriate error message
        if not candidates.exists():
            return Response({'error': 'No candidates found for the given college ID'}, status=404)
        
        # Aggregate the number of offers for all the candidates
        total_number_of_offers = sum(candidate.number_of_offers for candidate in candidates)

        return Response({'number_of_offers': total_number_of_offers})
    except Exception as e:
        return Response({'error': str(e)}, status=500)


#----Request count---college_id---------#


@api_view(['GET'])
def count_student_requests_college_id(request, college_id):
    try:
        # Filter candidate_master records by college_id
        candidate_ids = candidate_master.objects.filter(college_id=college_id).values_list('id', flat=True)
        
        # Count the number of student_request records with student_id in the filtered candidate ids
        request_count = student_request.objects.filter(student_id__in=candidate_ids).count()
        
        return Response({'request_count': request_count})
    except candidate_master.DoesNotExist:
        return Response({'error': 'College not found'}, status=404)






#--------------------------------------Student Attendance Table View-------------------------#


# def view_attendance_summary(request):
#     data = list(StudentAttendance.objects.values())
#     return JsonResponse(data, safe=False)




#----------------------student plan--------------------#





@api_view(['GET'])
def get_schedule_with_tests(request, dtm_start, student_id):
    try:
        # Convert dtm_start to a timezone-aware datetime
        dtm_start = timezone.make_aware(timezone.datetime.strptime(dtm_start, '%Y-%m-%d'))

        # Get the college_id, department_id, and year for the given student_id from candidate_master
        candidate = candidate_master.objects.get(id=student_id)
        print('candidate: ', candidate)

        college_id = candidate.college_id
        print('college_id: ', college_id)

        department_id = candidate.department_id
        print('department_id: ', department_id)

        year = candidate.year
        print('year: ', year)

        # Filter course_schedule based on college_id, department_id, year, and dtm_start
        schedule_with_tests = course_schedule.objects.filter(
            college_id=college_id,
            department_id=department_id,
            year=year,
            dtm_start_student__date=dtm_start.date()
        ).select_related('college_id', 'department_id')
        print('schedule_with_tests: ', schedule_with_tests)

        # Subquery to get test names and dtm_start for the given student_id
        test_names = tests_candidates_map.objects.filter(
            student_id=student_id,
            college_id=college_id,
            department_id=department_id,
            dtm_start__date=dtm_start.date()
        ).values('test_name', 'dtm_start')
        print('test_names: ', test_names)

        # Get attendance data grouped by course_schedule_id
        attendance_data = training_attendance_sheet.objects.filter(
            course_schedule_id__in=schedule_with_tests
        ).values('course_schedule_id').annotate(
            total_present=Count('present'),
            total_absent=Count('absent')
        )
        attendance_dict = {attendance['course_schedule_id']: attendance for attendance in attendance_data}

        # Compile results
        results = []
        if not schedule_with_tests:
            # If schedule_with_tests is empty, return test_names directly
            for test in test_names:
                results.append({
                    'college_name': candidate.college_id.college,
                    'department_name': candidate.department_id.department,
                    'year': candidate.year,
                    'topic_name': '',
                    'subtopic_name': '',
                    'test_name': test['test_name'],
                    'dtm_start': test['dtm_start'],
                    'total_present': 0,
                    'total_absent': 0
                })
        else:
            # Iterate through schedule_with_tests and test_names
            for schedule in schedule_with_tests:
                total_present = attendance_dict.get(schedule.id, {}).get('total_present', 0)
                total_absent = attendance_dict.get(schedule.id, {}).get('total_absent', 0)
                for test in test_names:
                    results.append({
                        'college_name': schedule.college_id.college,
                        'department_name': schedule.department_id.department,
                        'year': schedule.year,
                        'topic_name': schedule.topic,
                        'subtopic_name': schedule.sub_topic,
                        'test_name': test['test_name'],
                        'dtm_start': test['dtm_start'],
                        'total_present': total_present,
                        'total_absent': total_absent
                    })

        return Response(results)
    except Exception as e:
        return Response({'error': str(e)}, status=500)





@api_view(['GET'])
def get_tests_candidates_map_MCQ(request, user_name):
    # Get the current date
    current_date = datetime.now().date()
    print('current_date: ', current_date)

    # Apply the filters
    tests_candidates = tests_candidates_map.objects.filter(
        deleted=0,
        student_id__user_name=user_name,
        question_id__test_type='MCQ Test',
        is_active=False,
       # dtm_start__date__gte=current_date
    ).filter(
        Q(dtm_start__date__gte=current_date) | Q(dtm_start__date__lt=current_date, dtm_end__date__gte=current_date)
    )
    
    # Initialize logger info if needed
    # logger.info('Fetching test-candidate-map')

    test_candidate_map_data = []
    for testing in tests_candidates:
        question_id = None  # Initialize question_type variable
        student_id = None
        student_name = None
        user_name = None
        college_id = None
        college_id_id = None
        department_id = None
        department_id_id = None
        department_name = None
        test_type = None
        rules_id = None
        rules_id_id = None
        instruction = None
        skill_type = None
        question_paper_name = None
        test_type = None

        if testing.rules_id:
            rules_id_id = testing.rules_id.id
            rules_id = testing.rules_id.rule_name
            instruction = testing.rules_id.instruction

        if testing.department_id:
            department_id = testing.department_id.department
            department_id_id = testing.department_id.id

        if testing.question_id:
            question_id = testing.question_id.id
            question_paper_name = testing.question_id.question_paper_name
            test_type = testing.question_id.test_type

        if testing.student_id:
            student_id = testing.student_id.id
            student_name = testing.student_id.students_name
            user_name = testing.student_id.user_name

        if testing.college_id:
            college_id = testing.college_id.college
            college_id_id = testing.college_id.id

        test_candidate_map_data.append({
            'id': testing.id,
            'test_name': testing.test_name,
            'college_id_id': college_id_id,
            'college_id': college_id,
            'department_id_id': department_id_id,
            'department_id': department_id,
            'question_id': question_id,
            'question_paper_name': question_paper_name,
            'test_type': test_type,
            'student_id': student_id,
            'student_name': student_name,
            'user_name': user_name,
            'dtm_start': testing.dtm_start,
            'dtm_end': testing.dtm_end,
            'attempt_count': testing.attempt_count,
            'is_actual_test': testing.is_actual_test,
            'is_active': testing.is_active,
            'duration': testing.duration,
            'duration_type': testing.duration_type,
            'year': testing.year,
            'rules_id': rules_id_id,
            'rules': rules_id,
            'instruction': instruction,
            'need_candidate_info': testing.need_candidate_info,
            'total_score': testing.total_score,
            'avg_mark': testing.avg_mark
        })

    return Response(test_candidate_map_data)




@api_view(['GET'])
def get_tests_candidates_map_Coding(request, user_name):
    # Get the current date
    current_date = datetime.now().date()
    print('current_date: ', current_date)

    # Apply the filters
    tests_candidates = tests_candidates_map.objects.filter(
        deleted=0,
        student_id__user_name=user_name,
        question_id__test_type='Coding Test',
        is_active=False,
       # dtm_start__date__gte=current_date
    ).filter(
        Q(dtm_start__date__gte=current_date) | Q(dtm_start__date__lt=current_date, dtm_end__date__gte=current_date)
    )
    
    # Initialize logger info if needed
    # logger.info('Fetching test-candidate-map')

    test_candidate_map_data = []
    for testing in tests_candidates:
        question_id = None  # Initialize question_type variable
        student_id = None
        student_name = None
        user_name = None
        college_id = None
        college_id_id = None
        department_id = None
        department_id_id = None
        department_name = None
        test_type = None
        rules_id = None
        rules_id_id = None
        instruction = None
        skill_type = None
        question_paper_name = None
        test_type = None

        if testing.rules_id:
            rules_id_id = testing.rules_id.id
            rules_id = testing.rules_id.rule_name
            instruction = testing.rules_id.instruction

        if testing.department_id:
            department_id = testing.department_id.department
            department_id_id = testing.department_id.id

        if testing.question_id:
            question_id = testing.question_id.id
            question_paper_name = testing.question_id.question_paper_name
            test_type = testing.question_id.test_type

        if testing.student_id:
            student_id = testing.student_id.id
            student_name = testing.student_id.students_name
            user_name = testing.student_id.user_name

        if testing.college_id:
            college_id = testing.college_id.college
            college_id_id = testing.college_id.id

        test_candidate_map_data.append({
            'id': testing.id,
            'test_name': testing.test_name,
            'college_id_id': college_id_id,
            'college_id': college_id,
            'department_id_id': department_id_id,
            'department_id': department_id,
            'question_id': question_id,
            'question_paper_name': question_paper_name,
            'test_type': test_type,
            'student_id': student_id,
            'student_name': student_name,
            'user_name': user_name,
            'dtm_start': testing.dtm_start,
            'dtm_end': testing.dtm_end,
            'attempt_count': testing.attempt_count,
            'is_actual_test': testing.is_actual_test,
            'is_active': testing.is_active,
            'duration': testing.duration,
            'duration_type': testing.duration_type,
            'year': testing.year,
            'rules_id': rules_id_id,
            'rules': rules_id,
            'instruction': instruction,
            'need_candidate_info': testing.need_candidate_info,
            'total_score': testing.total_score,
            'avg_mark': testing.avg_mark
        })

    return Response(test_candidate_map_data)




#----------------------PLACEMENT-----------------------#



@api_view(['GET'])
def get_tests_Reports_placement(request, college_id):
    tests_candidates = tests_candidates_map.objects.filter(deleted=0, college_id=college_id)
     ##logger.info('Fetching test-candidate-map')
    
    test_candidate_map_data = []
    for testing in tests_candidates:
        question_name = None
        student_id = None
        student_name = None
        user_name = None
        college_id = None
        department_id = None
        test_type = None
        rules_id = None
        instruction = None
        email_id = None
        mobile_number = None
        gender = None
        registration_number = None

        if testing.rules_id:
            rules_id = testing.rules_id.rule_name
            instruction = testing.rules_id.instruction

        if testing.department_id:
            department_id = testing.department_id.department

        if testing.question_id:
            question_name = testing.question_id.question_paper_name

        if testing.student_id:
            student_id = testing.student_id.id
            student_name = testing.student_id.students_name
            user_name = testing.student_id.user_name
            email_id = testing.student_id.email_id
            mobile_number = testing.student_id.mobile_number
            gender = testing.student_id.gender
            registration_number = testing.student_id.registration_number

        if testing.college_id:
            college_id = testing.college_id.college

        # Format dates
        dtm_start_formatted = django_format_date(localtime(testing.dtm_start), 'd-m-Y h:i A')
        dtm_end_formatted = django_format_date(localtime(testing.dtm_end), 'd-m-Y h:i A')

        test_candidate_map_data.append({
            'id': testing.id,
            'test_name': testing.test_name,
            'college_id': college_id,
            'department_id': department_id,
            'question_id': question_name,
            'student_id': student_id,
            'registration_number': registration_number,
            'email_id': email_id,
            'mobile_number': mobile_number,
            'gender': gender,
            'student_name': student_name,
            'user_name': user_name,
            'dtm_start': dtm_start_formatted,
            'dtm_end': dtm_end_formatted,
            'attempt_count': testing.attempt_count,
            'is_actual_test': testing.is_actual_test,
            'is_active': testing.is_active,
            'duration': testing.duration,
            'duration_type': testing.duration_type,
            'year': testing.year,
            'rules': rules_id,
            'instruction': instruction,
            'need_candidate_info': testing.need_candidate_info,
            'total_score': testing.total_score,
            'avg_mark': testing.avg_mark
        })
    return Response(test_candidate_map_data)





@api_view(['GET'])
def get_group_test_name_placement(request, college_id):
    # Join tests_candidates_map with test_master and filter by test_name and deleted=0
    tests_candidates = (
        tests_candidates_map.objects.filter(deleted=0, college_id=college_id)
        .values('test_name', 'question_id', 'question_id__question_paper_name', 'dtm_start', 'dtm_end', 'dtm_created')
        .annotate(student_count=Count('student_id'))
        .order_by(
            Case(
                When(dtm_created__isnull=True, then=1),
                default=0,
                output_field=DateTimeField()
            ),
            '-dtm_created'
        )
    )

    test_candidate_map_data = []
    for testing in tests_candidates:
        # Format dates
        dtm_start_formatted = django_format_date(localtime(testing['dtm_start']), 'd-m-Y h:i A')
        dtm_end_formatted = django_format_date(localtime(testing['dtm_end']), 'd-m-Y h:i A')

        test_candidate_map_data.append({
            'test_name': testing['test_name'],
            'question_id': testing['question_id'],
            'question_paper_name': testing['question_id__question_paper_name'],
            'dtm_start': dtm_start_formatted,
            'dtm_end': dtm_end_formatted,
            'student_count': testing['student_count'],
            'dtm_created': testing['dtm_created']
        })

    return Response(test_candidate_map_data)



@api_view(['GET'])
def get_tests_Reports_placement_Candidates(request, college_id, test_name):
    tests_candidates = tests_candidates_map.objects.filter(deleted=0, college_id=college_id, test_name=test_name)
     ##logger.info('Fetching test-candidate-map')
    
    test_candidate_map_data = []
    for testing in tests_candidates:
        question_name = None
        student_id = None
        student_name = None
        user_name = None
        college_id = None
        department_id = None
        test_type = None
        rules_id = None
        instruction = None
        email_id = None
        mobile_number = None
        gender = None
        registration_number = None

        if testing.rules_id:
            rules_id = testing.rules_id.rule_name
            instruction = testing.rules_id.instruction

        if testing.department_id:
            department_id = testing.department_id.department

        if testing.question_id:
            question_name = testing.question_id.question_paper_name

        if testing.student_id:
            student_id = testing.student_id.id
            student_name = testing.student_id.students_name
            user_name = testing.student_id.user_name
            email_id = testing.student_id.email_id
            mobile_number = testing.student_id.mobile_number
            gender = testing.student_id.gender
            registration_number = testing.student_id.registration_number

        if testing.college_id:
            college_id = testing.college_id.college

        # Format dates
        dtm_start_formatted = django_format_date(localtime(testing.dtm_start), 'd-m-Y h:i A')
        dtm_end_formatted = django_format_date(localtime(testing.dtm_end), 'd-m-Y h:i A')

        test_candidate_map_data.append({
            'id': testing.id,
            'test_name': testing.test_name,
            'college_id': college_id,
            'department_id': department_id,
            'question_id': question_name,
            'student_id': student_id,
            'registration_number': registration_number,
            'email_id': email_id,
            'mobile_number': mobile_number,
            'gender': gender,
            'student_name': student_name,
            'user_name': user_name,
            'dtm_start': dtm_start_formatted,
            'dtm_end': dtm_end_formatted,
            'attempt_count': testing.attempt_count,
            'is_actual_test': testing.is_actual_test,
            'is_active': testing.is_active,
            'duration': testing.duration,
            'duration_type': testing.duration_type,
            'year': testing.year,
            'rules': rules_id,
            'instruction': instruction,
            'need_candidate_info': testing.need_candidate_info,
            'total_score': testing.total_score,
            'avg_mark': testing.avg_mark
        })
    return Response(test_candidate_map_data)




@api_view(['GET'])
def get_candidate_all_job_master1(request):
    try:
        # Fetch the last added job from job_master
        last_job = job_master.objects.filter(deleted=0).order_by('-id').first()
        print('last job: ', last_job)
        print('last_job.college_id: ', last_job.college_id)
        print('last_job.department_id: ', last_job.department_id)
        print('last_job.gender: ', last_job.gender)
        print('last_job.marks_10th: ', last_job.marks_10th)
        print('last_job.marks_12th: ', last_job.marks_12th)
        print('last_job.cgpa: ', last_job.cgpa)
        print('last_job.history_of_arrears: ', last_job.history_of_arrears)
        print('last_job.standing_arrears: ', last_job.standing_arrears)
        print('last_job.skill_id: ', last_job.skill_id)

        if not last_job:
            return Response({'error': 'No job found'}, status=404)
        
        # Fetch candidates based on the last job criteria
        candidates = candidate_master.objects.filter(
            deleted=0,
            college_id=last_job.college_id,
            department_id=last_job.department_id,
            gender=last_job.gender,
            marks_10th=last_job.marks_10th,
            marks_12th=last_job.marks_12th,
            cgpa=last_job.cgpa,
            history_of_arrears=str(last_job.history_of_arrears) if last_job.history_of_arrears else None,
            standing_arrears=last_job.standing_arrears,
            skill_id__in=[last_job.skill_id.id] if last_job.skill_id else []
        ).distinct()

        print('candidates: ', candidates)

        # Create a list to store the final output
        candidate_data = []

        for candidate in candidates:
            skills = candidate.skill_id.all()
            skill_ids = [skill.id for skill in skills]
            candidate_info = {
                'id': candidate.id,
                'college_id': candidate.college_id.college if candidate.college_id else None,
                'students_name': candidate.students_name,
                'user_name': candidate.user_name,
                'registration_number': candidate.registration_number,
                'gender': candidate.gender,
                'email_id': candidate.email_id,
                'mobile_number': candidate.mobile_number,
                'year': candidate.year,
                'cgpa': candidate.cgpa,
                'department_id': candidate.department_id.department if candidate.department_id else None,
                'marks_10th': candidate.marks_10th,
                'marks_12th': candidate.marks_12th,
                'marks_semester_wise': candidate.marks_semester_wise,
                'history_of_arrears': candidate.history_of_arrears,
                'standing_arrears': candidate.standing_arrears,
                'number_of_offers': candidate.number_of_offers,
                'text': candidate.text,
                'it_of_offers': candidate.it_of_offers,
                'core_of_offers': candidate.core_of_offers,
                'skill_id': skill_ids
            }
            candidate_data.append(candidate_info)

        print('candidate_data: ', candidate_data)

        return Response(candidate_data)

    except Exception as e:
        return Response({'error': str(e)}, status=500)



@api_view(['GET'])
def get_candidate_all_job_master_AND(request):
    try:
        # Fetch the last added job from job_master
        last_job = job_master.objects.filter(deleted=0).order_by('-id').first()
        if not last_job:
            return Response({'error': 'No job found'}, status=404)
        
        print('last job: ', last_job)
        print('last_job.college_id: ', last_job.college_id)
        print('last_job.department_id: ', last_job.department_id)
        print('last_job.gender: ', last_job.gender)
        print('last_job.marks_10th: ', last_job.marks_10th)
        print('last_job.marks_12th: ', last_job.marks_12th)
        print('last_job.cgpa: ', last_job.cgpa)
        print('last_job.history_of_arrears: ', last_job.history_of_arrears)
        print('last_job.standing_arrears: ', last_job.standing_arrears)
        print('last_job.skill_id: ', last_job.skill_id.all())

        # Check each filter step-by-step
        candidates = candidate_master.objects.filter(deleted=0)
        print('Initial candidates count: ', candidates.count())

        candidates = candidates.filter(college_id=last_job.college_id)
        print('After college_id filter count: ', candidates.count())

        candidates = candidates.filter(department_id=last_job.department_id)
        print('After department_id filter count: ', candidates.count())

        candidates = candidates.filter(gender=last_job.gender)
        print('After gender filter count: ', candidates.count())

        candidates = candidates.filter(marks_10th=last_job.marks_10th)
        print('After marks_10th filter count: ', candidates.count())

        candidates = candidates.filter(marks_12th=last_job.marks_12th)
        print('After marks_12th filter count: ', candidates.count())

        candidates = candidates.filter(cgpa=last_job.cgpa)
        print('After cgpa filter count: ', candidates.count())

        candidates = candidates.filter(history_of_arrears=last_job.history_of_arrears)
        print('After history_of_arrears filter count: ', candidates.count())

        candidates = candidates.filter(standing_arrears=last_job.standing_arrears)
        print('After standing_arrears filter count: ', candidates.count())

        if last_job.skill_id.exists():
            skill_ids = last_job.skill_id.values_list('id', flat=True)
            candidates = candidates.filter(skill_id__in=skill_ids).distinct()
            print('After skill_id filter count: ', candidates.count())

        # Create a list to store the final output
        candidate_data = []

        for candidate in candidates:
            skills = candidate.skill_id.all()
            skill_ids = [skill.id for skill in skills]
            candidate_info = {
                'id': candidate.id,
                'college_id': candidate.college_id.college if candidate.college_id else None,
                'students_name': candidate.students_name,
                'user_name': candidate.user_name,
                'registration_number': candidate.registration_number,
                'gender': candidate.gender,
                'email_id': candidate.email_id,
                'mobile_number': candidate.mobile_number,
                'year': candidate.year,
                'cgpa': candidate.cgpa,
                'department_id': candidate.department_id.department if candidate.department_id else None,
                'marks_10th': candidate.marks_10th,
                'marks_12th': candidate.marks_12th,
                'marks_semester_wise': candidate.marks_semester_wise,
                'history_of_arrears': candidate.history_of_arrears,
                'standing_arrears': candidate.standing_arrears,
                'number_of_offers': candidate.number_of_offers,
                'text': candidate.text,
                'it_of_offers': candidate.it_of_offers,
                'core_of_offers': candidate.core_of_offers,
                'skill_id': skill_ids
            }
            candidate_data.append(candidate_info)

        print('candidate_data: ', candidate_data)

        return Response(candidate_data)

    except Exception as e:
        print('Exception: ', str(e))
        return Response({'error': str(e)}, status=500)



@api_view(['GET'])
def get_candidate_all_job_master(request):
    try:
        # Fetch the last added job from job_master
        last_job = job_master.objects.filter(deleted=0).order_by('-id').first()
        if not last_job:
            return Response({'error': 'No job found'}, status=404)
        
        print('last job: ', last_job)
        print('last_job.college_id: ', last_job.college_id)
        print('last_job.department_id: ', last_job.department_id)
        print('last_job.gender: ', last_job.gender)
        print('last_job.marks_10th: ', last_job.marks_10th)
        print('last_job.marks_12th: ', last_job.marks_12th)
        print('last_job.cgpa: ', last_job.cgpa)
        print('last_job.history_of_arrears: ', last_job.history_of_arrears)
        print('last_job.standing_arrears: ', last_job.standing_arrears)
        print('last_job.skill_id: ', last_job.skill_id.all())

        # Check each filter step-by-step
        candidates = candidate_master.objects.filter(deleted=0, college_id=last_job.college_id)
        print('After filtering by college_id: ', candidates.count())

        if last_job.department_id is not None:
            candidates = candidates.filter(department_id=last_job.department_id)
            print('After filtering by department_id: ', candidates.count())
        
        if last_job.gender is not None:
            candidates = candidates.filter(gender=last_job.gender)
            print('After filtering by gender: ', candidates.count())
        
        if last_job.marks_10th is not None:
            candidates = candidates.filter(marks_10th=last_job.marks_10th)
            print('After filtering by marks_10th: ', candidates.count())
        
        if last_job.marks_12th is not None:
            candidates = candidates.filter(marks_12th=last_job.marks_12th)
            print('After filtering by marks_12th: ', candidates.count())
        
        if last_job.cgpa is not None:
            candidates = candidates.filter(cgpa=last_job.cgpa)
            print('After filtering by cgpa: ', candidates.count())
        
        if last_job.history_of_arrears is not None:
            candidates = candidates.filter(history_of_arrears=last_job.history_of_arrears)
            print('After filtering by history_of_arrears: ', candidates.count())
        
        if last_job.standing_arrears is not None:
            candidates = candidates.filter(standing_arrears=last_job.standing_arrears)
            print('After filtering by standing_arrears: ', candidates.count())

        if last_job.skill_id.exists():
            skill_ids = last_job.skill_id.values_list('id', flat=True)
            candidates = candidates.filter(skill_id__in=skill_ids)
            print('After filtering by skill_id: ', candidates.count())

        # Create a list to store the final output
        candidate_data = []

        for candidate in candidates:
            skills = candidate.skill_id.all()
            skill_ids = [skill.id for skill in skills]
            candidate_info = {
                'id': candidate.id,
                'college_id': candidate.college_id.college if candidate.college_id else None,
                'students_name': candidate.students_name,
                'user_name': candidate.user_name,
                'registration_number': candidate.registration_number,
                'gender': candidate.gender,
                'email_id': candidate.email_id,
                'mobile_number': candidate.mobile_number,
                'year': candidate.year,
                'cgpa': candidate.cgpa,
                'department_id': candidate.department_id.department if candidate.department_id else None,
                'marks_10th': candidate.marks_10th,
                'marks_12th': candidate.marks_12th,
                'marks_semester_wise': candidate.marks_semester_wise,
                'history_of_arrears': candidate.history_of_arrears,
                'standing_arrears': candidate.standing_arrears,
                'number_of_offers': candidate.number_of_offers,
                'text': candidate.text,
                'it_of_offers': candidate.it_of_offers,
                'core_of_offers': candidate.core_of_offers,
                'skill_id': skill_ids
            }
            candidate_data.append(candidate_info)

        print('candidate_data: ', candidate_data)

        return Response(candidate_data)

    except Exception as e:
        print('Exception: ', str(e))
        return Response({'error': str(e)}, status=500)




#-----------Students Test Schedule-----------------------#

@api_view(['GET'])
def get_students_test_schedule(request, college_id, user_name):
    tests_candidates = tests_candidates_map.objects.filter(deleted=0, college_id=college_id, student_id__user_name=user_name)
    ##logger.info('Fetching test-candidate-map')
   
    test_candidate_map_data = []
    for testing in tests_candidates:
        question_id = None  # Initialize question_type variable
        topic_name = None
        student_id = None
        student_name = None
        user_name = None
        college_id = None
        college_id_id = None
        department_id = None
        department_id_id = None
        department_name = None
        test_type = None
        rules_id = None
        rules_id_id = None
        instruction = None
        skill_type = None
        question_paper_name = None
        test_type = None

        if testing.rules_id:
            rules_id_id = testing.rules_id.id
            rules_id = testing.rules_id.rule_name
            instruction = testing.rules_id.instruction
        
        if testing.department_id:
            department_id = testing.department_id.department
            department_id_id = testing.department_id.id
      
        if testing.question_id:
            question_id = testing.question_id.id
            question_paper_name = testing.question_id.question_paper_name
            test_type = testing.question_id.test_type

        if testing.student_id:
            student_id = testing.student_id.id
            student_name = testing.student_id.students_name
            user_name = testing.student_id.user_name
        if testing.college_id:
            college_id = testing.college_id.college
            college_id_id = testing.college_id.id

        test_candidate_map_data.append({
            'id': testing.id,
            'test_name': testing.test_name,
            'college_id_id': college_id_id,
            'college_id': college_id,
            'department_id_id': department_id_id,
            'department_id': department_id,
            'question_id': question_id, 
            'question_paper_name': question_paper_name,
            'test_type': test_type,
            'student_id': student_id,
            'student_name': student_name,
            'user_name': user_name,
            'dtm_start': testing.dtm_start,
            'dtm_end': testing.dtm_end,
            'attempt_count': testing.attempt_count,
            'is_actual_test': testing.is_actual_test,
            'is_active': testing.is_active,
            'duration': testing.duration,
            'duration_type': testing.duration_type,
            'year': testing.year,
            'rules_id': rules_id_id,
            'rules': rules_id,
            'instruction': instruction,
            'need_candidate_info': testing.need_candidate_info,
            'total_score': testing.total_score,
            'avg_mark': testing.avg_mark
        })
    
    return Response(test_candidate_map_data)



#---------------------Questions Master with images--------------#
def get_questions_IO_filter(request, question_id):
    questionset = question_master.objects.filter(deleted=0, question_name_id=question_id)
    #logger.info("Fetching Questions where deleted=0")
   
    question_data = []
    for question in questionset:
        question_name_id = None
        question_paper_name = None
      
        if question.question_name_id:
            question_name_id = question.question_name_id.id           
            question_paper_name = question.question_name_id.question_paper_name

        question_image_data = None
        if question.question_image_data:
            question_image_data = base64.b64encode(question.question_image_data).decode('utf-8')
        
        option_a_image_data = None
        if question.option_a_image_data:
            option_a_image_data = base64.b64encode(question.option_a_image_data).decode('utf-8')

        option_b_image_data = None
        if question.option_b_image_data:
            option_b_image_data = base64.b64encode(question.option_b_image_data).decode('utf-8')

        option_c_image_data = None
        if question.option_c_image_data:
            option_c_image_data = base64.b64encode(question.option_c_image_data).decode('utf-8')

        option_d_image_data = None
        if question.option_d_image_data:
            option_d_image_data = base64.b64encode(question.option_d_image_data).decode('utf-8')

        question_data.append({
            'id': question.id,
            'question_name_id': question_name_id,
            'question_paper_name': question_paper_name,
            'question_text': question.question_text,
            'question_image_data': question_image_data,
            'option_a_image_data': option_a_image_data,
            'option_b_image_data': option_b_image_data,
            'option_c_image_data': option_c_image_data,
            'option_d_image_data': option_d_image_data,
            'option_a': question.option_a,
            'option_b': question.option_b,
            'option_c': question.option_c,
            'option_d': question.option_d,
            'view_hint': question.view_hint,
            'mark': question.mark,
            'explain_answer': question.explain_answer,
            'answer': question.answer,
            'negative_mark': question.negative_mark,
            'input_format': question.input_format,
            
        })
    random.shuffle(question_data)
    return JsonResponse(question_data, safe=False)




@api_view(['GET'])
def get_tests_candidates_map_testName(request, test_name):
    tests_candidates=tests_candidates_map.objects.filter(deleted=0, test_name=test_name)
    ##logger.info('Fetching test-candidate-map')
   
    test_candidate_map_data=[]
    for testing in  tests_candidates:
        question_id= None  # Initialize question_type variable
        topic_name=None
        student_id=None
        student_name=None
        user_name=None
        college_id=None
        college_id_id=None
        department_id=None
        department_id_id=None
        department_name=None
        test_type=None
        rules_id=None
        rules_id_id = None
        instruction=None
        skill_type=None
        question_paper_name=None
        test_type=None

        if testing.rules_id:
           rules_id_id=testing.rules_id.id
           rules_id=testing.rules_id.rule_name
           instruction=testing.rules_id.instruction
        
        if testing.department_id:
           department_id=testing.department_id.department
           department_id_id=testing.department_id.id
      
        if testing.question_id:
           question_id=testing.question_id.id
           question_paper_name=testing.question_id.question_paper_name
           test_type=testing.question_id.test_type


        if testing.student_id:
           student_id=testing.student_id.id
           student_name=testing.student_id.students_name
           user_name=testing.student_id.user_name
        if testing.college_id:
            college_id= testing.college_id.college
            college_id_id= testing.college_id.id

        test_candidate_map_data.append({
            'id': testing.id,
            'test_name':testing.test_name,
            'college_id_id':college_id_id,
            'college_id':college_id,
            'department_id_id':department_id_id,
            'department_id':department_id,
            'question_id':question_id, 
            'question_paper_name':question_paper_name,
            'test_type': test_type,
            'student_id':student_id,
            'student_name':student_name,
            'user_name':user_name,
            'dtm_start': testing.dtm_start,
            'dtm_end': testing.dtm_end,
           
            'attempt_count':testing.attempt_count,
            'is_actual_test':testing.is_actual_test,
            'is_active': testing.is_active,
            'duration':testing.duration,
            'duration_type':testing.duration_type,
            'year':testing.year,
            'rules_id':rules_id_id,
            'rules':rules_id,
            'instruction':instruction,
            'need_candidate_info':testing.need_candidate_info,
            'total_score': testing.total_score,
            'avg_mark': testing.avg_mark,
            'dtm_created': testing.dtm_created,
       
       
          

            })
    return Response(test_candidate_map_data)




@api_view(['GET'])
def get_tests_Reports_candidates(request, test_name):
    tests_candidates = tests_candidates_map.objects.filter(deleted=0, test_name=test_name)
    ##logger.info('Fetching test-candidate-map')
    
    test_candidate_map_data = []
    for testing in tests_candidates:
        question_name = None
        student_id = None
        student_name = None
        user_name = None
        college_id = None
        department_id = None
        test_type = None
        rules_id = None
        instruction = None
        email_id = None
        mobile_number = None
        gender = None
        registration_number = None

        if testing.rules_id:
            rules_id = testing.rules_id.rule_name
            instruction = testing.rules_id.instruction

        if testing.department_id:
            department_id = testing.department_id.department

        if testing.question_id:
            question_name = testing.question_id.question_paper_name

        if testing.student_id:
            student_id = testing.student_id.id
            student_name = testing.student_id.students_name
            user_name = testing.student_id.user_name
            email_id = testing.student_id.email_id
            mobile_number = testing.student_id.mobile_number
            gender = testing.student_id.gender
            registration_number = testing.student_id.registration_number

        if testing.college_id:
            college_id = testing.college_id.college

        # Format dates
        dtm_start_formatted = django_format_date(localtime(testing.dtm_start), 'd-m-Y h:i A')
        dtm_end_formatted = django_format_date(localtime(testing.dtm_end), 'd-m-Y h:i A')

        test_candidate_map_data.append({
            'id': testing.id,
            'test_name': testing.test_name,
            'college_id': college_id,
            'department_id': department_id,
            'question_id': question_name,
            'student_id': student_id,
            'registration_number': registration_number,
            'email_id': email_id,
            'mobile_number': mobile_number,
            'gender': gender,
            'student_name': student_name,
            'user_name': user_name,
            'dtm_start': dtm_start_formatted,
            'dtm_end': dtm_end_formatted,
            'attempt_count': testing.attempt_count,
            'is_actual_test': testing.is_actual_test,
            'is_active': testing.is_active,
            'duration': testing.duration,
            'duration_type': testing.duration_type,
            'year': testing.year,
            'rules': rules_id,
            'instruction': instruction,
            'need_candidate_info': testing.need_candidate_info,
            'total_score': testing.total_score,
            'avg_mark': testing.avg_mark
        })
    return Response(test_candidate_map_data)




class CourseScheduleMapView(APIView):
    def post(self, request, format=None):
        print('request.data: ', request.data)

        college_id = request.data.get('college_id', [])
        department_id = request.data.get('department_id', [])
        year = request.data.get('year')
        topic_ids = request.data.get('topic_id', [])  # Assuming 'topic_ids' is provided as a list in the request

        if not college_id or not department_id or not year or not topic_ids:
            return Response({'error': 'college_id, department_id, year, and topic_ids are required fields.'}, status=status.HTTP_400_BAD_REQUEST)

        # Filter candidates based on multiple college_ids and department_ids
        students = candidate_master.objects.filter(
            college_id__in=college_id,
            department_id__in=department_id,
            year=year
        )

        data = []
        current_date_and_time = datetime.now()  # Get the current date and time

        for student in students:
            for topic_id in topic_ids:
                test_candidate_data = {
                    'topic_id': topic_id,  # Use the topic_id directly
                    'student_id': student.id,
                    'college_id': student.college_id.id,  # Use student.college_id.id for pk
                    'department_id': student.department_id.id,  # Use student.department_id.id for pk
                    'dtm_start_student': request.data.get('dtm_start_student'),
                    'dtm_end_student': request.data.get('dtm_end_student'),
                    'dtm_start_trainer': request.data.get('dtm_start_trainer'),
                    'dtm_end_trainer': request.data.get('dtm_end_trainer'),
                    'dtm_of_training': request.data.get('dtm_of_training'),
                    'year': year,
                    'trainer_id': request.data.get('trainer_id'),
                    'dtm_created': current_date_and_time  # Add current date and time
                }

                serializer = courseScheduleSerializer(data=test_candidate_data)
                if serializer.is_valid():
                    serializer.save()
                    data.append(serializer.data)
                else:
                    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        return Response(data, status=status.HTTP_201_CREATED)



@api_view(['GET'])
def get_content_master_subTopic(request, topic_name):
    if not topic_name:
        return Response({'error': 'Topic name is required.'}, status=status.HTTP_400_BAD_REQUEST)

    # Fetch the id and sub_topic based on the topic name
    topics = content_master.objects.filter(topic=topic_name).values('id', 'sub_topic')
    if not topics.exists():
        return Response({'error': f'No topics found for the given topic name: {topic_name}'}, status=status.HTTP_404_NOT_FOUND)

    return Response(list(topics), status=status.HTTP_200_OK)




@api_view(['GET'])
def students_course_content_view(request):
    user_name = request.query_params.get('user_name')
    current_date_time = datetime.now()

    if not user_name:
        return Response({'error': 'user_name is required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        candidate = get_object_or_404(candidate_master, user_name=user_name)
        student_id = candidate.id

        course_schedules = course_schedule.objects.filter(
            deleted=0,
            student_id=student_id,
            dtm_start_student__lte=current_date_time,
            dtm_end_student__gte=current_date_time
        ).select_related('topic_id')

        data = []
        for cs in course_schedules:
            cm = cs.topic_id
            data.append({
                'topic': cm.topic,
                'sub_topic': cm.sub_topic,
                'Content_URL': cm.content_url,
                'Actual_Content': cm.actual_content,
                'Start_Date': cs.dtm_start_student,
                'End_Date': cs.dtm_end_student,
            })

        return Response(data, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)




@api_view(['DELETE'])
def delete_question_paper_by_name(request, question_paper_name):
    try:
        question_paper = get_object_or_404(question_paper_master, question_paper_name=question_paper_name)
        question_paper.delete()
        return Response({'message': 'Question paper deleted successfully'}, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'error': f'Error occurred while trying to delete the question paper: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



@api_view(['GET'])
def Trainer_course_content_view(request):
    user_name = request.query_params.get('user_name')
    current_date_time = datetime.now()

    if not user_name:
        return Response({'error': 'user_name is required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        trainer = get_object_or_404(trainer_master, trainer_name=user_name)
        trainer_id = trainer.id

        course_schedules = course_schedule.objects.filter(
            deleted=0,
            trainer_id=trainer_id,
            dtm_start_trainer__lte=current_date_time,
            dtm_end_trainer__gte=current_date_time
        ).select_related('topic_id')

        unique_data = {}
        for cs in course_schedules:
            cm = cs.topic_id
            key = (cm.topic, cm.sub_topic, cm.content_url, cm.actual_content, cs.dtm_start_trainer, cs.dtm_end_trainer)
            if key not in unique_data:
                unique_data[key] = {
                    'topic': cm.topic,
                    'sub_topic': cm.sub_topic,
                    'Content_URL': cm.content_url,
                    'Actual_Content': cm.actual_content,
                    'Start_Date': cs.dtm_start_trainer,
                    'End_Date': cs.dtm_end_trainer,
                }

        data = list(unique_data.values())

        return Response(data, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



@api_view(['GET'])
def get_candidate_user_name(request, user_name):
    try:
        # Fetch candidates with related college and department, and prefetch skills
        candidatelist = candidate_master.objects.filter(deleted=0, user_name=user_name).select_related('college_id', 'department_id').prefetch_related('skill_id').values(
            'id', 
            'college_id',
            'college_id__college',  # Get college name
            'students_name', 
            'user_name', 
            'registration_number', 
            'gender',
            'email_id', 
            'mobile_number', 
            'year', 
            'cgpa', 
            'department_id',
            'department_id__department',  # Get department name
            'marks_10th',
            'marks_12th', 
            'marks_semester_wise', 
            'history_of_arrears', 
            'standing_arrears',
            'number_of_offers', 
            'text', 
            'it_of_offers', 
            'core_of_offers'
        ).distinct()

        # Create a list to store the final output
        candidate_data = []
        
        for candidate in candidatelist:
            skills = candidate_master.objects.get(id=candidate['id']).skill_id.all()
            skill_ids = [skill.id for skill in skills]
            candidate['skill_id'] = skill_ids
            candidate_data.append(candidate)
        
        return Response(candidate_data)
    except Exception as e:
        return Response({'error': str(e)}, status=500)



@api_view(['GET'])
def candidate_info(request):
    user_name = request.GET.get('user_name', None)

    if not user_name:
        return Response({'error': 'user_name parameter is required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        candidate = candidate_master.objects.get(user_name=user_name)
        need_candidate_info = candidate.need_candidate_info

        return Response({'need_candidate_info': need_candidate_info}, status=status.HTTP_200_OK)

    except candidate_master.DoesNotExist:
        return Response({'error': 'Candidate not found'}, status=status.HTTP_404_NOT_FOUND)




@api_view(['PUT'])
def update_candidate_info(request):
    user_name = request.data.get('user_name')
    need_candidate_info = request.data.get('need_candidate_info', False)  # Default to False if not provided

    if not user_name:
        return Response({'error': 'user_name is required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        candidate = candidate_master.objects.get(user_name=user_name)
        candidate.need_candidate_info = need_candidate_info
        candidate.save()

        return Response({'success': f'Updated need_candidate_info for user {user_name}'}, status=status.HTTP_200_OK)

    except candidate_master.DoesNotExist:
        return Response({'error': 'Candidate not found'}, status=status.HTTP_404_NOT_FOUND)


@api_view(['GET'])
def get_course_scheduleold(request):
    candidatelist = course_schedule.objects.filter(deleted=0)
    candidate_data = []

    for candidate in candidatelist:
        student_id = None
        college_id = None
        college_name = None
        department_id = None
        department_name = None
        topic_id = None
        sub_topic = None
        trainer_id = None

        if candidate.student_id:
            student = candidate.student_id
            student_id = student.students_name
            college_id = student.college_id.id if student.college_id else None
            college_name = student.college_id.college_name if student.college_id else None
            department_id = student.department_id.id if student.department_id else None
            department_name = student.department_id.department_name if student.department_id else None

        if candidate.topic_id:
            topic_id = candidate.topic_id.topic
            sub_topic = candidate.topic_id.sub_topic

        if candidate.trainer_id:
            trainer_id = candidate.trainer_id.trainer_name

        candidate_data.append({
            'id': candidate.id,
            'college_id': college_id,
            'college_name': college_name,
            'student_id': student_id,
            'year': candidate.year,
            'department_id': department_id,
            'department_name': department_name,
            'sub_topic': sub_topic,
            'topic_id': topic_id,
            'trainer_id': trainer_id,
            'dtm_start_student': candidate.dtm_start_student,
            'dtm_end_student': candidate.dtm_end_student,
            'dtm_start_trainer': candidate.dtm_start_trainer,
            'dtm_end_trainer': candidate.dtm_end_trainer,
            'dtm_of_training': candidate.dtm_of_training,
        })

    return Response(candidate_data)






