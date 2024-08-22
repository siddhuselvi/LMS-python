from django.shortcuts import render
from rest_framework.response import Response
from rest_framework import generics
from .models import appinfo,course_schedule,user_profile,course_trainer_feedback,eligible_student_list,student_request,training_attendance_sheet, course_content_feedback, attendance_master,job_master,event_master,company_master, announcement_master, test_question_map,skill_type,login,tests_candidates_answers,tests_candidates_map,test_master,question_master,candidate_master,content_master,skills_master,topic_master,test_type,question_type,college_master,department_master,trainer_master, rules, question_paper_master, compiler_output, job_master, event_master, company_master,  student_request, training_attendance_sheet, course_trainer_feedback
from .serializers import candidateSerializerImport,UserProfileSerializer,eligible_studentSerializer,UserSerializer,studentRequestSerializer,trainerfeedbackSerializer,trainingreportsheetSerializer,loginSerializer,eventSerializer,companySerializer,jobSerializer,appinfoSerializer,tests_candidates_answerSerializer,courseScheduleSerializer, courseContentFeedbackSerializer, attendanceMasterSerializer, announcementSerializer,  testQuestionMapSerializer,testcandidatemapSerializers,testsSerializers,questionsSerializer,skilltypeSerializer,contentSerializers,collegeSerializers,departmentSerializers,questiontypeSerializers,testtypeSerializers,topicSerializers,skillSerializer,candidatesSerializer,trainerSerializer, testsSerializersImport, questionsSerializerImport, ruleSerializers, testsSerializersAddUpdate, questionsSerializerMasterData, loginSerializerStu, questionsSerializerCodeImport, candidatesoneSerializer, questionsPaperSerializer, questionsSerializer_IO, candidateuserSerializerImport, loginSerializerStuser, questionsSerializer_code, NonDbTestAssignSerializer, compilerSerializer, jobSerializer, eventSerializer, companySerializer, studentRequestSerializer, trainingreportsheetSerializer, trainerfeedbackSerializer
from rest_framework.decorators import api_view
from .forms import QuestionForm, QuestionCodeForm, CollegeForm, QuestionFormMCQ
from django.http import HttpResponse 

from django.db.models import Q
from django.shortcuts import render, redirect
from .forms import QuestionImportForm
from .models import question_master_temp
import docx
import re

from .forms import DocumentForm
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt

import base64
from django.utils import timezone
from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from django.http.response import JsonResponse
from rest_framework.views import APIView
from rest_framework import status
import pandas as pd
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Avg, Sum, Subquery, OuterRef
from docx.oxml.ns import qn
from .models import question_master, question_paper_master

import requests
import io

import json


from django.core.files.uploadedfile import InMemoryUploadedFile
from django.db.models import Q
from thefuzz import fuzz
import urllib.parse
import math
import random
from django.db.models import Count, Value, F, DateTimeField
from django.db.models import Avg, Sum, Subquery, OuterRef, Max

from django.db.models import Count
import io,math,random,datetime,json,string,collections,itertools,functools
import sys
import builtins
from django.core.exceptions import ObjectDoesNotExist
from datetime import date
from io import StringIO
import datetime
import logging
import subprocess
import os,re,shutil
import uuid
import io,math,random,datetime,json,string,collections,itertools,functools
import sys
from django.conf import settings
from datetime import datetime
import json
from django.contrib.auth.models import User
from .models import user_profile
from .serializers import UserProfileSerializer, TestCandidateMapSerializerNeedInfo, testsSerializersAdd, testcandidatemapSerializersupdate
from django.utils.timezone import localtime
from django.utils.dateformat import format as django_format_date
from django.db.models.functions import ExtractMonth
import calendar
import pytz
# from .models import StudentAttendance
from django.utils import timezone
from django.db.models.functions import Coalesce
from django.db.models import Case, When
from django.utils.timezone import make_aware

# Get an instance of a logger
logger = logging.getLogger('app')
from django.db import transaction
from django.db.models.functions import Length
from django.db import connection
from django.core.cache import cache
from django.core.files.base import ContentFile 


# Create your views here.


class appinfogetAPIView(generics.ListAPIView):
    queryset = appinfo.objects.filter(deleted=0)

    
    serializer_class = appinfoSerializer


class appinfoAPIView(generics.CreateAPIView):
    queryset = appinfo.objects.all()
    serializer_class = appinfoSerializer

class appinfoRetrieveUpdateDestroyAPIView(generics.RetrieveUpdateDestroyAPIView):
    queryset = appinfo.objects.all()
    serializer_class = appinfoSerializer
#_____________________________Master Tables_____________________________________________#

class skilltypegetAPIView(generics.ListAPIView):
    queryset = skill_type.objects.filter(deleted=0)
    serializer_class = skilltypeSerializer

    def get(self, request, *args, **kwargs):
       
        return super().get(request, *args, **kwargs)


class skilltypecreateAPIView(generics.CreateAPIView):
    queryset = skill_type.objects.all()
    serializer_class = skilltypeSerializer

    def create(self, request, *args, **kwargs):
       
        return super().create(request, *args, **kwargs)


class skilltypeRetrieveUpdateDestroyAPIView(generics.RetrieveUpdateAPIView):
    queryset = skill_type.objects.all()
    serializer_class = skilltypeSerializer

    def update(self, request, *args, **kwargs):
       # ##logger.debug('Updating skill type with id: %s', kwargs.get('pk'))
        return super().update(request, *args, **kwargs)

    def retrieve(self, request, *args, **kwargs):
       # ##logger.debug('Retrieving skill type with id: %s', kwargs.get('pk'))
        return super().retrieve(request, *args, **kwargs)


@api_view(['PUT', 'PATCH'])
def delete_skill_type(request, pk):
   # ##logger.debug('Entering delete_skill_type function with id: %s', pk)
    try:
        skilltype = skill_type.objects.get(id=pk)
      #  ##logger.debug('Skill type found: %s', skilltype)
    except skill_type.DoesNotExist:
        logger.error('Skill type with id %s not found', pk)
        return JsonResponse("Skill type not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    skilltype.deleted = 1
    skilltype.save()
   # ##logger.debug('Skill type marked as deleted: %s', skilltype)

    return JsonResponse("Skill 'deleted' field updated successfully", safe=False)


#------------------------------Test-type-------------------------------------------------------------#


class test_type_listView(generics.ListAPIView):
    queryset = test_type.objects.filter(deleted=0)
    serializer_class = testtypeSerializers

    def get(self, request, *args, **kwargs):
       # #logger.info("Fetching test types where deleted=0")
        return super().get(request, *args, **kwargs)

class test_type_create(generics.CreateAPIView):
    queryset = test_type.objects.all()
    serializer_class = testtypeSerializers

    def post(self, request, *args, **kwargs):
      #  #logger.info("Creating a new test type")
        return super().post(request, *args, **kwargs)

class test_type_Update(generics.RetrieveUpdateAPIView):
    queryset = test_type.objects.all()
    serializer_class = testtypeSerializers

    def put(self, request, *args, **kwargs):
      #  #logger.info(f"Updating test type with id {kwargs.get('pk')}")
        return super().put(request, *args, **kwargs)

    def patch(self, request, *args, **kwargs):
       # #logger.info(f"Partially updating test type with id {kwargs.get('pk')}")
        return super().patch(request, *args, **kwargs)

@api_view(['PUT', 'PATCH'])
def delete_test_type(request, pk):
    try:
       # #logger.info(f"Attempting to mark test type with id {pk} as deleted")
        test_typevar = test_type.objects.get(id=pk)
    except test_type.DoesNotExist:
        logger.error(f"Test type with id {pk} not found")
        return JsonResponse("test_type not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    test_typevar.deleted = 1
    test_typevar.save()

   # #logger.info(f"Marked test type with id {pk} as deleted")
    return JsonResponse("test_type 'deleted' field updated successfully", safe=False)

#------------------------------------Question_type--------------------------------------#

class question_type_listView(generics.ListAPIView):
    serializer_class = questiontypeSerializers

    def get_queryset(self):
       # #logger.info("Fetching question types where deleted=0")
        queryset = question_type.objects.filter(deleted=0)
        ##logger.info("Fetched question types successfully")
        return queryset

class question_type_create(generics.CreateAPIView):
    queryset = question_type.objects.all()
    serializer_class = questiontypeSerializers

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new question type")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new question type successfully")
        return response

class question_type_Update(generics.RetrieveUpdateAPIView):
    queryset = question_type.objects.all()
    serializer_class = questiontypeSerializers

    def put(self, request, *args, **kwargs):
       # #logger.info(f"Updating question type with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
      #  #logger.info(f"Updated question type with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
       # #logger.info(f"Partially updating question type with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
       # #logger.info(f"Partially updated question type with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_question_type(request, pk):
   # #logger.info(f"Attempting to mark question type with id {pk} as deleted")
    try:
        question_typevar = question_type.objects.get(id=pk)
    except question_type.DoesNotExist:
        logger.error(f"Question type with id {pk} not found")
        return JsonResponse("question_type not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    question_typevar.deleted = 1
    question_typevar.save()

   # #logger.info(f"Marked question type with id {pk} as deleted successfully")
    return JsonResponse("question_type 'deleted' field updated successfully", safe=False)


#------------------------------college----------------------------------------#


class CollegeListView(generics.ListAPIView):
    serializer_class = collegeSerializers

    def get_queryset(self):
        
        queryset = college_master.objects.filter(deleted=0)
       
        return queryset

class collegeCreateView(generics.CreateAPIView):
    queryset = college_master.objects.all()
    serializer_class = collegeSerializers

    def post(self, request, *args, **kwargs):
        
        response = super().post(request, *args, **kwargs)
       
        return response

class collegeUpdateView(generics.RetrieveUpdateAPIView):
    queryset = college_master.objects.all()
    serializer_class = collegeSerializers

    def put(self, request, *args, **kwargs):
       
        response = super().put(request, *args, **kwargs)
       
        return response

    def patch(self, request, *args, **kwargs):
        
        response = super().patch(request, *args, **kwargs)
       
        return response

@api_view(['PUT', 'PATCH'])
def delete_college(request, pk):
   
    try:
        college = college_master.objects.get(id=pk)
    except college_master.DoesNotExist:
        logger.error(f"College with id {pk} not found")
        return Response("college not found", status=404)

    college.deleted = 1
    college.save()

   

    return Response("college 'deleted' field updated successfully")

#----------------------------------------department---------------------------------------------#

class DepartmentListView(generics.ListAPIView):
    serializer_class = departmentSerializers

    def get_queryset(self):
        
        return department_master.objects.filter(deleted=0)

class departmentCreateView(generics.CreateAPIView):
    queryset = department_master.objects.all()
    serializer_class = departmentSerializers

    def post(self, request, *args, **kwargs):
       
        response = super().post(request, *args, **kwargs)
      
        return response

class departmentUpdateView(generics.RetrieveUpdateAPIView):
    queryset = department_master.objects.all()
    serializer_class = departmentSerializers

    def put(self, request, *args, **kwargs):
       
        response = super().put(request, *args, **kwargs)
       
        return response

    def patch(self, request, *args, **kwargs):
        
        response = super().patch(request, *args, **kwargs)
       
        return response

@api_view(['PUT', 'PATCH'])
def delete_department(request, pk):
   
    try:
        department = department_master.objects.get(id=pk)
    except department_master.DoesNotExist:
        logger.error(f"department with id {pk} not found")
        return Response("department not found", status=404)

    department.deleted = 1
    department.save()

   


    return Response("department 'deleted' field updated successfully")


#__________________________topic_master___________________________________#


class topicListView(generics.ListAPIView):
    serializer_class = topicSerializers

    def get_queryset(self):
       
        return topic_master.objects.filter(deleted=0)

class topicCreateView(generics.CreateAPIView):
    queryset = topic_master.objects.all()
    serializer_class = topicSerializers

    def post(self, request, *args, **kwargs):
       
        response = super().post(request, *args, **kwargs)
       
        return response

class topicUpdateView(generics.RetrieveUpdateAPIView):
    queryset = topic_master.objects.all()
    serializer_class = topicSerializers

    def put(self, request, *args, **kwargs):
       
        response = super().put(request, *args, **kwargs)
        
        return response

    def patch(self, request, *args, **kwargs):
       
        response = super().patch(request, *args, **kwargs)
       
        return response

@api_view(['PUT', 'PATCH'])
def delete_topic(request, pk):
   
    try:
        topic = topic_master.objects.get(id=pk)
    except topic_master.DoesNotExist:
        return Response("topic not found", status=404)

    topic.deleted = 1
    topic.save()

   


    return Response("topic 'deleted' field updated successfully")
#___________________________SKILL_master_________________________________________
class skillsAPIView(generics.ListAPIView):
    queryset = skills_master.objects.all()
    serializer_class = skillSerializer

   

class skillscreateAPIView(generics.CreateAPIView):
    queryset = skills_master.objects.all()
    serializer_class = skillSerializer

    def post(self, request, *args, **kwargs):
        
        response = super().post(request, *args, **kwargs)
        
        return response

class skillsRetrieveUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = skills_master.objects.all()
    serializer_class = skillSerializer

    def put(self, request, *args, **kwargs):
       
        response = super().put(request, *args, **kwargs)
       
        return response

    def patch(self, request, *args, **kwargs):
       
        response = super().patch(request, *args, **kwargs)
       
        return response

@api_view(['PUT', 'PATCH'])
def delete_skills(request, pk):
   
    try:
        print("Entering Function..")
        skills = skills_master.objects.get(id=pk)

        print("skill: ", skills)
    except skills_master.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    skills.deleted = 1
    skills.save()

   

    print("skill: ", skills)

    return JsonResponse("skill 'deleted' field updated successfully", safe=False)


#___________________________________candidate_master________________________________________

@api_view(['GET'])
def get_candidate_OLD(request):
    try:
        # Fetch all candidates with related data in one query
        candidate_data = candidate_master.objects.filter(deleted=0).select_related(
            'college_id', 'department_id'
        ).values(
            'id',
            'students_name',
            'user_name',
            'registration_number',
            'gender',
            'email_id',
            'mobile_number',
            'year',
            'cgpa',
            'marks_10th',
            'marks_12th',
            'marks_semester_wise',
            'history_of_arrears',
            'standing_arrears',
            'number_of_offers',
            'text',
            'it_of_offers',
            'core_of_offers',
            'college_id__college',   # Direct field access
            'college_id__id',        # Direct field access
            'department_id__department',  # Direct field access
            'department_id__id'      # Direct field access
        )

        # Convert the queryset into a list of dictionaries
        candidate_list = list(candidate_data)

        return Response(candidate_list)
    except Exception as e:
        return Response({'error': str(e)}, status=500)



@api_view(['GET'])
def get_candidate(request):
    try:
        # Create a cache key that uniquely identifies this query
        cache_key = 'candidate_master_data'
        
        # Try to get the data from the cache
        candidate_list = cache.get(cache_key)
        
        # If the data is not in the cache, fetch it from the database
        if not candidate_list:
            # Fetch all candidates with related data in one query
            candidate_data = candidate_master.objects.filter(deleted=0).select_related(
                'college_id', 'department_id'
            ).values(
                'id',
                'students_name',
                'user_name',
                'registration_number',
                'gender',
                'email_id',
                'mobile_number',
                'year',
                'cgpa',
                'marks_10th',
                'marks_12th',
                'marks_semester_wise',
                'history_of_arrears',
                'standing_arrears',
                'number_of_offers',
                'text',
                'it_of_offers',
                'core_of_offers',
                'college_id__college',   # Direct field access
                'college_id__id',        # Direct field access
                'department_id__department',  # Direct field access
                'department_id__id'      # Direct field access
            )

            # Convert the queryset into a list of dictionaries
            candidate_list = list(candidate_data)
            
            # Store the fetched data in the cache with a timeout
            cache.set(cache_key, candidate_list, timeout=3600)  # Timeout is set to 1 hour (3600 seconds)
        
        # Return the cached or freshly fetched data
        return Response(candidate_list)
    except Exception as e:
        return Response({'error': str(e)}, status=500)


@api_view(['GET'])
def get_candidate_all(request):
    try:
        # Fetch candidates with related college and department, and prefetch skills
        candidatelist = candidate_master.objects.filter(deleted=0).select_related('college_id', 'department_id').prefetch_related('skill_id').values(
            'id', 
            'college_id__college',  # Get college name
            'students_name', 
            'user_name', 
            'registration_number', 
            'gender',
            'email_id', 
            'mobile_number', 
            'year', 
            'cgpa', 
            'department_id__department',  # Get department name
            'marks_10th',
            'marks_12th', 
            'marks_semester_wise', 
            'history_of_arrears', 
            'standing_arrears',
            'number_of_offers', 
            'text', 
            'it_of_offers', 
            'core_of_offers'
        ).distinct()

        # Create a list to store the final output
        candidate_data = []
        
        for candidate in candidatelist:
            skills = candidate_master.objects.get(id=candidate['id']).skill_id.all()
            skill_ids = [skill.id for skill in skills]
            candidate['skill_id'] = skill_ids
            candidate_data.append(candidate)
        
        return Response(candidate_data)
    except Exception as e:
        return Response({'error': str(e)}, status=500)



class candidatescreateAPIView(generics.CreateAPIView):
    queryset = candidate_master.objects.all()
    serializer_class = candidatesSerializer

    def perform_create(self, serializer):
        serializer.save(is_database=True)

    def post(self, request, *args, **kwargs):        
        response = super().post(request, *args, **kwargs)       
        return response
class candidates_Select_Update(generics.RetrieveUpdateAPIView):
    queryset = candidate_master.objects.all()
    serializer_class = candidatesSerializer

    def put(self, request, *args, **kwargs):
        
        response = super().put(request, *args, **kwargs)
        
        return response

    def patch(self, request, *args, **kwargs):
       
        response = super().patch(request, *args, **kwargs)
        
        return response

@api_view(['PUT', 'PATCH'])
def delete_candidates(request, pk):
    
    try:
        print("Entering Function..")
        trainees = candidate_master.objects.get(id=pk)

        print("tests: ", trainees)
    except candidate_master.DoesNotExist:
        return JsonResponse("trainees not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    trainees.deleted = 1
    trainees.save()

    

    print("tests: ", trainees)

    return JsonResponse("trainees 'deleted' field updated successfully", safe=False)

#____________________________________content_master________________________________________



@api_view(['GET'])
def get_content(request):
    # Fetch data with ORM and annotate related fields directly
    content_data = content_master.objects.filter(deleted=0).select_related('question_type_id', 'skill_type_id').values(
        'id',
        'topic',
        'content_type',
        'question_type_id__question_type',  # Annotate the related question_type field
        'skill_type_id__skill_type',        # Annotate the related skill_type field
        'content_url',
        'actual_content',
        'status',
        'sub_topic',
        'dtm_active_from',
        'dtm_validity',
    )

    # Format dtm_validity for each item
    for content in content_data:
        content['dtm_validity'] = django_format_date(localtime(content['dtm_validity']), 'd-m-Y h:i A')

    return Response(content_data)




class contentcreateAPIView(generics.CreateAPIView):
    queryset = content_master.objects.all()
    serializer_class =contentSerializers

    def post(self, request, *args, **kwargs):
       
        response = super().post(request, *args, **kwargs)
       
        return response

class contentUpdateAPIView(generics.RetrieveUpdateDestroyAPIView):
    queryset = content_master.objects.all()
    serializer_class =contentSerializers

    def put(self, request, *args, **kwargs):
        
        response = super().put(request, *args, **kwargs)
       
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating Content with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated Content with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_content(request, pk):
    ##logger.info(f"Attempting to mark Content with id {pk} as deleted")
    try:
        print("Entering Function..")
        content=content_master.objects.get(id=pk)

        print("content: ",content)
    except content_master.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    content.deleted = 1
    content.save()

    ##logger.info(f"Marked Content with id {pk} as deleted successfully")

    print("content: ",content)

    return JsonResponse("content 'deleted' field updated successfully", safe=False)
#_______________________________________trainer_master__________________________________________________________

@api_view(['GET'])
def get_trainer_all(request):
    try:
        # Fetch candidates with related skills
        candidatelist = trainer_master.objects.filter(deleted=0).prefetch_related('skill_id')

        # Serialize the data
        serializer = trainerSerializer(candidatelist, many=True)

        return Response(serializer.data)
    except Exception as e:
        return Response({'error': str(e)}, status=500)



class TrainerCreateAPIView(generics.CreateAPIView):
    queryset = trainer_master.objects.all()
    serializer_class = trainerSerializer

    def post(self, request, *args, **kwargs):
        # Extract and decode the Base64 encoded resume
        resume_data = request.data.get('resume')
        if resume_data:
            # Remove the data URL scheme part
            resume_data = resume_data.split(',')[1]
            resume_content = base64.b64decode(resume_data)
            # Use the trainer_name to name the file
            fileName = str(uuid.uuid4())  #unique file name
            resume_file = ContentFile(resume_content, name=f'{fileName}.pdf')

            # Create a new instance of the serializer with the modified data
            data = request.data.copy()  # Create a mutable copy of request.data
            data['resume'] = resume_file  # Replace the 'resume' field with the new file
            
            # Instantiate the serializer with the modified data
            serializer = self.get_serializer(data=data)
            if serializer.is_valid():
                # Save the object if the data is valid
                self.perform_create(serializer)
                headers = self.get_success_headers(serializer.data)
                return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        else:
            return Response({"error": "No resume data provided"}, status=status.HTTP_400_BAD_REQUEST)


class TrainerRetrieveUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = trainer_master.objects.all()
    serializer_class = trainerSerializer

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating Trainer with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated Trainer with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating Trainer with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated Trainer with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_trainer(request, pk):
    ##logger.info(f"Attempting to mark Trainer with id {pk} as deleted")
    try:
        print("Entering Function..")
        trainer=trainer_master.objects.get(id=pk)

        print("trainer: ",trainer)
    except trainer_master.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    trainer.deleted = 1
    trainer.save()

    ##logger.info(f"Marked Trainer with id {pk} as deleted successfully")

    print("content: ",trainer)

    return JsonResponse("trainer 'deleted' field updated successfully", safe=False)

#_____________________________________Test_________________________________________________________#


@api_view(['GET'])
def get_test(request):
    # Create a cache key that uniquely identifies this query
    cache_key = 'test_master_data'
    
    # Try to get the data from the cache
    test_data = cache.get(cache_key)
    
    # If the data is not in the cache, fetch it from the database
    if not test_data:
        test_data = list(test_master.objects.filter(deleted=0)
                        .select_related('test_type_id', 'question_type_id', 'skill_type_id')
                        .values(
                            'id',
                            'test_name',
                            'test_type_id__test_type',               # Fetch the related test_type field
                            'test_type_id__test_type_categories',    # Fetch the related test_type_categories field
                            'question_type_id__question_type',       # Fetch the related question_type field
                            'skill_type_id__skill_type'              # Fetch the related skill_type field
                        ))

        # Store the fetched data in the cache with a timeout
        cache.set(cache_key, test_data, timeout=3600)  # Timeout is set to 1 hour (3600 seconds)
    
    # Return the cached or freshly fetched data
    return Response(test_data)



class testcreateAPIView(generics.CreateAPIView):
    queryset = test_master.objects.all()
    serializer_class = testsSerializersAdd

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new Test")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new Test successfully")
        return response

class testsRetrieveUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = test_master.objects.all()
    serializer_class = testsSerializersAddUpdate

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating Test with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated Test with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating Test with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated Test with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_tests(request, pk):
    ##logger.info(f"Attempting to mark Test with id {pk} as deleted")
    try:
        print("Entering Function..")
        test = test_master.objects.get(id=pk)

      
    except test_master.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    test.deleted = 1
    test.save()

    ##logger.info(f"Marked Test with id {pk} as deleted successfully")

    print("skill1: ",test)

    return JsonResponse("skill 'deleted' field updated successfully", safe=False)
#______________________________________Question_master________________________________________

@api_view(['GET'])
def get_questions(request):
    # Fetch data using ORM with related fields
    question_data = question_master.objects.filter(deleted=0).select_related('question_name_id').values(
        'id',
        'question_name_id__question_paper_name',  # Fetch the related question paper name
        'question_text',
        'option_a',
        'option_b',
        'option_c',
        'option_d',
        'view_hint',
        'mark',
        'explain_answer',
        'answer',
        'negative_mark',
        'input_format'
    )

    # Rename the 'question_name_id__question_paper_name' to 'question_id' in the resulting data
    question_data = [
        {
            **question,
            'question_id': question.pop('question_name_id__question_paper_name')
        } 
        for question in question_data
    ]

    return Response(question_data)


class questionscreateAPIView(generics.CreateAPIView):
    queryset = question_master.objects.all()
    serializer_class = questionsSerializer

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new Questions")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new Questions successfully")
        return response

class questionsRetrieveUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = question_master.objects.all()
    serializer_class = questionsSerializer 

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating Questions with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated Questions with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating Questions with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated Questions with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_question(request, pk):
    ##logger.info(f"Attempting to mark Questions with id {pk} as deleted")
    try:
        print("Entering Function..")
        question = question_master.objects.get(id=pk)

        print("skill: ", question)
    except question_master.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    question.deleted = 1
    question.save()

    ##logger.info(f"Marked Questions with id {pk} as deleted successfully")

    print("skill: ", question)

    return JsonResponse("skill 'deleted' field updated successfully", safe=False)



#____________________________test_candidate-map________________________________#



@api_view(['GET'])
def get_tests_candidates_map(request):
    try:
        # Fetch all test-candidate mappings with related data in one query
        tests_candidates = tests_candidates_map.objects.filter(deleted=0).select_related(
            'college_id', 'department_id', 'question_id', 'student_id', 'rules_id'
        ).values(
            'id', 'test_name',
            'college_id', 'college_id__college',  # Directly access related model fields
            'department_id', 'department_id__department',
            'question_id', 'question_id__question_paper_name', 'question_id__test_type',
            'student_id', 'student_id__students_name', 'student_id__user_name',
            'rules_id', 'rules_id__rule_name', 'rules_id__instruction',
            'dtm_start', 'dtm_end', 'attempt_count', 'is_camera_on', 'is_active',
            'duration', 'duration_type', 'year', 'need_candidate_info',
            'total_score', 'avg_mark', 'dtm_created','dtm_login','dtm_submit','dtm_start_test','status',
        )

        # Convert the queryset into a list of dictionaries
        test_candidate_map_data = list(tests_candidates)

        return Response(test_candidate_map_data)
    except Exception as e:
        return Response({'error': str(e)}, status=500)



@api_view(['GET'])
def get_tests_candidates_map_Update_ID(request, id):
    tests_candidates = tests_candidates_map.objects.filter(deleted=0, id=id).select_related(
        'rules_id', 'department_id', 'question_id', 'student_id', 'college_id'
    ).annotate(
        rules_id_value=F('rules_id__id'),
        question_id_value=F('question_id__id'),
    ).values(
        'id', 'test_name', 'question_id_value', 
        'dtm_start', 'dtm_end', 'duration', 'duration_type',
        'rules_id_value'
    )

    if tests_candidates.exists():
        testing = tests_candidates.first()

        # Format the datetime fields
        dtm_start_formatted = django_format_date(localtime(testing['dtm_start']), 'd-m-Y h:i A')
        dtm_end_formatted = django_format_date(localtime(testing['dtm_end']), 'd-m-Y h:i A')

        test_candidate_map_data = {
            'id': testing['id'],
            'test_name': testing['test_name'],
            'question_id': testing['question_id_value'], 
            'dtm_start': dtm_start_formatted,
            'dtm_end': dtm_end_formatted,
            'duration': testing['duration'],
            'duration_type': testing['duration_type'],
            'rules_id': testing['rules_id_value'],
        }
        
        return Response([test_candidate_map_data])  # Return as a list
    
    return Response({'error': 'Test candidate map not found'}, status=404)










@api_view(['GET'])
def get_tests_candidates_map_Reports(request):
    try:
        print('Entering reports function......')

        tests_candidates = tests_candidates_map.objects.filter(deleted=0).select_related(
            'college_id', 'department_id', 'question_id', 'student_id', 'rules_id'
        ).annotate(
            college_name=F('college_id__college'),
            department_name=F('department_id__department'),
            question_name=F('question_id__question_paper_name'),
            student_name=F('student_id__students_name'),
            user_name=F('student_id__user_name'),
            email_id=F('student_id__email_id'),
            mobile_number=F('student_id__mobile_number'),
            gender=F('student_id__gender'),
            registration_number=F('student_id__registration_number'),
            rules_name=F('rules_id__rule_name'),
            instruction=F('rules_id__instruction')
        ).values(
            'id', 'test_name', 'college_name', 'department_name', 'question_name', 'student_id', 'registration_number',
            'email_id', 'mobile_number', 'gender', 'student_name', 'user_name', 'dtm_start', 'dtm_end',
            'attempt_count', 'is_camera_on', 'is_active', 'duration', 'duration_type', 'year', 'rules_name', 'instruction', 
            'need_candidate_info', 'total_score', 'avg_mark'
        )

        test_candidate_map_data = list(tests_candidates)

        # Format dates and handle total_score
        for item in test_candidate_map_data:
            item['dtm_start'] = django_format_date(localtime(item['dtm_start']), 'd-m-Y h:i A')
            item['dtm_end'] = django_format_date(localtime(item['dtm_end']), 'd-m-Y h:i A')
            item['total_score'] = item['total_score'] if item['total_score'] is not None else 'AA'

        print('test_candidate_map_data: ', test_candidate_map_data)
        return Response(test_candidate_map_data)
    except Exception as e:
        print('Error:', e)
        return Response({'error': str(e)}, status=500)


@api_view(['GET'])
def get_tests_Reports(request):
    tests_candidates = tests_candidates_map.objects.filter(deleted=0).select_related(
        'rules_id', 'department_id', 'question_id', 'student_id', 'college_id'
    )

    test_candidate_map_data = []
    for testing in tests_candidates:
        rules_id = testing.rules_id.rule_name if testing.rules_id else None
        instruction = testing.rules_id.instruction if testing.rules_id else None
        department_id = testing.department_id.department if testing.department_id else None
        question_name = testing.question_id.question_paper_name if testing.question_id else None
        student_info = testing.student_id if testing.student_id else None
        college_id = testing.college_id.college if testing.college_id else None

        student_data = {
            'student_id': student_info.id if student_info else None,
            'student_name': student_info.students_name if student_info else None,
            'user_name': student_info.user_name if student_info else None,
            'email_id': student_info.email_id if student_info else None,
            'mobile_number': student_info.mobile_number if student_info else None,
            'gender': student_info.gender if student_info else None,
            'registration_number': student_info.registration_number if student_info else None,
        } if student_info else {}

        dtm_start_formatted = django_format_date(localtime(testing.dtm_start), 'd-m-Y h:i A')
        dtm_end_formatted = django_format_date(localtime(testing.dtm_end), 'd-m-Y h:i A')
        total_score_display = testing.total_score if testing.total_score is not None else 'AA'

        test_candidate_map_data.append({
            'id': testing.id,
            'test_name': testing.test_name,
            'college_id': college_id,
            'department_id': department_id,
            'question_id': question_name,
            'dtm_start': dtm_start_formatted,
            'dtm_end': dtm_end_formatted,
            'attempt_count': testing.attempt_count,
            'is_camera_on': testing.is_camera_on,
            'is_active': testing.is_active,
            'duration': testing.duration,
            'duration_type': testing.duration_type,
            'year': testing.year,
            'rules': rules_id,
            'instruction': instruction,
            'need_candidate_info': testing.need_candidate_info,
            'total_score': total_score_display,
            'avg_mark': testing.avg_mark,
            **student_data
        })

    return Response(test_candidate_map_data)




class testcandidatemapcreateAPIView(generics.CreateAPIView):
    queryset = tests_candidates_map.objects.all()
    serializer_class =testcandidatemapSerializers

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new test-candidate-map")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new test-candidate-map successfully")
        return response



class testcandidatemapUpdateAPIView_OLD(generics.UpdateAPIView):
    serializer_class = testcandidatemapSerializers

    def get_queryset(self):
        test_name = self.request.data.get('testName')
        print(f"Querying for test_name: {test_name}")  # Add logging
        return tests_candidates_map.objects.filter(test_name=test_name)

    def put(self, request, *args, **kwargs):
        print(f"Request data: {request.data}")  # Log the entire request data
        queryset = self.get_queryset()
        if not queryset.exists():
            print("Test not found")  # Add logging
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)
        
        response_data = []
        for instance in queryset:
            serializer = self.get_serializer(instance, data=request.data, partial=False)
            serializer.is_valid(raise_exception=True)
            self.perform_update(serializer)
            response_data.append(serializer.data)
        
        print("Update successful")  # Add logging
        return Response(response_data, status=status.HTTP_200_OK)

    def patch(self, request, *args, **kwargs):
        print(f"Request data: {request.data}")  # Log the entire request data
        queryset = self.get_queryset()
        if not queryset.exists():
            print("Test not found")  # Add logging
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)
        
        response_data = []
        for instance in queryset:
            serializer = self.get_serializer(instance, data=request.data, partial=True)
            serializer.is_valid(raise_exception=True)
            self.perform_update(serializer)
            response_data.append(serializer.data)
        
        print("Partial update successful")  # Add logging
        return Response(response_data, status=status.HTTP_200_OK)

    def perform_update(self, serializer):
        # Parse the datetime fields to ensure correct timezone handling
        dtm_start = self.request.data.get('dtm_start')
        dtm_end = self.request.data.get('dtm_end')

        if dtm_start:
            dtm_start = datetime.fromisoformat(dtm_start)
            serializer.validated_data['dtm_start'] = dtm_start

        if dtm_end:
            dtm_end = datetime.fromisoformat(dtm_end)
            serializer.validated_data['dtm_end'] = dtm_end

        serializer.save()



class testcandidatemapUpdateAPIView(generics.UpdateAPIView):
    serializer_class = testcandidatemapSerializersupdate

    def get_queryset(self):
        test_name = self.request.data.get('testName')
        print(f"Querying for test_name: {test_name}")
        return tests_candidates_map.objects.filter(test_name=test_name)

    def update(self, request, *args, **kwargs):
        current_date_time = datetime.now()
        print(f"Request data: {request.data}")
        queryset = self.get_queryset()
        print('QuerySet: ', queryset)
        if not queryset.exists():
            print("Test not found")
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)

        response_data = []
        for instance in queryset:
            request.data['dtm_created'] = current_date_time
            serializer = self.get_serializer(instance, data=request.data, partial=False)
            try:
                serializer.is_valid(raise_exception=True)
                print(f"Validated data before saving: {serializer.validated_data}")
                serializer.save()
                response_data.append(serializer.data)
            except Exception as e:
                print(f"Validation error: {e}")
                return Response({"detail": str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        print("Update successful")
        return Response(response_data, status=status.HTTP_200_OK)

    def partial_update(self, request, *args, **kwargs):
        current_date_time = datetime.now()
        print(f"Request data: {request.data}")
        queryset = self.get_queryset()
        if not queryset.exists():
            print("Test not found")
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)

        response_data = []
        for instance in queryset:
            request.data['dtm_created'] = current_date_time
            serializer = self.get_serializer(instance, data=request.data, partial=True)
            try:
                serializer.is_valid(raise_exception=True)
                print(f"Validated data before saving: {serializer.validated_data}")
                serializer.save()
                response_data.append(serializer.data)
            except Exception as e:
                print(f"Validation error: {e}")
                return Response({"detail": str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        print("Partial update successful")
        return Response(response_data, status=status.HTTP_200_OK)


class test_master_UpdateAPIView(generics.UpdateAPIView):
    serializer_class = testsSerializersAddUpdate

    def get_queryset(self):
        test_name = self.request.data.get('testName')
        print(f"Querying for test_name: {test_name}")  # Add logging
        return test_master.objects.filter(test_name=test_name)  # Use correct model

    def put(self, request, *args, **kwargs):
        print(f"Request data: {request.data}")  # Log the entire request data
        queryset = self.get_queryset()
        if not queryset.exists():
            print("Test not found")  # Add logging
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)
        
        response_data = []
        for instance in queryset:
            serializer = self.get_serializer(instance, data=request.data, partial=False)
            serializer.is_valid(raise_exception=True)
            self.perform_update(serializer)
            response_data.append(serializer.data)
        
        print("Update successful")  # Add logging
        return Response(response_data, status=status.HTTP_200_OK)

    def patch(self, request, *args, **kwargs):
        print(f"Request data: {request.data}")  # Log the entire request data
        queryset = self.get_queryset()
        if not queryset.exists():
            print("Test not found")  # Add logging
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)
        
        response_data = []
        for instance in queryset:
            serializer = self.get_serializer(instance, data=request.data, partial=True)
            serializer.is_valid(raise_exception=True)
            self.perform_update(serializer)
            response_data.append(serializer.data)
        
        print("Partial update successful")  # Add logging
        return Response(response_data, status=status.HTTP_200_OK)




@api_view(['PUT', 'PATCH'])
def delete_testcandidatemap(request, test_name):
    try:
        print("Entering Function..")
        tests_candidates = tests_candidates_map.objects.filter(test_name=test_name)
        if not tests_candidates.exists():
            return JsonResponse({"error": "tests not found"}, status=404)

        for candidate in tests_candidates:
            candidate.deleted = 1
            candidate.save()
            print("Updated candidate: ", candidate)

        return JsonResponse("tests_candidates_map 'deleted' field updated successfully", safe=False)
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)


@api_view(['PUT', 'PATCH'])
def update_testcandidatemap_is_active(request, pk):
    try:
        print("Entering Function..")
        tests_candidates=tests_candidates_map.objects.get(id=pk)
        ##logger.info('Fetching test-candidate-map ')

        print("tests_candidates: ",tests_candidates)
    except tests_candidates_map.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    tests_candidates.is_active = True
    tests_candidates.save()
    ##logger.info('test-candidates updated')
    print("tests_candidates: ",tests_candidates)

    return JsonResponse("tests_candidates_map 'deleted' field updated successfully", safe=False)

@api_view(['PUT', 'PATCH'])
def update_testcandidatemap_status_login(request, student_id_value):
    print("Entering Function student database.")
    try:
        print("Entering Function..")
        # Fetch and update all test candidates related to the given student_id_value
        currentDate = timezone.now()  # Assuming 'curr' is the current datetime
        updated_count = tests_candidates_map.objects.filter(student_id=student_id_value).update(
            dtm_login=currentDate,
            status="Database Updated"  # Assuming you want to set status to this string
        )
        
        if updated_count == 0:
            return JsonResponse("Tests not found", status=404)

        print(f"{updated_count} tests_candidates_map entries updated.")

    except Exception as e:
        print("Error: ", str(e))
        return JsonResponse("An error occurred while updating tests_candidates_map.", status=500)

    return JsonResponse(f"All {updated_count} tests_candidates_map entries updated successfully", safe=False)


@api_view(['PUT', 'PATCH'])
def update_testcandidatemap_status_submit(request, pk):
    try:
        print("Entering Function..")
        tests_candidates = tests_candidates_map.objects.get(id=pk)
        print("tests_candidates: ", tests_candidates)
    except tests_candidates_map.DoesNotExist:
        return JsonResponse("Tests not found", status=404)

    # Update fields
    currentDate = datetime.now()  # Assuming 'curr' is the current datetime
    tests_candidates.dtm_submit = currentDate 
    tests_candidates.status = "Submitted"  # Assuming you want to set status to this string
    
    tests_candidates.save()
    print("tests_candidates updated: ", tests_candidates)

    return JsonResponse("tests_candidates_map 'deleted' field updated successfully", safe=False)
@api_view(['PUT', 'PATCH'])
def update_testcandidatemap_status_start_test(request, pk):
    try:
        print("Entering Function..")
        tests_candidates = tests_candidates_map.objects.get(id=pk)
        print("tests_candidates: ", tests_candidates)
    except tests_candidates_map.DoesNotExist:
        return JsonResponse("Tests not found", status=404)

    # Update fields
    currentDate = datetime.now()  # Assuming 'curr' is the current datetime
    tests_candidates.dtm_start_test = currentDate 
    tests_candidates.status = "Test Started"  # Assuming you want to set status to this string
    print("tests_candidates updated: ", tests_candidates)
    tests_candidates.save()
    print("tests_candidates updated: ", tests_candidates)

    return JsonResponse("tests_candidates_map 'deleted' field updated successfully", safe=False)

#__________________________test_candidate_answer____________________________


@api_view(['GET'])
def get_tests_candidates_answer_OLD(request):
    # Fetch data using ORM with related fields
    test_candidate_answer_data = tests_candidates_answers.objects.filter(deleted=0).select_related('student_id', 'question_id').order_by('-id').values(
        'id',
        'student_id__id',  # Fetch student ID
        'student_id__students_name',  # Fetch student name
        'question_id__id',  # Fetch question ID
        'question_id__question_text',  # Fetch question text
        'test_name',
        'answer',
        'result',
        'dtm_start',
        'dtm_end',
        'submission_compile_code',
        'compile_code_editor'
    )

    # Rename the fields to match your desired structure
    test_candidate_answer_data = [
        {
            'id': testans['id'],
            'student_id': testans['student_id__id'],
            'student_name': testans['student_id__students_name'],
            'question_id': testans['question_id__id'],
            'question_name': testans['question_id__question_text'],
            'test_name': testans['test_name'],
            'answer': testans['answer'],
            'result': testans['result'],
            'dtm_start': testans['dtm_start'],
            'dtm_end': testans['dtm_end'],
            'submission_compile_code': testans['submission_compile_code'],
            'compile_code_editor': testans['compile_code_editor'],
        }
        for testans in test_candidate_answer_data
    ]

    return Response(test_candidate_answer_data)


@api_view(['GET'])
def get_tests_candidates_answer(request):
    username = request.query_params.get('username') 
    testName = request.query_params.get('testName')

    # Fetch data using ORM with related fields
    test_candidate_answer_data = tests_candidates_answers.objects.filter(deleted=0, student_id__user_name=username, test_name=testName).select_related('student_id', 'question_id').order_by('-id').values(
        'id',
        'student_id__id',  # Fetch student ID
        'question_id__id',  # Fetch question ID
        'test_name',
        'answer',
        'result',
        'submission_compile_code',
    )

    return Response(test_candidate_answer_data)



@csrf_exempt
@api_view(['POST'])
def test_candidates_answer_view(request, format=None):
    ques_id =  request.data.get('question_id')
    ans =  request.data.get('ans')
    code =  request.data.get('code')

    # code = urllib.parse.unquote(code)
    print("Parameters: ", ques_id, ans, code)
    
    try:
        question = question_master.objects.get(id=ques_id)
    except question_master.DoesNotExist:
        return JsonResponse({'error': 'Question not found'}, status=404)

    if question.answer == ans:
        result = question.mark
    else:
        # Handling multi-line strings
        similarity_score = fuzz.ratio(question.explain_answer.strip(), code.strip())
        result = round((similarity_score / 100) * question.mark)


    # Round the result to the nearest integer
    result = round(result)
    print("Result: ", result)
    
    test_candidate_answer_data = {
        'test_name': request.data.get('test_name'),
        'question_id': ques_id,
        'student_id': request.data.get('student_id'),
        'submission_compile_code': ans,
        'compile_code_editor': code,
        'result': result,
        'dtm_start': request.data.get('dtm_start'),
        'dtm_end': request.data.get('dtm_end'),
    }

    serializer = tests_candidates_answerSerializer(data=test_candidate_answer_data)
    print("Serializer: ", serializer)
    
    if serializer.is_valid():
        serializer.save()
        print('Test Answer is added Successfully')
        return JsonResponse(serializer.data, status=201)
    else:
        print(serializer.errors)
        return JsonResponse(serializer.errors, status=400)




@csrf_exempt
@api_view(['POST'])
def test_candidates_answer_view_Submit(request, format=None):
    ques_id =  request.data.get('question_id')
    ans =  request.data.get('ans')
    code =  request.data.get('code')

    # code = urllib.parse.unquote(code)
    print("Parameters: ", ques_id, ans, code)
    
    try:
        question = question_master.objects.get(id=ques_id)
    except question_master.DoesNotExist:
        return JsonResponse({'error': 'Question not found'}, status=404)

    if question.answer == ans:
        result = question.mark
    else:
        # Handling multi-line strings
        similarity_score = fuzz.ratio(question.explain_answer.strip(), code.strip())
        result = round((similarity_score / 100) * question.mark)


    # Round the result to the nearest integer
    result = round(result)
    print("Result: ", result)
    
    test_candidate_answer_data = {
        'test_name': request.data.get('test_name'),
        'question_id': ques_id,
        'student_id': request.data.get('student_id'),
        'answer': ans,
        'compile_code_editor': code,
        'result': result,
        'dtm_start': request.data.get('dtm_start'),
        'dtm_end': request.data.get('dtm_end'),
    }

    serializer = tests_candidates_answerSerializer(data=test_candidate_answer_data)
    print("Serializer: ", serializer)
    
    if serializer.is_valid():
        serializer.save()
        print('Test Answer is added Successfully')
        return JsonResponse(serializer.data, status=201)
    else:
        print(serializer.errors)
        return JsonResponse(serializer.errors, status=400)




class testcandidateanscreateAPIView(generics.CreateAPIView):
    queryset = tests_candidates_answers.objects.all()
    serializer_class =tests_candidates_answerSerializer

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new test-candidate-answer")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new test-candidate-answer successfully")
        return response

class testcandidateansUpdateAPIView(generics.RetrieveUpdateDestroyAPIView):
    queryset = tests_candidates_answers.objects.all()
    serializer_class =tests_candidates_answerSerializer

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating test-candidate-answer with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated test-candidate-answer with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating test-candidate-answer with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated test-candidate-answer with id {kwargs.get('pk')} successfully")
        return response



@api_view(['PUT', 'PATCH'])
def delete_testcandidateanswer(request, pk):
    ##logger.info(f"Attempting to mark test-candidate-answer with id {pk} as deleted")

    try:
        print("Entering Function..")
        tests_candidates_ans=tests_candidates_answers.objects.get(id=pk)

        print("tests_candidates_ans: ",tests_candidates_ans)
    except tests_candidates_answers.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    tests_candidates_ans.deleted = 1
    tests_candidates_ans.save()

    ##logger.info(f"Marked test-candidate-answer with id {pk} as deleted successfully")

    print("tests_candidates: ",tests_candidates_ans)

    return JsonResponse("tests_candidates_ans 'deleted' field updated successfully", safe=False)



#-----------------------------Login-----------------------------------------#


class login_create(generics.CreateAPIView):
    queryset = login.objects.all()
    serializer_class = loginSerializer

    def create(self, request, *args, **kwargs):
        ##logger.debug('login_create view called')
        
        try:
            response = super().create(request, *args, **kwargs)
            #logger.debug(f'Login created with data: {request.data}')
            #logger.debug(f'Response: {response.data}')
            return response
        except Exception as e:
            logger.error(f'Error occurred in login_create view: {e}')
            return Response({'error': 'An error occurred while creating login'}, status=500)

 

@api_view(['GET'])
def get_login(request):
    try:
        # Fetch login records with related college data using ORM
        login_data = login.objects.select_related('college_id').values(
            'id',
            'email_id',
            'user_name',
            'password',
            'role',
            'college_id__id',  # Fetch the related college's ID
            'college_id__college'  # Fetch the related college's name
        )

        # Rename fields to match the desired output
        formatted_login_data = [
            {
                'id': login_entry['id'],
                'email_id': login_entry['email_id'],
                'user_name': login_entry['user_name'],
                'password': login_entry['password'],
                'role': login_entry['role'],
                'college_id': login_entry['college_id__id'],
                'college_name': login_entry['college_id__college']
            }
            for login_entry in login_data
        ]

        return Response(formatted_login_data)
    except Exception as e:
        logger.error(f'Error occurred in get_login function: {e}')
        return Response({'error': 'An error occurred'}, status=500)

class update_login(generics.RetrieveUpdateAPIView):
    queryset = login.objects.all()
    serializer_class = loginSerializer



@api_view(['PUT', 'PATCH'])
def delete_login(request, pk):
    try:
        print("Entering Function..")
        logins = login.objects.get(id=pk)

        print("skilltyp: ", logins)
    except login.DoesNotExist:
        return JsonResponse("login not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    logins.deleted = 1
    logins.save()
    print("logins: ", logins)

    return JsonResponse("logins 'deleted' field updated successfully", safe=False)


#-----------------------------Course Schedule-----------------------------------------#

class add_course_schedule(generics.CreateAPIView):
    queryset = course_schedule.objects.all()
    serializer_class = courseScheduleSerializer


@api_view(['GET'])
def get_course_schedule(request):
    try:
        # Fetch the required data with related fields using ORM
        course_schedule_data = course_schedule.objects.filter(deleted=0).select_related(
            'college_id', 'student_id', 'department_id', 'topic_id', 'trainer_id'
        ).values(
            'id',
            'college_id__college',
            'college_id__id',
            'student_id__students_name',
            'year',
            'department_id__department',
            'department_id__id',
            'topic_id__topic',
            'topic_id__sub_topic',
            'trainer_id__trainer_name',
            'dtm_start_student',
            'dtm_end_student',
            'dtm_start_trainer',
            'dtm_end_trainer',
            'dtm_of_training'
        )

        # Rename the fields as per the desired output format
        formatted_course_schedule_data = [
            {
                'id': course['id'],
                'college_id': course['college_id__college'],
                'college_name_id': course['college_id__id'],
                'student_id': course['student_id__students_name'],
                'year': course['year'],
                'department_id': course['department_id__department'],
                'department_name_id': course['department_id__id'],
                'topic_id': course['topic_id__topic'],
                'sub_topic': course['topic_id__sub_topic'],
                'trainer_id': course['trainer_id__trainer_name'],
                'dtm_start_student': course['dtm_start_student'],
                'dtm_end_student': course['dtm_end_student'],
                'dtm_start_trainer': course['dtm_start_trainer'],
                'dtm_end_trainer': course['dtm_end_trainer'],
                'dtm_of_training': course['dtm_of_training']
            }
            for course in course_schedule_data
        ]

        return Response(formatted_course_schedule_data)
    except Exception as e:
        logger.error(f'Error occurred in get_course_schedule function: {e}')
        return Response({'error': 'An error occurred'}, status=500)


class update_course_schedule(generics.UpdateAPIView):
    queryset = course_schedule.objects.all()
    serializer_class = courseScheduleSerializer



@api_view(['PUT', 'PATCH'])
def delete_course_schedule(request, pk):
    try:
        print("Entering Function..")
        course = course_schedule.objects.get(id=pk)

        print("course: ", course)
    except course_schedule.DoesNotExist:
        return JsonResponse("course not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    course.deleted = 1
    course.save()
    print("logins: ", course)

    return JsonResponse("course 'deleted' field updated successfully", safe=False)



#-----------------------------Course Content Feedback-----------------------------------------#

class add_course_content_feedback(generics.CreateAPIView):
    queryset = course_content_feedback.objects.all()
    serializer_class = courseContentFeedbackSerializer

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new Course content feedback")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new Course content feedback successfully")
        return response


@api_view(['GET'])
def get_course_content_feedback(request):
    try:
        # Fetch data with related fields using select_related and values
        course_content_feedback_data = course_content_feedback.objects.filter(deleted=0).select_related(
            'student_id__college_id', 'student_id__department_id', 'topic_id', 'trainer_id'
        ).values(
            'id',
            'student_id__students_name',
            'student_id__college_id__id',
            'student_id__college_id__college',
            'student_id__department_id__id',
            'student_id__department_id__department',
            'topic_id__topic',
            'topic_id__sub_topic',
            'trainer_id__trainer_name',
            'dtm_session',
            'feedback'
        )

        # Rename the fields as per the desired output format
        formatted_course_content_feedback_data = [
            {
                'id': content['id'],
                'student_id': content['student_id__students_name'],
                'topic_id': content['topic_id__topic'],
                'sub_topic': content['topic_id__sub_topic'],
                'college_id': content['student_id__college_id__id'],
                'college_name': content['student_id__college_id__college'],
                'department_id': content['student_id__department_id__id'],
                'department_name': content['student_id__department_id__department'],
                'dtm_session': content['dtm_session'],
                'trainer_id': content['trainer_id__trainer_name'],
                'feedback': content['feedback'],
            }
            for content in course_content_feedback_data
        ]

        return Response(formatted_course_content_feedback_data)
    except Exception as e:
        logger.error(f'Error occurred in get_course_content_feedback function: {e}')
        return Response({'error': 'An error occurred'}, status=500)


class update_course_content_feedback(generics.UpdateAPIView):
    queryset = course_content_feedback.objects.all()
    serializer_class = courseContentFeedbackSerializer

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating course content feedback with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        ##logger.info(f"Updated course content feedback with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        ##logger.info(f"Partially updating course content feedback with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        ##logger.info(f"Partially updated course content feedback with id {kwargs.get('pk')} successfully")
        return response



@api_view(['PUT', 'PATCH'])
def delete_course_content_feedback(request, pk):
    ##logger.info(f"Attempting to mark course content feedback with id {pk} as deleted")

    try:
        print("Entering Function..")
        course = course_content_feedback.objects.get(id=pk)

        print("course: ", course)
    except course_content_feedback.DoesNotExist:
        return JsonResponse("course not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    course.deleted = 1
    course.save()

    ##logger.info(f"Marked course content feedback with id {pk} as deleted successfully")

    print("course: ", course)

    return JsonResponse("course 'deleted' field updated successfully", safe=False)


#-----------------------------Attendance Master-----------------------------------------#

class add_attendance_master(generics.CreateAPIView):
    queryset = attendance_master.objects.all()
    serializer_class = attendanceMasterSerializer

    def post(self, request, *args, **kwargs):
        ##logger.info("Creating a new Attendance")
        response = super().post(request, *args, **kwargs)
        ##logger.info("Created a new Attendance successfully")
        return response


@api_view(['GET'])
def get_attendance_master(request):
    attend_set = attendance_master.objects.select_related('student_id', 'test_id').filter(deleted=0).annotate(
        student_name=F('student_id__students_name'),
       
        test_name=F('test_id__test_name')
    ).values(
        'id', 'student_name', 'course_name', 'test_name', 'dtm_attendance'
    )

    return Response(list(attend_set))

class update_attendance_master(generics.UpdateAPIView):
    queryset = attendance_master.objects.all()
    serializer_class = attendanceMasterSerializer

    def put(self, request, *args, **kwargs):
        ##logger.info(f"Updating Attendance with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        #logger.info(f"Updated Attendance with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        #logger.info(f"Partially updating Attendance with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        #logger.info(f"Partially updated Attendance with id {kwargs.get('pk')} successfully")
        return response



@api_view(['PUT', 'PATCH'])
def delete_attendance_master(request, pk):
    #logger.info(f"Attempting to mark Attendance with id {pk} as deleted")

    try:
        print("Entering Function..")
        attend = attendance_master.objects.get(id=pk)

        print("attend: ", attend)
    except attendance_master.DoesNotExist:
        return JsonResponse("course not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    attend.deleted = 1
    attend.save()

    #logger.info(f"Marked Attendance with id {pk} as deleted successfully")

    print("attend: ", attend)

    return JsonResponse("attend 'deleted' field updated successfully", safe=False)



#-----------------------------Announcement Master-----------------------------------------#

class add_announcement_master(generics.CreateAPIView):
    queryset = announcement_master.objects.all()
    serializer_class = announcementSerializer

class update_announcement_master(generics.UpdateAPIView):
    queryset = announcement_master.objects.all()
    serializer_class = announcementSerializer



@api_view(['PUT', 'PATCH'])
def delete_announcement_master(request, pk):
    try:
        print("Entering Function..")
        announce = announcement_master.objects.get(id=pk)

        print("announce: ", announce)
    except announcement_master.DoesNotExist:
        return JsonResponse("announce not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    announce.deleted = 1
    announce.save()
    print("announce: ", announce)

    return JsonResponse("announce 'deleted' field updated successfully", safe=False)


#-----------------------------Test Question Map-----------------------------------------#

class add_testQuestion_map(generics.CreateAPIView):
    queryset = test_question_map.objects.all()
    serializer_class = testQuestionMapSerializer


@api_view(['GET'])
def get_testQuestion_map(request):
    mapping_data = test_question_map.objects.select_related('test_id', 'question_id').filter(deleted=0).annotate(
        test_name=F('test_id__test_name'),
        question_name=F('question_id__question_name')
    ).values(
        'id', 'test_name', 'question_name'
    )

    return Response(list(mapping_data))

class update_testQuestion_map(generics.UpdateAPIView):
    queryset = test_question_map.objects.all()
    serializer_class = testQuestionMapSerializer



@api_view(['PUT', 'PATCH'])
def delete_testQuestion_map(request, pk):
    try:
        print("Entering Function..")
        mapping = test_question_map.objects.get(id=pk)

        print("mapping: ", mapping)
    except test_question_map.DoesNotExist:
        return JsonResponse("mapping not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    mapping.deleted = 1
    mapping.save()
    print("mapping: ", mapping)

    return JsonResponse("mapping 'deleted' field updated successfully", safe=False)





#--------------------------------------Import Questions-----------------------------------#


def get_college_id(college_name):
    college_instance = college_master.objects.filter(college=college_name).first()
    if not college_instance:
        raise ValueError(f"College '{college_name}' not found")
    return college_instance.id

def get_department_id(department_name):
    department_instance = department_master.objects.filter(department=department_name).first()
    if not department_instance:
        raise ValueError(f"Department '{department_name}' not found")
    return department_instance.id

# Function to handle conversion of float values to strings without .0
def clean_char_field(value):
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


class ExcelImportView_CandidateLAST(APIView):
    def post(self, request, format=None):
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        header_mapping = {
            'College Name**': 'college_id',
            'Student Name': 'students_name',
            'User Name**': 'user_name',
            'Reg No': 'registration_number',
            'Gender': 'gender',
            'Email ID': 'email_id',
            'Mobile Number': 'mobile_number',
            'Year**': 'year',
            'CGPA': 'cgpa',
            'Department**': 'department_id',
            '10th Mark': 'marks_10th',
            '12th Mark': 'marks_12th',
            'Semaster Wise': 'marks_semester_wise',
            'History Of Arrears': 'history_of_arrears',
            'Standing Arrears': 'standing_arrears',
            'No.Of.IT Offers': 'it_of_offers',
            'No.Of.Core Offers': 'core_of_offers',
            'No.Of.Offers': 'number_of_offers',
            'Password**': 'password'
        }
        
        df.rename(columns=header_mapping, inplace=True)

        char_fields = ['college_id', 'user_name', 'year', 'password']  # Add any other char fields as needed

        # Combine clean_char_field and strip operations in one loop
        for column in df.columns:
            if column in char_fields:
                df[column] = df[column].apply(clean_char_field)
            if df[column].dtype == 'object':
                df[column] = df[column].str.strip()

        # Ensure all IDs are mapped to instance IDs
        try:
            df['college_id'] = df['college_id'].apply(get_college_id)
            df['department_id'] = df['department_id'].apply(get_department_id)
        except ValueError as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)


        mandatory_columns = ['college_id', 'user_name', 'department_id', 'year', 'password']
        for column in mandatory_columns:
            if column not in df.columns:
                return Response({'error': f'Mandatory column {column} is missing'}, status=status.HTTP_400_BAD_REQUEST)
        
        if df[mandatory_columns].isnull().any().any():
            return Response({'error': 'Mandatory fields cannot be null'}, status=status.HTTP_400_BAD_REQUEST)

        email_regex = re.compile(r'^\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b$')

        user_names = df['user_name']
        existing_users = set(User.objects.filter(username__in=user_names).values_list('username', flat=True))

        skipped_usernames = []
        added_usernames = []

        processed_usernames = set()
        valid_rows = []
        
        for index, row in df.iterrows():
            username = row['user_name']

            if username in existing_users or username in processed_usernames:
                skipped_usernames.append(username)
                continue  # Skip this row if username already exists or is a duplicate within the file

            processed_usernames.add(username)
            
            if not row['college_id']:
                return Response({'error': f'College not found in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
            elif not row['department_id']:
                return Response({'error': f'Department not found in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            year_value = str(row['year'])
            print('year value: ', year_value)
            if year_value not in ['1', '2', '3', '4']:
                return Response({'error': f'Year is mismatch in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            if pd.notna(row['email_id']):
                if not isinstance(row['email_id'], str) or not email_regex.match(row['email_id']):
                    return Response({'error': f'Email ID is not in correct format in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
        
            if pd.notna(row['mobile_number']):
                if len(str(int(row['mobile_number']))) != 10:
                    return Response({'error': f'Mobile number is not 10 digits in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            integer_fields = ['history_of_arrears', 'standing_arrears', 'core_of_offers', 'it_of_offers', 'number_of_offers']
            for field in integer_fields:
                if pd.notna(row[field]):
                    try:
                        value = float(row[field])
                        if value.is_integer() and 0 <= value <= 100:
                            row[field] = int(value)
                        else:
                            return Response({'error': f'{field.replace("_", " ").title()} must be an integer between 0 and 100 in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
                    except ValueError:
                        return Response({'error': f'{field.replace("_", " ").title()} must be an integer between 0 and 100 in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)

            float_fields = ['cgpa', 'marks_10th', 'marks_12th']
            for field in float_fields:
                if pd.notna(row[field]):
                    try:
                        row[field] = float(row[field])
                    except ValueError:
                        return Response({'error': f'{field.replace("_", " ").title()} must be a number in row {index + 2}'}, status=status.HTTP_400_BAD_REQUEST)
    
            df.at[index, 'user_name'] = username

            valid_rows.append(row)
        
        df_valid = pd.DataFrame(valid_rows)
        df_valid['role'] = 'Student'
        df_valid['is_database'] = True

        print('*********df_valid**********: ', df_valid)
        
        optional_columns = ['students_name', 'registration_number', 'gender', 'email_id', 'mobile_number',
                            'cgpa', 'marks_10th', 'marks_12th', 'marks_semester_wise',
                            'history_of_arrears', 'standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers', 'text']
        
        for column in optional_columns:
            if column in df_valid.columns:
                if df_valid[column].dtype == 'float64':  # For float columns
                    df_valid[column] = df_valid[column].fillna(0.0)
                elif df_valid[column].dtype == 'int64':  # For integer columns
                    df_valid[column] = df_valid[column].fillna(0)
                else:  # For other columns
                    df_valid[column] = df_valid[column].fillna(pd.NA)
            else:
                if column in ['cgpa', 'marks_10th', 'marks_12th']:
                    df_valid[column] = 0.0
                elif column in ['history_of_arrears', 'standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers']:
                    df_valid[column] = 0
                else:
                    df_valid[column] = pd.NA
        
        df_valid['email_id'] = df_valid['email_id'].apply(lambda x: x if pd.notna(x) and '@' in str(x) else 'placeholder@example.com')
        
        df_valid = df_valid.where(pd.notnull(df_valid), None)

        # Prepare data for User and UserProfile serializers
        user_profile_data = []
        for index, row in df_valid.iterrows():
            user_data = {
                'username': row['user_name'],
                'email': row['email_id'],
                'password': row['password']
            }
            user_profile_data.append({
                'user': user_data,
                'role': row['role'],
                'college_id': row['college_id']
            })

        user_profile_serializer = UserProfileSerializer(data=user_profile_data, many=True)
        print('user_profile_serializer: ', user_profile_serializer)

        if not user_profile_serializer.is_valid():
            print('user_profile_serializer.errors: ', user_profile_serializer.errors)
            return Response(user_profile_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        
        candidate_data = df_valid[['college_id', 'students_name', 'registration_number', 'gender', 'email_id', 'mobile_number',
                             'department_id', 'year', 'cgpa', 'marks_10th', 'marks_12th', 'marks_semester_wise',
                             'history_of_arrears', 'standing_arrears', 'number_of_offers', 'it_of_offers', 'core_of_offers', 'user_name', 'text', 'is_database']]
        candidate_records = candidate_data.to_dict(orient='records')
        
        candidate_serializer = candidateSerializerImport(data=candidate_records, many=True)
        
        if not candidate_serializer.is_valid():
            print('candidate_serializer.errors: ', candidate_serializer.errors)
            return Response(candidate_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        with transaction.atomic():
            user_profile_serializer.save()
            candidate_serializer.save()

        # Return a summary message
        summary_message = f"Excel file imported successfully. Skipped {len(skipped_usernames)} existing users, added {len(valid_rows)} new users."
        return Response({'message': summary_message}, status=status.HTTP_201_CREATED)





class ExcelImportView_Candidateuser(APIView):
    def post(self, request, format=None):
        print("Starting Excel Import Process...")
        
        # Initialize current date and time
        current_date_time = datetime.now()
        print('current_date_time: ', current_date_time)
        
        # Check if file is present in the request
        if 'file' not in request.FILES:
            print("Error: No file uploaded.")
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        # Check if file is in Excel format
        if not file.name.endswith('.xlsx'):
            print("Error: File is not in Excel format.")
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print("File read successfully.")
        except Exception as e:
            print(f"Error reading file: {str(e)}")
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
        # Check required columns
        if 'user_name' not in df.columns or df['user_name'].isnull().any():
            print("Error: User name is missing or contains null values.")
            return Response({'error': 'User name cannot be null'}, status=status.HTTP_400_BAD_REQUEST)

        if 'password' not in df.columns:
            print("Error: Password column is missing.")
            return Response({'error': 'Password column is missing'}, status=status.HTTP_400_BAD_REQUEST)

        

        # Ensure all IDs are mapped to instance IDs
        try:
            df['college_id'] = df['college_id'].apply(get_college_id)
           
        except ValueError as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        # Prepare login data
        login_data = df[['user_name', 'password', 'college_id']]
        login_data['role'] = 'Student'  # Default role
        print("Login data prepared.")

        # Get existing users
        existing_users = set(User.objects.filter(username__in=login_data['user_name']).values_list('username', flat=True))
        print(f"Existing users fetched: {existing_users}")

        # Filter out existing users
       # new_login_data = login_data[~login_data['user_name'].isin(existing_users)]
        new_login_data = login_data[~login_data['user_name'].astype(str).isin(existing_users)]
       
        
        if not new_login_data.empty:
            print(f"New login data found: {len(new_login_data)} entries")
            new_login_records = new_login_data.to_dict(orient='records')

            # Create user profiles and users
            with transaction.atomic():
                for record in new_login_records:
                    print('Record: ', record)
                    # Convert password to string
                    password_str = str(record['password'])

                    # Create user
                    user = User.objects.create_user(
                        username=record['user_name'],
                        password=password_str  # Use the stringified password
                    )

                    # Fetch college instance using the college_id from record
                    college_instance = college_master.objects.filter(id=record['college_id']).first()
                    if not college_instance:
                        raise ValueError(f"College with ID {record['college_id']} not found")


                    # Create user profile
                    user_profile.objects.create(
                        user=user,
                        role=record['role'],
                        college_id=college_instance
                    )
                print("New login data saved successfully.")
        else:
            print("No new login data to save (all user_names already exist).")

        # Prepare candidate data
        candidate_data = df[['user_name', 'college_id']]
        candidate_data['is_database'] = False
        candidate_data['dtm_upload'] = current_date_time  
        print("Candidate data prepared.")

        # Get existing candidates
        existing_candidates = set(candidate_master.objects.filter(user_name__in=candidate_data['user_name']).values_list('user_name', flat=True))
        print(f"Existing candidates fetched: {existing_candidates}")

        # Debugging: Print candidate data for verification
        print("Candidate data from file:")
        print(candidate_data)

        # Filter out existing candidates
        new_candidate_data = candidate_data[~candidate_data['user_name'].astype(str).isin(existing_candidates)]
    
       # new_candidate_data = candidate_data[~candidate_data['user_name'].isin(existing_candidates)]
        print(f"Filtered new candidate data: {new_candidate_data}")

        if not new_candidate_data.empty:
            print(f"New candidate data found: {len(new_candidate_data)} entries")
            new_candidate_records = new_candidate_data.to_dict(orient='records')

            # Avoid validation errors by handling existing candidate names before calling the serializer
            candidate_serializer = candidateuserSerializerImport(data=new_candidate_records, many=True)

            if candidate_serializer.is_valid():
                candidate_serializer.save()
                print("New candidate data saved successfully.")
            else:
                print("Candidate data validation failed with errors:")
                print(candidate_serializer.errors)
        else:
            print("No new candidate data to save (all user_names already exist).")

        # Return success response with saved data
        print("Import process completed successfully.")
        return Response({
            'new_login_data': new_login_records if 'new_login_records' in locals() else [],
            'new_candidate_data': new_candidate_records if 'new_candidate_records' in locals() else []
        }, status=status.HTTP_201_CREATED)





class ExcelImportView_Questions(APIView):
    def post(self, request, format=None):

        # Extract data for the question_paper_master
        question_paper_name = request.data.get('question_paper_name')
        duration_of_test = request.data.get('duration_of_test')
        topic = request.data.get('topic')
        sub_topic = request.data.get('sub_topic')
        no_of_questions = request.data.get('no_of_questions')
        upload_type = request.data.get('upload_type')
        test_type = request.data.get('test_type')

        # Validate required fields
        if not all([question_paper_name, duration_of_test, topic, sub_topic, no_of_questions, upload_type, test_type]):
            return Response({'error': 'Missing fields for question_paper_master'}, status=status.HTTP_400_BAD_REQUEST)

        # Create a new question_paper_master entry
        question_paper_data = {
            'question_paper_name': question_paper_name,
            'duration_of_test': duration_of_test,
            'topic': topic,
            'sub_topic': sub_topic,
            'no_of_questions': no_of_questions,
            'upload_type': upload_type,
            'test_type': test_type
        }

        question_paper_serializer = questionsPaperSerializer(data=question_paper_data)
        if question_paper_serializer.is_valid():
            question_paper_instance = question_paper_serializer.save()
            question_paper_id = question_paper_instance.id
        else:
            return Response(question_paper_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        print('Files:', request.FILES)

        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)

        file = request.FILES['file']

        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())  # Print the first few rows of the DataFrame
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        # Mapping of Excel header names to actual column names
        header_mapping = {
            'Questions**': 'question_text',
            'Option A': 'option_a',
            'Option B': 'option_b',
            'Option C': 'option_c',
            'Option D': 'option_d',
            'Answer**': 'answer',
            'Mark**': 'mark',
            'Explain Answer': 'explain_answer',
        }
        
       
        df.rename(columns=header_mapping, inplace=True)
        

        # List of mandatory columns
        mandatory_columns = [
            'question_text',
            'answer',
            'mark'
        ]

        # Initialize error messages
        error_messages = []

        # Check for empty values in mandatory columns and record specific rows
        for col in mandatory_columns:
            if col in df.columns:
                missing_values = df[df[col].isnull()]
                if not missing_values.empty:
                    for index, row in missing_values.iterrows():
                        error_messages.append(f"Row {index + 1}: Column '{col}' is empty.")

        if error_messages:
            error_message = ' '.join(error_messages)
            return Response({'error': error_message}, status=status.HTTP_400_BAD_REQUEST)

        # Convert DataFrame to records and handle NaN values
        records = df.fillna('').to_dict(orient='records')

        for record in records:
            record['question_name_id'] = question_paper_id
            for key in record:
                if isinstance(record[key], str):
                    record[key] = record[key].strip()

        # Serialize and save the records
        serializer = questionsSerializerImport(data=records, many=True)
        print("Serializer: ", serializer)

        if serializer.is_valid():
            print("Before Saving..")
            serializer.save()
            print("After Saving..")
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        else:
            # Initialize a list to collect all errors
            detailed_errors = []
            print("Entering Detailed Error..")

            # Iterate over each error in the list of errors
            for idx, error_dict in enumerate(serializer.errors):
                for field, errors in error_dict.items():
                    if isinstance(errors, list):
                        for error in errors:
                            detailed_errors.append(f"Row {idx + 1}, Column '{field}': {error}")

            print('Deltailed Error: ', detailed_errors)
            # Return the collected errors
            return Response({'error': detailed_errors}, status=status.HTTP_400_BAD_REQUEST)

class ExcelImportView_Questions_Code(APIView):
    def post(self, request, format=None):
        # Extract data for the question_paper_master
        question_paper_name = request.data.get('question_paper_name')
        duration_of_test = request.data.get('duration_of_test')
        topic = request.data.get('topic')
        sub_topic = request.data.get('sub_topic')
        no_of_questions = request.data.get('no_of_questions')
        upload_type = request.data.get('upload_type')
        test_type = request.data.get('test_type')

        # Validate required fields
        if not all([question_paper_name, duration_of_test, topic, sub_topic, no_of_questions, upload_type]):
            return Response({'error': 'Missing fields for question_paper_master'}, status=status.HTTP_400_BAD_REQUEST)

        # Create a new question_paper_master entry
        question_paper_data = {
            'question_paper_name': question_paper_name,
            'duration_of_test': duration_of_test,
            'topic': topic,
            'sub_topic': sub_topic,
            'no_of_questions': no_of_questions,
            'upload_type': upload_type,
            'test_type': test_type
        }

        question_paper_serializer = questionsPaperSerializer(data=question_paper_data)
        if question_paper_serializer.is_valid():
            question_paper_instance = question_paper_serializer.save()
            question_paper_id = question_paper_instance.id
        else:
            return Response(question_paper_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        # Process the uploaded file
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)

        file = request.FILES['file']

        if not file.name.endswith('.xlsx'):
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            df = pd.read_excel(file)
            print('DataFrame contents:')
            print(df.head())
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        # Map Excel columns to the expected format
        header_mapping = {
            'Questions**': 'question_text',
            'Answer**': 'answer',
            'Mark**': 'mark',
            'Explain Answer**': 'explain_answer',
            'Input Format': 'input_format'
        }

        df.rename(columns=header_mapping, inplace=True)

        mandatory_columns = [
            'question_text',
            'answer',
            'mark',
            'explain_answer'
        ]

        empty_columns = [col for col in mandatory_columns if df[col].isnull().any()]

        if empty_columns:
            error_message = f'Mandatory columns have null values: {", ".join(empty_columns)}'
            return Response({'error': error_message}, status=status.HTTP_400_BAD_REQUEST)

        # Convert DataFrame to records and associate with the newly created question_paper_master
        records = df.to_dict(orient='records')

        for record in records:
            record['question_name_id'] = question_paper_id
            for key in record:
                if isinstance(record[key], str):
                    record[key] = record[key].strip()

        # Serialize and save the records
        serializer = questionsSerializerCodeImport(data=records, many=True)
        print("Serializer: ", serializer)

        if serializer.is_valid():
            print("Before Saving..")
            serializer.save()
            print("After Saving..")
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        else:
            print("Serializer Errors: ", serializer.errors)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

#------------------------------Rules-------------------------------------------------------------#


class rules_listView(generics.ListAPIView):
    queryset = rules.objects.filter(deleted=0)

    
    serializer_class = ruleSerializers

class rules_create(generics.CreateAPIView):
    queryset = rules.objects.all()
    serializer_class = ruleSerializers

class rules_Update(generics.RetrieveUpdateAPIView):
    queryset = rules.objects.all()
    serializer_class = ruleSerializers

@api_view(['PUT', 'PATCH'])
def delete_rules(request, pk):
    try:
        print("Entering Function..")
        rule = rules.objects.get(id=pk)

       
    except rules.DoesNotExist:
        return JsonResponse("rules not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    rule.deleted = 1
    rule.save()
 

    return JsonResponse("rules 'deleted' field updated successfully", safe=False)



#------------------------Dashboard---------------------------------#


def attendance_excel_data1(request):
    if request.method == 'GET':
        # Define URLs for each sheet
        sheet_urls = {
            'AIDS': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=257135364',
            'IT': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=747238312',
            'CSE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=1208914504',
            'ECE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=2128791655'
        }
        
        # Initialize an empty dictionary to store data for each sheet
        all_data = {}

        # Iterate over each sheet URL
        for sheet_name, url in sheet_urls.items():
            # Fetch the CSV data from the URL
            response = requests.get(url)

            # Read the CSV data into a Pandas DataFrame, skipping the first 8 rows
            df = pd.read_csv(io.StringIO(response.text), skiprows=7)

            # Transform the DataFrame into the desired format
            transformed_data = []
            for _, row in df.iterrows():
                transformed_row = {
                    'Department': row.get('DePArtment ', ''),  # Handle missing column gracefully
                    'Register Number': row.get('Register Number', ''),  # Handle missing column gracefully
                    'Name of the Student': row.get('NAme of the Student', '')
                }
                # Iterate over date columns and add FN and AN values
                for i in range(4, len(df.columns), 2):  # Start from index 4, increment by 2
                    col = df.columns[i]
                    date = col.split()[0]  # Extract the date
                    fn = row.get(col, '')  # Handle missing column gracefully
                    an_col_index = i + 1  # Get the index of the next column (AN)
                    if an_col_index < len(df.columns):
                        an = row.get(df.columns[an_col_index], '')  # Handle missing column gracefully
                    else:
                        an = ''  # If AN column does not exist, set it to empty string
                    transformed_row[date] = {'FN': fn, 'AN': an}
                transformed_data.append(transformed_row)

            # Store the transformed data in the dictionary
            all_data[sheet_name] = transformed_data

        return JsonResponse(all_data, safe=False)
    else:
        return JsonResponse({'error': 'Method not allowed'}, status=405)



#---------------------------------Password Reset----------------------------#


@api_view(['PUT'])
def update_login(request, user_name=None):
    login_to_update = login.objects.get(user_name=user_name)
    serializer = loginSerializer(instance=login_to_update, data=request.data, partial=True)
    
    if serializer.is_valid():
        serializer.save()
        print("Updated.")
        return JsonResponse("Login Updated Successfully", safe=False)
    return JsonResponse("Failed to Update Login")


@api_view(['PUT'])
def update_totalScore_test_candidate_map(request, pk=None):
    total_score_update = tests_candidates_map.objects.get(id=pk)
    serializer = testcandidatemapSerializers(instance=total_score_update, data=request.data, partial=True)

    if serializer.is_valid():
        serializer.save()
        print('Total score updated')
        return JsonResponse("Total score Updated Successfully", safe=False)
    return JsonResponse("Failed to Update total score")


@api_view(['PUT'])
def update_Avg_Marks_test_candidate_map(request, pk=None):
    avg_mark_update = tests_candidates_map.objects.get(id=pk)
    serializer = testcandidatemapSerializers(instance=avg_mark_update, data=request.data, partial=True)

    if serializer.is_valid():
        serializer.save()
        print('Avg mark is updated')
        return JsonResponse("Avg mark is  Updated Successfully", safe=False)
    return JsonResponse("Failed to Update Avg mark")




#---------------------------------Batch wise Test assign-------------------------------#


class TestCandidatesMapView_batch(APIView):
    def post(self, request, batch_no, format=None):
        print("Batch No: ", batch_no)
        students = candidate_master.objects.filter(batch_no=batch_no)
        print("Students: ", students)
        data = []
        for student in students:
            print("Student id: ", student.id)
            test_candidate_data = {
                'test_id': request.data.get('test_id'),  # Assuming 'test_id' is provided in the request
                'question_id': request.data.get('question_id'),  # Assuming 'question_id' is provided in the request
                'student_id': student.id,
                'college_id': request.data.get('college_id'),
                'department_id': request.data.get('department_id'),
                'dtm_start': request.data.get('dtm_start'),
                'dtm_end': request.data.get('dtm_end'),
                'is_camera_on': request.data.get('is_camera_on'),
                'duration': request.data.get('duration'),
                'year': request.data.get('year'),
                'rules_id': request.data.get('rules_id'),
                'need_candidate_info': request.data.get('need_candidate_info'),
                # Add other fields here
            }

            serializer = testcandidatemapSerializers(data=test_candidate_data)
            print("Serializer: ", serializer)
            if serializer.is_valid():
                serializer.save()
                data.append(serializer.data)
                print("Data.append: ", data)
            else:
                print(serializer.errors)
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        print("Data: ", data)
        return Response(data, status=status.HTTP_201_CREATED)



class TestCandidatesMapView(APIView):
    def post(self, request, format=None):

        test_name = request.data.get('test_name')
        test_type_id = request.data.get('test_type_id')
        question_type_id = request.data.get('question_type_id')
        skill_type_id = request.data.get('skill_type_id')

        if not all([test_name, test_type_id, question_type_id]):
            return Response({'error': 'Missing fields for test_master'}, status=status.HTTP_400_BAD_REQUEST)

        test_master_data = {
            'test_name': test_name,
            'test_type_id': test_type_id,
            'question_type_id': question_type_id,
            'skill_type_id': skill_type_id
        }

        test_master_serializer = testsSerializersAdd(data=test_master_data)
        
        if not test_master_serializer.is_valid():
            print(f"Test master serializer errors: {test_master_serializer.errors}")
            return Response(test_master_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        

        college_id = request.data.get('college_id', [])
        department_id = request.data.get('department_id', [])
        year = request.data.get('year')

        if not college_id or not department_id or not year:
            return Response({'error': 'college_id, department_id, and year are required fields.'}, status=status.HTTP_400_BAD_REQUEST)

        print("Batch No: ", year, college_id, department_id)

        # Filter candidates based on multiple college_ids and department_ids
        students = candidate_master.objects.filter(
            college_id__in=college_id,
            department_id__in=department_id,
            year=year
        )

        print("Students: ", students)
        data = []
        updated_candidates = []  # To store IDs of updated candidates

        current_date_and_time = datetime.now()  # Get the current date and time

        for student in students:
            print("Student id: ", student.id)
            test_candidate_data = {
                'test_name': request.data.get('test_name'),  # Assuming 'test_name' is provided in the request
                'question_id': request.data.get('question_id'),  # Assuming 'question_id' is provided in the request
                'student_id': student.id,
                'college_id': student.college_id.id,  # Use student.college_id.id for pk
                'department_id': student.department_id.id,  # Use student.department_id.id for pk
                'dtm_start': request.data.get('dtm_start'),
                'dtm_end': request.data.get('dtm_end'),
                'is_camera_on': request.data.get('is_camera_on'),
                'duration': request.data.get('duration'),
                'duration_type': request.data.get('duration_type'),
                'year': year,
                'rules_id': request.data.get('rules_id'),
                'need_candidate_info': request.data.get('need_candidate_info'),
                'dtm_created': current_date_and_time  # Add current date and time
            }

            serializer = testcandidatemapSerializers(data=test_candidate_data)
            
            print("Serializer: ", serializer)

            if not serializer.is_valid():
                print(f"serializer errors: {serializer.errors}")
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        


            with transaction.atomic():  # Use transaction to ensure atomicity
                test_master_serializer.save()
                print('Test master saved Successfully')
                serializer.save()
                data.append(serializer.data)
                print("Data.append: ", data)

                # Check and update candidate_master if needed
                if student.need_candidate_info is None and request.data.get('need_candidate_info') is True:
                    student.need_candidate_info = request.data.get('need_candidate_info')
                    student.save(update_fields=['need_candidate_info'])
                    updated_candidates.append(student.id)
        
        print("Data: ", data)
        return Response(data, status=status.HTTP_201_CREATED)


#---------------Non database test assign-------------------#

class NonDbTestAssign_OLD(APIView):
    def post(self, request):
        students = candidate_master.objects.filter(
            Q(college_id__isnull=True),
            Q(department_id__isnull=True) ,
            Q(year__isnull=True) | Q(year=''),
            Q(batch_no__isnull=True) | Q(batch_no=''),
            Q(students_name__isnull=True) | Q(students_name=''),
            Q(registration_number__isnull=True) | Q(registration_number=''),
            Q(gender__isnull=True) | Q(gender=''),
            Q(email_id__isnull=True) | Q(email_id=''),
            Q(mobile_number__isnull=True) | Q(mobile_number=''),
            Q(cgpa__isnull=True) | Q(cgpa=''),
            Q(marks_10th__isnull=True) | Q(marks_10th=''),
            Q(marks_12th__isnull=True) | Q(marks_12th=''),
            Q(marks_semester_wise__isnull=True) | Q(marks_semester_wise=''),
            Q(history_of_arrears__isnull=True) | Q(history_of_arrears=''),
            Q(standing_arrears__isnull=True),
            Q(number_of_offers__isnull=True),
            Q(text__isnull=True) | Q(text=''),
            ~Q(user_name__isnull=True) & ~Q(user_name=''),
        )
        
        print("Students: ", students)
        data = []
        updated_candidates = []  # To store IDs of updated candidates

        current_date_and_time = datetime.now() 

        for student in students:
            print("Student id: ", student.id)
            test_candidate_data = {
                'test_name': request.data.get('test_name'),  # Assuming 'test_id' is provided in the request
                'question_id': request.data.get('question_id'),  # Assuming 'question_id' is provided in the request
                'student_id': student.id,
                'dtm_start': request.data.get('dtm_start'),
                'dtm_end': request.data.get('dtm_end'),
                'is_camera_on': request.data.get('is_camera_on'),
                'duration': request.data.get('duration'),
                'duration_type': request.data.get('duration_type'),
                'rules_id': request.data.get('rules_id'),
                'need_candidate_info': request.data.get('need_candidate_info'),
                'dtm_created': current_date_and_time  # Add current date and time
                # Add other fields here
            }

            serializer = NonDbTestAssignSerializer(data=test_candidate_data)
            print("Serializer: ", serializer)
            if serializer.is_valid():
                with transaction.atomic():  # Use transaction to ensure atomicity
                    serializer.save()
                    data.append(serializer.data)
                    print("Data.append: ", data)

                    # Check and update candidate_master if needed
                    if student.need_candidate_info is None and request.data.get('need_candidate_info') is True:
                        student.need_candidate_info = request.data.get('need_candidate_info')
                        student.save(update_fields=['need_candidate_info'])
                        updated_candidates.append(student.id)
            else:
                print(serializer.errors)
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        print("Data: ", data)
        return Response(data, status=status.HTTP_201_CREATED)


class NonDbTestAssign(APIView):
    def post(self, request):

        print('*****Request.data*****: ', request.data)

        test_name = request.data.get('test_name')
        test_type_id = request.data.get('test_type_id')
        question_type_id = request.data.get('question_type_id')
        skill_type_id = request.data.get('skill_type_id')

        college_id = request.data.get('college_id')
        dtm_upload = request.data.get('dtm_upload')

        if not all([test_name, test_type_id, question_type_id]):
            return Response({'error': 'Missing fields for test_master'}, status=status.HTTP_400_BAD_REQUEST)

        test_master_data = {
            'test_name': test_name,
            'test_type_id': test_type_id,
            'question_type_id': question_type_id,
            'skill_type_id': skill_type_id
        }

        test_master_serializer = testsSerializersAdd(data=test_master_data)
        print('test_master_serializer: ', test_master_serializer)
        
        if not test_master_serializer.is_valid():
            print(f"Test master serializer errors: {test_master_serializer.errors}")
            return Response(test_master_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        


        # Filter the students based on the constructed query
        students = candidate_master.objects.filter(
            is_database = False,
            college_id = college_id,
            dtm_upload = dtm_upload
        )
        print('students: ', students)
        
        data = []
        updated_candidates = []  # To store IDs of updated candidates
        current_date_and_time = datetime.now()

        for student in students:
            test_candidate_data = {
                'test_name': request.data.get('test_name'),
                'question_id': request.data.get('question_id'),
                'student_id': student.id,
                'college_id': request.data.get('college_id'),
                'dtm_start': request.data.get('dtm_start'),
                'dtm_end': request.data.get('dtm_end'),
                'is_camera_on': request.data.get('is_camera_on'),
                'duration': request.data.get('duration'),
                'duration_type': request.data.get('duration_type'),
                'rules_id': request.data.get('rules_id'),
                'need_candidate_info': request.data.get('need_candidate_info'),
                'dtm_created': current_date_and_time
            }

            serializer = NonDbTestAssignSerializer(data=test_candidate_data)
            

            if not serializer.is_valid():
                print(f"serializer errors: {serializer.errors}")
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
            
            with transaction.atomic():  # Use transaction to ensure atomicity
                test_master_serializer.save()
                print('Test master saved Successfully')
                
                serializer.save()
                data.append(serializer.data)
                # Check and update candidate_master if needed
                if student.need_candidate_info is None and request.data.get('need_candidate_info') is True:
                    student.need_candidate_info = request.data.get('need_candidate_info')
                    student.save(update_fields=['need_candidate_info'])
                    updated_candidates.append(student.id)
            
        return Response(data, status=status.HTTP_201_CREATED)


#-----------------------Get Unique Batch---------------------#

@api_view(['GET'])
def batch_list(request):
    batch_numbers = candidate_master.objects.exclude(batch_no__isnull=True).values_list('batch_no', flat=True).distinct()
    return Response({'batch_numbers': list(batch_numbers)})

@api_view(['GET'])
def question_name_list(request):
    question_names = question_master.objects.exclude(question_name__isnull=True).values_list('question_name', flat=True).distinct()
    return Response({'question_new': list(question_names)})

@api_view(['GET'])
def topic_list(request):
    topics = content_master.objects.filter(deleted=0).exclude(topic__isnull=True).values_list('topic', flat=True).distinct()
    return Response({'topics': list(topics)})


@api_view(['GET'])
def unique_test_type(request):
    test_types = test_type.objects.filter(deleted=0).values_list('test_type', flat=True).distinct()
    return Response({'test_type': list(test_types)})


@api_view(['GET'])
def MCQ_test_type(request):
    test_types = test_type.objects.filter(deleted=0, test_type='MCQ Test').values_list('test_type_categories', flat=True).distinct()

    return Response({"test_type": list(test_types)})

@api_view(['GET'])
def Coding_test_type(request):
    test_types = test_type.objects.filter(deleted=0, test_type='Coding Test').values_list('test_type_categories', flat=True).distinct()

    return Response({"test_type": list(test_types)})


@api_view(['GET'])
def unique_question_type(request):
    question_types = question_type.objects.filter(deleted=0).values_list('question_type', flat=True).distinct()

    return Response({"question_types": list(question_types)})




class sidebarMenu(generics.ListAPIView):
    queryset = appinfo.objects.filter(deleted=0, info_code='MAIN_MENU')
    serializer_class = appinfoSerializer


@api_view(['GET'])
def main_menu_Training(request):
    menus = appinfo.objects.filter(deleted=0, info_code='MAIN_MENU').values_list('info_value', flat=True).distinct()

    return Response({"MAIN_MENU": list(menus)})




def fetch_csv_with_retries(url, retries=3, delay=5):
    for attempt in range(retries):
        try:
            response = requests.get(url)
            response.raise_for_status()  # Raise an HTTPError for bad responses
            return response.text
        except requests.exceptions.RequestException as e:
            if attempt < retries - 1:  # If not the last attempt
                time.sleep(delay)
            else:
                raise e  # Re-raise the last exception

def attendance_excel_data(request):
    if request.method == 'GET':
        # Define URLs for each sheet
        sheet_urls = {
            'AIDS': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=257135364',
            'IT': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=747238312',
            'CSE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=1208914504',
            'ECE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=2128791655'
        }
        
        # Initialize an empty dictionary to store data for each sheet
        all_data = {}

        # Iterate over each sheet URL
        for sheet_name, url in sheet_urls.items():
            try:
                # Fetch the CSV data with retries
                csv_data = fetch_csv_with_retries(url)
                # Read the CSV data into a Pandas DataFrame, skipping the first 8 rows
                df = pd.read_csv(io.StringIO(csv_data), skiprows=7)

                # Transform the DataFrame into the desired format
                transformed_data = []
                for _, row in df.iterrows():
                    transformed_row = {
                        'Department': row.get('DePArtment ', ''),  # Handle missing column gracefully
                        'Register Number': row.get('Register Number', ''),  # Handle missing column gracefully
                        'Name of the Student': row.get('NAme of the Student', '')
                    }
                    # Iterate over date columns and add FN and AN values
                    for i in range(4, len(df.columns), 2):  # Start from index 4, increment by 2
                        col = df.columns[i]
                        date = col.split()[0]  # Extract the date
                        fn = row.get(col, '')  # Handle missing column gracefully
                        an_col_index = i + 1  # Get the index of the next column (AN)
                        if an_col_index < len(df.columns):
                            an = row.get(df.columns[an_col_index], '')  # Handle missing column gracefully
                        else:
                            an = ''  # If AN column does not exist, set it to empty string
                        transformed_row[date] = {'FN': fn, 'AN': an}
                    transformed_data.append(transformed_row)

                # Store the transformed data in the dictionary
                all_data[sheet_name] = transformed_data

            except requests.exceptions.RequestException as e:
                # Log the error and add an error message to the response
                all_data[sheet_name] = {'error': str(e)}
                print(f"Error fetching data for sheet {sheet_name}: {e}")

        return JsonResponse(all_data, safe=False)
    else:
        return JsonResponse({'error': 'Method not allowed'}, status=405)


class candidates_one_Update(generics.RetrieveUpdateAPIView):
    queryset = candidate_master.objects.all()
    serializer_class = candidatesoneSerializer


#------------------------------------------Question Paper Master----------------------------------#

@api_view(['GET'])
def get_question_paper_OLD(request):
    questions_data = question_paper_master.objects.filter(deleted=0).order_by(
        Case(
            When(dtm_created__isnull=True, then=1),
            default=0,
            output_field=DateTimeField()
        ),
        '-dtm_created'
    ).values(
        'id', 'question_paper_name', 'topic', 'sub_topic', 'test_type',
        'duration_of_test', 'no_of_questions', 'upload_type', 'dtm_created'
    )

    return Response(list(questions_data))

@api_view(['GET'])
def get_question_paper(request):
    questions_data = question_paper_master.objects.filter(deleted=0).order_by(
        '-dtm_created'
    ).values(
        'id', 'question_paper_name', 'topic', 'sub_topic', 'test_type',
        'duration_of_test', 'no_of_questions', 'upload_type', 'dtm_created'
    )

    # Convert dtm_created to local time zone
    for question in questions_data:
        question['dtm_created'] = timezone.localtime(question['dtm_created'])

    return Response(list(questions_data))


class questionpapercreateAPIView(generics.CreateAPIView):
    queryset = question_paper_master.objects.all()
    serializer_class =questionsPaperSerializer

    def post(self, request, *args, **kwargs):
        print("Request data:", request.data)  
        #logger.info("Creating a new question paper")
        response = super().post(request, *args, **kwargs)
        #logger.info("Created a new question paper successfully")
        return response

class questionpaperUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = question_paper_master.objects.all()
    serializer_class =questionsPaperSerializer

    def put(self, request, *args, **kwargs):
        #logger.info(f"Updating Question Paper with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        #logger.info(f"Updated Question Paper with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        #logger.info(f"Partially updating Question Paper with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        #logger.info(f"Partially updated Question Paper with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_question_paper(request, pk):
    #logger.info(f"Attempting to mark Courses with id {pk} as deleted")
    try:
        print("Entering Function..")
        questions=question_paper_master.objects.get(id=pk)

        print("questions: ",questions)
    except question_paper_master.DoesNotExist:
        return JsonResponse("Question Paper not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    questions.deleted = 1
    questions.save()

    #logger.info(f"Marked Courses with id {pk} as deleted successfully")

    print("questions: ",questions)

    return JsonResponse("questions paper 'deleted' field updated successfully", safe=False)



def get_last_added_question_paper(request):
    try:
        last_added_paper = question_paper_master.objects.order_by('-id').first()
        serializer = questionsPaperSerializer(last_added_paper)
        return JsonResponse(serializer.data)
    except question_paper_master.DoesNotExist:
        return JsonResponse({'message': 'No question papers available'}, status=404)



#---------------------Questions Master with images--------------#



class questionscreateAPIView_IO_OLD(generics.CreateAPIView):
    queryset = question_master.objects.all()
    serializer_class = questionsSerializer_IO

    def post(self, request, *args, **kwargs):
        #logger.info("Creating a new Questions")
        response = super().post(request, *args, **kwargs)
        #logger.info("Created a new Questions successfully")
        return response

class questionsRetrieveUpdateAPIView_IO_OLD(generics.RetrieveUpdateAPIView):
    queryset = question_master.objects.all()
    serializer_class = questionsSerializer_IO


    def update(self, request, *args, **kwargs):
        instance = self.get_object()
        response = super().update(request, *args, **kwargs)
        
        # Update the dtm_created field of the related question_paper_master
        question_paper = instance.question_name_id  # Adjust this based on your ForeignKey field name
        question_paper.dtm_created = datetime.now()
        question_paper.save()

        return response

    def put(self, request, *args, **kwargs):
        # logger.info(f"Updating Questions with id {kwargs.get('pk')}")
        response = self.update(request, *args, **kwargs)
        # logger.info(f"Updated Questions with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        # logger.info(f"Partially updating Questions with id {kwargs.get('pk')}")
        response = self.update(request, *args, **kwargs)
        # logger.info(f"Partially updated Questions with id {kwargs.get('pk')} successfully")
        return response



#---------------------Questions Master with images--------------#


@api_view(['GET'])
def get_questions_IO(request):
    # Define a unique cache key
    cache_key = 'questions_IO_data'

    # Try to retrieve the data from the cache
    question_data = cache.get(cache_key)

    if not question_data:
        # Fetch the questions from the database
        questionset = question_master.objects.filter(deleted=0).select_related('question_name_id').values(
            'id',
            'question_name_id__id',
            'question_name_id__question_paper_name',
            'question_text',
            'question_image_data',
            'option_a_image_data',
            'option_b_image_data',
            'option_c_image_data',
            'option_d_image_data',
            'option_a',
            'option_b',
            'option_c',
            'option_d',
            'mark',
            'explain_answer',
            'answer',
        )

        # Process the data to encode images and shuffle the list
        def process_question(question):
            return {
                'id': question['id'],
                'question_name_id': question['question_name_id__id'],
                'question_paper_name': question['question_name_id__question_paper_name'],
                'question_text': question['question_text'],
                'question_image_data': base64.b64encode(question['question_image_data']).decode('utf-8') if question['question_image_data'] else None,
                'option_a_image_data': base64.b64encode(question['option_a_image_data']).decode('utf-8') if question['option_a_image_data'] else None,
                'option_b_image_data': base64.b64encode(question['option_b_image_data']).decode('utf-8') if question['option_b_image_data'] else None,
                'option_c_image_data': base64.b64encode(question['option_c_image_data']).decode('utf-8') if question['option_c_image_data'] else None,
                'option_d_image_data': base64.b64encode(question['option_d_image_data']).decode('utf-8') if question['option_d_image_data'] else None,
                'option_a': question['option_a'],
                'option_b': question['option_b'],
                'option_c': question['option_c'],
                'option_d': question['option_d'],
                'mark': question['mark'],
                'explain_answer': question['explain_answer'],
                'answer': question['answer'],
            }

        # Map the processing function over the queryset results and shuffle
        question_data = list(map(process_question, questionset))
        random.shuffle(question_data)

        # Cache the data with a timeout of 1 hour (3600 seconds)
        cache.set(cache_key, question_data, timeout=3600)

    # Return the cached or freshly generated data
    return Response(question_data)




@csrf_exempt
def upload_question(request):
    print('***1')
    if request.method == 'POST':
        print('****2')
        print('POST data:', request.POST)
        print('FILES data:', request.FILES)

        form = QuestionForm(request.POST, request.FILES)
        print('****3')
        if form.is_valid():
            print('*****4')
            question = form.save(commit=False)
            if 'question_image_data' in request.FILES:
                question.question_image_data = request.FILES['question_image_data'].read()
            if 'option_a_image_data' in request.FILES:
                question.option_a_image_data = request.FILES['option_a_image_data'].read()
            if 'option_b_image_data' in request.FILES:
                question.option_b_image_data = request.FILES['option_b_image_data'].read()
            if 'option_c_image_data' in request.FILES:
                question.option_c_image_data = request.FILES['option_c_image_data'].read()
            if 'option_d_image_data' in request.FILES:
                question.option_d_image_data = request.FILES['option_d_image_data'].read()
            question.save()
            print('Question Created Successfully')
            return HttpResponse("Question created successfully")
    else:
        print('***1')
        form = QuestionForm()
        print('******2')
    return render(request, 'upload_image.html', {'form': form})



def update_question_OLD(request, question_id):
    question = get_object_or_404(question_master, id=question_id, deleted=0)

    if request.method == 'POST':
        form = QuestionForm(request.POST, request.FILES, instance=question)
        if form.is_valid():
            question = form.save(commit=False)
            if 'question_image_data' in request.FILES:
                question.question_image_data = request.FILES['question_image_data'].read()
            if 'option_a_image_data' in request.FILES:
                question.option_a_image_data = request.FILES['option_a_image_data'].read()
            if 'option_b_image_data' in request.FILES:
                question.option_b_image_data = request.FILES['option_b_image_data'].read()
            if 'option_c_image_data' in request.FILES:
                question.option_c_image_data = request.FILES['option_c_image_data'].read()
            if 'option_d_image_data' in request.FILES:
                question.option_d_image_data = request.FILES['option_d_image_data'].read()
            else:
                question.question_image_data = None
            question.save()
            return HttpResponse("Question updated successfully")
    else:
        form = QuestionForm(instance=question)

    return render(request, 'update_image.html', {'form': form})


@csrf_exempt
def update_question(request, question_id):
    # Retrieve the question object, or return a 404 error if not found
    question = get_object_or_404(question_master, id=question_id)

    if request.method == 'POST':
        # Bind the form to the POST data and files
        form = QuestionFormMCQ(request.POST, request.FILES, instance=question)
        if form.is_valid():
            question = form.save(commit=False)
            
            # Handle file uploads for the image data fields
            if 'question_image_data' in request.FILES:
                question.question_image_data = request.FILES['question_image_data'].read()
            if 'option_a_image_data' in request.FILES:
                question.option_a_image_data = request.FILES['option_a_image_data'].read()
            if 'option_b_image_data' in request.FILES:
                question.option_b_image_data = request.FILES['option_b_image_data'].read()
            if 'option_c_image_data' in request.FILES:
                question.option_c_image_data = request.FILES['option_c_image_data'].read()
            if 'option_d_image_data' in request.FILES:
                question.option_d_image_data = request.FILES['option_d_image_data'].read()

            # Save the updated question object
            question.save()

            # Retrieve the related question_paper_master instance
            question_paper = question.question_name_id
            print('question_paper: ', question_paper)
            if question_paper:

                # Update the dtm_created field to the current date and time
                question_paper.dtm_created = datetime.now()
                print('question_paper.dtm_created: ', question_paper.dtm_created)
                question_paper.save()
                print('Question Paper DTM Updated Succesfully')

            # Return a success response
            return HttpResponse("Question updated successfully")
    else:
        # Display the form with the current question data
        form = QuestionFormMCQ(instance=question)

    return render(request, 'update_question.html', {'form': form})





def delete_question_IO(request, pk):
    #logger.info(f"Attempting to mark Questions with id {pk} as deleted")
    try:
        question = question_master.objects.get(id=pk)
    except question_master.DoesNotExist:
        return JsonResponse("Question not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    question.deleted = 1
    question.save()

    #logger.info(f"Marked Questions with id {pk} as deleted successfully")

    return JsonResponse("Question 'deleted' field updated successfully", safe=False)

def get_last_added_question_paper(request):
    try:
        last_added_paper = question_paper_master.objects.order_by('-id').first()
        serializer = questionsPaperSerializer(last_added_paper)
        return JsonResponse(serializer.data)
    except question_paper_master.DoesNotExist:
        return JsonResponse({'message': 'No question papers available'}, status=404)
    #QUESTION_Code_master_#

def get_questions_Code(request):
    try:
        # Fetch data with related fields using select_related and values
        questionset = question_master.objects.filter(deleted=0).select_related('question_name_id').annotate(
            question_name_id_value=Case(
                When(question_name_id__isnull=False, then='question_name_id__id'),
                default=Value(None),
            ),
            question_paper_name_value=Case(
                When(question_name_id__isnull=False, then='question_name_id__question_paper_name'),
                default=Value(None),
            ),
            question_image_data_value=Case(
                When(question_image_data__isnull=False, then=Cast('question_image_data', output_field=Value(base64.b64encode))),
                default=Value(None),
            ),
        ).values(
            'id',
            'question_name_id_value',
            'question_paper_name_value',
            'question_text',
            'view_hint',
            'mark',
            'explain_answer',
            'answer',
            'negative_mark',
            'input_format',
            'question_image_data_value',
        )

        # Process the data to ensure the image is properly encoded in base64
        question_data = [
            {
                'id': question['id'],
                'question_name_id': question['question_name_id_value'],
                'question_paper_name': question['question_paper_name_value'],
                'question_text': question['question_text'],
                'question_image_data': base64.b64encode(question['question_image_data_value']).decode('utf-8') if question['question_image_data_value'] else None,
                'view_hint': question['view_hint'],
                'mark': question['mark'],
                'explain_answer': question['explain_answer'],
                'answer': question['answer'],
                'negative_mark': question['negative_mark'],
                'input_format': question['input_format'],
            }
            for question in questionset
        ]

        random.shuffle(question_data)

        return JsonResponse(question_data, safe=False)
    except Exception as e:
        logger.error(f'Error occurred in get_questions_Code function: {e}')
        return JsonResponse({'error': 'An error occurred'}, status=500)
def upload_question_code(request):
    if request.method == 'POST':
        form = QuestionCodeForm(request.POST, request.FILES)
        if form.is_valid():
            question = form.save(commit=False)
            if 'question_image_data' in request.FILES:
                question.question_image_data = request.FILES['question_image_data'].read()
            else:
                question.question_image_data = None
                question.save()
            return HttpResponse("Question created successfully")
    else:
        form = QuestionCodeForm()
    return render(request, 'upload_image.html', {'form': form})

def update_question_code(request, question_id):
    question = get_object_or_404(question_master, id=question_id, deleted=0)

    if request.method == 'POST':
        form = QuestionCodeForm(request.POST, request.FILES, instance=question)
        if form.is_valid():
            question = form.save(commit=False)
            if 'question_image_data' in request.FILES:
                question.question_image_data = request.FILES['question_image_data'].read()
            question.save()
            return HttpResponse("Question updated successfully")
    else:
        form = QuestionCodeForm(instance=question)

    return render(request, 'update_image.html', {'form': form})



#-----------------------Getting All questions where question_paper_id--------------------------#


@api_view(['GET'])
def get_questions_Qp_id(request, question_name_id):
    # Fetch the required data using select_related and values
    questionset = question_master.objects.filter(deleted=0, question_name_id=question_name_id).select_related('question_name_id').values(
        'id',
        'question_name_id__id',
        'question_name_id__question_paper_name',
        'question_text',
        'question_image_data',
        'option_a_image_data',
        'option_b_image_data',
        'option_c_image_data',
        'option_d_image_data',
        'option_a',
        'option_b',
        'option_c',
        'option_d',
        'view_hint',
        'mark',
        'explain_answer',
        'answer',
        'negative_mark',
        'input_format'
    )

    question_data = []
    for question in questionset:
        question_image_data = base64.b64encode(question['question_image_data']).decode('utf-8') if question['question_image_data'] else None
        option_a_image_data = base64.b64encode(question['option_a_image_data']).decode('utf-8') if question['option_a_image_data'] else None
        option_b_image_data = base64.b64encode(question['option_b_image_data']).decode('utf-8') if question['option_b_image_data'] else None
        option_c_image_data = base64.b64encode(question['option_c_image_data']).decode('utf-8') if question['option_c_image_data'] else None
        option_d_image_data = base64.b64encode(question['option_d_image_data']).decode('utf-8') if question['option_d_image_data'] else None

        question_data.append({
            'id': question['id'],
            'question_name_id': question['question_name_id__id'],
            'question_paper_name': question['question_name_id__question_paper_name'],
            'question_text': question['question_text'],
            'question_image_data': question_image_data,
            'option_a_image_data': option_a_image_data,
            'option_b_image_data': option_b_image_data,
            'option_c_image_data': option_c_image_data,
            'option_d_image_data': option_d_image_data,
            'option_a': question['option_a'],
            'option_b': question['option_b'],
            'option_c': question['option_c'],
            'option_d': question['option_d'],
            'view_hint': question['view_hint'],
            'mark': question['mark'],
            'explain_answer': question['explain_answer'],
            'answer': question['answer'],
            'negative_mark': question['negative_mark'],
            'input_format': question['input_format'],
        })

    return JsonResponse(question_data, safe=False)

class questionsRetrieveUpdateAPIView_IO_code(generics.RetrieveUpdateAPIView):
    queryset = question_master.objects.all()
    serializer_class = questionsSerializer_code


    def update(self, request, *args, **kwargs):
        instance = self.get_object()
        response = super().update(request, *args, **kwargs)
        
        # Update the dtm_created field of the related question_paper_master
        question_paper = instance.question_name_id  # Adjust this based on your ForeignKey field name
        print('question_paper: ', question_paper)
        question_paper.dtm_created = datetime.now()
        question_paper.save()
        print('question_paper_saved....')

        return response

    def put(self, request, *args, **kwargs):
        # logger.info(f"Updating Questions with id {kwargs.get('pk')}")
        response = self.update(request, *args, **kwargs)
        # logger.info(f"Updated Questions with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        # logger.info(f"Partially updating Questions with id {kwargs.get('pk')}")
        response = self.update(request, *args, **kwargs)
        # logger.info(f"Partially updated Questions with id {kwargs.get('pk')} successfully")
        return response



class questionscreateAPIView_IO_code(generics.CreateAPIView):
    queryset = question_master.objects.all()
    serializer_class = questionsSerializer_code

    def post(self, request, *args, **kwargs):
        #logger.info("Creating a new Questions")
        response = super().post(request, *args, **kwargs)
        #logger.info("Created a new Questions successfully")
        return response




@api_view(['PUT'])
def update_Need_Candidate_info(request, studentID=None):
    n_info_list = tests_candidates_map.objects.filter(student_id=studentID)
    print('request.data: ', request.data)

    if not n_info_list.exists():
        return JsonResponse("No candidate info found for the given student ID", safe=False)

    for n_info in n_info_list:
        serializer = testcandidatemapSerializers(instance=n_info, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            print('serializer: ', serializer)
        else:
            return JsonResponse("Failed to Update Need Candidate info, clg, dept, year", safe=False)

    print('Need Candidate info, clg, dept, year updated')
    return JsonResponse("Need Candidate info, clg, dept, year Updated Successfully", safe=False)


@api_view(['PUT'])
def update_clg_login(request, userName=None):
    try:
        # Fetch the user profile instance
        profile = user_profile.objects.get(user__username=userName)
    except user_profile.DoesNotExist:
        return JsonResponse({"error": "User profile not found."}, status=status.HTTP_404_NOT_FOUND)

    print('request.data: ', request.data)

    # Initialize the serializer with instance and request data
    serializer = UserProfileSerializer(instance=profile, data=request.data, partial=True)

    if serializer.is_valid():
        serializer.save()
        print('serializer: ', serializer)
        print('College updated')
        return JsonResponse("College Updated Successfully", safe=False)
    return JsonResponse("Failed to Update College. ")


@api_view(['GET'])
def get_candidate_login(request):
    # Subquery to fetch the password from the user_profile table based on the username
    user_profile_subquery = user_profile.objects.filter(
        user__username=OuterRef('user_name')
    ).values('user__password')[:1]

    # Fetch the candidates and annotate the password from the user_profile table
    candidatelist = candidate_master.objects.filter(deleted=0).annotate(
        password=Subquery(user_profile_subquery)
    ).values(
        'id', 'user_name', 'students_name', 'password'
    )

    candidate_data = [
        {
            'student_id': candidate['id'],
            'user_name': candidate['user_name'],
            'student_name': candidate['students_name'],
            'password': candidate['password']
        }
        for candidate in candidatelist
    ]

    return Response(candidate_data)

class TestAssignFor_Selected(APIView):
    def post(self, request):
        stu_id = request.data.get('stu_id', [])  # Fetching stu_id from request data
        if not isinstance(stu_id, list):
            return Response({"error": "stu_id must be a list"}, status=status.HTTP_400_BAD_REQUEST)

        data = []
        for s_id in stu_id:
            try:
                student = candidate_master.objects.get(id=s_id)
            except candidate_master.DoesNotExist:
                return Response({"error": f"Student with id {s_id} does not exist"}, status=status.HTTP_400_BAD_REQUEST)

            print("Students: ", student)
            print("Student id: ", student.id)
            
            # Preparing test candidate data
            test_candidate_data = {
                'test_name': request.data.get('test_name'),  
                'question_id': request.data.get('question_id'), 
                'college_id': student.college_id.id if student.college_id else None,
                'department_id': student.department_id.id if student.department_id else None,
                'year': student.year if student.year else None, 
                'student_id': student.id,
                'dtm_start': request.data.get('dtm_start'),
                'dtm_end': request.data.get('dtm_end'),
                'is_camera_on': request.data.get('is_camera_on'),
                'duration': request.data.get('duration'),
                'duration_type': request.data.get('duration_type'),
                'rules_id': request.data.get('rules_id'),
                'need_candidate_info': student.college_id is None or student.department_id is None or student.year is None,
                'dtm_created': request.data.get('dtm_created')
            }

            serializer = testcandidatemapSerializers(data=test_candidate_data)
            print("Serializer: ", serializer)
            if serializer.is_valid():
                serializer.save()
                data.append(serializer.data)
                print("Data.append: ", data)
            else:
                print(serializer.errors)
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        print("Data: ", data)
        return Response(data, status=status.HTTP_201_CREATED)



@api_view(['GET'])
def get_group_test_name(request):
    # Join tests_candidates_map with test_master and filter by test_name and deleted=0
    tests_candidates = (
        tests_candidates_map.objects.filter(deleted=0)
        .values('test_name', 'question_id', 'question_id__question_paper_name', 'dtm_start', 'dtm_end', 'dtm_created')
        .annotate(student_count=Count('student_id'))
        .order_by(
            Case(
                When(dtm_created__isnull=True, then=1),
                default=0,
                output_field=DateTimeField()
            ),
            '-dtm_created'
        )
    )

    test_candidate_map_data = []
    for testing in tests_candidates:
        # Format dates
        dtm_start_formatted = django_format_date(localtime(testing['dtm_start']), 'd-m-Y h:i A')
        dtm_end_formatted = django_format_date(localtime(testing['dtm_end']), 'd-m-Y h:i A')

        test_candidate_map_data.append({
            'test_name': testing['test_name'],
            'question_id': testing['question_id'],
            'question_paper_name': testing['question_id__question_paper_name'],
            'dtm_start': dtm_start_formatted,
            'dtm_end': dtm_end_formatted,
            'student_count': testing['student_count'],
            'dtm_created': testing['dtm_created']
        })

    return Response(test_candidate_map_data)






#_____________________________________________Company_master______________________________________________


class company_listView(generics.ListAPIView):
    queryset = company_master.objects.filter(deleted=0)
    serializer_class = companySerializer

    def get(self, request, *args, **kwargs):
        #logger.info("Fetching test types where deleted=0")
        return super().get(request, *args, **kwargs)

class company_create(generics.CreateAPIView):
    queryset = company_master.objects.all()
    serializer_class = companySerializer

    def post(self, request, *args, **kwargs):
        #logger.info("Creating a new test type")
        return super().post(request, *args, **kwargs)
class company_master_delete(generics.RetrieveDestroyAPIView):
    queryset = company_master.objects.all()
    serializer_class = companySerializer

   


class company_master_Update(generics.RetrieveUpdateAPIView):
    queryset = company_master.objects.all()
    serializer_class = companySerializer

    def put(self, request, *args, **kwargs):
        #logger.info(f"Updating test type with id {kwargs.get('pk')}")
        return super().put(request, *args, **kwargs)

    def patch(self, request, *args, **kwargs):
        #logger.info(f"Partially updating test type with id {kwargs.get('pk')}")
        return super().patch(request, *args, **kwargs)

@api_view(['PUT', 'PATCH'])
def delete_company_master(request, pk):
    try:
        #logger.info(f"Attempting to mark test type with id {pk} as deleted")
        company_mastervar = company_master.objects.get(id=pk)
    except company_master.DoesNotExist:
        logger.error(f"Test type with id {pk} not found")
        return JsonResponse("company_master not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    company_mastervar.deleted = 1
    company_mastervar.save()

    #logger.info(f"Marked test type with id {pk} as deleted")
    return JsonResponse("company_master 'deleted' field updated successfully", safe=False)

#______________________________________________________job_master______________________________________________#


@api_view(['GET'])
def get_job(request):
    job_data = job_master.objects.filter(deleted=0).values(
        'id',
        'company_name',
        'company_profile',
        'college_id__college',  # Access related object's field
        'department_id__department',  # Access related object's field
        'skill_id__skill_name',  # Access related object's field
        'intern_fulltime',
        'on_off_campus',
        'marks_10th',
        'marks_12th',
        'cgpa',
        'gender',
        'year',
        'interview_date',
        'history_of_arrears',
        'standing_arrears',
        'location',
    )

    # Convert the QuerySet to a list of dictionaries
    job_list = list(job_data)

    return Response(job_list)



class jobcreateAPIView(generics.CreateAPIView):
    queryset = job_master.objects.all()
    serializer_class = jobSerializer

    def perform_create(self, serializer):
        job_instance = serializer.save()
        self.create_eligible_students(job_instance)

    def create_eligible_students(self, job_instance):
        try:
            candidates = self.filter_candidates(job_instance)
            eligible_students = [
                {
                    'students_id': candidate.id,
                    'announcement': 'Job Announcement Text',
                    'job_id': job_instance.id,
                    'round_of_interview': 'Interview Date',
                    'whatsapp_text': 'Sample WhatsApp Text'
                }
                for candidate in candidates
            ]
            # Bulk create eligible students
            eligible_serializer = eligible_studentSerializer(data=eligible_students, many=True)
            if eligible_serializer.is_valid():
                eligible_serializer.save()
            else:
                raise ValueError("Eligible student data is not valid")
        except Exception as e:
            print('Exception:', str(e))

    def filter_candidates(self, job):
        # Build Q objects for dynamic filtering
        filters = Q(deleted=0)

        # Use dictionary to dynamically add filters
        criteria = {
            'college_id': job.college_id,
            'department_id': job.department_id,
            'year':job.year,
            'marks_10th__gte': job.marks_10th,
            'marks_12th__gte': job.marks_12th,
            'cgpa__gte': job.cgpa,
            'history_of_arrears__lte': job.history_of_arrears,
            'standing_arrears__lte': job.standing_arrears
        }

        # Update filters dynamically
        for field, value in criteria.items():
            if value is not None:  # Ensure value is not None
                filters &= Q(**{field: value})

        # Handle gender separately to allow both genders if not specified
        if job.gender is not None:
            filters &= Q(gender=job.gender)

        # Filter candidates based on dynamic filters
        candidates = candidate_master.objects.filter(filters)

        # Skill filtering can be added if needed
        # if job.skill_id.exists():
        #     skill_ids = job.skill_id.values_list('id', flat=True)
        #     candidates = candidates.filter(skill_id__in=skill_ids)

        return candidates

class jobUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = job_master.objects.all()
    serializer_class =jobSerializer

    def put(self, request, *args, **kwargs):
        #logger.info(f"Updating job with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        #logger.info(f"Updated job with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        #logger.info(f"Partially updating job with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        #logger.info(f"Partially updated job with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_job(request, pk):
    #logger.info(f"Attempting to mark job with id {pk} as deleted")
    try:
        print("Entering Function..")
        jobs=job_master.objects.get(id=pk)

        print("job: ",jobs)
    except job_master.DoesNotExist:
        return JsonResponse("job not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    jobs.deleted = 1
    jobs.save()

    #logger.info(f"Marked job with id {pk} as deleted successfully")

    print("course: ",jobs)

    return JsonResponse("course 'deleted' field updated successfully", safe=False)


#______________________________________________________Event_master______________________________________________#


@api_view(['GET'])
def get_event(request):
    try:
        # Fetch all events with related data in one query
        event_data = event_master.objects.filter(deleted=0).select_related(
            'department_id', 'college_id'
        ).values(
            'id',
            'event_name',
            'event_desc',
            'dtm_start',
            department_name=F('department_id__department'),
            college_name=F('college_id__college')
        )
        
        # Convert the queryset into a list of dictionaries
        event_list = list(event_data)

        return Response(event_list)
    except Exception as e:
        return Response({'error': str(e)}, status=500)



class eventcreateAPIView(generics.CreateAPIView):
    queryset = event_master.objects.all()
    serializer_class =eventSerializer

    def post(self, request, *args, **kwargs):
        #logger.info("Creating a new event")
        response = super().post(request, *args, **kwargs)
        #logger.info("Created a new event successfully")
        return response

class eventUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = event_master.objects.all()
    serializer_class =eventSerializer

    def put(self, request, *args, **kwargs):
        #logger.info(f"Updating events with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        #logger.info(f"Updated events with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        #logger.info(f"Partially updating events with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        #logger.info(f"Partially updated events with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_events(request, pk):
    #logger.info(f"Attempting to mark events with id {pk} as deleted")
    try:
        print("Entering Function..")
        eventss=event_master.objects.get(id=pk)

        print("events: ",eventss)
    except event_master.DoesNotExist:
        return JsonResponse("events not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    eventss.deleted = 1
    eventss.save()

    #logger.info(f"Marked event with id {pk} as deleted successfully")

    print("course: ",eventss)

    return JsonResponse("event 'deleted' field updated successfully", safe=False)

#----------------------------Training admin Dashboard datas-------------------#

# Get Total Test Count

def get_distinct_test_name_count(request, college_id):
    # Perform the query to count distinct test names for the given college_id
    distinct_test_name_count = tests_candidates_map.objects.filter(college_id=college_id).values('test_name').distinct().count()
    
    # Return the count as a JSON response
    return JsonResponse({'distinct_test_name_count': distinct_test_name_count})



# Avg Score of Aptitude

@api_view(['GET'])
def get_avg_score_by_department(request, college_id, dtm_start):
    try:
        print('college_id: ', college_id)
        print('dtm_start: ', dtm_start)
        # Convert dtm_start to a datetime object
        dtm_start_date = timezone.make_aware(timezone.datetime.strptime(dtm_start, '%Y-%m-%d'))
        print('dtm_start: ', dtm_start_date)

        # Get all departments
        departments = department_master.objects.filter(deleted=0).values('id', 'department')

        # Calculate the average score for each department
        results = []
        for department in departments:
            avg_score = tests_candidates_map.objects.filter(
                college_id=college_id,
                dtm_start__date=dtm_start_date,
                question_id__test_type='MCQ Test',
                department_id=department['id']
            ).aggregate(avg_score=Avg('total_score'))['avg_score'] or 0

            results.append({
                'department_name': department['department'],
                'avg_score': avg_score
            })

        # Return the results as a JSON response
        return JsonResponse(results, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


# Avg Score of Coding
@api_view(['GET'])
def avg_score_by_department_Coding(request, college_id, dtm_start):
    try:
        print('college_id: ', college_id)
        print('dtm_start: ', dtm_start)
        # Convert dtm_start to a datetime object
        dtm_start_date = timezone.make_aware(timezone.datetime.strptime(dtm_start, '%Y-%m-%d'))
        print('dtm_start: ', dtm_start_date)

        # Get all departments
        departments = department_master.objects.filter(deleted=0).values('id', 'department')

        # Calculate the average score for each department
        results = []
        for department in departments:
            avg_score = tests_candidates_map.objects.filter(
                college_id=college_id,
                dtm_start__date=dtm_start_date,
                question_id__test_type='Coding Test',
                department_id=department['id']
            ).aggregate(avg_score=Avg('total_score'))['avg_score'] or 0

            results.append({
                'department_name': department['department'],
                'avg_score': avg_score
            })

        # Return the results as a JSON response
        return JsonResponse(results, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


#  College Topper


def get_max_score_by_department(request, college_id):
    try:
        # First, create a subquery to get the maximum total score per department
        subquery = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='MCQ Test',
            total_score__isnull=False,
            department_id=OuterRef('department_id')
        ).values(
            'department_id'
        ).annotate(
            max_total_score=Max('total_score')
        ).values('max_total_score')

        # Then, filter the original query using this subquery to get the corresponding student names
        results = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='MCQ Test',
            total_score__isnull=False,
            total_score=Subquery(subquery)  # Match the max score in each department
        ).values(
            'student_id__students_name',
            'department_id__department',
            'total_score'
        ).distinct()  # Ensure unique records in case of ties

        # Format the results as a list of dictionaries
        data = [
            {
                'student_name': result['student_id__students_name'],
                'department': result['department_id__department'],
                'max_total_score': result['total_score']
            }
            for result in results
        ]

        # Return the results as a JSON response
        return JsonResponse(data, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)



def get_max_score_by_department_coding(request, college_id):
    try:
        # First, create a subquery to get the maximum total score per department
        subquery = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='Coding Test',
            total_score__isnull=False,
            department_id=OuterRef('department_id')
        ).values(
            'department_id'
        ).annotate(
            max_total_score=Max('total_score')
        ).values('max_total_score')

        # Then, filter the original query using this subquery to get the corresponding student names
        results = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='Coding Test',
            total_score__isnull=False,
            total_score=Subquery(subquery)  # Match the max score in each department
        ).values(
            'student_id__students_name',
            'department_id__department',
            'total_score'
        ).distinct()  # Ensure unique records in case of ties

        # Format the results as a list of dictionaries
        data = [
            {
                'student_name': result['student_id__students_name'],
                'department': result['department_id__department'],
                'max_total_score': result['total_score']
            }
            for result in results
        ]

        # Return the results as a JSON response
        return JsonResponse(data, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


# Count of company


def count_company_names(request):
    try:
        # Perform the query to count company_name entries
        result = company_master.objects.aggregate(
            count_company_name=Count('company_name')
        )

        # Return the count as a JSON response
        return JsonResponse(result, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


#------------------------------Placement admin Dashboard datas-----------------------#


def get_distinct_test_name_count_today(request):
    try:
        # Get today's date
        today_date = date.today()

        # Perform the query to count distinct test_name where dtm_start is today's date
        distinct_test_name_count = tests_candidates_map.objects.filter(
            dtm_start__date=today_date
        ).values('test_name').distinct().count()

        # Return the count as JSON response
        return JsonResponse({'distinct_test_name_count': distinct_test_name_count}, status=200)
    
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)







# URLs for the CSV files
sheet_urls = {
    'AIDS': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=257135364',
    'IT': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=747238312',
    'CSE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=1208914504',
    'ECE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=2128791655'
}

# Function to process each sheet
def process_sheet(url):
    response = requests.get(url)
    csv_data = response.content.decode('utf-8')
    df = pd.read_csv(StringIO(csv_data), skiprows=7)  # Skip first 7 rows to reach the headers
    df = df[2:]  # Skip the next row (8th and 9th rows)
    df.reset_index(drop=True, inplace=True)  # Reset index after slicing
    df.columns = df.columns.str.strip()  # Strip any leading/trailing spaces from columns
    return df

@api_view(['GET'])
def get_avg_total_present(request):
    try:
        department_data = {}

        for department, url in sheet_urls.items():
            df = process_sheet(url)
            print(f"Headers for {department}: {df.columns}")  # Print headers for debugging

            # Check if 'Total Present' column exists
            if 'Total Present' in df.columns:
                # Calculate the average of 'Total Present'
                avg_total_present = df['Total Present'].astype(float).mean()  # Ensure conversion to float
                department_data[department] = round(avg_total_present, 1)
            else:
                department_data[department] = 'Total Present column not found'

        return Response(department_data)
    except Exception as e:
        return Response({'error': str(e)}, status=400)



@api_view(['GET'])
def get_avg_total_absent(request):
    try:
        department_data = {}

        for department, url in sheet_urls.items():
            df = process_sheet(url)
            print(f"Headers for {department}: {df.columns}")  # Print headers for debugging

            # Check if 'Total Present' column exists
            if 'Total Present' in df.columns:
                # Calculate the average of 'Total Present'
                avg_total_absent = df['Total Abscent'].astype(float).mean()  # Ensure conversion to float
                department_data[department] = round(avg_total_absent, 1)
            else:
                department_data[department] = 'Total absent column not found'

        return Response(department_data)
    except Exception as e:
        return Response({'error': str(e)}, status=400)







#---------------------------------------Students Dashboards-----------------------------------#



def get_events_by_college_and_department(request, college_id, department_id):
    try:
        # Filter the events based on college_id and department_id
        events = event_master.objects.filter(
            college_id=college_id,
            department_id=department_id
        ).values('event_name', 'event_desc', 'dtm_start')

        # Convert the queryset to a list of dictionaries
        data = list(events)

        # Return the results as a JSON response
        return JsonResponse(data, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)



#__________________________________IMport MCQ   Questions____________________________#


def extract_text_and_images_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = []
    images_binary = {}
    image_count = 0

    for para in doc.paragraphs:
        text.append(para.text)

    for rel in doc.part.rels:
        target = doc.part.rels[rel].target_ref
        if "image" in target:
            image_count = str(target).split('.')[0][-1]
            image = doc.part.rels[rel].target_part.blob
            images_binary[f"image_{image_count}"] = base64.b64encode(image).decode('utf-8')

    return '\n'.join(text), images_binary

def remove_serial_number(question_text):
    return re.sub(r'^\d+\s?\)\s*', '', question_text)

def remove_option_choice(option_text):
    return re.sub(r'^[a-d]\s*?\)\s*', '', option_text)


def create_json_structure(text, images_binary):
    print("Entering create_json_structure")
    print("Text received:", text[:100], "...")
    print("Images received:", images_binary)
    
    # Function logic
    questions = []
    lines = text.split('\n')
    current_question = {}
    options = {}
    image_keys = list(images_binary.keys())
    image_index = 1
    unmatched_lines = []

    # Regex patterns
    question_pattern = re.compile(r'^\s*\d+\s?\)')
    option_a_pattern = re.compile(r'^\s*a\s?\)')
    option_b_pattern = re.compile(r'^\s*b\s?\)')
    option_c_pattern = re.compile(r'^\s*c\s?\)')
    option_d_pattern = re.compile(r'^\s*d\s?\)')
    answer_pattern = re.compile(r'^\s*c_ans:')
    explanation_pattern = re.compile(r'^\s*e_ans:')
    mark_pattern = re.compile(r'^\s*mark:')
    negative_mark_pattern = re.compile(r'^\s*n_m:')

    for line in lines:
        line = line.strip()
        print("Processing line:", line)
        if question_pattern.match(line):
            if current_question:
                current_question['options'] = options
                questions.append(current_question)
                current_question = {}
                options = {}
            question_text = remove_serial_number(line)
            if question_text == '':
                if image_index <= len(image_keys):
                    current_question['question_image_data'] = images_binary[image_keys[image_index - 1]]
                    image_index += 1
                else:
                    unmatched_lines.append((line, "Image not found for question"))
            else:
                current_question['question_text'] = question_text
        elif option_a_pattern.match(line):
            option_a_text = remove_option_choice(line)
            if option_a_text == '':
                if image_index <= len(image_keys):
                    options['option_a_image_data'] = images_binary[image_keys[image_index - 1]]
                    image_index += 1
                else:
                    unmatched_lines.append((line, "Image not found for option A"))
            else:
                options['option_a'] = option_a_text
        elif option_b_pattern.match(line):
            option_b_text = remove_option_choice(line)
            if option_b_text == '':
                if image_index <= len(image_keys):
                    options['option_b_image_data'] = images_binary[image_keys[image_index - 1]]
                    image_index += 1
                else:
                    unmatched_lines.append((line, "Image not found for option B"))
            else:
                options['option_b'] = option_b_text
        elif option_c_pattern.match(line):
            option_c_text = remove_option_choice(line)
            if option_c_text == '':
                if image_index <= len(image_keys):
                    options['option_c_image_data'] = images_binary[image_keys[image_index - 1]]
                    image_index += 1
                else:
                    unmatched_lines.append((line, "Image not found for option C"))
            else:
                options['option_c'] = option_c_text
        elif option_d_pattern.match(line):
            option_d_text = remove_option_choice(line)
            if option_d_text == '':
                if image_index <= len(image_keys):
                    options['option_d_image_data'] = images_binary[image_keys[image_index - 1]]
                    image_index += 1
                else:
                    unmatched_lines.append((line, "Image not found for option D"))
            else:
                options['option_d'] = option_d_text
        elif answer_pattern.match(line):
            current_question['answer'] = line.split(':', 1)[1].strip()
        elif explanation_pattern.match(line):
            current_question['explanation'] = line.split(':', 1)[1].strip()
        elif mark_pattern.match(line):
            current_question['marks'] = int(line.split(':', 1)[1].strip())
        elif negative_mark_pattern.match(line):
            current_question['negative_marks'] = int(line.split(':', 1)[1].strip())
        else:
            if line:
                unmatched_lines.append((line, "Unmatched pattern"))

    if current_question:
        current_question['options'] = options
        current_question['image_based'] = any(isinstance(opt[0], str) and 'data:image' in opt[0] for opt in options.values())
        questions.append(current_question)

    print("Questions parsed:", questions)
    print("Unmatched lines:", unmatched_lines)
    return {"questions": questions, "unmatched_lines": unmatched_lines}

@csrf_exempt
def import_questions_from_word(request):
    if request.method == 'POST':
        form = QuestionImportForm(request.POST, request.FILES)
        if form.is_valid():
            docx_file = request.FILES['docx_file']
            try:
                # Extract text and images from the Word document
                text, images_binary = extract_text_and_images_from_docx(docx_file)

                # Assume `text` is a list where each question is [question_text, some_boolean]
                questions = text
                data = create_json_structure(questions, images_binary)

                # Create a new question_paper_master entry
                question_paper_name = request.POST.get('question_paper_name')
                duration_of_test = request.POST.get('duration_of_test')
                topic = request.POST.get('topic')
                sub_topic = request.POST.get('sub_topic')
                no_of_questions = request.POST.get('no_of_questions')
                upload_type = request.POST.get('upload_type')
                test_type = request.POST.get('test_type')
                current_date_and_time = timezone.now()  # Use timezone-aware datetime

                # Validate required fields for question_paper_master
                if not all([question_paper_name, duration_of_test, topic, sub_topic, no_of_questions, upload_type]):
                    return HttpResponse("Missing fields for question_paper_master", status=400)

                question_paper_data = {
                    'question_paper_name': question_paper_name,
                    'duration_of_test': duration_of_test,
                    'topic': topic,
                    'sub_topic': sub_topic,
                    'no_of_questions': no_of_questions,
                    'upload_type': upload_type,
                    'test_type': test_type,
                    'dtm_created': current_date_and_time
                }

                question_paper_instance = question_paper_master.objects.create(**question_paper_data)

                # Temporary storage of questions
                for question_data in data['questions']:
                    try:
                        question_text = question_data.get('question_text')
                        question_image_data = question_data.get('question_image_data')
                        option_a_image_data = question_data.get('options', {}).get('option_a_image_data')
                        option_b_image_data = question_data.get('options', {}).get('option_b_image_data')
                        option_c_image_data = question_data.get('options', {}).get('option_c_image_data')
                        option_d_image_data = question_data.get('options', {}).get('option_d_image_data')

                        def sanitize_option(option):
                            return option if isinstance(option, str) else ''

                        temp_question = question_master_temp.objects.create(
                            question_name_id=question_paper_instance,
                            question_text=question_text,
                            question_image_data=base64.b64decode(question_image_data) if question_image_data else None,
                            option_a=sanitize_option(question_data.get('options', {}).get('option_a')),
                            option_a_image_data=base64.b64decode(option_a_image_data) if option_a_image_data else None,
                            option_b=sanitize_option(question_data.get('options', {}).get('option_b')),
                            option_b_image_data=base64.b64decode(option_b_image_data) if option_b_image_data else None,
                            option_c=sanitize_option(question_data.get('options', {}).get('option_c')),
                            option_c_image_data=base64.b64decode(option_c_image_data) if option_c_image_data else None,
                            option_d=sanitize_option(question_data.get('options', {}).get('option_d')),
                            option_d_image_data=base64.b64decode(option_d_image_data) if option_d_image_data else None,
                            answer=question_data.get('answer', ''),
                            explain_answer=question_data.get('explanation', ''),
                            mark=question_data.get('marks', '')
                        )
                    except Exception as e:
                        print("Error saving to temporary table:", e)
                        return HttpResponse(f"Error saving to temporary table: {e}")


                # Move from temporary to main table
                for temp_question in question_master_temp.objects.filter(question_name_id=question_paper_instance):
                    try:
                        main_question = question_master.objects.create(
                            question_name_id=temp_question.question_name_id,
                            question_text=temp_question.question_text,
                            question_image_data=temp_question.question_image_data,
                            option_a=temp_question.option_a,
                            option_a_image_data=temp_question.option_a_image_data,
                            option_b=temp_question.option_b,
                            option_b_image_data=temp_question.option_b_image_data,
                            option_c=temp_question.option_c,
                            option_c_image_data=temp_question.option_c_image_data,
                            option_d=temp_question.option_d,
                            option_d_image_data=temp_question.option_d_image_data,
                            answer=temp_question.answer,
                            explain_answer=temp_question.explain_answer,
                            mark=temp_question.mark,
                        )
                        # Optionally, delete the temporary record if it's no longer needed
                        temp_question.delete()
                        print('main_questions: ', main_question)

                    except Exception as e:
                        print("Error moving to main table:", e)
                        return HttpResponse(f"Error moving to main table: {e}")

                return HttpResponse("Questions imported and moved to main table successfully.")
            except Exception as e:
                print("Error processing file:", e)
                return HttpResponse(f"Error processing file: {e}")
        else:
            return HttpResponse("Form is not valid.")
    else:
        form = QuestionImportForm()

    return render(request, 'import_questions.html', {'form': form})


#___________________________________IMPORT CODING QUESTION___________________________________________________#


def extract_text_and_images_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = []
    images_binary = {}
    image_count = 0

    for para in doc.paragraphs:
        text.append(para.text)

    for rel in doc.part.rels:
        target = doc.part.rels[rel].target_ref
        if "image" in target:
            image_count += 1
            image = doc.part.rels[rel].target_part.blob
            images_binary[f"image_{image_count}"] = base64.b64encode(image).decode('utf-8')

    return '\n'.join(text), images_binary

def remove_serial_number(question_text):
    return re.sub(r'^\d+\s?\)\s*', '', question_text)

def create_json_structures(text, images_binary):
    questions = []
    lines = text.split('\n')
    current_question = {}
    image_index = 1
    unmatched_lines = []
    input_format_lines = []

    question_pattern = re.compile(r'^question:\s*')
    answer_pattern = re.compile(r'^c_ans:\s*')
    explanation_pattern = re.compile(r'^e_ans:\s*')
    mark_pattern = re.compile(r'^mark:\s*')
    input_format_pattern = re.compile(r'^input_format:\s*')

    def append_current_question():
        if 'question' in current_question:
            if input_format_lines:
                current_question['input_format'] = ' '.join(input_format_lines).strip()
                input_format_lines.clear()
            questions.append(current_question.copy())
            current_question.clear()

    for line in lines:
        line = line.strip()
        if question_pattern.match(line):
            append_current_question()
            current_question['question'] = line[len('question:'):].strip()
            if f'image_{image_index}' in images_binary:
                current_question['question_image_data'] = images_binary[f'image_{image_index}']
                image_index += 1
            current_question['answers'] = []
        elif answer_pattern.match(line):
            current_question['answers'].append(line[len('c_ans:'):].strip())
        elif explanation_pattern.match(line):
            current_question['explanation'] = line[len('e_ans:'):].strip()
        elif mark_pattern.match(line):
            current_question['marks'] = line[len('mark:'):].strip()
        elif input_format_pattern.match(line):
            input_format_lines.append(line[len('input_format:'):].strip())
        elif input_format_lines:
            input_format_lines.append(line)
        else:
            unmatched_lines.append(line)

    append_current_question()
    return {"questions": questions, "unmatched_lines": unmatched_lines}

@csrf_exempt
def import_questions(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            docfile = request.FILES['docfile']
            try:
                text, images_binary = extract_text_and_images_from_docx(docfile)
                data = create_json_structures(text, images_binary)

                # Create a new question_paper_master entry
                question_paper_name = request.POST.get('question_paper_name')
                duration_of_test = request.POST.get('duration_of_test')
                topic = request.POST.get('topic')
                sub_topic = request.POST.get('sub_topic')
                no_of_questions = request.POST.get('no_of_questions')
                upload_type = request.POST.get('upload_type')
                test_type = request.POST.get('test_type')
                current_date_and_time = timezone.now()

                if not all([question_paper_name, duration_of_test, topic, sub_topic, no_of_questions, upload_type]):
                    return HttpResponse("Missing fields for question_paper_master", status=400)

                question_paper_data = {
                    'question_paper_name': question_paper_name,
                    'duration_of_test': duration_of_test,
                    'topic': topic,
                    'sub_topic': sub_topic,
                    'no_of_questions': no_of_questions,
                    'upload_type': upload_type,
                    'test_type': test_type,
                    'dtm_created': current_date_and_time
                }

                question_paper_instance = question_paper_master.objects.create(**question_paper_data)

                for ques in data["questions"]:
                    if 'question' in ques:
                        # Debugging: Print the data being processed
                        question_text = ques['question']
                        input_format = ques.get('input_format', '')
                        answers = ques.get('answers', [])
                        marks = ques.get('marks', '')
                        explain_answer = ques.get('explanation', '')
                        question_image_data = ques.get('question_image_data', '')

                        print(f"Processing question: {question_text}")
                        print(f"Input format: {input_format}")
                        print(f"Answers: {answers}")
                        print(f"Marks: {marks}")
                        print(f"Explanation: {explain_answer}")
                        print(f"Image data: {question_image_data}")

                        try:
                            # Save to temporary table
                            temp_question = question_master_temp.objects.create(
                                question_name_id=question_paper_instance,
                                question_text=question_text,
                                question_image_data=base64.b64decode(question_image_data) if question_image_data else None,
                                input_format=input_format,
                                answer=', '.join(answers),
                                mark=marks,
                                explain_answer=explain_answer
                            )

                            # Move to main table
                            main_question = question_master.objects.create(
                                question_name_id=temp_question.question_name_id,
                                question_text=temp_question.question_text,
                                question_image_data=temp_question.question_image_data,
                                input_format=temp_question.input_format,
                                answer=temp_question.answer,
                                mark=temp_question.mark,
                                explain_answer=temp_question.explain_answer
                            )

                            # Optionally delete the temporary question
                            temp_question.delete()

                        except Exception as e:
                            return HttpResponse(f"Error saving or moving question: {e}")

                return HttpResponse("Questions imported and moved to main table successfully.")
            except Exception as e:
                return HttpResponse(f"Error processing file: {e}")
        else:
            return HttpResponse("Form is not valid.")
    else:
        form = DocumentForm()
        return render(request, 'import_coding.html', {'form': form})


#_________________________Training_schedule_sheet________________________________________________________#


@api_view(['GET'])
def get_training_report(request):
    try:
        # Fetch all training attendance sheets with related data
        course_data = training_attendance_sheet.objects.filter(deleted=0).select_related(
            'course_schedule_id__college_id',
            'course_schedule_id__department_id',
            'course_schedule_id__trainer_id'
        ).values(
            'id',
            'dtm_attendance',
            'present',
            'absent',
            year=F('course_schedule_id__year'),
            dtm_start=F('course_schedule_id__dtm_start_trainer'),  # Check your actual field name
            dtm_end=F('course_schedule_id__dtm_end_trainer'),      # Check your actual field name
            college_name=F('course_schedule_id__college_id__college'),
            department_name=F('course_schedule_id__department_id__department'),
            trainer_name=F('course_schedule_id__trainer_id__trainer_name'),
            topic=F('course_schedule_id__topic_id__topic'),        # Adjust field name accordingly
            sub_topic=F('course_schedule_id__topic_id__sub_topic') # Adjust field name accordingly
        )
        
        # Convert the queryset into a list of dictionaries
        course_list = list(course_data)

        return Response(course_list)
    except Exception as e:
        return Response({'error': str(e)}, status=500)

class trainingreportcreateAPIView(generics.CreateAPIView):
    queryset = training_attendance_sheet.objects.all()
    serializer_class =trainingreportsheetSerializer

    def post(self, request, *args, **kwargs):
        #logger.info("Creating a new training_report")
        response = super().post(request, *args, **kwargs)
        #logger.info("Created a new training_report successfully")
        return response

class trainingreportUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = training_attendance_sheet.objects.all()
    serializer_class =trainingreportsheetSerializer

    def put(self, request, *args, **kwargs):
        #logger.info(f"Updating training_report with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        #logger.info(f"Updated training_report with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        #logger.info(f"Partially updating training_report with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        #logger.info(f"Partially updated training_report with id {kwargs.get('pk')} successfully")
        return response

@api_view(['PUT', 'PATCH'])
def delete_training_report(request, pk):
    #logger.info(f"Attempting to mark training_report with id {pk} as deleted")
    try:
        print("Entering Function..")
        training_report=training_attendance_sheet.objects.get(id=pk)

        print("trainingreport: ",training_report)
    except training_attendance_sheet.DoesNotExist:
        return JsonResponse("training_report not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    training_report.deleted = 1
    training_report.save()

    #logger.info(f"Marked training_report with id {pk} as deleted successfully")

    print("training_report: ",training_report)

    return JsonResponse("training_report 'deleted' field updated successfully", safe=False)



#-----------------------------Course _ trainer_feedback-----------------------------------------#

class add_trainer_feedback(generics.CreateAPIView):
    queryset = course_trainer_feedback.objects.all()
    serializer_class = trainerfeedbackSerializer


@api_view(['GET'])
def get_trainer_feedback(request):
    # Use the `values` method to fetch only the needed fields
    feedback_data = course_trainer_feedback.objects.select_related(
        'college_id', 'topic_id', 'trainer_id', 'department_id'
    ).filter(deleted=0).values(
        'id',
        'college_id__college',
        'department_id__department',
        'topic_id__topic',
        'trainer_id__trainer_name',
        'dtm_complete',
        'completion_status',
        'feedback'
    )

    return Response(feedback_data)
class update_trainer_feedback(generics.UpdateAPIView):
    queryset = course_trainer_feedback.objects.all()
    serializer_class = trainerfeedbackSerializer



@api_view(['PUT', 'PATCH'])
def delete_trainer_feedback(request, pk):
    try:
        print("Entering Function..")
        course = course_trainer_feedback.objects.get(id=pk)

        print("course: ", course)
    except course_trainer_feedback.DoesNotExist:
        return JsonResponse("course not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    course.deleted = 1
    course.save()
    print("logins: ", course)

    return JsonResponse("course 'deleted' field updated successfully", safe=False)

#------------------------Compiler---------------------#



# List of forbidden file operation functions
FORBIDDEN_FUNCTIONS = [
    r'\bfopen\b', r'\bfclose\b', r'\bfread\b', r'\bfwrite\b', r'\bfscanf\b', 
    r'\bfprintf\b', r'\bfgetc\b', r'\bfputc\b', r'\bfgets\b', r'\bfputs\b', 
    r'\bremove\b', r'\brename\b',r'\bfopen\b',r'\bfread\b',r'\bfwrite\b',r'\bstd::ifstream\b',r'\bstd::ofstream\b',r'\bstd::fstream\b'
]

def contains_forbidden_functions(code):
    """Check if the code contains any forbidden functions."""
    for func in FORBIDDEN_FUNCTIONS:
        if re.search(func, code):
            return True
    return False




def execute_python_code(code,inputs=""):

    try:
        

        addon="""
import builtins
# Define safe import and restricted open functions
def safe_import(name, globals=None, locals=None, fromlist=(), level=0):
    if name not in ['os',"globe","shutill",subprocess]:
        return original_import(name, globals, locals, fromlist, level)
    raise ImportError(f"Importing {name} is not allowed")

def restricted_open(*args, **kwargs):
    raise IOError("File opening is not allowed.")

# Save the original built-in functions
original_import = builtins.__import__
original_open = builtins.open

# Replace built-in functions with restricted versions
builtins.__import__ = safe_import
builtins.open = restricted_open

"""
        code=addon+"\n"+code
        #use subprocess to execute the code at terminal
        result = subprocess.run(
            ['python', '-c', code],
            capture_output=True,
            input=inputs,
            text=True,
            timeout=settings.CODE_EXECUTION_TIMEOUT #the program will raise timeout error if it exceeds settings.CODE_EXECUTION_TIMEOUTs of execution
        )
        return result.stdout or result.stderr
    except subprocess.TimeoutExpired:
        #time error is handled and the message is returned
        return {'error':'Execution timed out'}
    except Exception as e:
        return {'error': str(e)}




def execute_java_code(code,inputs=""):
    asset_path = settings.ASSET_DIR
    # Extract class name
    class_match = re.search(r'public\s+class\s+(\w+)', code)
    if not class_match:
        return 'No public class found'
    
    class_name = class_match.group(1)
    dir_path = os.path.abspath(os.path.join(asset_path,str(uuid.uuid4())))
    os.makedirs(dir_path)
    #file_path = os.path.join(dir_path, f'{class_name}.java')

    #using uuid to create temporary c file for execution
    temp_java_path = os.path.abspath(os.path.join(dir_path, f'{class_name}.java'))
    # print('temp_java_path_created...: ', temp_java_path)
    # logger.error("Temp Java PAth Created...: %s", temp_java_path)

    temp_exec_path = f'{class_name}'
    # print('temp_exec_path_created...: ', temp_exec_path)
    # logger.error("Temp Java Excecute PAth Created...: %s", temp_exec_path)

    #writing the code into temporary file
    with open(temp_java_path, 'w') as f:
        f.write(code)
    try:
        #compiling the file using subprocess at terminal
        compile_result = subprocess.run(
            ['javac', temp_java_path],
            capture_output=True,
            cwd=dir_path,
            text=True
        )

        # print('*******compile_result******')
        # logger.error("Compile Result...: %s", compile_result)
        #if program has syntax error return the error message
        if compile_result.returncode != 0:
            print("syntax error occured")
            return compile_result.stderr
        #return the programs output
        run_result = subprocess.run(
            ['java', '-Djava.security.manager', '-Djava.security.policy==java.policy', temp_exec_path],
            capture_output=True,
            input=inputs,
            text=True,
            cwd=dir_path,
            timeout=settings.CODE_EXECUTION_TIMEOUT
        )
        # print('*******Run Result*********')
        # logger.error("Run Result...: %s", run_result)
        return run_result.stdout or run_result.stderr
    except subprocess.TimeoutExpired:
        return 'Execution timed out'
    except Exception as e:
         return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    finally:
        shutil.rmtree(dir_path)
     

   
def execute_c_code(code,inputs=""):
    file_uuid=uuid.uuid4()
    asset_path = settings.ASSET_DIR
    #using uuid to create temporary c file for execution
    temp_c_path = os.path.abspath(os.path.join(asset_path, f'{file_uuid}.c'))
   # print('temp_path_created...: ', temp_c_path)
    #logger.error("Temp C Path Created...: %s", temp_c_path)

    temp_exec_path = os.path.abspath(os.path.join(asset_path, f'{file_uuid}'))
    # print('temp_exec_path_created...: ', temp_exec_path)
    #logger.error("Temp C Execute Path Created...: %s", temp_exec_path)

    # Check for forbidden file operations
    if contains_forbidden_functions(code):
        return "Error: Code contains forbidden file operations."
    #writing the code into temporary file
    with open(temp_c_path, 'w') as f:
        f.write(code)
    try:
        #compiling the file using subprocess at terminal
        compile_result = subprocess.run(
            ['gcc', temp_c_path, '-o', temp_exec_path],
            capture_output=True,
            #inputs=inputs,
            text=True
        )
        # print('compile_result: ', compile_result)
        #logger.error("Compile Result...: %s", compile_result)

        #if program has syntax error return the error message
        if compile_result.returncode != 0:
            print("syntax error occured")
            return compile_result.stderr
        #return the programs output
        # run_result = subprocess.run(
        #     [temp_exec_path],
        #     capture_output=True,
        #     text=True,
        #     timeout=settings.CODE_EXECUTION_TIMEOUT
        # )
        # Execute the compiled  program with inputs
        process = subprocess.Popen(
                [temp_exec_path],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
        )

            # Communicate with the process: send inputs and read outputs
        stdout, stderr = process.communicate(inputs)
        return stdout or stderr
    except subprocess.TimeoutExpired:
        return 'Execution timed out'
    except Exception as e:
         return f'error:{str(e)}'
    finally:
        if os.path.exists(temp_c_path):
            os.remove(temp_c_path)
        if os.path.exists(f"{temp_exec_path}"):
            os.remove(f"{temp_exec_path}")
        # print('files are deleted....')
        #logger.error("Files Delete in Assets")
        #subprocess.run(['rm', temp_c_path, temp_exec_path])

def execute_cpp_code(code,inputs=""):
    file_uuid=uuid.uuid4()
    asset_path = settings.ASSET_DIR
    file_path = os.path.abspath(os.path.join(asset_path, f'{file_uuid}.cpp'))
    # print('file_path_created...: ', file_path)
    logger.error("C++ File Path Created...: %s", file_path)

    exec_path = os.path.abspath(os.path.join(asset_path, f'{file_uuid}'))
    # print('exec_path_created...: ', exec_path)
    logger.error("C++ File Execute Path Created...: %s", exec_path)

    # Check for forbidden file operations
    if contains_forbidden_functions(code):
        return "Error: Code contains forbidden file operations."
   
    try:
        with open(file_path, 'w') as f:
            f.write(code)
        compile_result = subprocess.run(
            ['g++', file_path, '-o', exec_path],
            capture_output=True,
            text=True
        )
        # print('**********compiled**********')
        logger.error("Compile Result...: %s", compile_result)
        
        if compile_result.returncode != 0:
            return compile_result.stderr
        
        process = subprocess.Popen(
                [exec_path],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True)
        # print('*****run process***********')
        logger.error("Process Result...: %s", process)

            # Communicate with the process: send inputs and read outputs
        stdout, stderr = process.communicate(inputs)
        return stdout or stderr
    except subprocess.TimeoutExpired:
        return 'Execution timed out'
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)
        if os.path.exists(f"{exec_path}"):
            os.remove(f"{exec_path}")

        # os.remove(file_path)
        # os.remove(f"{exec_path}.exe")
        # print('files are deleted....')
        logger.error("Files Delete in Assets")
        # subprocess.run(['rm', file_path, exec_path])






@api_view(['POST'])
def program_compiler(request):

    language_dict={
        "python":execute_python_code,
        "c":execute_c_code,
        "cpp":execute_cpp_code,
        "java":execute_java_code,
    }
    
    code = request.data.get('code')
    p_type = request.data.get('p_type')
    inputs=request.data.get('inputs')
    if p_type not in language_dict.keys():
        return Response({"result":"Unsupported language"})
    else:
        print(f"executing {p_type} code")
        return Response({"result":language_dict[p_type](code,inputs)})




def execute_code(p_type, code, inputs):
    language_dict = {
        "python": execute_python_code,
        "c": execute_c_code,
        "cpp": execute_cpp_code,
        "java": execute_java_code,
    }

    if p_type not in language_dict.keys():
        raise ValueError("Unsupported language")
    
    print('output: ', language_dict[p_type](code,inputs))

    return language_dict[p_type](code,inputs)


@csrf_exempt
@api_view(['POST'])
def test_candidates_answer_view_Com(request):
    print('Request.Data......: ', request.data)

    code = request.data.get('code')
    p_type = request.data.get('p_type')
    inputs = request.data.get('inputs')

    try:
        output = execute_code(p_type, code, inputs)  # This should return the code output
        print("Code Output: ", output)
        # Return the output in the response
        return JsonResponse({'message': 'Successfully updated output', 'output': output}, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)



@csrf_exempt 
@api_view(['POST']) 
def test_candidates_answer_view_Submit_Com(request, format=None): 
    print('Request.Data......: ', request.data)

    stu_id = request.data.get('student_id') 
    ques_id = request.data.get('question_id') 
    code = request.data.get('code') 
    p_type = request.data.get('p_type') 
    inputs = request.data.get('inputs') 
    explain_ans = request.data.get('explain_answer')
    question_mark = request.data.get('mark')
    question_answer = request.data.get('answer')

    try: 
        output = execute_code(p_type, code, inputs)  # Consider using Celery for async processing 
        result = question_mark if question_answer == output else round((fuzz.ratio(explain_ans.strip(), code.strip()) / 100) * question_mark) 
        print('Result: ', result)

        test_candidate_answer_data = { 
            'test_name': request.data.get('test_name'), 
            'question_id': ques_id, 
            'student_id': stu_id, 
            'answer': output, 
            'compile_code_editor': code, 
            'result': result, 
            'dtm_start': request.data.get('dtm_start'), 
            'dtm_end': request.data.get('dtm_end'), 
        } 

        serializer = tests_candidates_answerSerializer(data=test_candidate_answer_data) 
        
        if serializer.is_valid(): 
            serializer.save() 
            print('....Data Saved Successfully....')
            return JsonResponse(serializer.data, status=201) 
        else: 
            return JsonResponse(serializer.errors, status=400) 
    except Exception as e: 
        return JsonResponse({'error': str(e)}, status=500) 

    


@api_view(['GET'])
def get_last_compiler_output(request, student_id):
    last_output = compiler_output.objects.filter(student_id=student_id).last()
    output = last_output.output if last_output else ""
    
    return Response({"last_output": output})




@api_view(['POST'])
def insert_empty_output_view(request, student_id):
    if not student_id:
        print('student_id is required')
        return Response({'error': 'student_id is required'}, status=400)
    
    try:
        student = candidate_master.objects.get(id=student_id)
        print('student id: ', student)
    except candidate_master.DoesNotExist:
        return Response({'error': 'CandidateMaster with the given id does not exist'}, status=404)
    
    existing_record = compiler_output.objects.filter(student_id=student)
    if not existing_record:
        compiler_output.objects.create(student_id=student, output='')
        print('Record created successfully')
        return Response({'message': 'Record created successfully'})
    else:
        print('Record already exists')
        return Response({'message': 'Record already exists'})


#----------------------------Training admin Dashboard datas-------------------#



def get_max_score_by_department(request, college_id):
    try:
        # First, create a subquery to get the maximum total score per department
        subquery = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='MCQ Test',
            total_score__isnull=False,
            department_id=OuterRef('department_id'),
            deleted=0
        ).values(
            'department_id'
        ).annotate(
            max_total_score=Max('total_score')
        ).values('max_total_score')

        # Then, filter the original query using this subquery to get the corresponding student names
        results = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='MCQ Test',
            total_score__isnull=False,
            total_score=Subquery(subquery),
            deleted=0
        ).values(
            'student_id__students_name',
            'department_id__department',
            'total_score'
        ).distinct()  # Ensure unique records in case of ties

        # Format the results as a list of dictionaries
        data = [
            {
                'student_name': result['student_id__students_name'],
                'department': result['department_id__department'],
                'max_total_score': result['total_score']
            }
            for result in results
        ]

        # Return the results as a JSON response
        return JsonResponse(data, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)



def get_max_score_by_department_coding(request, college_id):
    try:
        # First, create a subquery to get the maximum total score per department
        subquery = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='Coding Test',
            total_score__isnull=False,
            department_id=OuterRef('department_id'),
            deleted=0
        ).values(
            'department_id'
        ).annotate(
            max_total_score=Max('total_score')
        ).values('max_total_score')

        # Then, filter the original query using this subquery to get the corresponding student names
        results = tests_candidates_map.objects.filter(
            college_id=college_id,
            question_id__test_type='Coding Test',
            total_score__isnull=False,
            total_score=Subquery(subquery),
            deleted=0
        ).values(
            'student_id__students_name',
            'department_id__department',
            'total_score'
        ).distinct()  # Ensure unique records in case of ties

        # Format the results as a list of dictionaries
        data = [
            {
                'student_name': result['student_id__students_name'],
                'department': result['department_id__department'],
                'max_total_score': result['total_score']
            }
            for result in results
        ]

        # Return the results as a JSON response
        return JsonResponse(data, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


# Count of company


def count_company_names(request):
    try:
        # Perform the query to count company_name entries
        result = company_master.objects.aggregate(
            count_company_name=Count('company_name')
        )

        # Return the count as a JSON response
        return JsonResponse(result, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


#------------------------------Placement admin Dashboard datas-----------------------#

# URLs for the CSV files
sheet_urls = {
    'AIDS': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=257135364',
    'IT': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=747238312',
    'CSE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=1208914504',
    'ECE': 'https://docs.google.com/spreadsheets/d/1WZ4R-VJmYy21muG8Rvd2DFrULJHzICf_/export?format=csv&gid=2128791655'
}

# Function to process each sheet
def process_sheet(url):
    response = requests.get(url)
    csv_data = response.content.decode('utf-8')
    df = pd.read_csv(StringIO(csv_data), skiprows=7)  # Skip first 7 rows to reach the headers
    df = df[2:]  # Skip the next row (8th and 9th rows)
    df.reset_index(drop=True, inplace=True)  # Reset index after slicing
    df.columns = df.columns.str.strip()  # Strip any leading/trailing spaces from columns
    return df

@api_view(['GET'])
def get_avg_total_present(request):
    try:
        department_data = []

        for department, url in sheet_urls.items():
            df = process_sheet(url)
            # print(f"Headers for {department}: {df.columns}")  # Print headers for debugging

            # Check if 'Total Present' column exists
            if 'Total Present' in df.columns:
                # Calculate the average of 'Total Present'
                avg_total_present = df['Total Present'].astype(float).mean()  # Ensure conversion to float
                department_data.append({
                    "department_name": department,
                    "total_present_avg": round(avg_total_present, 1)
                })
            else:
                department_data.append({
                    "department_name": department,
                    "total_present_avg": 0
                })

        return Response(department_data)
    except Exception as e:
        return Response({'error': str(e)}, status=400)



@api_view(['GET'])
def get_avg_total_absent(request):
    try:
        department_data = []

        for department, url in sheet_urls.items():
            df = process_sheet(url)
            # print(f"Headers for {department}: {df.columns}")  # Print headers for debugging

            # Check if 'Total Present' column exists
            if 'Total Abscent' in df.columns:
                # Calculate the average of 'Total Present'
                avg_total_present = df['Total Abscent'].astype(float).mean()  # Ensure conversion to float
                department_data.append({
                    "department_name": department,
                    "total_present_avg": round(avg_total_present, 1)
                })
            else:
                department_data.append({
                    "department_name": department,
                    "total_absent_avg": 0
                })

        return Response(department_data)
    except Exception as e:
        return Response({'error': str(e)}, status=400)







#---------------------------------------Students Dashboards-----------------------------------#



def get_events_by_college_and_department(request, college_id, department_id):
    try:
        # Filter the events based on college_id and department_id
        events = event_master.objects.filter(
            college_id=college_id,
            department_id=department_id,
            deleted=0
        ).values('event_name', 'event_desc', 'dtm_start')

        # Convert the queryset to a list of dictionaries and format the datetime
        data = []
        for event in events:
            dtm_start = event['dtm_start']
            formatted_dtm_start = dtm_start.strftime('%A %I:%M %p')
            data.append({
                'event_name': event['event_name'],
                'event_desc': event['event_desc'],
                'dtm_start': dtm_start,
                'day_time': formatted_dtm_start
            })

        # Return the results as a JSON response
        return JsonResponse(data, safe=False, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


#---------------Student Course Progress----------------------#


@api_view(['GET'])
def course_progress(request, student_id):
    # Count distinct topics in course_content_feedback for the specified student
    course_completed_count = course_content_feedback.objects.filter(student_id=student_id).values('topic_id').distinct().count()
    
    # Count distinct topics in content_master
    content_master_count = content_master.objects.values('id').distinct().count()
    
    # Calculate course in progress
    course_inprogress_count = content_master_count - course_completed_count
    
    results = {
        "Course_Inprogress": course_inprogress_count,
        "Course_Completed": course_completed_count
    }

    return Response(results)




@api_view(['GET'])
def get_tests_by_student(request, student_id):
    # Filter tests_candidates_map by student_id and select required fields
    tests = tests_candidates_map.objects.filter(student_id=student_id).values('test_name', 'total_score', 'is_active')
    
    return Response(list(tests))


#----------------Performance of students-------------------#


@api_view(['GET'])
def get_avg_total_score_by_month(request, student_id):
    # Perform the query
    result = tests_candidates_map.objects.filter(
        question_id__test_type='MCQ Test', 
        student_id=student_id  # Filter by student_id
    ).annotate(
        month=ExtractMonth('dtm_start')
    ).values('month').annotate(
        avg_total_score=Avg('total_score')
    ).order_by('month')
    
    # Convert the QuerySet to a dictionary with months as keys
    result_dict = {item['month']: item['avg_total_score'] for item in result}

    # Create the final list with all months
    result_list = [
        {
            'month': calendar.month_abbr[month],
            'avg_total_score': result_dict.get(month, 0)
        }
        for month in range(1, 13)
    ]

    return Response(result_list)

#-----------------------newly-- student attendance performance-------------------------------#



@api_view(['GET'])
def attendance_report_new(request, reg_no):
    try:
        department_data = {}

        for department, url in sheet_urls.items():
            df = process_sheet(url)
            
            # Filter data by registration number for the current department
            filtered_data = df[df['Register Number'] == reg_no]
            
            # Group by weekdays and calculate average performance for the current department
            weekdays = group_by_weekdays_NEW(filtered_data)
            avg_performance = calculate_avg_performance_NEW(weekdays)
            
            # Store results for each department
            department_data[department] = {
                'weeks': weekdays,
                'avg_performance': avg_performance
            }
        
        return Response(department_data)
    
    except Exception as e:
        return Response({'error': str(e)}, status=400)

def group_by_weekdays_NEW(data):
    weekdays = {'Monday': [], 'Tuesday': [], 'Wednesday': [], 'Thursday': [], 'Friday': [], 'Saturday': [], 'Sunday': []}
    for idx, row in data.iterrows():
        for column_name in data.columns[4:]:  # Assuming attendance starts from the 5th column
            value = row[column_name]
            if value == 'P':
                try:
                    date_obj = datetime.strptime(column_name, '%d-%b')
                    weekday = date_obj.strftime('%A')
                    weekdays[weekday].append(value)
                except ValueError:
                    # Handle invalid or empty dates if needed
                    pass
    return weekdays

def calculate_avg_performance_NEW(weekdays):
    avg_performance = {}
    for day, attendance in weekdays.items():
        total = len(attendance)
        present = attendance.count('P')
        avg_performance[day] = (present / total) * 100 if total else 0
    return avg_performance


#------------------------Django Login----------------------#

@api_view(['POST'])
def add_user_profile(request):
    if request.method == 'POST':
        serializer = UserProfileSerializer(data=request.data)
        print('serializer: ', serializer)

        if serializer.is_valid():
            print('Before Saving....')
            serializer.save()
            print('Profile Add successfully in User_Profile')
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        else:
            print('serializer.errors: ', serializer.errors)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
def view_user_profile(request):
    try:
        user_profiles = user_profile.objects.all()  # Renamed variable to user_profiles
    except user_profile.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        serializer = UserProfileSerializer(user_profiles, many=True, context={'request': request})  # Updated to handle multiple objects
        return Response(serializer.data)


@api_view(['PUT'])
def update_user_profile(request, pk):
    try:
        user_profiles = user_profile.objects.get(pk=pk)
    except user_profile.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)

    if request.method == 'PUT':
        serializer = UserProfileSerializer(user_profiles, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@api_view(['DELETE'])
def delete_user_profile(request, pk):
    try:
        user_profile = user_profile.objects.get(pk=pk)
    except user_profile.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)

    if request.method == 'DELETE':
        user_profile.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)





@api_view(['GET'])
def custom_user_profiles(request):
    try:
        user_profiles = user_profile.objects.all()
    except user_profile.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)

    data = []
    for profile in user_profiles:
        college = None
        if profile.college_id:
            college = profile.college_id.college
        
        serialized_data = {
            "id": profile.id,
            "username": profile.user.username,
            "email": profile.user.email,
            "password": "****",  # Avoid exposing passwords directly
            "role": profile.role,
            "college_id": profile.college_id.id if profile.college_id else None,  # Use college_id instead of the object
            "college": college if college else None  # Serialize college_master object
        }
        data.append(serialized_data)

    return Response(data)


from django.contrib.auth import authenticate, login, logout

@api_view(['POST'])
def login_view(request):
    username = request.data.get('username')
    password = request.data.get('password')

    print('Request.data: ', request.data)
    
    user = authenticate(username=username, password=password)
    print('user: ', user)
    
    if user is not None:
        try:
            profile = user_profile.objects.get(user=user)
            role = profile.role
            college_id = profile.college_id.id
            college_name = profile.college_id.college if profile.college_id else None

            # Log the user in (this creates the session)
            login(request, user)

            response_data = {
                'message': 'Login successful',
                'role': role,
                'college_id': college_id,
                'college_name': college_name,
            }
            print('Response Data: ', response_data)

            return Response(response_data, status=status.HTTP_200_OK)
        
        except user_profile.DoesNotExist:
            return Response({'message': 'User profile not found'}, status=status.HTTP_404_NOT_FOUND)
    
    return Response({'message': 'Invalid credentials'}, status=status.HTTP_400_BAD_REQUEST)


@csrf_exempt
def logout_view(request):
    if request.method == 'POST':
        # Log the user out and clear the session
        logout(request)

        # Debugging information (if needed)
        print('Session Data After Logout: ', request.session.items())

        return JsonResponse({'message': 'Logout successful'}, status=200)
    
    return JsonResponse({'message': 'Invalid method'}, status=405)

#-------------------------------Student Request------------------------------------#


@api_view(['GET'])
def get_student_request(request):
    try:
        # Fetch all student requests that are not deleted, using select_related to fetch related student data
        stu_queries = student_request.objects.filter(deleted=0).select_related('student_id')
        
        # Annotate the queryset with related fields
        query_data = stu_queries.values(
            'id',
            'student_id__id',  # Fetch student_id from related student model
            'student_id__students_name',  # Fetch students_name from related student model
            'dtm_request',
            'student_query',
            'approved_by',
            'dtm_approved',
            'dtm_student_update',
            'status'
        )
        
        # Prepare the response data
        formatted_data = [
            {
                'id': item['id'],
                'student_id': item['student_id__id'],
                'student_name': item['student_id__students_name'],
                'dtm_request': item['dtm_request'],
                'student_query': item['student_query'],
                'approved_by': item['approved_by'],
                'dtm_approved': item['dtm_approved'],
                'dtm_student_update': item['dtm_student_update'],
                'status': item['status']
            }
            for item in query_data
        ]
        
        return Response(formatted_data)
    
    except Exception as e:
        return Response({'error': str(e)}, status=500)

class student_request_createAPIView(generics.CreateAPIView):
    queryset = student_request.objects.all()
    serializer_class =studentRequestSerializer

    def post(self, request, *args, **kwargs):
        # logger.info("Creating a new student_request")
        response = super().post(request, *args, **kwargs)
        # logger.info("Created a new student_request successfully")
        return response

class student_request_UpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = student_request.objects.all()
    serializer_class =studentRequestSerializer

    def put(self, request, *args, **kwargs):
        # logger.info(f"Updating student_request with id {kwargs.get('pk')}")
        response = super().put(request, *args, **kwargs)
        # logger.info(f"Updated student_request with id {kwargs.get('pk')} successfully")
        return response

    def patch(self, request, *args, **kwargs):
        # logger.info(f"Partially updating student_request with id {kwargs.get('pk')}")
        response = super().patch(request, *args, **kwargs)
        # logger.info(f"Partially updated student_request with id {kwargs.get('pk')} successfully")
        return response



#----------------get LMS Updated id data-------------------#


@api_view(['GET'])
def update_LMS_id(request, id):
    # Get the content_master object by ID
    try:
        lms = content_master.objects.get(id=id)
    except content_master.DoesNotExist:
        return Response(status=404)

    # Serialize the content_master object
    serializer = contentSerializers(lms)
    
    return Response(serializer.data)



@api_view(['GET'])
def update_LMS_Topic_id(request, id):
    # Get the content_master object by ID
    try:
        lms = topic_master.objects.get(id=id)
    except topic_master.DoesNotExist:
        return Response(status=404)

    # Serialize the content_master object
    serializer = topicSerializers(lms)
    
    return Response(serializer.data)




@api_view(['GET'])
def get_test_type_category(request, test_name):
    try:
        # Create a cache key based on the test_name
        cache_key = f'test_type_category_{test_name}'
        
        # Try to get the data from the cache
        test_type_category = cache.get(cache_key)
        
        # If the data is not in the cache, fetch it from the database
        if not test_type_category:
            test_master_instance = test_master.objects.get(test_name=test_name)
            test_type_category = test_master_instance.test_type_id.test_type_categories
            
            # Store the fetched data in the cache with a timeout (e.g., 1 hour)
            cache.set(cache_key, test_type_category, timeout=3600)
        
        # Return the cached or freshly fetched data
        return Response({'test_type_category': test_type_category})
    except test_master.DoesNotExist:
        return Response({'error': 'Test name does not exist'}, status=404)




#------------Need Candidate info-----------------------#


@api_view(['GET'])
def get_test_candidates_NeedInfo(request, username):
    candidates = tests_candidates_map.objects.filter(student_id__user_name=username).select_related('student_id')
    serializer = TestCandidateMapSerializerNeedInfo(candidates, many=True)
    return Response(serializer.data)


#---------------------Student Dasboard New--------------------------#

#--Total Test Taken--------#

@api_view(['GET'])
def active_tests_count(request, student_id):
    count = tests_candidates_map.objects.filter(is_active=True, student_id=student_id).count()
    return Response({'count': count})


#---Total no.of offers------#


@api_view(['GET'])
def get_number_of_offers(request, candidate_id):
    try:
        candidate = candidate_master.objects.get(id=candidate_id)
        number_of_offers = candidate.number_of_offers
    except candidate_master.DoesNotExist:
        return Response({'error': 'Candidate not found'}, status=404)
    return Response({'number_of_offers': number_of_offers})


#----Request count---------#


@api_view(['GET'])
def count_student_requests(request, student_id):
    try:
        request_count = student_request.objects.filter(student_id=student_id).count()
    except student_request.DoesNotExist:
        return Response({'error': 'Student not found'}, status=404)
    return Response({'request_count': request_count})


#------Apditute Avg Score------#

@api_view(['GET'])
def get_monthly_avg_total_score_apditute(request, student_id):
    try:
        # Subquery to get the question_type id where question_type is 'Apditute'
        question_type_subquery = question_type.objects.filter(
            question_type='Aptitude'
        ).values('id')
        print('question_type_subquery: ', question_type_subquery)

        # Subquery to get the test_type id where test_type is 'MCQ Test'
        test_type_subquery = test_type.objects.filter(
            test_type='MCQ Test'
        ).values('id')
        print('test_type_subquery: ', test_type_subquery)

        # Subquery to get the test_name from test_master where the question_type and test_type match
        test_name_subquery = test_master.objects.filter(
            question_type_id__in=Subquery(question_type_subquery),
            test_type_id__in=Subquery(test_type_subquery)
        ).values('test_name')
        print('test_name_subquery: ', test_name_subquery)

        # Main query to get the average total score grouped by month
        monthly_avg_scores = tests_candidates_map.objects.filter(
            test_name__in=Subquery(test_name_subquery),
            student_id=student_id
        ).annotate(month=ExtractMonth('dtm_start')).values('month').annotate(avg_total_score=Avg('total_score')).order_by('month')

        print('monthly_avg_scores: ', monthly_avg_scores)

        # Convert month numbers to month names and create a full list with 0 defaults
        full_months = {i: {"month": calendar.month_abbr[i], "avg_score": 0} for i in range(1, 13)}
        
        print('full months: ', full_months)

        for entry in monthly_avg_scores:
            full_months[entry["month"]]["avg_score"] = entry["avg_total_score"]

        return Response(list(full_months.values()))
    except Exception as e:
        return Response({'error': str(e)}, status=500)


#------SoftSkills Avg Score------#

@api_view(['GET'])
def get_monthly_avg_total_score_softskill(request, student_id):
    try:
        # Subquery to get the question_type id where question_type is 'Apditute'
        question_type_subquery = question_type.objects.filter(
            question_type='Softskills'
        ).values('id')
        print('question_type_subquery: ', question_type_subquery)

        # Subquery to get the test_type id where test_type is 'MCQ Test'
        test_type_subquery = test_type.objects.filter(
            test_type='MCQ Test'
        ).values('id')
        print('test_type_subquery: ', test_type_subquery)

        # Subquery to get the test_name from test_master where the question_type and test_type match
        test_name_subquery = test_master.objects.filter(
            question_type_id__in=Subquery(question_type_subquery),
            test_type_id__in=Subquery(test_type_subquery)
        ).values('test_name')
        print('test_name_subquery: ', test_name_subquery)

        # Main query to get the average total score grouped by month
        monthly_avg_scores = tests_candidates_map.objects.filter(
            test_name__in=Subquery(test_name_subquery),
            student_id=student_id
        ).annotate(month=ExtractMonth('dtm_start')).values('month').annotate(avg_total_score=Avg('total_score')).order_by('month')

        print('monthly_avg_scores: ', monthly_avg_scores)

        # Convert month numbers to month names and create a full list with 0 defaults
        full_months = {i: {"month": calendar.month_abbr[i], "avg_score": 0} for i in range(1, 13)}
        
        print('full months: ', full_months)

        for entry in monthly_avg_scores:
            full_months[entry["month"]]["avg_score"] = entry["avg_total_score"]

        return Response(list(full_months.values()))
    except Exception as e:
        return Response({'error': str(e)}, status=500)




#------Technical Avg Score------#

@api_view(['GET'])
def get_monthly_avg_total_score_technical(request, student_id):
    try:
        # Subquery to get the question_type id where question_type is 'Apditute'
        question_type_subquery = question_type.objects.filter(
            question_type='Technical'
        ).values('id')
        print('question_type_subquery: ', question_type_subquery)

        # Subquery to get the test_type id where test_type is 'MCQ Test'
        test_type_subquery = test_type.objects.filter(
            test_type='MCQ Test'
        ).values('id')
        print('test_type_subquery: ', test_type_subquery)

        # Subquery to get the test_name from test_master where the question_type and test_type match
        test_name_subquery = test_master.objects.filter(
            question_type_id__in=Subquery(question_type_subquery),
            test_type_id__in=Subquery(test_type_subquery)
        ).values('test_name')
        print('test_name_subquery: ', test_name_subquery)

        # Main query to get the average total score grouped by month
        monthly_avg_scores = tests_candidates_map.objects.filter(
            test_name__in=Subquery(test_name_subquery),
            student_id=student_id
        ).annotate(month=ExtractMonth('dtm_start')).values('month').annotate(avg_total_score=Avg('total_score')).order_by('month')

        print('monthly_avg_scores: ', monthly_avg_scores)

        # Convert month numbers to month names and create a full list with 0 defaults
        full_months = {i: {"month": calendar.month_abbr[i], "avg_score": 0} for i in range(1, 13)}
        
        print('full months: ', full_months)

        for entry in monthly_avg_scores:
            full_months[entry["month"]]["avg_score"] = entry["avg_total_score"] if entry["avg_total_score"] is not None else 0

        return Response(list(full_months.values()))
    except Exception as e:
        return Response({'error': str(e)}, status=500)




#------Coding Avg Score------#


@api_view(['GET'])
def get_monthly_avg_total_score_Coding(request, student_id):
    try:
        # Subquery to get the test_type id where test_type is 'CodingTest'
        test_type_subquery = test_type.objects.filter(
            test_type='Coding Test'
        ).values('id')

        # Subquery to get the test_name from test_master where the test_type matches
        test_name_subquery = test_master.objects.filter(
            test_type_id__in=Subquery(test_type_subquery)
        ).values('test_name')

        # Main query to get the average total score grouped by month
        monthly_avg_scores = tests_candidates_map.objects.filter(
            test_name__in=Subquery(test_name_subquery),
            student_id=student_id
        ).annotate(month=ExtractMonth('dtm_start')).values('month').annotate(avg_total_score=Avg('total_score')).order_by('month')

        # Convert month numbers to abbreviated month names and create a full list with 0 defaults
        full_months = {i: {"month": calendar.month_abbr[i], "avg_score": 0} for i in range(1, 13)}

        for entry in monthly_avg_scores:
            full_months[entry["month"]]["avg_score"] = entry["avg_total_score"]

        return Response(list(full_months.values()))
    except Exception as e:
        return Response({'error': str(e)}, status=500)

#------------------Offer Chart where college id


@api_view(['GET'])
def get_candidate_all_college_id(request, college_id):
    try:
        candidatelist = candidate_master.objects.filter(deleted=0, college_id=college_id).select_related('college_id', 'department_id').values(
            'id', 
            'college_id__college',  # Get college name
            'students_name', 
            'user_name', 
            'registration_number', 
            'gender',
            'email_id', 
            'mobile_number', 
            'year', 
            'cgpa', 
            'department_id__department',  # Get department name
            'marks_10th',
            'marks_12th', 
            'marks_semester_wise', 
            'history_of_arrears', 
            'standing_arrears',
            'number_of_offers', 
            'text', 
            'it_of_offers', 
            'core_of_offers', 
            'skill_id',
        )
        candidate_data = list(candidatelist)
        return Response(candidate_data)
    except Exception as e:
        return Response({'error': str(e)}, status=500)

#------------------number of offers --college_id--------

@api_view(['GET'])
def get_number_of_offers_college_id(request, college_id):
    try:
        # Filter candidate_master records by college_id
        candidates = candidate_master.objects.filter(college_id=college_id)
        
        # Aggregate the number of offers for all the candidates
        total_number_of_offers = sum(candidate.number_of_offers for candidate in candidates)

        return Response({'number_of_offers': total_number_of_offers})
    except Exception as e:
        return Response({'error': str(e)}, status=500)


#----Request count---college_id---------#


@api_view(['GET'])
def count_student_requests_college_id(request, college_id):
    try:
        # Filter candidate_master records by college_id
        candidate_ids = candidate_master.objects.filter(college_id=college_id).values_list('id', flat=True)
        
        # Count the number of student_request records with student_id in the filtered candidate ids
        request_count = student_request.objects.filter(student_id__in=candidate_ids).count()
        
        return Response({'request_count': request_count})
    except candidate_master.DoesNotExist:
        return Response({'error': 'College not found'}, status=404)



#----------------------student plan--------------------#

@api_view(['GET'])
def get_schedule_with_tests(request, dtm_start, student_id):
    try:
        # Convert dtm_start to a timezone-aware datetime
        dtm_start = timezone.make_aware(timezone.datetime.strptime(dtm_start, '%Y-%m-%d'))

        # Get the college_id, department_id, and year for the given student_id from candidate_master
        candidate = candidate_master.objects.get(id=student_id)
        print('candidate: ', candidate)

        college_id = candidate.college_id
        print('college_id: ', college_id)

        department_id = candidate.department_id
        print('department_id: ', department_id)

        year = candidate.year
        print('year: ', year)

        # Filter course_schedule based on college_id, department_id, year, and dtm_start
        schedule_with_tests = course_schedule.objects.filter(
            college_id=college_id,
            department_id=department_id,
            year=year,
            dtm_start_student__date=dtm_start.date()
        ).select_related('college_id', 'department_id')
        print('schedule_with_tests: ', schedule_with_tests)

        # Subquery to get test names and dtm_start for the given student_id
        test_names = tests_candidates_map.objects.filter(
            student_id=student_id,
            college_id=college_id,
            department_id=department_id,
            dtm_start__date=dtm_start.date()
        ).values('test_name', 'dtm_start')
        print('test_names: ', test_names)

        # Get attendance data grouped by course_schedule_id
        attendance_data = training_attendance_sheet.objects.filter(
            course_schedule_id__in=schedule_with_tests
        ).values('course_schedule_id').annotate(
            total_present=Count('present'),
            total_absent=Count('absent')
        )
        attendance_dict = {attendance['course_schedule_id']: attendance for attendance in attendance_data}

        # Compile results
        results = []
        if not schedule_with_tests:
            # If schedule_with_tests is empty, return test_names directly
            for test in test_names:
                results.append({
                    'college_name': candidate.college_id.college,
                    'department_name': candidate.department_id.department,
                    'year': candidate.year,
                    'topic_name': '',
                    'subtopic_name': '',
                    'test_name': test['test_name'],
                    'dtm_start': test['dtm_start'],
                    'total_present': 0,
                    'total_absent': 0
                })
        else:
            # Iterate through schedule_with_tests and test_names
            for schedule in schedule_with_tests:
                total_present = attendance_dict.get(schedule.id, {}).get('total_present', 0)
                total_absent = attendance_dict.get(schedule.id, {}).get('total_absent', 0)
                for test in test_names:
                    results.append({
                        'college_name': schedule.college_id.college,
                        'department_name': schedule.department_id.department,
                        'year': schedule.year,
                        'topic_name': schedule.topic,
                        'subtopic_name': schedule.sub_topic,
                        'test_name': test['test_name'],
                        'dtm_start': test['dtm_start'],
                        'total_present': total_present,
                        'total_absent': total_absent
                    })

        return Response(results)
    except Exception as e:
        return Response({'error': str(e)}, status=500)


@api_view(['GET'])
def get_tests_candidates_map_MCQ(request, user_name):
    # Get the current date
    current_date = datetime.now().date()

    # Apply the filters and fetch data
    tests_candidates = tests_candidates_map.objects.filter(
        Q(deleted=0),
        Q(student_id__user_name=user_name),
        Q(question_id__test_type='MCQ Test'),
        Q(is_active=False),
        Q(dtm_start__date__gte=current_date) | Q(dtm_start__date__lt=current_date, dtm_end__date__gte=current_date)
    ).select_related(
        'college_id', 'department_id', 'question_id', 'student_id', 'rules_id'
    ).values(
        'id',
        'test_name',
        'college_id__id',
        'college_id__college',
        'department_id__id',
        'department_id__department',
        'question_id__id',
        'question_id__question_paper_name',
        'question_id__test_type',
        'student_id__id',
        'student_id__students_name',
        'student_id__user_name',
        'dtm_start',
        'dtm_end',
        'attempt_count',
        'is_camera_on',
        'is_active',
        'duration',
        'duration_type',
        'year',
        'rules_id__id',
        'rules_id__rule_name',
        'rules_id__instruction',
        'need_candidate_info',
        'total_score',
        'avg_mark'
    )

    # Rename the keys to match the expected output
    test_candidate_map_data = [
        {
            'id': testing['id'],
            'test_name': testing['test_name'],
            'college_id_id': testing['college_id__id'],
            'college_id': testing['college_id__college'],
            'department_id_id': testing['department_id__id'],
            'department_id': testing['department_id__department'],
            'question_id': testing['question_id__id'],
            'question_paper_name': testing['question_id__question_paper_name'],
            'test_type': testing['question_id__test_type'],
            'student_id': testing['student_id__id'],
            'student_name': testing['student_id__students_name'],
            'user_name': testing['student_id__user_name'],
            'dtm_start': testing['dtm_start'],
            'dtm_end': testing['dtm_end'],
            'attempt_count': testing['attempt_count'],
            'is_camera_on': testing['is_camera_on'],
            'is_active': testing['is_active'],
            'duration': testing['duration'],
            'duration_type': testing['duration_type'],
            'year': testing['year'],
            'rules_id': testing['rules_id__id'],
            'rules': testing['rules_id__rule_name'],
            'instruction': testing['rules_id__instruction'],
            'need_candidate_info': testing['need_candidate_info'],
            'total_score': testing['total_score'],
            'avg_mark': testing['avg_mark']
        }
        for testing in tests_candidates
    ]


    # Return the cached or freshly generated data
    return Response(test_candidate_map_data)



@api_view(['GET']) 

def get_tests_candidates_map_Coding(request, user_name): 
    current_date = datetime.now().date() 

    # Optimize database queries by using select_related 

    tests_candidates = tests_candidates_map.objects.filter( 

        Q(deleted=0), 

        Q(student_id__user_name=user_name), 

        Q(question_id__test_type='Coding Test'), 

        Q(is_active=False), 

        Q(dtm_start__date__gte=current_date) | Q(dtm_start__date__lt=current_date, dtm_end__date__gte=current_date) 

    ).select_related( 

        'question_id', 'student_id', 'rules_id' 

    ).values( 

        'id', 'test_name', 'question_id__id',  'student_id__id', 

        'dtm_start', 'dtm_end', 'duration', 'duration_type', 

        'rules_id__instruction',

    )   

    # Return cached or freshly generated data 

    return Response(list(tests_candidates)) 



#----------------------PLACEMENT-----------------------#

@api_view(['GET'])
def get_tests_Reports_placement(request, college_id):
    try:
        # Fetch all test candidates for the given college that are not deleted
        tests_candidates = tests_candidates_map.objects.filter(
            deleted=0, college_id=college_id
        ).select_related(
           
            'department_id',
           
            'student_id',
            'college_id'
        ).values(
            'id', 'test_name',  'is_active',
            'year', 
            'total_score', 'avg_mark',
          
            'department_id__department', 
            'student_id__id', 'student_id__students_name', 'student_id__user_name',
            'student_id__email_id', 'student_id__mobile_number', 'student_id__gender',
            'student_id__registration_number',
            'college_id__college',
            'dtm_start', 'dtm_end'
        )
        
        # Format and prepare the response data
        test_candidate_map_data = [
            {
                'id': item['id'],
                'test_name': item['test_name'],
                'college_id': item['college_id__college'],
                'department_id': item['department_id__department'],
               
                
                'registration_number': item['student_id__registration_number'],
                'email_id': item['student_id__email_id'],
                'mobile_number': item['student_id__mobile_number'],
                'gender': item['student_id__gender'],
                'student_name': item['student_id__students_name'],
                'user_name': item['student_id__user_name'],
                'dtm_start': django_format_date(localtime(item['dtm_start']), 'd-m-Y h:i A'),
                'dtm_end': django_format_date(localtime(item['dtm_end']), 'd-m-Y h:i A'),
               
                'is_active': item['is_active'],
               
                'year': item['year'],
                
                'total_score': item['total_score'],
                'avg_mark': item['avg_mark']
            }
            for item in tests_candidates
        ]
        
        return Response(test_candidate_map_data)

    except Exception as e:
        return Response({'error': str(e)}, status=500)



@api_view(['GET'])
def get_group_test_name_placement(request, college_id):
    # Join tests_candidates_map with test_master and filter by test_name and deleted=0
    tests_candidates = (
        tests_candidates_map.objects.filter(deleted=0, college_id=college_id)
        .values('test_name', 'question_id', 'question_id__question_paper_name', 'dtm_start', 'dtm_end', 'dtm_created')
        .annotate(student_count=Count('student_id'))
        .order_by(
            Case(
                When(dtm_created__isnull=True, then=1),
                default=0,
                output_field=DateTimeField()
            ),
            '-dtm_created'
        )
    )

    test_candidate_map_data = []
    for testing in tests_candidates:
        # Format dates
        dtm_start_formatted = django_format_date(localtime(testing['dtm_start']), 'd-m-Y h:i A')
        dtm_end_formatted = django_format_date(localtime(testing['dtm_end']), 'd-m-Y h:i A')

        test_candidate_map_data.append({
            'test_name': testing['test_name'],
            'question_id': testing['question_id'],
            'question_paper_name': testing['question_id__question_paper_name'],
            'dtm_start': dtm_start_formatted,
            'dtm_end': dtm_end_formatted,
            'student_count': testing['student_count'],
            'dtm_created': testing['dtm_created']
        })

    return Response(test_candidate_map_data)


@api_view(['GET'])
def get_tests_Reports_placement_Candidates(request, college_id, test_name):
    # Fetch related data using select_related for efficiency
    tests_candidates = tests_candidates_map.objects.filter(
        deleted=0,
        college_id=college_id,
        test_name=test_name
    ).select_related(
        'rules_id', 
        'department_id', 
        'question_id', 
        'student_id', 
        'college_id'
    ).values(
        'id',
        'test_name',
        'college_id__college',
        'department_id__department',
        'question_id__question_paper_name',
        'student_id__id',
        'student_id__students_name',
        'student_id__user_name',
        'student_id__email_id',
        'student_id__mobile_number',
        'student_id__gender',
        'student_id__registration_number',
        'rules_id__rule_name',
        'rules_id__instruction',
        'dtm_start',
        'dtm_end',
        'attempt_count',
        'is_camera_on',
        'is_active',
        'duration',
        'duration_type',
        'year',
        'need_candidate_info',
        'total_score',
        'avg_mark'
    )
    
    # Convert the QuerySet to a list of dictionaries
    test_candidate_map_data = []
    for testing in tests_candidates:
        dtm_start_formatted = django_format_date(localtime(testing['dtm_start']), 'd-m-Y h:i A')
        dtm_end_formatted = django_format_date(localtime(testing['dtm_end']), 'd-m-Y h:i A')
        
        test_candidate_map_data.append({
            'id': testing['id'],
            'test_name': testing['test_name'],
            'college_id': testing['college_id__college'],
            'department_id': testing['department_id__department'],
            'question_id': testing['question_id__question_paper_name'],
            'student_id': testing['student_id__id'],
            'registration_number': testing['student_id__registration_number'],
            'email_id': testing['student_id__email_id'],
            'mobile_number': testing['student_id__mobile_number'],
            'gender': testing['student_id__gender'],
            'student_name': testing['student_id__students_name'],
            'user_name': testing['student_id__user_name'],
            'dtm_start': dtm_start_formatted,
            'dtm_end': dtm_end_formatted,
            'attempt_count': testing['attempt_count'],
            'is_camera_on': testing['is_camera_on'],
            'is_active': testing['is_active'],
            'duration': testing['duration'],
            'duration_type': testing['duration_type'],
            'year': testing['year'],
            'rules': testing['rules_id__rule_name'],
            'instruction': testing['rules_id__instruction'],
            'need_candidate_info': testing['need_candidate_info'],
            'total_score': testing['total_score'],
            'avg_mark': testing['avg_mark']
        })
    
    return Response(test_candidate_map_data)

@api_view(['GET'])
def get_candidate_all_job_master(request):
    try:
        # Fetch all jobs that are not deleted and include related fields
        jobs = job_master.objects.filter(deleted=0).select_related(
            'college_id',
            'department_id'
        ).prefetch_related(
            'skill_id'
        )

        if not jobs.exists():
            return Response({'error': 'No jobs found'}, status=404)

        all_candidates_data = []

        # Using list comprehension and Django ORM's conditional aggregation
        job_candidates = []
        for job in jobs:
            filters = Q(deleted=0)
            
            # Apply filters based on job attributes
            filters &= Q(college_id=job.college_id) if job.college_id else filters
            filters &= Q(department_id=job.department_id) if job.department_id else filters
            filters &= Q(gender=job.gender) if job.gender else filters
            filters &= Q(marks_10th__gte=job.marks_10th) if job.marks_10th else filters
            filters &= Q(marks_12th__gte=job.marks_12th) if job.marks_12th else filters
            filters &= Q(cgpa__gte=job.cgpa) if job.cgpa else filters
            filters &= Q(history_of_arrears__lte=job.history_of_arrears) if job.history_of_arrears is not None else filters
            filters &= Q(standing_arrears__lte=job.standing_arrears) if job.standing_arrears is not None else filters
            filters &= Q(skill_id__in=job.skill_id.values_list('id', flat=True)) if job.skill_id.exists() else filters

            candidates = candidate_master.objects.filter(filters).select_related(
                'college_id',
                'department_id'
            ).prefetch_related(
                'skill_id'
            )
            
            for candidate in candidates:
                skill_ids = list(candidate.skill_id.values_list('id', flat=True))
                job_candidates.append({
                    'job_id': job.id,
                    'student_id': candidate.id,
                    'college_id': candidate.college_id.college if candidate.college_id else None,
                    'students_name': candidate.students_name,
                    'user_name': candidate.user_name,
                    'registration_number': candidate.registration_number,
                    'gender': candidate.gender,
                    'email_id': candidate.email_id,
                    'mobile_number': candidate.mobile_number,
                    'year': candidate.year,
                    'cgpa': candidate.cgpa,
                    'department_id': candidate.department_id.department if candidate.department_id else None,
                    'marks_10th': candidate.marks_10th,
                    'marks_12th': candidate.marks_12th,
                    'marks_semester_wise': candidate.marks_semester_wise,
                    'history_of_arrears': candidate.history_of_arrears,
                    'standing_arrears': candidate.standing_arrears,
                    'number_of_offers': candidate.number_of_offers,
                    'text': candidate.text,
                    'it_of_offers': candidate.it_of_offers,
                    'core_of_offers': candidate.core_of_offers,
                    'skill_id': skill_ids
                })

        return Response(job_candidates)

    except Exception as e:
        return Response({'error': str(e)}, status=500)
#-----------Students Test Schedule-----------------------#


@api_view(['GET'])
def get_students_test_schedule(request, user_name):
    # Fetch related data using select_related for efficiency
    tests_candidates = tests_candidates_map.objects.filter(
        deleted=0,
        student_id__user_name=user_name
    ).select_related(
        'question_id', 
        'student_id', 
    ).values(
        'id',
        'test_name',
        'question_id__id',
        'question_id__question_paper_name',
        'question_id__test_type',
        'dtm_start',
        'dtm_end',
        
        'is_active',
        'need_candidate_info',
        'total_score',
    )
    
    return Response(list(tests_candidates))
#---------------------Questions Master with images--------------#
def get_questions_IO_filter(request, question_id):
    # Fetch the required data using select_related and values
    questionset = question_master.objects.filter(deleted=0, question_name_id=question_id).select_related('question_name_id').values(
        'id',
        'question_name_id__id',
        'question_name_id__question_paper_name',
        'question_text',
        'question_image_data',
        'option_a_image_data',
        'option_b_image_data',
        'option_c_image_data',
        'option_d_image_data',
        'option_a',
        'option_b',
        'option_c',
        'option_d',
        'view_hint',
        'mark',
        'explain_answer',
        'answer',
        'negative_mark',
        'input_format'
    )

    question_data = []
    for question in questionset:
        question_image_data = base64.b64encode(question['question_image_data']).decode('utf-8') if question['question_image_data'] else None
        option_a_image_data = base64.b64encode(question['option_a_image_data']).decode('utf-8') if question['option_a_image_data'] else None
        option_b_image_data = base64.b64encode(question['option_b_image_data']).decode('utf-8') if question['option_b_image_data'] else None
        option_c_image_data = base64.b64encode(question['option_c_image_data']).decode('utf-8') if question['option_c_image_data'] else None
        option_d_image_data = base64.b64encode(question['option_d_image_data']).decode('utf-8') if question['option_d_image_data'] else None

        question_data.append({
            'id': question['id'],
            'question_name_id': question['question_name_id__id'],
            'question_paper_name': question['question_name_id__question_paper_name'],
            'question_text': question['question_text'],
            'question_image_data': question_image_data,
            'option_a_image_data': option_a_image_data,
            'option_b_image_data': option_b_image_data,
            'option_c_image_data': option_c_image_data,
            'option_d_image_data': option_d_image_data,
            'option_a': question['option_a'],
            'option_b': question['option_b'],
            'option_c': question['option_c'],
            'option_d': question['option_d'],
            'view_hint': question['view_hint'],
            'mark': question['mark'],
            'explain_answer': question['explain_answer'],
            'answer': question['answer'],
            'negative_mark': question['negative_mark'],
            'input_format': question['input_format'],
        })

    random.shuffle(question_data)
    return JsonResponse(question_data, safe=False)
@api_view(['GET'])
def get_tests_candidates_map_testName(request, test_name):
    # Fetch related data using select_related for efficiency
    tests_candidates = tests_candidates_map.objects.filter(
        deleted=0,
        test_name=test_name
    ).select_related(
        'rules_id', 
        'department_id', 
        'question_id', 
        'student_id', 
        'college_id'
    ).values(
        'id',
        'test_name',
        'college_id__college',
        'college_id__id',
        'department_id__department',
        'department_id__id',
        'question_id__id',
        'question_id__question_paper_name',
        'question_id__test_type',
        'student_id__id',
        'student_id__students_name',
        'student_id__user_name',
        'rules_id__id',
        'rules_id__rule_name',
        'rules_id__instruction',
        'dtm_start',
        'dtm_end',
        'attempt_count',
        'is_camera_on',
        'is_active',
        'duration',
        'duration_type',
        'year',
        'need_candidate_info',
        'total_score',
        'avg_mark',
        'dtm_created'
    )
    
    # Convert the QuerySet to a list of dictionaries
    test_candidate_map_data = []
    for testing in tests_candidates:
        test_candidate_map_data.append({
            'id': testing['id'],
            'test_name': testing['test_name'],
            'college_id_id': testing['college_id__id'],
            'college_id': testing['college_id__college'],
            'department_id_id': testing['department_id__id'],
            'department_id': testing['department_id__department'],
            'question_id': testing['question_id__id'],
            'question_paper_name': testing['question_id__question_paper_name'],
            'test_type': testing['question_id__test_type'],
            'student_id': testing['student_id__id'],
            'student_name': testing['student_id__students_name'],
            'user_name': testing['student_id__user_name'],
            'dtm_start': testing['dtm_start'],
            'dtm_end': testing['dtm_end'],
            'attempt_count': testing['attempt_count'],
            'is_camera_on': testing['is_camera_on'],
            'is_active': testing['is_active'],
            'duration': testing['duration'],
            'duration_type': testing['duration_type'],
            'year': testing['year'],
            'rules_id': testing['rules_id__id'],
            'rules': testing['rules_id__rule_name'],
            'instruction': testing['rules_id__instruction'],
            'need_candidate_info': testing['need_candidate_info'],
            'total_score': testing['total_score'],
            'avg_mark': testing['avg_mark'],
            'dtm_created': testing['dtm_created'],
        })

    return Response(test_candidate_map_data)
@api_view(['GET'])
def get_tests_Reports_candidates(request, test_name):
    # Fetch related data using select_related for efficiency
    tests_candidates = tests_candidates_map.objects.filter(
        deleted=0,
        test_name=test_name
    ).select_related(
        'rules_id', 
        'department_id', 
        'question_id', 
        'student_id', 
        'college_id'
    ).values(
        'id',
        'test_name',
        'college_id__college',
        'department_id__department',
        'question_id__question_paper_name',
        'student_id__id',
        'student_id__students_name',
        'student_id__user_name',
        'student_id__email_id',
        'student_id__mobile_number',
        'student_id__gender',
        'student_id__registration_number',
        'rules_id__rule_name',
        'rules_id__instruction',
        'dtm_start',
        'dtm_end',
        'attempt_count',
        'is_camera_on',
        'is_active',
        'duration',
        'duration_type',
        'year',
        'need_candidate_info',
        'total_score',
        'avg_mark'
    )
    
    # Format the datetime fields
    test_candidate_map_data = []
    for testing in tests_candidates:
        dtm_start_formatted = django_format_date(localtime(testing['dtm_start']), 'd-m-Y h:i A')
        dtm_end_formatted = django_format_date(localtime(testing['dtm_end']), 'd-m-Y h:i A')

        test_candidate_map_data.append({
            'id': testing['id'],
            'test_name': testing['test_name'],
            'college_id': testing['college_id__college'],
            'department_id': testing['department_id__department'],
            'question_id': testing['question_id__question_paper_name'],
            'student_id': testing['student_id__id'],
            'registration_number': testing['student_id__registration_number'],
            'email_id': testing['student_id__email_id'],
            'mobile_number': testing['student_id__mobile_number'],
            'gender': testing['student_id__gender'],
            'student_name': testing['student_id__students_name'],
            'user_name': testing['student_id__user_name'],
            'dtm_start': dtm_start_formatted,
            'dtm_end': dtm_end_formatted,
            'attempt_count': testing['attempt_count'],
            'is_camera_on': testing['is_camera_on'],
            'is_active': testing['is_active'],
            'duration': testing['duration'],
            'duration_type': testing['duration_type'],
            'year': testing['year'],
            'rules': testing['rules_id__rule_name'],
            'instruction': testing['rules_id__instruction'],
            'need_candidate_info': testing['need_candidate_info'],
            'total_score': testing['total_score'],
            'avg_mark': testing['avg_mark']
        })

    return Response(test_candidate_map_data)


class CourseScheduleMapView(APIView):
    def post(self, request, format=None):
        print('request.data: ', request.data)

        college_id = request.data.get('college_id', [])
        department_id = request.data.get('department_id', [])
        year = request.data.get('year')
        topic_ids = request.data.get('topic_id', [])  # Assuming 'topic_ids' is provided as a list in the request
        trainer_ids = request.data.get('trainer_id', [])

        if not college_id or not department_id or not year or not topic_ids:
            return Response({'error': 'college_id, department_id, year, and topic_ids are required fields.'}, status=status.HTTP_400_BAD_REQUEST)

        # Filter candidates based on multiple college_ids and department_ids
        students = candidate_master.objects.filter(
            college_id__in=college_id,
            department_id__in=department_id,
            year=year
        )

        data = []
        current_date_and_time = datetime.now()  # Get the current date and time

        for student in students:
            for topic_id in topic_ids:
                for trainer_id in trainer_ids:
                    test_candidate_data = {
                        'topic_id': topic_id,  # Use the topic_id directly
                        'student_id': student.id,
                        'college_id': student.college_id.id,  # Use student.college_id.id for pk
                        'department_id': student.department_id.id,  # Use student.department_id.id for pk
                        'dtm_start_student': request.data.get('dtm_start_student'),
                        'dtm_end_student': request.data.get('dtm_end_student'),
                        'dtm_start_trainer': request.data.get('dtm_start_trainer'),
                        'dtm_end_trainer': request.data.get('dtm_end_trainer'),
                        'dtm_of_training': request.data.get('dtm_of_training'),
                        'year': year,
                        'trainer_id': trainer_id,  # Use individual trainer_id
                        'dtm_created': current_date_and_time  # Add current date and time
                    }

                    serializer = courseScheduleSerializer(data=test_candidate_data)
                    if serializer.is_valid():
                        serializer.save()
                        data.append(serializer.data)
                    else:
                        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        return Response(data, status=status.HTTP_201_CREATED)



@api_view(['GET'])
def get_content_master_subTopic(request, topic_name):
    if not topic_name:
        return Response({'error': 'Topic name is required.'}, status=status.HTTP_400_BAD_REQUEST)

    # Fetch the id and sub_topic based on the topic name
    topics = content_master.objects.filter(topic=topic_name).values('id', 'sub_topic')
    if not topics.exists():
        return Response({'error': f'No topics found for the given topic name: {topic_name}'}, status=status.HTTP_404_NOT_FOUND)

    return Response(list(topics), status=status.HTTP_200_OK)

@api_view(['GET'])
def students_course_content_view(request):
    user_name = request.query_params.get('user_name')
    current_date_time = datetime.now()

    if not user_name:
        return Response({'error': 'user_name is required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        candidate = get_object_or_404(candidate_master, user_name=user_name)
        student_id = candidate.id

        course_schedules = course_schedule.objects.filter(
            deleted=0,
            student_id=student_id,
            dtm_start_student__lte=current_date_time,
            dtm_end_student__gte=current_date_time
        ).select_related('topic_id')

        data = []
        for cs in course_schedules:
            cm = cs.topic_id
            data.append({
                'topic': cm.topic,
                'sub_topic': cm.sub_topic,
                'Content_URL': cm.content_url,
                'Actual_Content': cm.actual_content,
                'Start_Date': cs.dtm_start_student,
                'End_Date': cs.dtm_end_student,
            })

        return Response(data, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['DELETE'])
def delete_question_paper_by_name(request, question_paper_name):
    try:
        question_paper = get_object_or_404(question_paper_master, question_paper_name=question_paper_name)
        question_paper.delete()
        return Response({'message': 'Question paper deleted successfully'}, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'error': f'Error occurred while trying to delete the question paper: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
def Trainer_course_content_view(request):
    user_name = request.query_params.get('user_name')
    current_date_time = datetime.now()

    if not user_name:
        return Response({'error': 'user_name is required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        trainer = get_object_or_404(trainer_master, trainer_name=user_name)
        trainer_id = trainer.id

        course_schedules = course_schedule.objects.filter(
            deleted=0,
            trainer_id=trainer_id,
            dtm_start_trainer__lte=current_date_time,
            dtm_end_trainer__gte=current_date_time
        ).select_related('topic_id')

        unique_data = {}
        for cs in course_schedules:
            cm = cs.topic_id
            key = (cm.topic, cm.sub_topic, cm.content_url, cm.actual_content, cs.dtm_start_trainer, cs.dtm_end_trainer)
            if key not in unique_data:
                unique_data[key] = {
                    'topic': cm.topic,
                    'sub_topic': cm.sub_topic,
                    'Content_URL': cm.content_url,
                    'Actual_Content': cm.actual_content,
                    'Start_Date': cs.dtm_start_trainer,
                    'End_Date': cs.dtm_end_trainer,
                }

        data = list(unique_data.values())

        return Response(data, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



@api_view(['GET'])
def get_candidate_user_name(request):
    user_name = request.query_params.get('user_name')
    
    if not user_name:
        return Response({'error': 'user_name is required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        # Fetch candidates with related college and department, and prefetch skills
        candidatelist = candidate_master.objects.filter(
            deleted=0,
            user_name=user_name
        ).select_related('college_id', 'department_id').prefetch_related('skill_id').values(
            'id', 
            'college_id',
            'college_id__college',  # Get college name
            'students_name', 
            'user_name', 
            'registration_number', 
            'gender',
            'email_id', 
            'mobile_number', 
            'year', 
            'cgpa', 
            'department_id',
            'department_id__department',  # Get department name
            'marks_10th',
            'marks_12th', 
            'marks_semester_wise', 
            'history_of_arrears', 
            'standing_arrears',
            'number_of_offers', 
            'text', 
            'it_of_offers', 
            'core_of_offers'
        ).distinct()

        # Create a list to store the final output
        candidate_data = []
        
        for candidate in candidatelist:
            skills = candidate_master.objects.get(id=candidate['id']).skill_id.all()
            skill_ids = [skill.id for skill in skills]
            candidate['skill_id'] = skill_ids
            candidate_data.append(candidate)
        
        return Response(candidate_data, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def candidate_info_OLD(request):
    user_name = request.GET.get('user_name', None)

    if not user_name:
        return Response({'error': 'user_name parameter is required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        candidate = candidate_master.objects.get(user_name=user_name)
        need_candidate_info = candidate.need_candidate_info

        return Response({'need_candidate_info': need_candidate_info}, status=status.HTTP_200_OK)

    except candidate_master.DoesNotExist:
        return Response({'error': 'Candidate not found'}, status=status.HTTP_404_NOT_FOUND)



@api_view(['GET'])
def candidate_info(request):
    user_name = request.GET.get('user_name', None)

    if not user_name:
        return Response({'error': 'user_name parameter is required'}, status=status.HTTP_400_BAD_REQUEST)

    # Define cache key
    cache_key = f'candidate_info_{user_name}'
    # Attempt to retrieve data from cache
    cached_data = cache.get(cache_key)

    if cached_data:
        # If data is found in cache, return it
        return Response(cached_data, status=status.HTTP_200_OK)

    try:
        # Fetch candidate from database
        candidate = candidate_master.objects.get(user_name=user_name)
        need_candidate_info = candidate.need_candidate_info

        # Prepare response data
        if not need_candidate_info:
            # Special handling for when need_candidate_info is false
            response_data = {'message': 'Candidate info is not required'}
            # Cache the result for future requests
            cache.set(cache_key, response_data, timeout=900)  # 15 mins
        else:
            # Normal handling when need_candidate_info is true
            response_data = {'need_candidate_info': need_candidate_info}

        return Response(response_data, status=status.HTTP_200_OK)

    except candidate_master.DoesNotExist:
        return Response({'error': 'Candidate not found'}, status=status.HTTP_404_NOT_FOUND)

    except Exception as e:
        # Handle unexpected errors
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



@api_view(['PUT'])
def update_candidate_info(request):
    user_name = request.data.get('user_name')
    need_candidate_info = request.data.get('need_candidate_info', False)  # Default to False if not provided

    if not user_name:
        return Response({'error': 'user_name is required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        candidate = candidate_master.objects.get(user_name=user_name)
        candidate.need_candidate_info = need_candidate_info
        candidate.save()

        return Response({'success': f'Updated need_candidate_info for user {user_name}'}, status=status.HTTP_200_OK)

    except candidate_master.DoesNotExist:
        return Response({'error': 'Candidate not found'}, status=status.HTTP_404_NOT_FOUND)





@api_view(['GET'])
def get_department_info(request):
    college_ids = request.query_params.getlist('college_id')  # Retrieve college_ids from query parameters
   # print('college_ids: ', college_ids)

    if not college_ids:
        return Response({"error": "No college_ids provided"}, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        # Perform the query using Django's ORM
        results = candidate_master.objects.filter(
            college_id__in=college_ids
        ).values(
            department_id_value=F('department_id__id'),
            department_name_value=F('department_id__department')
        ).distinct()

        return Response(results, status=status.HTTP_200_OK)

    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)




@api_view(['GET'])
def get_department_info_LMS(request):
    college_ids = request.query_params.getlist('college_id')  # Retrieve college_ids from query parameters
   # print('college_ids: ', college_ids)

    if not college_ids:
        return Response({"error": "No college_ids provided"}, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        # Perform the query using Django's ORM
        results = candidate_master.objects.filter(
            college_id__in=college_ids
        ).values(
            department_id_value=F('department_id__id'),
            department_name_value=F('department_id__department')
        ).distinct()

        return Response(results, status=status.HTTP_200_OK)

    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



@api_view(['GET'])
def get_candidate_details(request):
    user_name = request.query_params.get('user_name')
    if user_name:
        candidates = candidate_master.objects.filter(user_name=user_name).values('id', 'college_id', 'department_id')
        return Response(candidates)
    return Response({"error": "user_name parameter is required"}, status=400)


#------------------------college with Logo------------------------------#
@api_view(['GET'])
def get_colleges(request):
    # Query all colleges with their logos
    collegeset = college_master.objects.all().values(
        'id',
        'college',
        'college_logo'  # Fetch the binary data for the logo
    )

    # Prepare the data with Base64 encoding for logos
    college_data = []
    for college in collegeset:
        college_logo = None
        if college['college_logo']:
            college_logo = base64.b64encode(college['college_logo']).decode('utf-8')
        
        college_data.append({
            'id': college['id'],
            'college': college['college'],
            'college_logo': college_logo,
        })
    
    return JsonResponse(college_data, safe=False)
@csrf_exempt
def upload_college(request):
    if request.method == 'POST':
        form = CollegeForm(request.POST, request.FILES)
        if form.is_valid():
            college = form.save(commit=False)
            if 'college_logo' in request.FILES:
                college.college_logo = request.FILES['college_logo'].read()
            college.save()
            logger.info("College Created Successfully")
            return HttpResponse("College created successfully")
    else:
        form = CollegeForm()
    return render(request, 'upload_college.html', {'form': form})

@csrf_exempt
def update_college(request, college_id):
    college = get_object_or_404(college_master, id=college_id)

    if request.method == 'POST':
        form = CollegeForm(request.POST, request.FILES, instance=college)
        if form.is_valid():
            college = form.save(commit=False)
            if 'college_logo' in request.FILES:
                college.college_logo = request.FILES['college_logo'].read()
            else:
                college.college_logo = None

            college.save()
            logger.info("College Updated Successfully")
            return HttpResponse("College updated successfully")
    else:
        form = CollegeForm(instance=college)

    return render(request, 'update_college.html', {'form': form})

@csrf_exempt
def delete_college(request, pk):
    logger.info(f"Attempting to delete college with id {pk}")
    try:
        college = college_master.objects.get(id=pk)
    except college_master.DoesNotExist:
        return JsonResponse("College not found", status=404)

    college.delete()
    logger.info(f"Deleted college with id {pk} successfully")
    return JsonResponse("College deleted successfully", safe=False)


#___________________________________________Eligible_students_____________________________________#

@api_view(['GET'])
def get_eligible_studentsall(request):
    # Filter the eligible students based on your criteria
    eligible_students_data = eligible_student_list.objects.all().values(
        'id',
        'students_id__id',  
        'students_id__students_name',  
        'students_id__user_name',
        'students_id__department_id__department',
        'students_id__mobile_number',
        'students_id__registration_number',
        'students_id__year',
        'students_id__email_id',  # Access related object's field (candidate_master id)
       # 'students_id__skill_id__skill_name',  # Access related object's field (candidate_master name)
        'students_id__gender',
        'students_id__college_id__college',
        'students_id__cgpa',
        'announcement',
        'job_id__id',  # Access related object's field (job_master id)
        'job_id__company_name',  # Access related object's field (job_master company name)
        'job_id__company_profile',
        'job_id__location',
        'job_id__interview_date',
        'round_of_interview',
        'whatsapp_text',
        'is_accept'
    )
    
    # Convert the QuerySet to a list of dictionaries
    eligible_students_list = list(eligible_students_data)

    return Response(eligible_students_list)
@api_view(['GET'])
def get_registered_count(request):
    # Annotate each job with the count of accepted students
    eligible_students_data = eligible_student_list.objects.filter(is_accept=True).values(
        'job_id__id',  # Job ID
        'job_id__company_name',  # Job Company Name
        'job_id__company_profile',
        'job_id__location',
        'job_id__interview_date'
    ).annotate(
        accepted_students_count=Count('id')  # Count of students who accepted the job
    ).distinct()

    # Convert the QuerySet to a list of dictionaries
    eligible_students_list = list(eligible_students_data)

    return Response(eligible_students_list)

@api_view(['GET'])
def get_eligible_students(request):
    job_id = request.query_params.get('job_id', None)

    if not job_id:
        return Response({'error': 'job_id is required'}, status=400)

    try:
        # Filter the eligible students based on job_id
        eligible_students_data = eligible_student_list.objects.filter(job_id_id=job_id).values(
            'id',
            'students_id__id',  # Access related object's field (candidate_master id)
            'students_id__students_name',  # Access related object's field (candidate_master name)
            'students_id__department_id__department',
            'students_id__mobile_number',
            'students_id__registration_number',
            'students_id__year',
            'students_id__email_id',  # Access related object's field (candidate_master id)
            #'students_id__skill_id__skill_name',  # Access related object's field (candidate_master name)
            'students_id__gender',
            'students_id__college_id',
            'students_id__cgpa',
            'announcement',
            'job_id__id',  # Access related object's field (job_master id)
            'job_id__company_name',  # Access related object's field (job_master company name)
            'round_of_interview',
            'whatsapp_text'
        )
        
        # Convert the QuerySet to a list of dictionaries
        eligible_students_list = list(eligible_students_data)

        return Response(eligible_students_list)

    except Exception as e:
        print('Exception:', str(e))
        return Response({'error': str(e)}, status=500)

@api_view(['GET'])
def get_eligible_students_count(request):
    # Aggregate the eligible students by job_id and count the number of students
    eligible_students_data = eligible_student_list.objects.values(
        'job_id__id',  # Job ID
        'job_id__company_name',  # Company Name
      
    ).annotate(
        student_count=Count('students_id')  # Count of students for each job
    ).values(
        'job_id__id',
        'job_id__company_name',
      
        'student_count'
    )
    # Convert the QuerySet to a list of dictionaries
    eligible_students_list = list(eligible_students_data)

    return Response(eligible_students_list)


@api_view(['GET'])
def get_database_candidate(request):
    try:
        # Fetch candidates with related college and department, filtering by is_database=True
        candidatelist = candidate_master.objects.filter(deleted=0, is_database=True)\
            .select_related('college_id', 'department_id')\
            .values(
                'id', 
                'college_id__college',  # Get college name
                'students_name', 
                'user_name', 
                'year',
                'department_id__department',  # Get department name
            ).distinct()

        # Convert the queryset to a list to manipulate individual fields
        candidatelist = list(candidatelist)

        # Check and replace '0.0' values with empty strings
        for candidate in candidatelist:
            if candidate['students_name'] == '0.0':
                candidate['students_name'] = ''
        
        return Response(candidatelist)
    except Exception as e:
        return Response({'error': str(e)}, status=500)


@api_view(['GET'])
def get_nondb_candidate(request):
    # Create a subquery to fetch dtm_upload from candidate_master based on user_name
    subquery_dtm_upload = candidate_master.objects.filter(
        user_name=OuterRef('user__username')  # Assuming user_name is related to User model's username
    ).values('dtm_upload')[:1]  # Limiting to one record since we expect one match per user_name

    # Filter user_profile entries where user_name is in the list from the candidate_master
    nondatabase_candidates = user_profile.objects.filter(
        user__username__in=candidate_master.objects.filter(is_database=False).values_list('user_name', flat=True)
    ).annotate(
        dtm_upload=Subquery(subquery_dtm_upload)
    ).values('user__username', 'user__password', 'dtm_upload').order_by('-id')  # Order by id in descending order

    # Convert the queryset to a list of dictionaries
    result = list(nondatabase_candidates)
    
    return JsonResponse(result, safe=False)



#-----------------------------------------------------------------------------#

from .forms import EligibleStudentListForm

@csrf_exempt
def update_eligible_student(request, job_id_value,round_of_interview_value):
    print(f"Received request for updating students with Job ID: {job_id_value}")
    print(request)
    # Get the round_of_interview from the query parameters
   # round_of_interview = request.GET.get('round_of_interview')
   
    print("round_of_interview",round_of_interview_value)
    if not round_of_interview_value:
        return HttpResponse("Error: round_of_interview query parameter is required.", status=400)

    try:
        # Retrieve all eligible student records based on job_id and round_of_interview
        eligible_students = eligible_student_list.objects.filter(job_id_id=job_id_value, round_of_interview=round_of_interview_value)
        print("eligible_students",eligible_students)
        if not eligible_students.exists():
            return HttpResponse("Error: No eligible students found for the given criteria.", status=404)

        print(f"Retrieved {eligible_students.count()} eligible students")
    except Exception as e:
        print(f"Error occurred: {e}")
        return HttpResponse(f"Error: {e}", status=404)

    if request.method == 'POST':
        # Loop through each eligible student and update their details
        for eligible_student in eligible_students:
            form = EligibleStudentListForm(request.POST, request.FILES, instance=eligible_student)
            if form.is_valid():
                # Only update announcement if provided in the request
                if 'announcement' in request.POST:
                    eligible_student.announcement = request.POST.get('announcement')  # Use the input from the form

                # Only update the announcement image if provided in the request
                if 'announcement_image' in request.FILES:
                    file = request.FILES['announcement_image']
                    eligible_student.announcement_image = file.read()

                # Save the updated student data
                form.save()
            else:
                # If any form is invalid, return an error response with form errors
                return HttpResponse(f"Form errors: {form.errors}", status=400)

        print("All students updated successfully")
       
        
    else:
        # Prepopulate the form with data from the first student in the queryset
        form = EligibleStudentListForm(instance=eligible_students.first())

    return render(request, 'update_eligible_student.html', {'form': form, 'eligible_students': eligible_students})


@api_view(['GET'])
def get_eligible_students_round(request):
    round_of_interview = request.query_params.get('round_of_interview', None)

    if not round_of_interview:
        return Response({'error': 'round_of_interview is required'}, status=400)

    try:
        # Filter the eligible students based on round_of_interview
        eligible_students_data = eligible_student_list.objects.filter(round_of_interview=round_of_interview).values(
            'id',
            'students_id__id',  # Access related object's field (candidate_master id)
                 
        )
        
        # Convert the QuerySet to a list of dictionaries
        eligible_students_list = list(eligible_students_data)

        return Response(eligible_students_list)

    except Exception as e:
        print('Exception:', str(e))
        return Response({'error': str(e)}, status=500)

#---------------------Update Upload Shortlisted Cadidate----------------#

class UpdateEligibleStudentListView_REG_No(APIView):
    def post(self, request, format=None):
        if 'file' not in request.FILES:
            print("No file uploaded")
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        if not file.name.endswith('.xlsx'):
            print("File is not in Excel format")
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            df = pd.read_excel(file)
            print("Excel file read successfully")
        except Exception as e:
            print(f"Error reading Excel file: {str(e)}")
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        round_of_interview = request.data.get('round_of_interview')
        job_id_value = request.data.get('job_id')
        print('round_of_interview: ', round_of_interview)
        print('jobId: ', job_id_value)

        if not round_of_interview:
            return Response({'error': 'round_of_interview parameter is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        if not job_id_value:
            return Response({'error': 'jobId parameter is required'}, status=status.HTTP_400_BAD_REQUEST)

        for _, row in df.iterrows():
            reg_no = row.get('Reg No')
            if not reg_no:
                continue

            eligible_ids = eligible_student_list.objects.filter(
                students_id__registration_number=reg_no,
                job_id=job_id_value
            )
            if eligible_ids.exists():
                eligible_ids.update(round_of_interview=round_of_interview)
            else:
                # Handle case where no eligible student list record exists
                print(f"No eligible student list record found for candidate with reg_no {reg_no} and job_id {job_id_value}")
                continue

        return Response({'success': 'Eligible student list updated successfully'}, status=status.HTTP_200_OK)



class UpdateEligibleStudentListView(APIView):
    def post(self, request, format=None):
        print("Received POST request")

        if 'file' not in request.FILES:
            print("No file uploaded")
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)

        file = request.FILES['file']
        print(f"Uploaded file name: {file.name}")

        if not file.name.endswith('.xlsx'):
            print("File is not in Excel format")
            return Response({'error': 'File is not in Excel format'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            df = pd.read_excel(file)
            print("Excel file read successfully")
        except Exception as e:
            print(f"Error reading Excel file: {str(e)}")
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

        round_of_interview = request.data.get('round_of_interview')
        job_id_value = request.data.get('job_id')

        print(f"Round of Interview: {round_of_interview}")
        print(f"Job ID: {job_id_value}")

        if not round_of_interview:
            print("round_of_interview parameter is missing")
            return Response({'error': 'round_of_interview parameter is required'}, status=status.HTTP_400_BAD_REQUEST)

        if not job_id_value:
            print("jobId parameter is missing")
            return Response({'error': 'jobId parameter is required'}, status=status.HTTP_400_BAD_REQUEST)

        errors = []
        try:
            with transaction.atomic():
                for index, row in df.iterrows():
                    reg_no = row.get('Reg No**')

                    # Check if reg_no is NaN or an empty string
                    if pd.isna(reg_no) or str(reg_no).strip() == '':
                        error_message = f"Row {index + 2}: Reg No is empty"
                        print(error_message)
                        errors.append(error_message)
                        raise ValueError(error_message)  # Raise an exception to trigger rollback

                    dept = row.get('Department')
                    stud_name = row.get('Student Name')
                    email = row.get('Email')
                    mobile = row.get('Mobile No')

                    print(f"Processing row {index + 2}: Reg No: {reg_no}, Department: {dept}")

                    eligible_ids = eligible_student_list.objects.filter(
                        students_id__registration_number=reg_no,
                        job_id=job_id_value
                    )

                    if eligible_ids.exists():
                        eligible_ids.update(round_of_interview=round_of_interview)
                        print(f"Updated round_of_interview for Reg No: {reg_no}")

        except ValueError as ve:
            print(f"ValueError occurred: {str(ve)}")
            return Response({'error': str(ve)}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            print(f"An unexpected error occurred: {str(e)}")
            return Response({'error': f"An unexpected error occurred: {str(e)}"}, status=status.HTTP_400_BAD_REQUEST)

        if errors:
            print(f"Errors encountered: {errors}")
            return Response({'error': errors}, status=status.HTTP_400_BAD_REQUEST)

        print("Eligible student list updated successfully")
        return Response({'success': 'Eligible student list updated successfully'}, status=status.HTTP_200_OK)








@api_view(['PUT', 'PATCH'])
def update_is_accept(request, pk):
    try:
        print("Entering Function..")
        candidates=eligible_student_list.objects.get(id=pk)
        ##logger.info('Fetching test-candidate-map ')

        print("accept_candidates: ",candidates)
    except eligible_student_list.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    candidates.is_accept = True
    candidates.round_of_interview = 'Registered'
    candidates.save()
    ##logger.info('test-candidates updated')
    print("candidates_accepted: ",candidates)

    return JsonResponse("tests_candidates_map 'deleted' field updated successfully", safe=False)

@api_view(['PUT', 'PATCH'])
def update_is_decline(request, pk):
    try:
        print("Entering Function..")
        candidates=eligible_student_list.objects.get(id=pk)
        ##logger.info('Fetching test-candidate-map ')

        print("accept_candidates: ",candidates)
    except eligible_student_list.DoesNotExist:
        return JsonResponse("tests not found", status=404)

    # Update the 'deleted' field to 1 instead of deleting the object
    candidates.is_accept = False
    candidates.save()
    ##logger.info('test-candidates updated')
    print("candidates_accepted: ",candidates)

    return JsonResponse("tests_candidates_map 'deleted' field updated successfully", safe=False)




def get_questions_IO_filter_mcq(request, question_id):
    # Create a cache key based on the question_id
    cache_key = f'questions_IO_filter_mcq_{question_id}'
    
    # Try to get the data from the cache
    question_data = cache.get(cache_key)
    
    if not question_data:
        # Fetch the required data using select_related and values
        questionset = question_master.objects.filter(
            deleted=0, 
            question_name_id=question_id
        ).select_related('question_name_id').values(
            'id',
            'question_name_id__id',
            'question_name_id__question_paper_name',
            'question_text',
            'question_image_data',
            'option_a_image_data',
            'option_b_image_data',
            'option_c_image_data',
            'option_d_image_data',
            'option_a',
            'option_b',
            'option_c',
            'option_d',
            'mark',
            'explain_answer',
            'answer',
            'input_format'
        )

        # Process the data
        question_data = []
        for question in questionset:
            question_image_data = base64.b64encode(question['question_image_data']).decode('utf-8') if question['question_image_data'] else None
            option_a_image_data = base64.b64encode(question['option_a_image_data']).decode('utf-8') if question['option_a_image_data'] else None
            option_b_image_data = base64.b64encode(question['option_b_image_data']).decode('utf-8') if question['option_b_image_data'] else None
            option_c_image_data = base64.b64encode(question['option_c_image_data']).decode('utf-8') if question['option_c_image_data'] else None
            option_d_image_data = base64.b64encode(question['option_d_image_data']).decode('utf-8') if question['option_d_image_data'] else None

            question_data.append({
                'id': question['id'],
                'question_name_id': question['question_name_id__id'],
                'question_paper_name': question['question_name_id__question_paper_name'],
                'question_text': question['question_text'],
                'question_image_data': question_image_data,
                'option_a_image_data': option_a_image_data,
                'option_b_image_data': option_b_image_data,
                'option_c_image_data': option_c_image_data,
                'option_d_image_data': option_d_image_data,
                'option_a': question['option_a'],
                'option_b': question['option_b'],
                'option_c': question['option_c'],
                'option_d': question['option_d'],
                'mark': question['mark'],
                'explain_answer': question['explain_answer'],
                'answer': question['answer'],
                'input_format': question['input_format'],
            })

        # Shuffle the questions
        random.shuffle(question_data)

        # Store the processed data in the cache with a timeout (e.g., 1 hour)
        cache.set(cache_key, question_data, timeout=3600)
    
    # Return the cached or freshly processed data
    return JsonResponse(question_data, safe=False)


def get_questions_IO_filter_code(request, question_id):
    cache_key = f'question_data_code_{question_id}'
    logger.error("Catch Key...: %s", cache_key)

    question_data = cache.get(cache_key)
    logger.error("-----Data From Cache----")
    
    if not question_data:
        questionset = (
            question_master.objects
            .filter(deleted=0, question_name_id=question_id)
            .select_related('question_name_id')
            .values(
                'id',
                'question_text',
                'mark',
                'input_format',
                'explain_answer',
                'answer'
            )
        )
        question_data = list(questionset)
        cache.set(cache_key, question_data, timeout=3600)
        logger.error("-----Data From Database----")
    random.shuffle(question_data)
    return JsonResponse(question_data, safe=False)



@api_view(['GET'])
def get_tests_ansnwers_filter(request):
    # Directly apply filters
    student_id = request.GET.get('student_id')
    test_name = request.GET.get('test_name')
    question_id = request.GET.get('question_id')

    # Fetch data using ORM with related fields and filters
    test_candidate_answer_data = tests_candidates_answers.objects.annotate(
        answer_length=Length('answer')
    ).filter(
        deleted=0, 
        student_id=student_id, 
        test_name=test_name, 
        question_id=question_id, 
        submission_compile_code__isnull=True,
        answer_length__gt=0
    ).select_related('student_id', 'question_id').order_by('-id').values(
        'result',
        'question_id'
    )

    # Rename the fields to match your desired structure
    test_candidate_answer_data = [
        {
            'result': testans['result'],
            'question_id': testans['question_id']
        }
        for testans in test_candidate_answer_data
    ]

    return Response(test_candidate_answer_data)


@api_view(['GET'])
def get_total_marks(request, student_id_value, test_name_value):
    total_marks = tests_candidates_answers.objects.filter(
        student_id=student_id_value,
        test_name=test_name_value,
    ).aggregate(total_marks=Sum('result'))['total_marks']  # Sum of the 'result' field
    print('total_marks: ', total_marks)

    return Response({'total_marks': total_marks if total_marks is not None else 0})




@api_view(['DELETE'])
def delete_student_answers(request, student_id):
    try:
        # Query to delete the records for the given student_id
        count, _ = tests_candidates_answers.objects.filter(student_id=student_id).delete()

        if count > 0:
            return Response({'message': f'{count} records deleted.'}, status=status.HTTP_204_NO_CONTENT)
        else:
            return Response({'error': 'No records found for the given student_id.'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



@api_view(['GET'])
def get_trainer_by_userName(request, userName):
    try:
        # Fetch trainers with the given username and not deleted
        trainer = trainer_master.objects.filter(user_name=userName, deleted=0)
        # Serialize the data
        serializer = trainerSerializer(trainer, many=True)
        return Response(serializer.data)
    except Exception as e:
        return Response({'error': str(e)}, status=500)
		


@api_view(['GET'])
def get_distinct_dtm_uploads(request, college_id):
    try:
        # Filter the candidate_master table based on college_id and is_database = False
        distinct_dtm_uploads = candidate_master.objects.filter(
            Q(college_id=college_id) & Q(is_database=False)
        ).order_by('-dtm_upload').values_list('dtm_upload', flat=True).distinct()

        # Return the distinct dtm_upload values
        return Response({
            'distinct_dtm_uploads': list(distinct_dtm_uploads)
        }, status=status.HTTP_200_OK)

    except Exception as e:
        # Handle any errors
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)



@api_view(['GET'])
def get_tests_candidates_camera(request,id):
    try:
        # Fetch all test-candidate mappings with related data in one query
        tests_candidates = tests_candidates_map.objects.filter(deleted=0,id=id).values(
            'id', 
          'is_camera_on',
        )

        # Convert the queryset into a list of dictionaries
        test_candidate_map_data = list(tests_candidates)

        return Response(test_candidate_map_data)
    except Exception as e:
        return Response({'error': str(e)}, status=500)



@api_view(['GET'])
def get_eligible_student_count_rounds(request):
    # Get the round_of_interview and job_id from the request parameters
    round_of_interview = request.query_params.get('round_of_interview')
    job_name = request.query_params.get('job_name')

    if not round_of_interview or not job_name:
        return Response({"error": "round_of_interview and job_id are required parameters."}, status=status.HTTP_400_BAD_REQUEST)

    # Query the eligible_student_list model
    count = eligible_student_list.objects.filter(round_of_interview=round_of_interview, job_id__company_name=job_name).count()

    return Response({"count": count}, status=status.HTTP_200_OK)


@api_view(['GET'])
def get_candidate_offers_count(request, college_id):
    try:
        # Filter by college_id and count the number of offers
        count_offers = candidate_master.objects.filter(college_id=college_id).count()

        # Return the count in the response
        return Response({"total_offers": count_offers}, status=status.HTTP_200_OK)

    except candidate_master.DoesNotExist:
        # Handle the case where the college_id does not exist
        return Response({"error": "College ID not found."}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        # Handle any other exceptions
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)




@api_view(['GET'])
def get_eligible_students_job_id(request, job_id, round_of_interview):
    try:
        # Retrieve the relevant student information directly
        students_data = eligible_student_list.objects.filter(
            job_id=job_id,
            round_of_interview=round_of_interview
        ).select_related('students_id').values(
            'students_id__students_name',
            'students_id__registration_number',
            'students_id__department_id__department',
            'students_id__email_id',
            'students_id__mobile_number'
        )

        # Return the data as a response
        return Response(list(students_data), status=status.HTTP_200_OK)

    except Exception as e:
        # Handle any exceptions
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)




@api_view(['GET'])
def get_round_of_interview_count(request):
    try:
        
        college_id_value = request.query_params.get('college_id', None)
        company_name_value = request.query_params.get('company_name', None)

        # Group by 'round_of_interview' and count the number of records for each round status
        round_of_interview_counts = eligible_student_list.objects.filter(
            students_id__college_id=college_id_value, 
            job_id__company_name=company_name_value
        ).values(
            'round_of_interview'
        ).annotate(
            student_count=Count('id')  # Count the number of students for each 'round_of_interview' status
        ).order_by('round_of_interview')

        # Convert the QuerySet to a list of dictionaries
        round_of_interview_list = list(round_of_interview_counts)

        return Response(round_of_interview_list)

    except Exception as e:
        print(f"Error: {str(e)}")
        return Response({'error': str(e)}, status=500)


@api_view(['GET'])
def get_registered_count_by_company(request):
    try:
        college_id_value = request.query_params.get('college_id', None)
        department_name_value = request.query_params.get('department_name', None)

        # Group by company name (from job_id) and get the count of registered students (is_accept=True)
        company_student_counts = eligible_student_list.objects.filter(
            students_id__college_id=college_id_value,
            students_id__department_id__department=department_name_value,
            is_accept=True  # Only count registered students
        ).values(
            'job_id__company_name'  # Group by company name from job_id
        ).annotate(
            registered_count=Count('id')  # Count number of students for each company
        ).order_by('job_id__company_name')  # Optional: Order by company name

        # Convert the QuerySet to a list of dictionaries
        company_student_count_list = list(company_student_counts)

        return Response(company_student_count_list)

    except Exception as e:
        print(f"Error: {str(e)}")
        return Response({'error': str(e)}, status=500)














